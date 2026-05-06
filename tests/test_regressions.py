from adeu import AcceptChange, RedlineEngine
from adeu import ModifyText, RedlineEngine
from adeu import ModifyText, RedlineEngine, ReplyComment
from adeu.ingest import extract_text_from_stream
from adeu.mcp_components.tools.document import PROCESS_BATCH_OPERATIONS_DESC
from adeu.mcp_components.tools.sanitize import sanitize_docx
from adeu.models import ModifyText
from adeu.redline.engine import BatchValidationError
from adeu.redline.engine import BatchValidationError, RedlineEngine
from adeu.redline.engine import RedlineEngine
from adeu.sanitize.report import SanitizeReport
from adeu.sanitize.transforms import remove_all_comments
from adeu.utils.docx import _coalesce_runs_in_paragraph
from adeu.utils.docx import get_visible_runs
from docx import Document
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.part import XmlPart
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from io import BytesIO
from unittest.mock import AsyncMock, MagicMock, patch
from unittest.mock import MagicMock
import asyncio
import docx
import io
import pytest
import subprocess
import sys
import time
import zipfile


def test_batch_engine_accept_fake_id_attribute_error():
    """
    Reproduces a bug where accepting/rejecting a non-existent change ID
    crashes with a Pydantic AttributeError on `_match_start_index` rather
    than a clean missing target / validation error.
    """
    doc = docx.Document()
    doc.add_paragraph("Test")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="QA Bot")

    with pytest.raises(BatchValidationError):
        engine.process_batch([AcceptChange(target_id="Chg:999")])

def test_batch_engine_deletes_special_content():
    """
    Reproduces a bug where modifying text in a paragraph that contains
    special content (like images/drawings) deletes the drawing element,
    violating Safety Constraint #2.
    """
    # 1. Setup a doc with a paragraph containing text and an image.
    d = docx.Document()
    p = d.add_paragraph("Normal paragraph")
    r = p.add_run(" ")
    # Add a drawing element (use a dummy file or we can just inject XML directly to avoid file dependencies)
    from docx.oxml import parse_xml

    drawing_xml = parse_xml('<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    r._r.append(drawing_xml)

    stream = io.BytesIO()
    d.save(stream)
    stream.seek(0)

    # Verify the drawing is there before
    d_before = docx.Document(stream)
    assert (
        len(
            d_before.paragraphs[0]._element.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
            )
        )
        == 1
    )
    stream.seek(0)

    # 2. Modify the text preceding the drawing.
    engine = RedlineEngine(stream, author="QA")
    engine.apply_edits([ModifyText(target_text="Normal paragraph ", new_text="Modified ")])

    # 3. Verify the drawing was destroyed.
    out_stream = engine.save_to_stream()
    out_stream.seek(0)
    d_after = docx.Document(out_stream)

    drawings_after = d_after.paragraphs[0]._element.findall(
        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
    )
    assert len(drawings_after) == 1, "BUG: The drawing element was destroyed by the text modification!"

def test_batch_engine_heading_depth_enforcement():
    """
    Reproduces a bug where the Batch Engine fails to enforce the maximum
    heading depth limit of 6, silently allowing `#` * 7 to pass through
    and corrupt the Markdown to OOXML mapping.
    """
    doc = docx.Document()
    doc.add_paragraph("Target Text")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="QA Bot")

    with pytest.raises(BatchValidationError):
        engine.process_batch([ModifyText(target_text="Target Text", new_text="####### Heading 7")])

def test_batch_engine_reply_to_fake_comment():
    """
    Reproduces a bug where replying to a non-existent comment
    raises an unhandled AttributeError instead of a clean validation error.
    """
    doc = docx.Document()
    doc.add_paragraph("Target Text")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="QA Bot")

    with pytest.raises(BatchValidationError):
        engine.process_batch([ReplyComment(target_id="Com:999", text="Hello")])

def test_batch_engine_formatting_removal_fails():
    """
    Reproduces a bug where the Batch Engine fails to remove formatting.
    If the target_text has formatting (like **text**), but the new_text
    is plain ('text'), the inserted run inherits the old formatting
    instead of dropping it.
    """
    d = docx.Document()
    p = d.add_paragraph("Body ")
    r = p.add_run("text")
    r.bold = True
    p.add_run(" here.")
    stream = io.BytesIO()
    d.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="QA")
    engine.apply_edits([ModifyText(target_text="Body **text** here.", new_text="Body text here.")])

    out_stream = engine.save_to_stream()
    out_stream.seek(0)
    d2 = docx.Document(out_stream)

    # We look for the inserted run (w:ins) in the XML.
    xml = d2.paragraphs[0]._element.xml
    assert "<w:b/>" in xml.split("<w:ins")[1], "Bug fixed? The bold tag was properly stripped from the insertion."

def test_batch_engine_corrupts_multipara_insertion():
    """
    Reproduces a bug where replacing a space with multiple paragraphs
    (e.g., 'Cell Text' -> 'Cell\n\nNew\n\nText') causes the engine to
    swallow the trailing text and corrupt the OOXML structure, generating
    out-of-order text fragments.
    """
    d = docx.Document()
    d.add_paragraph("Cell Text")
    stream = io.BytesIO()
    d.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="QA")
    engine.process_batch([ModifyText(target_text="Cell Text", new_text="Cell\n\nNew\n\nText")])

    # We must accept all revisions here to properly inspect text structure
    # because docx.Paragraph.text natively ignores <w:ins> tag contents
    engine.accept_all_revisions()
    out_stream = engine.save_to_stream()
    out_stream.seek(0)
    d2 = docx.Document(out_stream)

    texts = [p.text for p in d2.paragraphs]
    assert texts == ["Cell", "New", "Text"], f"BUG: Text mapping corruption! Got {texts}"

def test_batch_engine_swallows_trailing_deletions():
    """
    Reproduces Bug 9: When replacing text that includes a deletion at the end
    of the target, AND a multi-paragraph insertion, the trailing deletion is
    silently swallowed (ignored) and left in the document.
    """
    d = docx.Document()
    d.add_paragraph("Party A; and")
    stream = io.BytesIO()
    d.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="QA")
    # We remove `; and` and insert a new paragraph.
    engine.apply_edits([ModifyText(target_text="Party A; and", new_text="Party A\n\nParty B")])

    out_stream = engine.save_to_stream()
    out_stream.seek(0)
    d2 = docx.Document(out_stream)

    # The first paragraph should have '; and' deleted. Let's check the XML.
    xml = d2.paragraphs[0]._element.xml
    assert "<w:del" in xml, "BUG: The trailing deletion of '; and' was swallowed entirely!"

def test_diff_engine_ignores_formatting_changes(tmp_path):
    """
    Reproduces a bug where `diff_docx_files` with compare_clean=True
    fails to detect formatting-only changes (e.g., Bold to Italic).
    """
    import docx

    d1_path = tmp_path / "format_1_pytest.docx"
    d2_path = tmp_path / "format_2_pytest.docx"

    d1 = docx.Document()
    p1 = d1.add_paragraph()
    r1 = p1.add_run("Silent Change")
    r1.bold = True
    d1.save(str(d1_path))

    d2 = docx.Document()
    p2 = d2.add_paragraph()
    r2 = p2.add_run("Silent Change")
    r2.italic = True
    d2.save(str(d2_path))

    import subprocess

    diff_output = subprocess.check_output(["uv", "run", "adeu", "diff", str(d1_path), str(d2_path)]).decode("utf-8")

    assert "[~]" in diff_output or "Found 1 changes:" in diff_output, (
        "Diff engine completely missed the bold -> italic formatting change."
    )

def test_diff_engine_splits_markdown_tokens(tmp_path):
    """
    Reproduces Bug 8: Diff engine breaks mid-markdown token.
    """
    d1_path = tmp_path / "diff_markdown_1.docx"
    d2_path = tmp_path / "diff_markdown_2.docx"

    d1 = docx.Document()
    d1.add_paragraph("ons)**  **(In millions)**          Dear")
    d1.save(str(d1_path))

    d2 = docx.Document()
    d2.add_paragraph("ons)**  **(In millions)**\n\n\n\n\n\n\n\n\nDear shareholders:")
    d2.save(str(d2_path))

    import subprocess

    diff_output = subprocess.check_output(["uv", "run", "adeu", "diff", str(d1_path), str(d2_path)]).decode("utf-8")

    # The bug produces `- *(In millions)**` which is invalid markdown.
    assert "- *(In millions)**" not in diff_output, "BUG: Diff engine split a markdown token!"

def test_extract_handles_tabs_and_breaks():
    """
    REPRO: Documents using tabs for spacing (e.g. 'Word<tab>Word')
    should extract as 'Word Word', not 'WordWord'.
    """
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run()

    # Manually inject w:t, w:tab, w:t to simulate "Word<tab>Word"
    # python-docx run.text doesn't support adding tabs easily, so we manipulate XML
    t1 = OxmlElement("w:t")
    t1.text = "Word"
    run._element.append(t1)

    tab = OxmlElement("w:tab")
    run._element.append(tab)

    t2 = OxmlElement("w:t")
    t2.text = "One"
    run._element.append(t2)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    text = extract_text_from_stream(stream)

    # EXPECTATION: "Word One", not "WordOne"
    assert "Word One" in text
    assert "WordOne" not in text

def test_heuristic_header_detection():
    """
    REPRO: 'Normal' style paragraphs that are BOLD and ALL-CAPS should
    be detected as headers (##) to give the LLM structural context.

    UPDATED: Now includes inline markdown **markers** for bold text.
    """
    doc = Document()
    p = doc.add_paragraph("LIABILITY CAP")
    p.style = doc.styles["Normal"]

    # Make it Bold
    p.runs[0].bold = True

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    text = extract_text_from_stream(stream)

    # EXPECTATION: "## **LIABILITY CAP**"
    # The header heuristic adds ##. The inline processor adds **...**.
    assert "## **LIABILITY CAP**" in text

def test_disk_engine_drops_comment_on_multipara_insertion():
    """
    Reproduces a bug where the Disk Engine (RedlineEngine) drops
    the `comment` parameter if the `new_text` contains a multi-paragraph
    insertion (\n\n).
    """
    from docx import Document

    doc = Document()
    doc.add_paragraph("This is paragraph 1. It is short.")
    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit = ModifyText(
        target_text="This is paragraph 1. It is short.",
        new_text="This is paragraph 1.\n\nIt is short and now has a second paragraph.",
        comment="This comment will be dropped by the disk engine!",
    )

    engine = RedlineEngine(stream, author="QA Disk Bot")
    engine.process_batch([edit])

    out_stream = engine.save_to_stream()

    import zipfile

    out_stream.seek(0)
    z = zipfile.ZipFile(out_stream)

    # Bug manifestation: word/comments.xml might not even exist!
    has_comments_part = any("word/comments" in f and f.endswith(".xml") for f in z.namelist())
    assert has_comments_part, "word/comments.xml was not created, comment was completely dropped."

    comments_xml = ""
    for f in z.namelist():
        if "word/comments" in f and f.endswith(".xml"):
            comments_xml = z.read(f).decode("utf-8")
            break

    assert "This comment will be dropped" in comments_xml, "Comment text not found in comments.xml"

def test_extractor_ignores_tracked_formatting(tmp_path):
    """
    Reproduces Bug 10: Extractor fails to generate CriticMarkup for <w:rPrChange>.
    """
    docx_path = tmp_path / "test_fmt_track_pytest.docx"
    d = docx.Document()
    p = d.add_paragraph("Test")
    r = p.runs[0]
    r.bold = True
    r._r.append(
        parse_xml(
            '<w:rPrChange xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'w:id="1" w:author="QA" w:date="2026-04-30T00:00:00Z"><w:rPr><w:b/></w:rPr></w:rPrChange>'
        )
    )
    d.save(str(docx_path))

    output = subprocess.check_output(["uv", "run", "adeu", "extract", str(docx_path)]).decode("utf-8")

    # Expected: {==**Test**==}{>>[Chg:1] QA<<} or {++**Test**++}
    # Actual: **Test** (no tracking markup)
    assert "{=" in output or "{+" in output or "{>" in output, (
        "BUG: Tracked formatting change was completely ignored by the extractor."
    )

def test_suppress_inherited_removes_complex_scripts():
    """
    Bug #12: Suppressing inherited formatting must strip <w:bCs/> and <w:iCs/>
    to prevent visual styling mismatches in modern Word.
    """
    r_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rPr.append(OxmlElement("w:b"))
    rPr.append(OxmlElement("w:bCs"))
    rPr.append(OxmlElement("w:i"))
    rPr.append(OxmlElement("w:iCs"))
    r_elem.append(rPr)

    engine = RedlineEngine.__new__(RedlineEngine)
    engine._apply_run_props(r_elem, props={}, suppress_inherited=True)

    assert len(rPr.findall(qn("w:b"))) == 0, "w:b should be stripped"
    assert len(rPr.findall(qn("w:i"))) == 0, "w:i should be stripped"
    assert len(rPr.findall(qn("w:bCs"))) == 0, "w:bCs should be stripped (Bug #12)"
    assert len(rPr.findall(qn("w:iCs"))) == 0, "w:iCs should be stripped (Bug #12)"

def test_process_document_batch_docstring_mentions_attribution():
    """
    OBS-01 Update: The process_document_batch docstring must explicitly state
    that author_name is correctly used for attribution, even in Live Word.
    """
    desc = PROCESS_BATCH_OPERATIONS_DESC

    assert "spoofed" not in desc.lower(), (
        "Docstring should no longer claim identities cannot be spoofed, as this was disproved."
    )
    assert "used for attribution" in desc.lower(), (
        "Docstring must explicitly state that author_name is used for attribution."
    )

@pytest.mark.skipif(sys.platform != "win32", reason="Live Word is Windows only")
def test_live_word_trims_common_context():
    """
    Bug #8: Live Word COM engine replaces the exact target substring wholesale,
    but it must trim common context like the disk engine does.
    """
    import adeu.mcp_components.tools.live_word as lw

    doc_mock = MagicMock()
    doc_mock.Content.Text = "The quick brown fox jumps."
    doc_mock.Revisions.Count = 0
    doc_mock.Content.End = 100

    rng_mock = MagicMock()
    rng_mock.Find.Execute.return_value = True
    rng_mock.Start = 4  # Index of 'quick'
    rng_mock.End = 23
    doc_mock.Range.return_value = rng_mock

    app_mock = MagicMock()
    app_mock.ActiveDocument = doc_mock

    import win32com.client

    original_get = getattr(win32com.client, "GetActiveObject", None)
    win32com.client.GetActiveObject = lambda name: app_mock
    try:
        changes = [ModifyText(target_text="The quick brown fox", new_text="The fast brown fox", comment="Speed up")]
        stats = lw._process_active_word_batch_core(changes, "Reviewer")

        assert stats["failed"] == 0, f"Batch failed with: {stats.get('skipped_details', [])}"

        # The real _apply_com_replacement will set target_rng.Text
        assert rng_mock.Text == "fast", f"Expected trimmed 'fast', got '{rng_mock.Text}'"
    finally:
        if original_get:
            win32com.client.GetActiveObject = original_get

def test_markdown_headers_leak_into_docx():
    doc = Document()
    p1 = doc.add_paragraph("Section 1.")
    p1.style = "Heading 1"
    p2 = doc.add_paragraph("Section 2.")
    p2.style = "Heading 1"

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    target_text = "# Section 1.\n\n# Section 2."
    new_text = "# Section 1.\n\n# New Section.\n\n# Section 2."

    edit = ModifyText(target_text=target_text, new_text=new_text)

    engine = RedlineEngine(stream)

    # Unit Test the Helper directly
    clean, style = engine._parse_markdown_style("# New Section.")
    print(f"DEBUG: Parse '# New Section.' -> '{clean}', '{style}'")
    assert clean == "New Section."
    assert style == "Heading 1"

    engine.apply_edits([edit])

    result_stream = engine.save_to_stream()
    doc_result = Document(result_stream)

    print("\n--- Result XML ---")
    print(doc_result.element.xml)
    print("------------------\n")

    assert len(doc_result.paragraphs) == 3
    p2_text = "".join(r.text for r in get_visible_runs(doc_result.paragraphs[1]))
    assert "New Section." in p2_text

def test_paragraph_merge_on_newline_deletion():
    """
    Test Case: When replacing text that spans a paragraph boundary (`\\n\\n`)
    with text that does not have a paragraph boundary, the engine should
    merge the two paragraphs into one.
    """
    doc = Document()
    doc.add_paragraph("Paragraph 1 end.")
    doc.add_paragraph("Paragraph 2 start.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit = ModifyText(target_text="1 end.\n\nParagraph 2", new_text="1 end. Paragraph 2")

    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])
    assert applied == 1, "Edit should be applied"

    result_stream = engine.save_to_stream()
    doc_result = Document(result_stream)

    # Initially there were 2 paragraphs. They should be merged into 1.
    assert len(doc_result.paragraphs) == 1, "Paragraphs were not merged"

    visible_text = "".join(r.text for r in get_visible_runs(doc_result.paragraphs[0]))
    assert "Paragraph 1 end. Paragraph 2 start." in visible_text

def test_run_coalescing_ooxml_compliance():
    """
    Bug #5: _coalesce_runs_in_paragraph should concatenate <w:t> texts
    rather than appending multiple <w:t> siblings into a single <w:r>.
    """
    p_elem = OxmlElement("w:p")

    # Run 1: "Con"
    r1 = OxmlElement("w:r")
    t1 = OxmlElement("w:t")
    t1.text = "Con"
    r1.append(t1)
    p_elem.append(r1)

    # Run 2: "tract"
    r2 = OxmlElement("w:r")
    t2 = OxmlElement("w:t")
    t2.text = "tract"
    r2.append(t2)
    p_elem.append(r2)

    p = Paragraph(p_elem, None)

    _coalesce_runs_in_paragraph(p)

    runs = p_elem.findall(qn("w:r"))
    assert len(runs) == 1, "Should be coalesced into 1 run"

    t_nodes = runs[0].findall(qn("w:t"))
    assert len(t_nodes) == 1, f"Expected 1 <w:t> node, got {len(t_nodes)}! Invalid OOXML."
    assert t_nodes[0].text == "Contract", "Text should be concatenated"

@pytest.fixture
def anyio_backend():
    return "asyncio"

@pytest.mark.anyio
async def test_sanitize_docx_does_not_block_event_loop():
    """
    Bug #13 & #14: sanitize_docx is a heavy synchronous function that blocks the MCP
    event loop. It must be dispatched to a worker thread via asyncio.to_thread().
    """

    # Mock the synchronous sanitize core to sleep for 0.5 seconds
    def mock_sanitize_sync(*args, **kwargs):
        time.sleep(0.5)
        mock_result = MagicMock()
        mock_result.output_path = "out.docx"
        mock_result.status = "Success"
        mock_result.tracked_changes_found = 0
        mock_result.tracked_changes_accepted = 0
        mock_result.comments_removed = 0
        mock_result.comments_kept = 0
        mock_result.metadata_stripped = []
        mock_result.warnings = []
        mock_result.report_text = ""
        return mock_result

    ctx = MagicMock()
    ctx.info = AsyncMock()

    with (
        patch("adeu.mcp_components.tools.sanitize._sanitize", side_effect=mock_sanitize_sync, create=True),
        patch("adeu.sanitize.core.sanitize_docx", side_effect=mock_sanitize_sync),
        patch("pathlib.Path.exists", return_value=True),
    ):
        start = time.perf_counter()

        loop_ticks = []

        async def tick():
            for _ in range(5):
                loop_ticks.append(time.perf_counter())
                await asyncio.sleep(0.1)

        # Run concurrently
        await asyncio.gather(sanitize_docx("dummy.docx", ctx=ctx), tick())

        first_tick_delay = loop_ticks[0] - start
        assert first_tick_delay < 0.2, f"Event loop was blocked! First tick took {first_tick_delay}s"

def test_sanitize_purges_empty_comment_parts():
    """
    Reproduces a bug where `remove_all_comments` physically ejects empty comment parts
    (e.g., word/comments.xml) from the OPC package.
    According to AI_CONTEXT.md (Architectural Decisions #8), this is explicitly forbidden.
    """
    doc = docx.Document()
    doc.add_paragraph("Test paragraph.")

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)

    from adeu import ModifyText, RedlineEngine

    engine = RedlineEngine(stream, author="QA Bot")
    engine.process_batch([ModifyText(target_text="Test", new_text="Test", comment="This is a comment.")])

    commented_stream = engine.save_to_stream()
    commented_stream.seek(0)

    doc_obj = docx.Document(commented_stream)
    remove_all_comments(doc_obj)

    out_stream = BytesIO()
    doc_obj.save(out_stream)
    out_stream.seek(0)

    z = zipfile.ZipFile(out_stream)

    has_comments_part = any("word/comments" in f and f.endswith(".xml") for f in z.namelist())
    assert has_comments_part, (
        "Violation of Architectural Decision #8: Empty comment parts were violently ejected instead of left intact."
    )

def test_report_routes_comment_lines_to_structural():
    """
    Reproduces a bug where the SanitizeReport heuristic incorrectly routes
    comment detail lines (like `[Open] "..."`) into the STRUCTURAL section
    because the string doesn't contain the word "comment".
    """
    report = SanitizeReport("test.docx")

    # This is what transforms.remove_all_comments() yields
    lines = ["Comments removed: 1 (0 resolved, 1 open)", '  [Open] "Updated term." (Counterparty)']

    report.add_transform_lines(lines)

    # The first line correctly goes to removed_comment_lines
    assert "Comments removed: 1 (0 resolved, 1 open)" in report.removed_comment_lines

    # The detail line should be in removed_comment_lines, not structural_lines
    assert '  [Open] "Updated term." (Counterparty)' in report.removed_comment_lines, (
        "BUG: The comment line was not routed to removed_comment_lines."
    )
    assert '  [Open] "Updated term." (Counterparty)' not in report.structural_lines, (
        "BUG: The comment line incorrectly landed in structural_lines."
    )

def test_repro_split_insertion_coalescing():
    """
    Scenario: Word stores "rapala" and " " as two separate runs inside ONE w:ins tag.

    XML:
      <w:ins w:id="415">
        <w:r><w:t>rapala</w:t></w:r>
        <w:r><w:t> </w:t></w:r>
      </w:ins>

    Current Behavior: {++rapala++}{++ ++}
    Desired Behavior: {++rapala ++}
    """
    doc = Document()
    p = doc.add_paragraph()

    # Manually construct the split insertion XML
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "415")
    ins.set(qn("w:author"), "Mikko")

    # Run 1: "rapala"
    r1 = OxmlElement("w:r")
    t1 = OxmlElement("w:t")
    t1.text = "rapala"
    r1.append(t1)
    ins.append(r1)

    # Run 2: " "
    r2 = OxmlElement("w:r")
    t2 = OxmlElement("w:t")
    t2.text = " "
    t2.set(qn("xml:space"), "preserve")
    r2.append(t2)
    ins.append(r2)

    p._element.append(ins)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Act
    text = extract_text_from_stream(stream)

    # Assert
    # We want them merged
    assert "{++rapala ++}" in text
    # We do NOT want them split
    assert "{++rapala++}" not in text

def test_engine_init_does_not_strip_proof_err():
    """
    Bug #11: RedlineEngine should not perform global document normalization
    (which strips proofErr tags) during initialization or batch processing.
    It should operate in Surgical Mode.
    """
    doc = Document()
    p = doc.add_paragraph("Some text ")
    proof_err = OxmlElement("w:proofErr")
    proof_err.set(qn("w:type"), "spellStart")
    p._element.append(proof_err)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Verify it exists before engine initialization
    test_doc = Document(stream)
    assert len(test_doc.element.xpath("//w:proofErr")) == 1
    stream.seek(0)

    # Init engine (this currently calls normalize_docx and strips the tag)
    engine = RedlineEngine(stream)

    surviving = engine.doc.element.xpath("//w:proofErr")
    assert len(surviving) == 1, "proofErr was stripped! Engine init should not trigger global normalization."

def test_cross_table_cell_edit_validation_error():
    """
    Test Case: When an edit tries to replace text that spans across table cells
    (which ingest formats with ` | `), and the new_text does not have the same
    number of separators, the engine should raise a BatchValidationError.
    """
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "CellA"
    table.cell(0, 1).text = "CellB"

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Ingest output: "CellA | CellB"
    edit = ModifyText(target_text="CellA | CellB", new_text="CellC")
    engine = RedlineEngine(stream)

    with pytest.raises(BatchValidationError) as exc_info:
        engine.apply_edits([edit])

    assert "Target text spans 2 table cells, but replacement provides 1" in str(exc_info.value)

def test_repro_unbound_local_curr_ins_id_failure():
    """
    Scenario: The FIRST run in a paragraph is empty (no text).

    This caused an UnboundLocalError in mapper.py because the 'curr_ins_id'
    variable initialization was skipped inside the 'if full_seg_text:' block,
    but the variable was accessed later in the loop for lookahead logic.
    """
    doc = Document()
    p = doc.add_paragraph()

    # 1. Empty Run FIRST
    # CRITICAL: We must give it a distinct property (like font name) so it is NOT
    # coalesced with the next run during normalize_docx().
    # We avoid Bold/Italic because those generate Markdown text ("**"), masking the bug.
    r1 = p.add_run()
    r1.font.name = "Arial"

    # 2. Subsequent content
    r2 = p.add_run("Subsequent text")
    r2.font.name = "Times New Roman"

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)
    # We define an edit just to trigger the mapping engine
    edit = ModifyText(target_text="Subsequent", new_text="Changed")

    try:
        # This triggers engine._apply_single_edit_heuristic -> mapper._build_map
        engine.apply_edits([edit])
    except UnboundLocalError as e:
        pytest.fail(f"Regression: UnboundLocalError raised! The fix is missing. Details: {e}")
    except Exception as e:
        # If the fix works, we might get other errors (e.g. not found), but NOT UnboundLocal
        if "local variable 'curr_ins_id' referenced before assignment" in str(e):
            pytest.fail(f"Regression: UnboundLocalError raised (wrapped)! Details: {e}")

def test_dk1_cell_split_empty_cell_placement():
    """
    Tests that a cell-spanning edit properly targets an empty cell,
    even if the LLM omits the trailing space in the target string.
    """
    # 1. Setup: 1x2 table, Cell 1 is empty.
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Site Organization"

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)

    # 2. Apply Edit
    engine = RedlineEngine(stream, author="Test AI")
    edit = ModifyText(
        type="modify",
        target_text="Site Organization |",  # Notice the missing trailing space
        new_text="Site Organization | 10%",
    )

    stats = engine.process_batch([edit])

    # Verify the engine didn't skip it
    assert stats["edits_applied"] == 1
    assert stats["edits_skipped"] == 0

    out_stream = engine.save_to_stream()
    out_doc = Document(out_stream)

    # 3. Verify XML directly (The <w:ins> should be in Cell 1)
    cell_0_xml = out_doc.tables[0].cell(0, 0)._tc.xml
    cell_1_xml = out_doc.tables[0].cell(0, 1)._tc.xml

    assert "10%" not in cell_0_xml, "The text '10%' was wrongly inserted into the first cell!"
    assert "<w:ins" in cell_1_xml, "The insertion tag is missing from the second cell!"
    assert "10%" in cell_1_xml, "The text '10%' is missing from the second cell!"

    # 4. Verify Final Accepted State
    out_stream.seek(0)
    final_engine = RedlineEngine(out_stream)
    final_engine.accept_all_revisions()
    final_doc = Document(final_engine.save_to_stream())

    assert final_doc.tables[0].cell(0, 0).text == "Site Organization"
    assert final_doc.tables[0].cell(0, 1).text == "10%"

def test_markdown_numbered_list_leak():
    """
    Test Case: Injecting a new numbered list item using Markdown syntax (`\\n\\n1. Numbered Item`)
    should trigger a paragraph split and create a proper `<w:numPr>` list item,
    without leaking the literal `1. ` into the text run.
    """
    doc = Document()
    doc.add_paragraph("Reference text.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit = ModifyText(target_text="Reference text.", new_text="Reference text.\n\n1. Numbered Item")

    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])
    assert applied == 1, "Edit should be applied"

    result_stream = engine.save_to_stream()
    doc_result = Document(result_stream)

    assert len(doc_result.paragraphs) >= 2, "Expected a new paragraph to be created."

    p_new = doc_result.paragraphs[1]
    visible_text = "".join(r.text for r in get_visible_runs(p_new))

    # 1. The literal '1. ' should NOT be in the text.
    assert not visible_text.startswith("1. "), f"Numbered list Markdown leaked into text: '{visible_text}'"
    assert "Numbered Item" in visible_text

    # 2. It should have list properties (either numPr or a List style).
    pPr = p_new._element.pPr
    has_numPr = False
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            has_numPr = True

    is_list_style = p_new.style is not None and "List" in p_new.style.name

    assert has_numPr or is_list_style, "Paragraph is not formatted as a list."

def test_markdown_bullet_leak():
    """
    Test Case: VAL-OBS-8
    Injecting a new list item using Markdown syntax (`\\n\\n* New Bullet`)
    should trigger a paragraph split and create a proper `<w:numPr>` list item,
    without leaking the literal `*` into the text run.
    """
    doc = Document()
    doc.add_paragraph("Reference is also made to this Section 2 for further detail.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Simulating the change that failed in Phase 2
    edit = ModifyText(
        target_text="Reference is also made to this Section 2 for further detail.",
        new_text="Reference is also made to this Section 2 for further detail.\n\n* New Regression Test Bullet",
    )

    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])
    assert applied == 1, "Edit should be applied"

    result_stream = engine.save_to_stream()
    doc_result = Document(result_stream)

    # We expect 2 paragraphs.
    assert len(doc_result.paragraphs) >= 2, "Expected a new paragraph to be created."

    # Check the newly inserted paragraph
    p_new = doc_result.paragraphs[1]
    visible_text = "".join(r.text for r in get_visible_runs(p_new))

    print(f"DEBUG: Inserted paragraph visible text: '{visible_text}'")

    # 1. The literal '* ' should NOT be in the text.
    assert not visible_text.startswith("* "), f"Bullet Markdown leaked into text: '{visible_text}'"
    assert "New Regression Test Bullet" in visible_text

    # 2. It should have list properties (either numPr or a List style).
    pPr = p_new._element.pPr
    has_numPr = False
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            has_numPr = True

    is_list_style = p_new.style is not None and "List" in p_new.style.name

    assert has_numPr or is_list_style, "Paragraph is not formatted as a list."

def test_multiline_insert_does_not_create_nested_paragraphs():
    """
    Validates Issue Fix: Paragraph merge during wholesale replacement.
    Simulates anchoring a multiline insert onto a run that is already wrapped
    inside a w:ins tag to ensure we don't accidentally create <w:p><w:p> nested structures.
    """
    # 1. Setup Document
    doc = Document()
    doc.add_paragraph()
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # 2. Inject a mocked "Word-style" inline insertion
    engine = RedlineEngine(stream, author="TestAuthor")

    ins_tag = OxmlElement("w:ins")
    ins_tag.set(qn("w:id"), "99")

    r_tag = OxmlElement("w:r")
    t_tag = OxmlElement("w:t")
    t_tag.text = "10. Force Majeure"

    r_tag.append(t_tag)
    ins_tag.append(r_tag)

    p_element = engine.doc.paragraphs[0]._element
    p_element.append(ins_tag)

    # 3. Act: Apply a multiline text edit targeting the tracked text
    # This will trigger track_delete_run (splitting the ins) and track_insert
    edit = ModifyText(
        target_text="10. Force Majeure",
        new_text="10. Force Majeure\n\n11. Entire Agreement",
    )

    applied, _ = engine.apply_edits([edit])
    assert applied == 1, "Edit should apply successfully"

    # 4. Assert: Validate XML structure
    doc_xml = engine.doc.element.xml

    # Word strictly forbids a w:p element being a direct child of another w:p.
    # The new paragraph (11. Entire Agreement) must be a sibling to the original paragraph.
    nested_p_check = engine.doc.element.xpath("//w:p//w:p")
    assert len(nested_p_check) == 0, "FATAL: Nested <w:p> tags detected. Word will merge these!"

    # Ensure the new paragraph was successfully inserted
    assert "Entire Agreement" in doc_xml

def test_val_obs_new_7_paragraph_break_tracking():
    """
    VAL-OBS-NEW-7: When a multi-line string is inserted, the paragraph break
    itself must be tracked inside the <w:pPr><w:rPr> of the newly created paragraph.
    """
    doc = Document()
    doc.add_paragraph("First paragraph")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="TestAuthor")
    edit = ModifyText(target_text="paragraph", new_text="paragraph\n\nSecond paragraph")

    engine.apply_edits([edit])

    # Find the newly created paragraph
    p_elements = engine.doc.element.xpath("//w:p")
    assert len(p_elements) == 2, "Should have exactly 2 paragraphs"

    # Assert the break is tracked: <w:pPr><w:rPr><w:ins/></w:rPr></w:pPr>
    new_p = p_elements[1]
    ins_marker = new_p.xpath("./w:pPr/w:rPr/w:ins")
    assert len(ins_marker) > 0, "Paragraph break must be tracked with an <w:ins> inside <w:pPr>"

def test_external_relationship_does_not_crash_comments_manager():
    """
    Validates the fix for:
    'ValueError: target_part property on _Relationship is undefined when target mode is External'

    This bug occurred when `_ensure_xml_part` looped through `doc.part.rels` looking for
    internal part targets, but encountered an external link (which has no `target_part`).
    """
    # 1. Create a base document
    doc = Document()
    doc.add_paragraph("Visit our website.")
    stream1 = io.BytesIO()
    doc.save(stream1)
    stream1.seek(0)

    # 2. Add an edit with a comment to force the creation of a Comments XML part.
    # The bug only triggers when CommentsManager tries to upgrade an EXISTING part.
    engine = RedlineEngine(stream1)
    engine.apply_edits([ModifyText(target_text="website", new_text="portal", comment="Update wording")])
    stream2 = engine.save_to_stream()

    # 3. Reload the document and inject an external relationship (e.g., a Hyperlink).
    # This perfectly mimics real-world documents with web links.
    doc_with_comments = Document(stream2)
    doc_with_comments.part.relate_to("https://kempower.com", RT.HYPERLINK, is_external=True)

    stream3 = io.BytesIO()
    doc_with_comments.save(stream3)
    stream3.seek(0)

    # 4. Trigger the bug via ingest
    # extract_text_from_stream initializes CommentsManager, which finds the existing comments part,
    # calls _ensure_xml_part, which loops over rels, hitting the new external HYPERLINK.
    try:
        text = extract_text_from_stream(stream3)
        assert "portal" in text
    except ValueError as e:
        if "target mode is External" in str(e):
            pytest.fail(f"Regression: Failed to handle external relationship in _ensure_xml_part: {e}")
        raise e

def test_repro_comments_namespace_xml_syntax_error():
    """
    Regression test for a bug where _ensure_namespaces constructed a malformed
    <w:comments> tag (missing closing '>') when patching namespaces.
    """
    doc = Document()
    doc.add_paragraph("Test content")

    # 1. Manually create a 'defective' comments part (missing w14/w15/Ignorable)
    # This forces _ensure_namespaces to trigger its regex replacement logic.
    # We use a minimal XML that lacks the modern namespaces Adeu requires.
    comments_xml = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        b'<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">\n'
        b'  <w:comment w:id="1" w:author="Tester" w:date="2026-01-26T10:00:00Z">\n'
        b"    <w:p><w:r><w:t>Comment</w:t></w:r></w:p>\n"
        b"  </w:comment>\n"
        b"</w:comments>"
    )

    # Inject into package manually to bypass python-docx defaults
    partname = doc.part.package.next_partname("/word/comments%d.xml")
    part = XmlPart(partname, CT.WML_COMMENTS, parse_xml(comments_xml), doc.part.package)
    doc.part.package.parts.append(part)
    doc.part.relate_to(part, RT.COMMENTS)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # 2. Initialize Engine/Manager
    # This calls CommentsManager.__init__ -> _ensure_namespaces
    # Before the fix, this raised:
    # XMLSyntaxError: Couldn't find end of Start Tag comments line 1, line 2, column 1
    try:
        engine = RedlineEngine(stream)
    except Exception as e:
        pytest.fail(f"CommentsManager crashed on init during namespace patching: {e}")

    # 3. Verify the patch was applied correctly
    # The root element should now have the namespaces and valid syntax
    comments_part = engine.comments_manager.comments_part
    xml_str = comments_part.blob.decode("utf-8")

    # Check for presence of injected attributes
    assert 'xmlns:w15="' in xml_str
    assert 'mc:Ignorable="w14 w15 w16cid w16cex"' in xml_str

    # Ensure the tag was closed correctly (no syntax error implies this, but good to check string)
    # We expect the start tag to end with >
    # A rough check:
    assert 'mc:Ignorable="w14 w15 w16cid w16cex">' in xml_str or 'mc:Ignorable="w14 w15 w16cid w16cex" >' in xml_str


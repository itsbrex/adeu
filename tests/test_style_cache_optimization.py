# FILE: tests/test_style_cache_optimization.py
"""
Unit tests for the Fast OXML Style Cache optimization.

Covers
------
1. _get_style_cache: built once, cached on main_document_part, resolves
   w:basedOn inheritance for outline_level and bold.
2. get_paragraph_prefix: correct prefix from each of the three detection
   paths (explicit outline_level, named Heading/Title style, heuristic
   all-caps + bold fallback).
3. get_run_style_markers: suppresses ** for runs inside named-style
   Heading/Title paragraphs, preserves _italic_ in those paragraphs, and
   does NOT suppress ** in heuristic-promoted heading paragraphs (this is
   the documented legacy behaviour we confirmed in session — the heuristic
   path doesn't propagate to bold suppression).

These tests build minimal in-memory DOCX packages to give us full control
over the styles.xml inheritance chain and the run/paragraph property mix.
"""

import io
import zipfile

from docx import Document

from adeu.utils.docx import (
    _get_style_cache,
    get_paragraph_prefix,
    get_run_style_markers,
)

W_NS_DECL = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'


def _wrap_docx(document_xml: str, styles_xml: str) -> io.BytesIO:
    """Bundles a document.xml and styles.xml into a minimal valid DOCX zip."""
    rels = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        b'<Relationship Id="rId1"'
        b' Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument"'
        b' Target="word/document.xml"/>'
        b"</Relationships>"
    )

    doc_rels = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        b'<Relationship Id="rId1"'
        b' Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles"'
        b' Target="styles.xml"/>'
        b"</Relationships>"
    )

    ct = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        b'<Default Extension="rels"'
        b' ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        b'<Override PartName="/word/document.xml"'
        b' ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        b'<Override PartName="/word/styles.xml"'
        b' ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>'
        b"</Types>"
    )

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as z:
        z.writestr("[Content_Types].xml", ct)
        z.writestr("_rels/.rels", rels)
        z.writestr("word/_rels/document.xml.rels", doc_rels)
        z.writestr("word/document.xml", document_xml.encode("utf-8"))
        z.writestr("word/styles.xml", styles_xml.encode("utf-8"))
    buf.seek(0)
    return buf


# --- Style cache mechanics ---


def test_style_cache_is_memoized_on_main_document_part():
    """Two calls to _get_style_cache should return the SAME dict object,
    proving the result is cached on main_document_part rather than rebuilt."""
    styles = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f"<w:styles {W_NS_DECL}>"
        f'<w:style w:type="paragraph" w:styleId="Heading1">'
        f'<w:name w:val="Heading 1"/>'
        f'<w:pPr><w:outlineLvl w:val="0"/></w:pPr>'
        f"</w:style>"
        f"</w:styles>"
    )
    document_xml = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f"<w:document {W_NS_DECL}><w:body><w:p/><w:sectPr/></w:body></w:document>"
    )
    doc = Document(_wrap_docx(document_xml, styles))

    cache_a, default_a = _get_style_cache(doc.part)
    cache_b, default_b = _get_style_cache(doc.part)

    assert cache_a is cache_b, "style cache must be memoized, not rebuilt"
    assert default_a == default_b


def test_style_cache_resolves_basedon_inheritance_for_outline_level():
    """A child style with no own outlineLvl must inherit from its w:basedOn parent."""
    styles = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f"<w:styles {W_NS_DECL}>"
        # Parent: Heading 2 with outlineLvl=1
        f'<w:style w:type="paragraph" w:styleId="Heading2">'
        f'<w:name w:val="Heading 2"/>'
        f'<w:pPr><w:outlineLvl w:val="1"/></w:pPr>'
        f"</w:style>"
        # Child: derived from Heading 2, no own outlineLvl
        f'<w:style w:type="paragraph" w:styleId="MyChildHeading">'
        f'<w:name w:val="My Child Heading"/>'
        f'<w:basedOn w:val="Heading2"/>'
        f"</w:style>"
        f"</w:styles>"
    )
    document_xml = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f"<w:document {W_NS_DECL}><w:body><w:p/><w:sectPr/></w:body></w:document>"
    )
    doc = Document(_wrap_docx(document_xml, styles))

    cache, _ = _get_style_cache(doc.part)
    child = cache.get("MyChildHeading")
    assert child is not None
    assert child["outline_level"] == 1, child


def test_style_cache_resolves_basedon_inheritance_for_bold():
    """Bold property must propagate down a basedOn chain."""
    styles = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f"<w:styles {W_NS_DECL}>"
        f'<w:style w:type="paragraph" w:styleId="BoldBase">'
        f'<w:name w:val="Bold Base"/>'
        f"<w:rPr><w:b/></w:rPr>"
        f"</w:style>"
        f'<w:style w:type="paragraph" w:styleId="BoldChild">'
        f'<w:name w:val="Bold Child"/>'
        f'<w:basedOn w:val="BoldBase"/>'
        f"</w:style>"
        f"</w:styles>"
    )
    document_xml = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f"<w:document {W_NS_DECL}><w:body><w:p/><w:sectPr/></w:body></w:document>"
    )
    doc = Document(_wrap_docx(document_xml, styles))

    cache, _ = _get_style_cache(doc.part)
    assert cache["BoldBase"]["bold"] is True
    assert cache["BoldChild"]["bold"] is True, "bold should inherit via basedOn"


# --- get_paragraph_prefix paths ---


def test_get_paragraph_prefix_outline_level_direct():
    """Paragraph with explicit outlineLvl on its own pPr should produce '#'*N+' '."""
    styles = f'<?xml version="1.0"?><w:styles {W_NS_DECL}/>'
    document_xml = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f"<w:document {W_NS_DECL}><w:body>"
        f'<w:p><w:pPr><w:outlineLvl w:val="2"/></w:pPr>'
        f"<w:r><w:t>A heading</w:t></w:r></w:p>"
        f"<w:sectPr/></w:body></w:document>"
    )
    doc = Document(_wrap_docx(document_xml, styles))
    p = doc.paragraphs[0]
    assert get_paragraph_prefix(p) == "### "  # outlineLvl=2 -> 3 hashes


def test_get_paragraph_prefix_named_style():
    """Paragraph using Heading 1 style should produce '# ' via the cache lookup."""
    styles = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f"<w:styles {W_NS_DECL}>"
        f'<w:style w:type="paragraph" w:styleId="Heading1">'
        f'<w:name w:val="Heading 1"/>'
        f'<w:pPr><w:outlineLvl w:val="0"/></w:pPr>'
        f"</w:style>"
        f"</w:styles>"
    )
    document_xml = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f"<w:document {W_NS_DECL}><w:body>"
        f'<w:p><w:pPr><w:pStyle w:val="Heading1"/></w:pPr>'
        f"<w:r><w:t>A heading</w:t></w:r></w:p>"
        f"<w:sectPr/></w:body></w:document>"
    )
    doc = Document(_wrap_docx(document_xml, styles))
    p = doc.paragraphs[0]
    assert get_paragraph_prefix(p) == "# "


def test_get_paragraph_prefix_heuristic_all_caps_bold():
    """Short, all-caps, bolded run in a Normal-styled paragraph should
    promote to '## ' via the heuristic fallback."""
    styles = f'<?xml version="1.0"?><w:styles {W_NS_DECL}/>'
    document_xml = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f"<w:document {W_NS_DECL}><w:body>"
        f"<w:p>"
        f"<w:r><w:rPr><w:b/></w:rPr><w:t>STANDARD PROCUREMENT DOCUMENT</w:t></w:r>"
        f"</w:p>"
        f"<w:sectPr/></w:body></w:document>"
    )
    doc = Document(_wrap_docx(document_xml, styles))
    p = doc.paragraphs[0]
    assert get_paragraph_prefix(p) == "## "


def test_get_paragraph_prefix_normal_paragraph_returns_empty():
    """A regular sentence in a Normal style with mixed case should NOT
    be promoted to a heading."""
    styles = f'<?xml version="1.0"?><w:styles {W_NS_DECL}/>'
    document_xml = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f"<w:document {W_NS_DECL}><w:body>"
        f"<w:p>"
        f"<w:r><w:t>This is a normal sentence in body text.</w:t></w:r>"
        f"</w:p>"
        f"<w:sectPr/></w:body></w:document>"
    )
    doc = Document(_wrap_docx(document_xml, styles))
    p = doc.paragraphs[0]
    assert get_paragraph_prefix(p) == ""


# --- get_run_style_markers suppression ---


def test_run_markers_bold_suppressed_inside_named_heading():
    """A bold run inside a Heading 1 paragraph must NOT emit '**' markers
    (because Heading 1 is natively bold; emitting markers would double up)."""
    styles = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f"<w:styles {W_NS_DECL}>"
        f'<w:style w:type="paragraph" w:styleId="Heading1">'
        f'<w:name w:val="Heading 1"/>'
        f'<w:pPr><w:outlineLvl w:val="0"/></w:pPr>'
        f"<w:rPr><w:b/></w:rPr>"
        f"</w:style>"
        f"</w:styles>"
    )
    document_xml = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f"<w:document {W_NS_DECL}><w:body>"
        f'<w:p><w:pPr><w:pStyle w:val="Heading1"/></w:pPr>'
        f"<w:r><w:rPr><w:b/></w:rPr><w:t>Heading text</w:t></w:r></w:p>"
        f"<w:sectPr/></w:body></w:document>"
    )
    doc = Document(_wrap_docx(document_xml, styles))
    run = doc.paragraphs[0].runs[0]
    prefix, suffix = get_run_style_markers(run)
    assert prefix == "", f"expected bold suppression in heading, got prefix={prefix!r}"
    assert suffix == "", f"expected bold suppression in heading, got suffix={suffix!r}"


def test_run_markers_italic_preserved_inside_named_heading():
    """An italic run inside a Heading 1 paragraph SHOULD still emit '_'
    markers — only bold is suppressed in heading paragraphs."""
    styles = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f"<w:styles {W_NS_DECL}>"
        f'<w:style w:type="paragraph" w:styleId="Heading1">'
        f'<w:name w:val="Heading 1"/>'
        f'<w:pPr><w:outlineLvl w:val="0"/></w:pPr>'
        f"<w:rPr><w:b/></w:rPr>"
        f"</w:style>"
        f"</w:styles>"
    )
    document_xml = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f"<w:document {W_NS_DECL}><w:body>"
        f'<w:p><w:pPr><w:pStyle w:val="Heading1"/></w:pPr>'
        f"<w:r><w:rPr><w:i/></w:rPr><w:t>Heading text</w:t></w:r></w:p>"
        f"<w:sectPr/></w:body></w:document>"
    )
    doc = Document(_wrap_docx(document_xml, styles))
    run = doc.paragraphs[0].runs[0]
    prefix, suffix = get_run_style_markers(run)
    assert prefix == "_", f"expected italic preserved in heading, got prefix={prefix!r}"
    assert suffix == "_", f"expected italic preserved in heading, got suffix={suffix!r}"


def test_run_markers_bold_emitted_inside_heuristic_heading():
    """Documented legacy behaviour: bold markers are NOT suppressed inside
    heuristic-promoted headings (the heuristic path lives in
    get_paragraph_prefix only; get_run_style_markers checks the named-style
    cache only). This test pins the legacy behaviour so a future refactor
    that *changes* it gets noticed."""
    styles = f'<?xml version="1.0"?><w:styles {W_NS_DECL}/>'
    document_xml = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f"<w:document {W_NS_DECL}><w:body>"
        f"<w:p>"
        f"<w:r><w:rPr><w:b/></w:rPr><w:t>STANDARD PROCUREMENT DOCUMENT</w:t></w:r>"
        f"</w:p>"
        f"<w:sectPr/></w:body></w:document>"
    )
    doc = Document(_wrap_docx(document_xml, styles))
    run = doc.paragraphs[0].runs[0]
    prefix, suffix = get_run_style_markers(run)
    # Legacy: heuristic heading still emits '**' because run-side suppression
    # is only checked against the named-style cache.
    assert prefix == "**", f"expected legacy '**' on heuristic heading, got {prefix!r}"
    assert suffix == "**", f"expected legacy '**' on heuristic heading, got {suffix!r}"


def test_run_markers_bold_emitted_in_normal_paragraph():
    """Sanity check: bold markers ARE emitted for normal body paragraphs."""
    styles = f'<?xml version="1.0"?><w:styles {W_NS_DECL}/>'
    document_xml = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f"<w:document {W_NS_DECL}><w:body>"
        f"<w:p>"
        f"<w:r><w:t>Some plain words </w:t></w:r>"
        f"<w:r><w:rPr><w:b/></w:rPr><w:t>and bold ones</w:t></w:r>"
        f"<w:r><w:t> follow.</w:t></w:r>"
        f"</w:p>"
        f"<w:sectPr/></w:body></w:document>"
    )
    doc = Document(_wrap_docx(document_xml, styles))
    runs = doc.paragraphs[0].runs

    plain_prefix, plain_suffix = get_run_style_markers(runs[0])
    bold_prefix, bold_suffix = get_run_style_markers(runs[1])

    assert plain_prefix == "" and plain_suffix == ""
    assert bold_prefix == "**" and bold_suffix == "**"

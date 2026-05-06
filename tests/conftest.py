import io
import sys
import pytest
from docx import Document

@pytest.fixture
def simple_docx_stream():
    """Returns a BytesIO stream containing a simple DOCX."""
    doc = Document()
    doc.add_heading("Contract Agreement", 0)
    doc.add_paragraph("This is a simple contract.")
    doc.add_paragraph("The party of the first part shall be known as the Seller.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream

# Only define COM fixtures on Windows
if sys.platform == "win32":
    import pythoncom
    import win32com.client

    @pytest.fixture
    def active_word_app():
        """
        Creates an ephemeral, visible MS Word instance with a fresh document.
        Ensures it is torn down properly after the test.
        """
        pythoncom.CoInitialize()

        app = None
        try:
            # Dispatch starts a new background instance if one doesn't exist.
            # GetActiveObject will then be able to hook into it in the tool.
            app = win32com.client.Dispatch("Word.Application")
            app.Visible = True  # Needs to be visible/active for GetActiveObject sometimes
            doc = app.Documents.Add()

            # Bring to front so GetActiveObject definitely binds to this instance
            app.Activate()

            # Seed initial content
            doc.Range(0, 0).Text = "Hello world! This is a live testing document.\n"

            yield app, doc

        except Exception as e:
            pytest.skip(f"Could not initialize Word COM for testing: {e}")

        finally:
            if app:
                try:
                    doc.Close(0)  # 0 = wdDoNotSaveChanges
                except Exception:
                    pass
                # We intentionally omit app.Quit() and pythoncom.CoUninitialize()
                # to avoid Windows Access Violations (0x800706be) when Pytest holds COM locals.

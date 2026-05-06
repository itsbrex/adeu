import io
import re

from docx import Document

from adeu.ingest import extract_text_from_stream
from adeu.models import ModifyText, ReplyComment
from adeu.redline.engine import RedlineEngine


def test_reply_creates_new_comment_entry():
    """
    Verifies that replying to a comment creates a separate comment entry
    targeting the same text range, rather than appending text to the old comment.
    """
    doc = Document()
    doc.add_paragraph("Text with comment.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # 1. Create initial comment via modification
    # We change "Text" to "TextModified" to ensure the engine processes it and attaches the comment.
    engine = RedlineEngine(stream, author="Author1")
    edit = ModifyText(target_text="Text", new_text="TextModified", comment="Initial Comment")
    engine.apply_edits([edit])

    stream_mid = engine.save_to_stream()

    # Verify initial state
    text_mid = extract_text_from_stream(stream_mid)

    # Extract Comment ID. Expect [Com:1] or similar.
    match = re.search(r"\[Com:(\d+)\]", text_mid)
    assert match, f"Initial comment not found in text: {text_mid}"
    com_id = match.group(1)

    # 2. Reply to the comment
    engine2 = RedlineEngine(stream_mid, author="Author2")
    action = ReplyComment(target_id=f"Com:{com_id}", text="This is a reply.")

    applied, skipped = engine2.apply_review_actions([action])

    assert applied == 1, "Reply action should be applied"
    assert skipped == 0

    stream_final = engine2.save_to_stream()
    text_final = extract_text_from_stream(stream_final)

    print(f"Final Text:\n{text_final}")

    # Expectation: TWO distinct comment IDs in the output
    # We look for [Com:X] patterns.
    com_ids = re.findall(r"\[Com:(\d+)\]", text_final)
    unique_ids = set(com_ids)

    assert len(unique_ids) == 2, f"Should have 2 distinct comments, found: {unique_ids}\nText: {text_final}"

    # Verify content
    assert "Initial Comment" in text_final
    assert "This is a reply." in text_final
    assert "Author1" in text_final
    assert "Author2" in text_final


def test_threaded_comment_structure():
    """
    Verifies that replying to a comment creates a valid threaded structure
    compatible with Word (w15:p attribute and w15 namespace declaration).
    """
    doc = Document()
    doc.add_paragraph("Threaded conversation anchor.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # 1. Create Parent Comment (Force edit to ensure comment is attached)
    engine = RedlineEngine(stream, author="UserA")
    # Change "anchor" -> "Anchor" to force a tracked change with comment
    edit = ModifyText(target_text="anchor", new_text="Anchor", comment="Parent Topic")
    engine.apply_edits([edit])
    stream_mid = engine.save_to_stream()

    # Verify Parent Exists
    text_mid = extract_text_from_stream(stream_mid)
    match = re.search(r"\[Com:(\d+)\]", text_mid)
    assert match, f"Parent comment not created. Text: {text_mid}"
    parent_id = match.group(1)

    # 2. Create Reply (Action: REPLY)
    engine2 = RedlineEngine(stream_mid, author="UserB")
    action = ReplyComment(target_id=f"Com:{parent_id}", text="Reply Content")
    applied, skipped = engine2.apply_review_actions([action])

    assert applied == 1
    assert skipped == 0

    stream_final = engine2.save_to_stream()

    # 3. XML Inspection (The Critical Check)
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    doc_final = Document(stream_final)
    comments_part = None
    for rel in doc_final.part.rels.values():
        if rel.reltype == RT.COMMENTS:
            comments_part = rel.target_part
            break

    assert comments_part, "Comments part missing"

    xml = comments_part.blob.decode("utf-8")

    # Check 1: Namespace Declaration in Root
    assert 'xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"' in xml, (
        "Missing w15 namespace declaration in comments.xml root"
    )

    # Check 2: Parent Attribute on the Reply
    # Search for w15:p="{parent_id}"
    # NOTE: Adeu suppresses w15:p if commentsExtended (Modern Comments) is used.
    # So we check for EITHER w15:p in comments.xml OR w15:paraIdParent in commentsExtended.xml
    expected_attr = f'w15:p="{parent_id}"'

    if expected_attr in xml:
        pass  # Found legacy threading
    else:
        # Check Extended Part
        extended_part = None
        rel_extended = "http://schemas.microsoft.com/office/2011/relationships/commentsExtended"
        for rel in doc_final.part.rels.values():
            if rel.reltype == rel_extended:
                extended_part = rel.target_part
                break

        assert extended_part, "Missing w15:p AND missing commentsExtended part"
        ext_xml = extended_part.blob.decode("utf-8")
        # We can't easily check paraIdParent mapping without parsing IDs,
        # but we can check if the part exists and has content.
        assert "w15:paraIdParent" in ext_xml, "commentsExtended missing threading info"

    # 4. Ingestion Inspection
    text_final = extract_text_from_stream(stream_final)

    assert "Parent Topic" in text_final
    assert "Reply Content" in text_final

    sigs = re.findall(r"\[Com:(\d+)\]", text_final)
    # Depending on ingest implementation, reply might reuse parent ID in signature or have its own.
    # Current implementation gives every comment its own ID signature.
    assert len(set(sigs)) == 2, f"Should have 2 unique comments visible. Found: {sigs}"


def test_threaded_rendering_order():
    """
    Ensures that when ingesting text, replies are rendered correctly.
    """
    doc = Document()
    doc.add_paragraph("Target")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Setup Engine with forced edit
    engine = RedlineEngine(stream, author="A")
    engine.apply_edits([ModifyText(target_text="Target", new_text="TargetModified", comment="Root")])

    # Get ID of Root
    data = engine.comments_manager.extract_comments_data()
    assert data, "Root comment not created"
    root_id = list(data.keys())[0]

    # Reply 1
    engine.author = "B"
    engine.apply_review_actions([ReplyComment(target_id=f"Com:{root_id}", text="Reply1")])

    # Get ID of Reply1
    data = engine.comments_manager.extract_comments_data()
    # Find key that isn't root_id
    others = [k for k in data.keys() if k != root_id]
    assert others, "Reply1 not created"
    reply1_id = others[0]

    # Reply 2 (Reply to Reply1)
    engine.author = "C"
    engine.apply_review_actions([ReplyComment(target_id=f"Com:{reply1_id}", text="Reply2")])

    stream_final = engine.save_to_stream()
    text = extract_text_from_stream(stream_final)

    # Check that all 3 appear
    assert "Root" in text
    assert "Reply1" in text
    assert "Reply2" in text


def test_threading_creates_extended_part():
    """
    Verifies that adding comments to a clean doc creates commentsExtended.xml,
    which is required for visible threading in modern Word.
    """
    doc = Document()
    doc.add_paragraph("Content")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # 1. Add Root Comment
    engine = RedlineEngine(stream)
    # Note: Must change text so the edit is not skipped as no-op
    engine.apply_edits([ModifyText(target_text="Content", new_text="Content Modified", comment="Root")])
    stream_1 = engine.save_to_stream()

    # Check if Extended part exists immediately
    doc_1 = Document(stream_1)
    extended_rel = "http://schemas.microsoft.com/office/2011/relationships/commentsExtended"
    has_extended = any(rel.reltype == extended_rel for rel in doc_1.part.rels.values())
    assert has_extended, "commentsExtended.xml should be created with the first comment"

    # 2. Add Reply
    engine2 = RedlineEngine(stream_1)
    # Get root ID
    root_id = list(engine2.comments_manager.extract_comments_data().keys())[0]

    engine2.apply_review_actions([ReplyComment(target_id=f"Com:{root_id}", text="Reply")])
    stream_2 = engine2.save_to_stream()

    # 3. Inspect XML for Threading
    doc_2 = Document(stream_2)
    extended_part = None
    for rel in doc_2.part.rels.values():
        if rel.reltype == extended_rel:
            extended_part = rel.target_part
            break

    xml = extended_part.blob.decode("utf-8")
    print(xml)

    # Should contain paraIdParent linking
    assert "w15:paraIdParent" in xml


def test_full_modern_comments_triad_creation():
    doc = Document()
    doc.add_paragraph("Content")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)
    engine.apply_edits([ModifyText(target_text="Content", new_text="Content Changed", comment="Modern")])
    stream_out = engine.save_to_stream()

    doc_out = Document(stream_out)
    rels = [rel.reltype for rel in doc_out.part.rels.values()]

    rel_extended = "http://schemas.microsoft.com/office/2011/relationships/commentsExtended"
    rel_ids = "http://schemas.microsoft.com/office/2016/09/relationships/commentsIds"
    rel_extensible = "http://schemas.microsoft.com/office/2018/08/relationships/commentsExtensible"

    assert any(r == rel_extended for r in rels), "Missing commentsExtended"
    assert any(r == rel_ids for r in rels), "Missing commentsIds"
    assert any(r == rel_extensible for r in rels), "Missing commentsExtensible"

    # Check Extensible Content
    extensible_part = None
    for rel in doc_out.part.rels.values():
        if rel.reltype == rel_extensible:
            extensible_part = rel.target_part
            break

    xml = extensible_part.blob.decode("utf-8")
    print(xml)
    assert "w16cex:durableId" in xml
    assert "w16cex:dateUtc" in xml

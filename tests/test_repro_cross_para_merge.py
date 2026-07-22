import io
from docx import Document
from adeu.redline.engine import RedlineEngine
from adeu.models import ModifyText


def test_three_paragraph_merge():
    """
    BUG 3 REPRO: Tests that modifying text across 3 paragraphs (2 newline boundaries)
    correctly merges them into a single paragraph without leaving orphaned <w:p> elements.
    """
    doc = Document()
    doc.add_paragraph("Paragraph 1 ends here.")
    doc.add_paragraph("Paragraph 2 is in the middle.")
    doc.add_paragraph("Paragraph 3 begins here.")
    doc.add_paragraph("Paragraph 4 is untouched.")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    engine = RedlineEngine(buf, author="Test AI")
    engine.process_batch(
        [
            ModifyText(
                type="modify",
                target_text="ends here.\n\nParagraph 2 is in the middle.\n\nParagraph 3 begins",
                new_text="ends here. MERGED",
            )
        ]
    )

    engine.accept_all_revisions()

    out_buf = engine.save_to_stream()
    out_doc = Document(out_buf)

    paras = [p.text for p in out_doc.paragraphs if p.text.strip()]

    # If the bug exists, this will be 3 or 4 paragraphs instead of 2.
    assert len(paras) == 2
    assert paras[0] == "Paragraph 1 ends here. MERGED here."
    assert paras[1] == "Paragraph 4 is untouched."

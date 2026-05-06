import io
import os
import tempfile
from docx import Document
from docx.enum.section import WD_ORIENT

def save_to_temp_docx(doc, suffix=".docx"):
    """Saves a python-docx Document to a temporary file and returns the path."""
    fd, path = tempfile.mkstemp(suffix=suffix)
    os.close(fd)
    doc.save(path)
    return path

def make_doc_with_track_changes():
    doc = Document()
    doc.add_paragraph("Original text.")
    # Tracked changes are handled by the engine, but we can seed a doc with them if needed
    # via raw XML. For now, this helper just creates a base doc.
    return doc

def make_doc_with_comments():
    doc = Document()
    p = doc.add_paragraph("Text with a comment.")
    # Comments are also usually handled by the engine/COM.
    return doc

def make_doc_with_custom_props():
    doc = Document()
    doc.core_properties.title = "Test Doc"
    doc.core_properties.author = "Test Author"
    return doc

def make_doc_with_multiple_sections():
    doc = Document()
    doc.add_paragraph("Section 1")
    new_section = doc.add_section()
    new_section.orientation = WD_ORIENT.LANDSCAPE
    doc.add_paragraph("Section 2")
    return doc

import io

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from adeu.ingest import extract_text_from_stream
from adeu.models import ModifyText
from adeu.redline.engine import BatchValidationError, RedlineEngine


def setup_links_fixture() -> io.BytesIO:
    """
    Programmatically creates a DOCX with an external hyperlink and a cross-reference.
    """
    doc = Document()

    # 1. External Hyperlink
    p1 = doc.add_paragraph("Please visit ")
    rel_id = doc.part.relate_to(
        "https://adeu.com",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    r1 = OxmlElement("w:r")
    t1 = OxmlElement("w:t")
    t1.text = "Adeu HQ"
    r1.append(t1)
    hyperlink.append(r1)
    p1._element.append(hyperlink)

    # 2. Cross-Reference (Simple Field)
    p2 = doc.add_paragraph("As detailed in ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " REF _Ref12345 \\h ")
    r2 = OxmlElement("w:r")
    t2 = OxmlElement("w:t")
    t2.text = "Section 5"
    r2.append(t2)
    fld.append(r2)
    p2._element.append(fld)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def test_link_and_xref_extraction():
    """Verify that hyperlinks and cross-references project into the Adeu Markdown Dialect."""
    stream = setup_links_fixture()
    text = extract_text_from_stream(stream)

    assert "[Adeu HQ](https://adeu.com)" in text, "External hyperlink projection failed."
    assert "[~Section 5~](#_Ref12345)" in text, "Cross-reference projection failed."


def test_xref_modification_rejected():
    """Verify that attempting to modify a cross-reference text or hash directly is rejected."""
    stream = setup_links_fixture()
    engine = RedlineEngine(stream)

    # Attempting to edit the text of the cross-reference
    edit_text = ModifyText(target_text="[~Section 5~](#_Ref12345)", new_text="[~Section 6~](#_Ref12345)")

    try:
        engine.process_batch([edit_text])
        raise AssertionError("Engine failed to reject cross-reference text modification.")
    except BatchValidationError as e:
        assert "Cross-reference display text" in str(e) or "rejected" in str(e).lower()

    # Attempting to edit the target hash of the cross-reference
    edit_hash = ModifyText(target_text="[~Section 5~](#_Ref12345)", new_text="[~Section 5~](#_Ref99999)")

    try:
        engine.process_batch([edit_hash])
        raise AssertionError("Engine failed to reject cross-reference hash modification.")
    except BatchValidationError as e:
        assert "dependency corruption" in str(e) or "rejected" in str(e).lower()

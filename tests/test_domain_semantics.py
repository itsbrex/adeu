import io

import pytest
from docx import Document
from docx.opc.packuri import PackURI
from docx.opc.part import XmlPart
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from adeu.domain import build_structural_appendix
from adeu.ingest import extract_text_from_stream
from adeu.models import ModifyText
from adeu.redline.engine import BatchValidationError, RedlineEngine


def add_bookmark(paragraph, name: str, id_val: str = "0", text: str = "") -> None:
    """Helper to inject a w:bookmarkStart and w:bookmarkEnd around an optional run."""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:name"), name)
    start.set(qn("w:id"), id_val)
    paragraph._p.append(start)

    if text:
        paragraph.add_run(text)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), id_val)
    paragraph._p.append(end)


def add_cross_reference(paragraph, ref_name: str, text: str) -> None:
    """Helper to inject a w:fldSimple REF pointing to a bookmark."""
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f" REF {ref_name} \\h ")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def add_hyperlink(doc, paragraph, url: str, text: str) -> None:
    """Helper to inject a w:hyperlink."""
    rel_id = doc.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)


def setup_footnotes_part(doc) -> None:
    """Helper to fabricate a standalone footnotes.xml package part."""
    fn_xml = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
    <w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:footnote w:type="separator" w:id="-1">
            <w:p><w:r><w:separator/></w:r></w:p>
        </w:footnote>
        <w:footnote w:id="1">
            <w:p><w:r><w:t>Footnote content.</w:t></w:r></w:p>
        </w:footnote>
    </w:footnotes>
    """
    partname = PackURI("/word/footnotes.xml")
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
    rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"

    part = XmlPart.load(partname, content_type, fn_xml, doc.part.package)
    doc.part.package.parts.append(part)
    doc.part.relate_to(part, rel_type)


@pytest.fixture
def domain_semantics_stream() -> io.BytesIO:
    """
    Unified fixture containing all semantic elements:
    - Appendix / Definitions
    - Internal Anchors & Bookmarks
    - Cross-References
    - External Hyperlinks
    - Footnotes
    """
    doc = Document()

    # 1. Appendix / Definitions
    doc.add_heading("1. Definitions", level=1)
    doc.add_paragraph('"Affiliate" means any entity that controls, is controlled by, or is under common control.')
    doc.add_paragraph(
        "“Confidential Information” means all non-public information disclosed by one party to the other."
    )
    doc.add_paragraph("This paragraph does not define anything.")

    doc.add_heading("2. Obligations", level=1)
    doc.add_paragraph("The Affiliate shall protect the Confidential Information to the highest standard.")

    # 3. Bookmarks and Cross-References
    p3 = doc.add_paragraph("Subject to ")
    add_bookmark(p3, "MyBookmark_1", "1", "Anchored Clause")
    p3.add_run(", the parties agree to...")

    p4 = doc.add_paragraph("As strictly stated in ")
    add_cross_reference(p4, "MyBookmark_1", "Anchored Clause")
    p4.add_run(", either party may terminate.")

    # 4. Internal Anchors
    p_anchor = doc.add_paragraph("Section 5. Indemnification")
    add_bookmark(p_anchor, "_Ref12345", "0")

    p_noise = doc.add_paragraph("Some text.")
    add_bookmark(p_noise, "_GoBack", "2")
    add_bookmark(p_noise, "_Toc123456789", "3")

    # 5. Footnotes
    p_fn = doc.add_paragraph("Sentence with footnote")
    r_fn = OxmlElement("w:r")
    ref = OxmlElement("w:footnoteReference")
    ref.set(qn("w:id"), "1")
    r_fn.append(ref)
    p_fn._element.append(r_fn)
    setup_footnotes_part(doc)

    # 6. Links and Cross references
    p_link = doc.add_paragraph("Please visit ")
    add_hyperlink(doc, p_link, "https://adeu.com", "Adeu HQ")

    p_xref = doc.add_paragraph("As detailed in ")
    add_cross_reference(p_xref, "_Ref12345", "Section 5")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def test_domain_semantics_projection(domain_semantics_stream):
    """
    Ensures that domain heuristics extract definitions, anchors, footnotes, links,
    and cross-references, appending them safely to the bottom of the projection.
    """
    text = extract_text_from_stream(domain_semantics_stream)

    # Boundary constraints
    assert "<!-- READONLY_BOUNDARY_START -->" in text
    assert "# Document Structure (Read-Only)" in text

    # Definitions
    assert "## Defined Terms" in text
    assert '"Affiliate"' in text
    assert '"Confidential Information"' in text
    assert "used 1 times" in text

    # Named Anchors & Back-References
    assert "## Named Anchors" in text
    assert "MyBookmark_1" in text
    assert "Anchored to:" in text
    assert "Referenced from:" in text

    # Internal anchors & Noise suppression
    assert "{#_Ref12345}" in text
    assert "Section 5. Indemnification{#_Ref12345}" in text
    assert "{#_GoBack}" not in text
    assert "{#_Toc123456789}" not in text

    # Footnotes
    assert "[^fn-1]" in text
    assert "## Footnotes" in text
    assert "[^fn-1]: Footnote content." in text

    # Links
    assert "[Adeu HQ](https://adeu.com)" in text
    assert "[~Section 5~](#_Ref12345)" in text


@pytest.mark.parametrize(
    "target_text, new_text, error_checker",
    [
        # Appendix edits
        (
            "# Document Structure (Read-Only)",
            "# Modified Document Structure",
            lambda m: "read-only boundary" in m or "appendix" in m,
        ),
        # Footnotes
        (
            "Sentence with footnote[^fn-1]",
            "Sentence with footnote",
            lambda m: "footnote" in m and ("delete" in m or "remove" in m),
        ),
        (
            "Sentence with footnote",
            "Sentence with footnote[^fn-99]",
            lambda m: "footnote" in m and ("insert" in m or "create" in m),
        ),
        # Internal Anchors
        (
            "Some text.",
            "Some text.{#_Ref99999}",
            lambda m: "cannot modify or insert internal anchor markers" in m,
        ),
        (
            "Section 5. Indemnification{#_Ref12345}",
            "Section 5. Indemnification{#_Ref99999}",
            lambda m: "cannot modify or insert internal anchor markers" in m,
        ),
        # Links
        (
            "[~Section 5~](#_Ref12345)",
            "[~Section 6~](#_Ref12345)",
            lambda m: "cross-reference display text" in m or "rejected" in m,
        ),
        (
            "[~Section 5~](#_Ref12345)",
            "[~Section 5~](#_Ref99999)",
            lambda m: "dependency corruption" in m or "rejected" in m,
        ),
        (
            "As detailed in [~Section 5~](#_Ref12345)",
            "As detailed in [~Section 5~](#_Ref12345) and [~Section 6~](#_Ref999)",
            lambda m: "cross-reference" in m or "read-only" in m,
        ),
        (
            "As detailed in [~Section 5~](#_Ref12345)",
            "As detailed in nothing",
            lambda m: "cross-reference" in m or "delete" in m,
        ),
        (
            "Please visit [Adeu HQ](https://adeu.com)",
            "Please visit [Adeu HQ](https://adeu.com) and [Google](https://google.com)",
            lambda m: "hyperlink" in m or "dialects" in m,
        ),
        (
            "Please visit [Adeu HQ](https://adeu.com)",
            "Please visit nothing",
            lambda m: "hyperlink" in m or "delete" in m,
        ),
    ],
)
def test_domain_edit_rejection(domain_semantics_stream, target_text, new_text, error_checker):
    """
    Ensures that the engine strictly rejects any unsupported structural or semantic modifications.
    """
    engine = RedlineEngine(domain_semantics_stream)
    edit = ModifyText(type="modify", target_text=target_text, new_text=new_text)

    with pytest.raises(BatchValidationError) as exc:
        engine.process_batch([edit])

    error_msg = str(exc.value).lower()
    assert error_checker(error_msg), f"Error message '{error_msg}' did not match expectations."


def test_footnote_redline_edit(domain_semantics_stream):
    """Verify that text edits safely target the footnotes.xml payload structure via the main pipeline."""
    engine = RedlineEngine(domain_semantics_stream)

    edit = ModifyText(target_text="Footnote content.", new_text="This is an edited footnote.")
    stats = engine.process_batch([edit])
    assert stats["edits_applied"] == 1

    engine.accept_all_revisions()
    clean_text = extract_text_from_stream(engine.save_to_stream(), clean_view=True)
    assert "[^fn-1]: This is an edited footnote." in clean_text


def test_footnote_accept_changes(domain_semantics_stream):
    """Verify accept_all_changes processes footnotes.xml and leaves it clean."""
    engine = RedlineEngine(domain_semantics_stream)

    edit = ModifyText(target_text="Footnote content.", new_text="Edited content.")
    engine.process_batch([edit])
    engine.accept_all_revisions()

    fn_part = next(p for p in engine.doc.part.package.parts if "footnotes" in p.partname)

    if hasattr(fn_part, "_adeu_element"):
        import lxml.etree as etree

        xml_after = etree.tostring(fn_part._adeu_element).decode("utf-8")
    else:
        xml_after = fn_part.blob.decode("utf-8")

    assert "w:ins" not in xml_after
    assert "w:del" not in xml_after
    assert "Edited" in xml_after
    assert "content." in xml_after


# --- Defined Terms Extraction & Typos Tests ---


def test_defined_terms_extraction_and_diagnostics():
    """Validates structural extraction of defined terms, duplicates, and semantic diagnostics."""
    doc = Document()

    # 1. Glossary Definitions
    doc.add_heading("Definitions", 1)
    doc.add_paragraph('"Agreement" means this contract.')
    doc.add_paragraph("“Party” shall mean either side.")
    doc.add_paragraph('"Agreement" means another thing.')  # Duplicate

    # 2. Inline Definitions
    doc.add_paragraph('This contract (hereinafter, the "Contract") is valid.')

    # 3. Multilingual Typographic Definitions
    doc.add_heading("Miscellaneous", 1)
    doc.add_paragraph('"Confidential Information" on salainen asia.')
    doc.add_paragraph('1.1 "Affiliate" tarkoittaa osakkuusyhtiötä.')
    doc.add_paragraph('We will act as the disclosing party (jäljempänä "Discloser").')

    # 4. Phantom avoidance
    doc.add_paragraph('This is a syntax example: ("Heading*") and ("<Term>")')

    # 5. Usages, Typos
    doc.add_paragraph("The Agreement is binding. The Contract is signed.")
    doc.add_paragraph("There is an Agrement here.")
    doc.add_paragraph("We shared Confidential Information with the Affiliate. The Discloser is happy.")

    base_text = "\n".join(p.text for p in doc.paragraphs)
    appendix = build_structural_appendix(doc, base_text)

    # Check Definitions Symbol Table
    assert '"Agreement" — used' in appendix
    assert '"Contract" — used' in appendix
    assert '"Confidential Information" — used' in appendix
    assert '"Affiliate" — used' in appendix
    assert '"Discloser" — used' in appendix

    # Ensure pruned logic works (unused terms removed)
    assert '"Party"' not in appendix

    # Ensure phantom syntax terms are strictly ignored
    assert '"Heading*"' not in appendix
    assert '"<Term>"' not in appendix

    # Check Semantic Diagnostics Linter Rules
    assert "[Error] Duplicate Definition: 'Agreement' is defined multiple times." in appendix
    assert "[Info] Possible Typos for 'Agreement': Found 'Agrement'" in appendix


def test_acronym_typo_noise_reduction():
    """
    Ensures short acronyms don't trigger false positive typos due to Levenshtein proximity.
    """
    doc = Document()
    doc.add_heading("Definitions", 1)
    doc.add_paragraph('"PSUs" means power supply units.')
    doc.add_paragraph('"CPU" means central processing unit.')
    doc.add_paragraph('"Party" means the entity.')

    # Body with acronyms
    doc.add_paragraph("We rely on ESAs, LSPs, and GPUs for the servers.")
    doc.add_paragraph("The GPU is very fast.")

    # Valid short typo
    doc.add_paragraph("The Pary signed the contract.")

    # Usages to prevent pruning
    doc.add_paragraph("We bought PSUs and a CPU.")
    doc.add_paragraph("The Party begins today.")

    base_text = "\n".join(p.text for p in doc.paragraphs)
    appendix = build_structural_appendix(doc, base_text)

    # Valid short typo SHOULD be detected
    assert "[Info] Possible Typos for 'Party': Found 'Pary'" in appendix

    # Acronym noise SHOULD BE suppressed
    assert "'GPU'" not in appendix
    assert "'GPUs'" not in appendix
    assert "'ESAs'" not in appendix
    assert "'LSPs'" not in appendix

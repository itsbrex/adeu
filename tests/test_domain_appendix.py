import io

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from adeu.ingest import extract_text_from_stream
from adeu.models import ModifyText
from adeu.redline.engine import BatchValidationError, RedlineEngine


def add_bookmark(paragraph, name: str, id_val: str = "0") -> None:
    """Helper to inject a w:bookmarkStart and w:bookmarkEnd around a run."""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:name"), name)
    start.set(qn("w:id"), id_val)
    paragraph._p.append(start)

    paragraph.add_run("Anchored Clause")

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), id_val)
    paragraph._p.append(end)


def add_cross_reference(paragraph, ref_name: str) -> None:
    """Helper to inject a w:fldSimple REF pointing to a bookmark."""
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f" REF {ref_name} \\h ")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Anchored Clause"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


@pytest.fixture
def domain_docx_stream() -> io.BytesIO:
    doc = Document()

    # 1. Definitions Section
    doc.add_heading("1. Definitions", level=1)
    doc.add_paragraph('"Affiliate" means any entity that controls, is controlled by, or is under common control.')
    doc.add_paragraph(
        "“Confidential Information” means all non-public information disclosed by one party to the other."
    )
    doc.add_paragraph("This paragraph does not define anything.")

    # 2. Body Text using definitions
    doc.add_heading("2. Obligations", level=1)
    doc.add_paragraph("The Affiliate shall protect the Confidential Information to the highest standard.")

    # 3. Bookmarks and Back-References
    p3 = doc.add_paragraph("Subject to ")
    add_bookmark(p3, "MyBookmark_1", "1")
    p3.add_run(", the parties agree to...")

    doc.add_heading("3. Termination", level=1)
    p4 = doc.add_paragraph("As strictly stated in ")
    add_cross_reference(p4, "MyBookmark_1")
    p4.add_run(", either party may terminate.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def test_appendix_projection(domain_docx_stream):
    """
    Ensures that domain heuristics extract definitions and anchors,
    appending them safely to the bottom of the projection.
    """
    text = extract_text_from_stream(domain_docx_stream)

    # Boundary constraints
    assert "<!-- READONLY_BOUNDARY_START -->" in text
    assert "# Document Structure (Read-Only)" in text

    # Track 3.2: Definitions
    assert "## Defined Terms" in text
    assert '"Affiliate"' in text
    assert '"Confidential Information"' in text
    assert "used 2 times" in text  # "Affiliate" should appear 2 times (definition + usage)

    # Track 3.3: Named Anchors & Back-References
    assert "## Named Anchors" in text
    assert "MyBookmark_1" in text
    assert "Anchored to:" in text
    assert "Referenced from:" in text


def test_appendix_edit_rejection(domain_docx_stream):
    """
    Ensures that the engine strictly rejects any text modification targeting
    the read-only appendix boundary.
    """
    engine = RedlineEngine(domain_docx_stream)

    # Simulate LLM trying to edit the metadata appendix
    edit = ModifyText(
        type="modify", target_text="# Document Structure (Read-Only)", new_text="# Modified Document Structure"
    )

    with pytest.raises(BatchValidationError) as exc:
        engine.process_batch([edit])

    error_msg = str(exc.value)
    assert "read-only boundary" in error_msg.lower() or "appendix" in error_msg.lower()

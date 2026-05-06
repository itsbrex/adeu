# FILE: tests/test_repro_id_pool_renumber.py
"""
Regression tests for Bug 5: Live Word snapshot ID renumbering.

Background: Word allocates Chg and Com IDs from a single shared counter,
while the disk path uses two independent counters. Without renumbering,
an agent that reads via the disk path and writes via Live Word (or vice
versa) can target the wrong XML element because Com:N from one path may
not match Com:N from the other.

`renumber_snapshot_ids` rewrites the snapshot's w:id attributes so the
projection always emits the disk-style two-pool scheme regardless of how
the underlying IDs were originally allocated. These tests pin that
behavior so future refactors can't silently regress it.
"""

import re
from io import BytesIO

import pytest
from docx import Document
from docx.oxml.ns import qn

from adeu.ingest import _extract_text_from_doc, extract_text_from_stream
from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine
from adeu.redline.mapper import DocumentMapper, renumber_snapshot_ids

APPENDIX_MARKER = "<!-- READONLY_BOUNDARY_START -->"
COMMENTS_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"


def _strip_appendix(text: str) -> str:
    if APPENDIX_MARKER in text:
        return text[: text.find(APPENDIX_MARKER)].rstrip()
    return text


def _projection_ids(doc_or_bytes) -> tuple[set[str], set[str]]:
    """Returns (chg_ids_set, com_ids_set) extracted from the projection."""
    if isinstance(doc_or_bytes, bytes):
        text = extract_text_from_stream(BytesIO(doc_or_bytes), clean_view=False)
    else:
        text = _extract_text_from_doc(doc_or_bytes, clean_view=False)
    text = _strip_appendix(text)
    chg = set(re.findall(r"\bChg:(\d+)\b", text))
    com = set(re.findall(r"\bCom:(\d+)\b", text))
    return chg, com


def _comments_part(doc):
    for part in doc.part.package.parts:
        if part.content_type == COMMENTS_CONTENT_TYPE:
            return part
    return None


def _comments_root(doc):
    """Returns the parsed root element of the comments part, or None."""
    part = _comments_part(doc)
    if part is None:
        return None
    if hasattr(part, "element"):
        return part.element
    from docx.oxml import parse_xml

    if not hasattr(part, "_adeu_element"):
        part._adeu_element = parse_xml(part.blob)
    return part._adeu_element


def _shift_com_ids(doc, shift: int) -> None:
    """
    Mutates `doc` in place to shift all comment IDs by `shift`.
    Used to simulate Word's single-pool numbering, where Com IDs land in a
    range that doesn't overlap with Chg IDs.
    """
    comments_root = _comments_root(doc)
    if comments_root is not None:
        for c in comments_root.findall(qn("w:comment")):
            old = c.get(qn("w:id"))
            if old is not None:
                c.set(qn("w:id"), str(int(old) + shift))

    for tag in (
        qn("w:commentReference"),
        qn("w:commentRangeStart"),
        qn("w:commentRangeEnd"),
    ):
        for el in doc.element.iter(tag):
            old = el.get(qn("w:id"))
            if old is not None:
                el.set(qn("w:id"), str(int(old) + shift))


@pytest.fixture
def doc_bytes_with_changes_and_comments():
    """
    Builds a DOCX with two tracked-change modifications, each with a comment.
    Returns the saved bytes. The disk path emits Chg in {1,2,3,4} and Com in
    {1,2}.
    """
    doc = Document()
    doc.add_paragraph("Quarterly revenue rose by twelve percent.")
    doc.add_paragraph("The team launched three new products this year.")
    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="Reviewer")
    engine.process_batch(
        [
            ModifyText(
                type="modify",
                target_text="twelve",
                new_text="fifteen",
                comment="Verify with finance",
            ),
            ModifyText(
                type="modify",
                target_text="three",
                new_text="four",
                comment="Confirm product count",
            ),
        ]
    )
    return engine.save_to_stream().getvalue()


def test_renumber_keeps_disk_style_two_pools(doc_bytes_with_changes_and_comments):
    """
    A doc whose IDs already follow the disk-style two-pool scheme should be
    unchanged in shape after renumbering: Chg pool stays {1..N}, Com pool
    stays {1..M}.
    """
    chg_baseline, com_baseline = _projection_ids(doc_bytes_with_changes_and_comments)
    assert chg_baseline == {"1", "2", "3", "4"}
    assert com_baseline == {"1", "2"}

    doc = Document(BytesIO(doc_bytes_with_changes_and_comments))
    renumber_snapshot_ids(doc)

    chg_after, com_after = _projection_ids(doc)
    assert chg_after == chg_baseline
    assert com_after == com_baseline


def test_renumber_normalizes_simulated_word_pool(doc_bytes_with_changes_and_comments):
    """
    The core Bug 5 case: simulate Word's single-pool numbering by shifting
    comment IDs into a non-overlapping range. After renumbering, the
    projection must show the disk-style two-pool scheme.
    """
    chg_baseline, com_baseline = _projection_ids(doc_bytes_with_changes_and_comments)

    # Simulate Word: shift Com IDs into a high range that doesn't clash with Chg.
    doc = Document(BytesIO(doc_bytes_with_changes_and_comments))
    _shift_com_ids(doc, shift=100)

    # Sanity: the simulation actually shifted things.
    _, com_simulated = _projection_ids(doc)
    assert com_simulated == {"101", "102"}, "simulation precondition failed"

    # Apply the fix.
    renumber_snapshot_ids(doc)

    chg_after, com_after = _projection_ids(doc)
    assert chg_after == chg_baseline, (
        "renumber should not change which Chg IDs exist, only normalize them to the {1..N} pool"
    )
    assert com_after == com_baseline, (
        "renumber must restore Com IDs to the disk-style {1..M} pool, not leave them in the simulated Word range"
    )


def test_renumber_preserves_comment_text_linkage(doc_bytes_with_changes_and_comments):
    """
    After renumbering a simulated-Word snapshot, the comments_map must map
    the renumbered Com IDs to the SAME author/text as the baseline disk doc.
    Without this, the renumber would technically produce disk-style IDs but
    they'd point to the wrong comments — defeating the whole purpose.
    """
    baseline_doc = Document(BytesIO(doc_bytes_with_changes_and_comments))
    baseline_mapper = DocumentMapper(baseline_doc)
    baseline_map = dict(baseline_mapper.comments_map)

    doc = Document(BytesIO(doc_bytes_with_changes_and_comments))
    _shift_com_ids(doc, shift=100)
    renumber_snapshot_ids(doc)

    renumbered_mapper = DocumentMapper(doc)
    renumbered_map = dict(renumbered_mapper.comments_map)

    assert renumbered_map == baseline_map, (
        "Comments_map must be identical after renumber. If it differs, the "
        "renumber lost the linkage between commentReference IDs and their "
        "comment payloads — which means an agent that targets Com:N would "
        "hit the wrong comment."
    )


def test_renumber_is_deterministic(doc_bytes_with_changes_and_comments):
    """
    Two independent renumbers of the same input snapshot must produce
    byte-identical projections. Stability matters because agents may make
    multiple read calls within a session and must see consistent IDs.
    """
    doc1 = Document(BytesIO(doc_bytes_with_changes_and_comments))
    doc2 = Document(BytesIO(doc_bytes_with_changes_and_comments))
    _shift_com_ids(doc1, shift=100)
    _shift_com_ids(doc2, shift=100)

    renumber_snapshot_ids(doc1)
    renumber_snapshot_ids(doc2)

    text1 = _strip_appendix(_extract_text_from_doc(doc1, clean_view=False))
    text2 = _strip_appendix(_extract_text_from_doc(doc2, clean_view=False))

    assert text1 == text2


def test_renumber_handles_doc_with_no_comments():
    """
    A document with tracked changes but no comments must renumber cleanly:
    Chg pool gets normalized, Com pool stays empty.
    """
    base = Document()
    base.add_paragraph("The team launched three new products this year.")
    stream = BytesIO()
    base.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="Reviewer")
    engine.process_batch(
        [
            ModifyText(type="modify", target_text="three", new_text="four"),
        ]
    )
    saved_bytes = engine.save_to_stream().getvalue()

    doc = Document(BytesIO(saved_bytes))
    chg_remap, com_remap = renumber_snapshot_ids(doc)

    assert com_remap == {}, "no comments → no com remap"
    assert len(chg_remap) >= 1, "tracked-change edit should produce at least one chg id"

    chg_after, com_after = _projection_ids(doc)
    assert com_after == set()
    # Chg IDs should form a contiguous {1..N} pool.
    assert chg_after == {str(i) for i in range(1, len(chg_after) + 1)}


def test_renumber_handles_doc_with_no_tracked_changes():
    """
    Symmetric case: doc with comments but no tracked changes must renumber
    cleanly. Chg pool stays empty, Com pool gets normalized.
    """
    base = Document()
    base.add_paragraph("This is the body of the document.")
    stream = BytesIO()
    base.save(stream)
    stream.seek(0)

    # Add a comment without a content modification by attaching it to a
    # same-text "modify" edit.
    engine = RedlineEngine(stream, author="Reviewer")
    engine.process_batch(
        [
            ModifyText(
                type="modify",
                target_text="body",
                new_text="body",
                comment="Note: this is the body.",
            ),
        ]
    )
    saved_bytes = engine.save_to_stream().getvalue()

    doc = Document(BytesIO(saved_bytes))
    _shift_com_ids(doc, shift=50)
    chg_remap, com_remap = renumber_snapshot_ids(doc)

    assert com_remap, "comment IDs should be remapped"

    _, com_after = _projection_ids(doc)
    assert com_after == {str(i) for i in range(1, len(com_after) + 1)}


def test_renumber_returns_translation_dicts(doc_bytes_with_changes_and_comments):
    """
    renumber_snapshot_ids returns (chg_remap, com_remap) so that callers who
    need to translate IDs across the renumber can do so. These dicts must
    cover every original ID that was present in the snapshot.
    """
    doc = Document(BytesIO(doc_bytes_with_changes_and_comments))
    _shift_com_ids(doc, shift=100)

    # Snapshot the original IDs before renumber.
    original_chg_ids = set()
    for tag in (qn("w:ins"), qn("w:del")):
        for el in doc.element.iter(tag):
            cid = el.get(qn("w:id"))
            if cid is not None:
                original_chg_ids.add(cid)

    original_com_ids = set()
    comments_root = _comments_root(doc)
    if comments_root is not None:
        for c in comments_root.findall(qn("w:comment")):
            cid = c.get(qn("w:id"))
            if cid is not None:
                original_com_ids.add(cid)

    chg_remap, com_remap = renumber_snapshot_ids(doc)

    assert set(chg_remap.keys()) == original_chg_ids, "chg_remap must cover every original Chg ID"
    assert set(com_remap.keys()) == original_com_ids, "com_remap must cover every original Com ID"

    # Values must be unique within each pool — no two original IDs map to the same new ID.
    assert len(set(chg_remap.values())) == len(chg_remap)
    assert len(set(com_remap.values())) == len(com_remap)


def test_renumber_normalizes_both_pools_independently():
    """
    The Chg pool and Com pool are independent. Renumbering must produce
    {1..N} for Chg and {1..M} for Com regardless of any overlap or non-overlap
    between the original pools. Specifically, an originally-overlapping
    Chg/Com input (both in {1,2}) and a non-overlapping input (Chg in {1,2},
    Com in {100,101}) must produce identical post-renumber projections.
    """
    # Build two docs that have the same semantic content but different
    # pre-renumber ID layouts.
    base = Document()
    base.add_paragraph("Alpha bravo charlie delta.")
    stream = BytesIO()
    base.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="R")
    engine.process_batch(
        [
            ModifyText(
                type="modify",
                target_text="bravo",
                new_text="BRAVO",
                comment="Bold this",
            ),
            ModifyText(
                type="modify",
                target_text="delta",
                new_text="DELTA",
                comment="Bold this too",
            ),
        ]
    )
    saved = engine.save_to_stream().getvalue()

    # Variant A: leave IDs alone (overlapping pools, disk-style).
    doc_a = Document(BytesIO(saved))
    renumber_snapshot_ids(doc_a)
    chg_a, com_a = _projection_ids(doc_a)

    # Variant B: simulate Word (Com IDs shifted to non-overlapping range).
    doc_b = Document(BytesIO(saved))
    _shift_com_ids(doc_b, shift=100)
    renumber_snapshot_ids(doc_b)
    chg_b, com_b = _projection_ids(doc_b)

    # Both must converge on identical disk-style two-pool projections.
    assert chg_a == chg_b
    assert com_a == com_b

    # And specifically: Com pool should be {1..M}, NOT {101, 102} or anything
    # leftover from the simulation.
    assert com_b == {str(i) for i in range(1, len(com_b) + 1)}

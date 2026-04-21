import io

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from adeu.ingest import extract_text_from_stream
from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine


def test_bug1_hyperlink_survives_unrelated_edit():
    doc = Document()
    doc.add_paragraph("The parties agree.")
    p2 = doc.add_paragraph("Visit ")

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), "rId5")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "the portal"
    r.append(t)
    hyperlink.append(r)
    p2._element.append(hyperlink)

    p2.add_run(" today.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # 1. Ensure ingest extracts hyperlink text
    text = extract_text_from_stream(stream)
    assert "the portal" in text, "BUG 1: Hyperlink text is completely invisible to ingest!"

    # 2. Edit unrelated paragraph
    engine = RedlineEngine(stream)
    engine.apply_edits([ModifyText(target_text="The parties", new_text="All parties")])

    res = engine.save_to_stream()
    doc_res = Document(res)
    xml = doc_res.element.xml

    assert "the portal" in xml, "BUG 1: Hyperlink text destroyed in XML by unrelated edit!"
    assert "w:hyperlink" in xml, "BUG 1: Hyperlink element removed from XML!"


def test_bug2_footnotes_visible():
    doc = Document()
    p = doc.add_paragraph("Some text")

    r = OxmlElement("w:r")
    f_ref = OxmlElement("w:footnoteReference")
    f_ref.set(qn("w:id"), "1")
    r.append(f_ref)
    p._element.append(r)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    text = extract_text_from_stream(stream)
    assert "[^1]" in text or "footnote" in text.lower(), "BUG 2: Footnote reference is invisible!"


def test_bug3_whole_paragraph_deletion_mark():
    doc = Document()
    doc.add_paragraph("Delete me completely.")
    doc.add_paragraph("Keep me.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)
    engine.apply_edits([ModifyText(target_text="Delete me completely.", new_text="")])

    res = engine.save_to_stream()
    doc_res = Document(res)

    # The first paragraph should have a w:del inside w:pPr/w:rPr
    p1 = doc_res.paragraphs[0]._element
    pPr = p1.find(qn("w:pPr"))
    assert pPr is not None, "BUG 3: Missing pPr"
    rPr = pPr.find(qn("w:rPr"))
    assert rPr is not None, "BUG 3: Missing rPr in pPr"
    del_mark = rPr.find(qn("w:del"))
    assert del_mark is not None, "BUG 3: Paragraph mark not marked as deleted! Will leave empty list item."


def test_bug4_list_markers_visible():
    doc = Document()
    p = doc.add_paragraph("List item 1")
    pPr = OxmlElement("w:pPr")
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), "1")
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)
    p._element.insert(0, pPr)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    text = extract_text_from_stream(stream)
    assert text.strip().startswith("* ") or text.strip().startswith("1."), "BUG 4: List marker invisible!"

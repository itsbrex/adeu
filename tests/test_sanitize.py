"""Tests for adeu sanitize — DOCX metadata scrubber."""

import io
import os
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.opc.part import Part
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from adeu.sanitize import transforms
from adeu.sanitize.core import SanitizeError, sanitize_docx

from .docx_fixtures import save_to_temp_docx
from .verify_sanitized import (
    check_full_scrub,
    check_keep_markup,
)

FIXTURE_DIR = Path(__file__).parent / "fixtures"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


FIXTURE_DIR = Path(__file__).parent / "fixtures"

# ---------------------------------------------------------------------------
# Specific Helpers for Sanitize
# ---------------------------------------------------------------------------


def _make_doc_with_track_changes() -> io.BytesIO:
    """Create a DOCX with track changes (insertion + deletion)."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("The ")

    d = OxmlElement("w:del")
    d.set(qn("w:id"), "1")
    d.set(qn("w:author"), "Opposing Counsel")
    d.set(qn("w:date"), "2025-01-15T10:00:00Z")
    rd = OxmlElement("w:r")
    rt = OxmlElement("w:delText")
    rt.set(qn("xml:space"), "preserve")
    rt.text = "Vendor"
    rd.append(rt)
    d.append(rd)
    p._element.append(d)

    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "2")
    ins.set(qn("w:author"), "Opposing Counsel")
    ins.set(qn("w:date"), "2025-01-15T10:00:00Z")
    ri = OxmlElement("w:r")
    ti = OxmlElement("w:t")
    ti.set(qn("xml:space"), "preserve")
    ti.text = "Supplier"
    ri.append(ti)
    ins.append(ri)
    p._element.append(ins)

    p.add_run(" shall provide services.")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def _make_doc_with_multi_author_track_changes() -> io.BytesIO:
    """Create a DOCX with track changes from multiple authors."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("The ")

    d = OxmlElement("w:del")
    d.set(qn("w:id"), "1")
    d.set(qn("w:author"), "Adeu Reviewer")
    d.set(qn("w:date"), "2025-01-15T10:00:00Z")
    rd = OxmlElement("w:r")
    rt = OxmlElement("w:delText")
    rt.text = "Vendor"
    rd.append(rt)
    d.append(rd)
    p._element.append(d)

    ins2 = OxmlElement("w:ins")
    ins2.set(qn("w:id"), "3")
    ins2.set(qn("w:author"), "Sneaky Counterparty")
    ri2 = OxmlElement("w:r")
    ti2 = OxmlElement("w:t")
    ti2.text = "Supplier"
    ri2.append(ti2)
    ins2.append(ri2)
    p._element.append(ins2)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def _make_doc_with_rsids() -> io.BytesIO:
    doc = Document()
    p = doc.add_paragraph("Hello World")
    p._element.set(qn("w:rsidR"), "00A21F3B")
    p._element.set(qn("w:rsidRDefault"), "004C12DE")
    p._element.set(qn("w:rsidP"), "00B33E21")
    for run in p.runs:
        run._element.set(qn("w:rsidR"), "00A21F3B")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def _save_to_tmp(stream: io.BytesIO) -> str:
    path = save_to_temp_docx(Document(stream))
    return path


# ---------------------------------------------------------------------------
# Transform unit tests
# ---------------------------------------------------------------------------


class TestTransforms:
    def test_strip_rsid(self):
        doc = Document(_make_doc_with_rsids())
        transforms.strip_rsid(doc)
        for el in doc.element.iter():
            for attr in transforms.RSID_ATTRS:
                assert attr not in el.attrib

    def test_strip_para_ids(self):
        doc = Document()
        p = doc.add_paragraph("Test")
        p._element.set(f"{{{transforms.W14_NS}}}paraId", "3F2A91BC")
        transforms.strip_para_ids(doc)
        for el in doc.element.iter():
            for attr in transforms.W14_ATTRS:
                assert attr not in el.attrib

    def test_count_and_accept_tracked_changes(self):
        doc = Document(_make_doc_with_track_changes())
        ins, dels, fmt = transforms.count_tracked_changes(doc)
        assert ins == 1 and dels == 1
        transforms.accept_all_tracked_changes(doc)
        ins, dels, fmt = transforms.count_tracked_changes(doc)
        assert ins == 0 and dels == 0
        assert "Supplier" in doc.paragraphs[0].text

    def test_strip_hidden_text(self):
        doc = Document()
        run = doc.add_paragraph().add_run("Hidden")
        run._element.get_or_add_rPr().append(OxmlElement("w:vanish"))
        assert len(transforms.strip_hidden_text(doc)) > 0

    def test_scrub_doc_properties(self):
        doc = Document()
        app_part = next(p for p in doc.part.package.parts if str(p.partname).endswith("app.xml"))
        app_part._blob = (
            f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
            f'<Properties xmlns="{transforms.EXTENDED_NS}">'
            f"<TotalTime>15</TotalTime><Template>T.dotm</Template></Properties>"
        ).encode("utf-8")
        transforms.scrub_doc_properties(doc)
        tree = etree.fromstring(app_part.blob)
        ns = {"app": transforms.EXTENDED_NS}
        assert tree.find(".//app:TotalTime", ns).text == "0"
        assert not tree.find(".//app:Template", ns).text

    def test_strip_custom_xml(self):
        doc = Document()
        pkg = doc.part.package
        custom_part = Part(pkg.next_partname("/customXml/item%d.xml"), "application/xml", b"<t/>", pkg)
        pkg.parts.append(custom_part)

        # Add content control with binding
        sdt = OxmlElement("w:sdt")
        sdtPr = OxmlElement("w:sdtPr")
        binding = OxmlElement("w:dataBinding")
        binding.set(qn("w:xpath"), "/t")
        sdtPr.append(binding)
        sdt.append(sdtPr)
        doc.add_paragraph()._element.append(sdt)

        transforms.strip_custom_xml(doc)
        assert custom_part not in pkg.parts
        assert len(doc.element.findall(f".//{qn('w:dataBinding')}")) == 0


# ---------------------------------------------------------------------------
# Orchestrator integration tests
# ---------------------------------------------------------------------------


class TestSanitizeIntegration:
    def test_full_sanitize_flow(self, tmp_path):
        # 1. Clean doc
        clean_doc = Document()
        clean_doc.add_paragraph("Clean")
        input_path = save_to_temp_docx(clean_doc)
        result = sanitize_docx(input_path)
        assert result.status == "clean"
        os.unlink(input_path)

        # 2. Unresolved changes error
        input_path = _save_to_tmp(_make_doc_with_track_changes())
        with pytest.raises(SanitizeError, match="unresolved"):
            sanitize_docx(input_path)

        # 3. Accept all
        result = sanitize_docx(input_path, accept_all=True)
        assert result.tracked_changes_accepted > 0
        os.unlink(input_path)

    def test_full_sanitize_with_accept_all_warns_multi_author(self):
        """Full sanitize with --accept-all warns if multiple authors are detected (VAL-OBS-NEW-9)."""
        stream = _make_doc_with_multi_author_track_changes()
        input_path = _save_to_tmp(stream)

        result = sanitize_docx(input_path, accept_all=True)
        assert result.status == "clean_with_warnings"
        assert any("Multiple authors detected" in w for w in result.warnings)
        assert "Adeu Reviewer" in result.report_text
        assert "Sneaky Counterparty" in result.report_text
        os.unlink(input_path)

    def test_keep_markup_and_author_replace(self):
        input_path = _save_to_tmp(_make_doc_with_track_changes())
        result = sanitize_docx(input_path, keep_markup=True, author="Firm X")
        out_doc = Document(result.output_path)
        for tag in [qn("w:ins"), qn("w:del")]:
            for el in out_doc.element.findall(f".//{tag}"):
                assert el.get(qn("w:author")) == "Firm X"
                assert el.get(qn("w:date")) == "2025-01-01T00:00:00Z"
        os.unlink(input_path)

    def test_baseline_recompute(self):
        baseline_doc = Document()
        baseline_doc.add_paragraph("The Vendor shall provide services.")
        b_path = save_to_temp_docx(baseline_doc)

        working_doc = Document()
        working_doc.add_paragraph("The Supplier shall provide services.")
        w_path = save_to_temp_docx(working_doc)

        result = sanitize_docx(w_path, baseline_path=b_path, author="Tester")
        assert result.tracked_changes_found > 0
        os.unlink(b_path)
        os.unlink(w_path)

    def test_file_not_found(self):
        with pytest.raises(FileNotFoundError):
            sanitize_docx("/nonexistent/file.docx")

    def test_baseline_not_found(self):
        input_path = _save_to_tmp(_make_doc_with_track_changes())
        with pytest.raises(FileNotFoundError):
            sanitize_docx(input_path, baseline_path="/nonexistent/baseline.docx")
        os.unlink(input_path)

    def test_e2e_dirty_sample(self, tmp_path):
        dirty_doc = FIXTURE_DIR / "dirty_sample.docx"
        if not dirty_doc.exists():
            pytest.skip("Fixture missing")

        out_full = tmp_path / "full.docx"
        sanitize_docx(str(dirty_doc), str(out_full), accept_all=True)
        with zipfile.ZipFile(out_full, "r") as zf:
            check_full_scrub(zf)

        out_keep = tmp_path / "keep.docx"
        sanitize_docx(str(dirty_doc), str(out_keep), keep_markup=True, author="Firm")
        with zipfile.ZipFile(out_keep, "r") as zf:
            check_keep_markup(zf, "Firm")

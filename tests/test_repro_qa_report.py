import io
import sys
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from docx.oxml.ns import qn

from adeu.models import ModifyText
from adeu.redline.engine import BatchValidationError, RedlineEngine


def create_doc_with_bold_run():
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("This is bold text.")
    r.bold = True
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


class TestQaReportDefects:
    def test_h1_insertion_inherits_bold(self):
        """
        H1: Plain text insertion into a bold run explicitly emits <w:b w:val="0"/>
        instead of inheriting the bold context.
        """
        stream = create_doc_with_bold_run()
        engine = RedlineEngine(stream)
        # new_text starts with target_text to bypass trim_common_context string diffs
        edit = ModifyText(target_text="bold text.", new_text="bold text. Plain suffix.")
        engine.process_batch([edit])

        ins_elements = engine.doc.element.xpath("//w:ins")
        assert len(ins_elements) >= 1

        for ins in ins_elements:
            for run in ins.findall(qn("w:r")):
                rPr = run.find(qn("w:rPr"))
                if rPr is not None:
                    b_tag = rPr.find(qn("w:b"))
                    if b_tag is not None:
                        val = b_tag.get(qn("w:val"))
                        assert val != "0", "Insertion explicitly disabled bold inside a bold run (H1)"

    def test_h2_markdown_bold_inside_bold(self):
        """
        H2: Markdown **bold** emits bare <w:b/> inside already-bold context,
        which Word interprets as a toggle (turning bold OFF).
        """
        stream = create_doc_with_bold_run()
        engine = RedlineEngine(stream)
        edit = ModifyText(target_text="bold text.", new_text="bold text. **Spicy suffix**.")
        engine.process_batch([edit])

        ins_elements = engine.doc.element.xpath("//w:ins")
        for ins in ins_elements:
            for run in ins.findall(qn("w:r")):
                t_tag = run.find(qn("w:t"))
                if t_tag is not None and "Spicy suffix" in t_tag.text:
                    rPr = run.find(qn("w:rPr"))
                    if rPr is not None:
                        b_tag = rPr.find(qn("w:b"))
                        assert b_tag is not None, "Expected <w:b> tag for markdown bold"
                        val = b_tag.get(qn("w:val"))
                        assert val == "1", (
                            "Inside a bold context, **bold** must emit <w:b w:val='1'> to avoid toggle semantics (H2)"
                        )

    def test_h3_accept_all_removes_orphan_paragraphs(self):
        """
        H3: accept_all_changes leaves orphan blank paragraphs after full-paragraph deletions.
        """
        doc = Document()
        doc.add_heading("5. Exclusions", level=1)
        doc.add_paragraph("Following text.")
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        engine = RedlineEngine(stream)

        edit = ModifyText(target_text="# 5. Exclusions", new_text="## 5. Exclusions (Revised)\n\nNew paragraph.")
        engine.process_batch([edit])
        engine.accept_all_revisions()

        paragraphs = [p.text for p in engine.doc.paragraphs]
        # 1: Empty orphan (BUG), 2: Revised heading, 3: New paragraph, 4: Following text
        assert len(paragraphs) == 3, (
            f"Expected 3 paragraphs, but got {len(paragraphs)}. Orphan paragraph was left behind: {paragraphs}"
        )

    def test_h7_skipped_edits_report(self):
        """
        H7: Skipped-edit reporting is aggregate-only; callers cannot identify which edits were skipped.
        """
        doc = Document()
        doc.add_paragraph("Overlapping target.")
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        engine = RedlineEngine(stream)
        # Both edits target the same text. Edit 1 applies, Edit 2 is skipped due to overlap.
        # This passes validation but tests the post-validation skip path.
        edit1 = ModifyText(target_text="Overlapping target.", new_text="First edit.")
        edit2 = ModifyText(target_text="Overlapping target.", new_text="Second edit.")

        stats = engine.process_batch([edit1, edit2])

        assert stats["edits_applied"] == 1
        assert stats["edits_skipped"] == 1

        assert "skipped_details" in stats, "Engine must return detailed information about skipped edits (H7)"

    def test_h8_noop_modify_drops_comment(self):
        """
        H8: Silent comment drop. Comments attached to no-op modify edits are lost.
        """
        doc = Document()
        doc.add_paragraph("This is some text.")
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        engine = RedlineEngine(stream)
        edit = ModifyText(target_text="some text.", new_text="some text.", comment="This is a QA comment.")
        engine.process_batch([edit])

        comments = engine.comments_manager.extract_comments_data()
        assert len(comments) == 1, "Comment attached to no-op modify was silently dropped (H8)"

    def test_h9_table_row_multi_paragraph(self):
        """
        H9: Silent multi-paragraph-insert drop in table row contexts.
        """
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Cell content."
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        engine = RedlineEngine(stream)
        edit = ModifyText(target_text="Cell content.", new_text="Cell content.\n\nNew paragraph.")

        stats = engine.process_batch([edit])

        if stats["edits_skipped"] > 0:
            pytest.fail("Edit was silently skipped. Must either apply or raise a validation error. (H9)")

        engine.accept_all_revisions()

        table_text = [p.text for row in engine.doc.tables[0].rows for cell in row.cells for p in cell.paragraphs]
        assert "New paragraph." in table_text, "Multi-paragraph insert into table silently failed (H9)"

    def test_m3_empty_comment_parts_not_created(self):
        """
        M3: Empty comment XML parts created on every modify edit, even when no comments are attached.
        """
        doc = Document()
        doc.add_paragraph("No comments here.")
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        engine = RedlineEngine(stream)
        edit = ModifyText(target_text="here.", new_text="here in this document.")
        engine.process_batch([edit])

        parts = engine.doc.part.package.parts
        comment_parts = [p for p in parts if "comments" in p.partname]
        assert len(comment_parts) == 0, "Comment XML parts should not be created if no comments are attached (M3)"

    def test_tech2_heading_depth_validation(self):
        """
        TECH-2: Heading levels above 6 should be rejected by validation to prevent
        silent template breakage (e.g., generating Heading7).
        """
        doc = Document()
        doc.add_paragraph("Replace me.")
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        engine = RedlineEngine(stream)
        edit = ModifyText(target_text="Replace me.", new_text="####### Heading 7")

        with pytest.raises(BatchValidationError, match="[Hh]eading"):
            engine.process_batch([edit])

    def test_tech3_table_cell_comment_anchoring(self):
        """
        TECH-3: Comments for cell edits should anchor to the specific edited cell,
        not the first cell of the row.
        """
        doc = Document()
        table = doc.add_table(rows=1, cols=3)
        row = table.rows[0]
        row.cells[0].text = "First"
        row.cells[1].text = "Second"
        row.cells[2].text = "Third"
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        engine = RedlineEngine(stream)
        # Edit only the 3rd cell and attach a comment
        edit = ModifyText(
            target_text="First | Second | Third", new_text="First | Second | Third Modified", comment="Review this cell"
        )
        engine.process_batch([edit])

        cells = engine.doc.tables[0].rows[0].cells
        cell1_xml = cells[0]._tc.xml
        cell3_xml = cells[2]._tc.xml

        assert "commentRangeStart" not in cell1_xml, "Comment incorrectly anchored to the first cell (TECH-3)"
        assert "commentRangeStart" in cell3_xml, "Comment must anchor to the modified cell (TECH-3)"

    def test_tech6_table_error_wording(self):
        """
        TECH-6: Error message for structural table changes should explicitly say "rows or columns",
        not just "columns".
        """
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        row = table.rows[0]
        row.cells[0].text = "Left"
        row.cells[1].text = "Right"
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        engine = RedlineEngine(stream)
        # Attempt a structural column insertion via text replace
        edit = ModifyText(target_text="Left | Right", new_text="Left | Right | Extra")
        stats = engine.process_batch([edit])

        assert stats["edits_skipped"] == 1
        skipped_msg = stats["skipped_details"][0].lower()

        assert "rows or columns" in skipped_msg, f"Error message lacks 'rows or columns'. Got: {skipped_msg}"

    @pytest.mark.skipif(sys.platform != "win32", reason="Live Word only supported on Windows")
    def test_tech1_live_com_author_warning(self):
        """
        TECH-1: Live COM mode silently ignores author_name overrides for tracked changes.
        The tool must surface this warning in the string response so the agent is aware.
        """
        import asyncio

        from fastmcp import Context

        from adeu.mcp_components.tools.live_word import process_active_word_batch

        ctx = MagicMock(spec=Context)

        with patch("adeu.mcp_components.tools.live_word._process_active_word_batch_core") as mock_core:
            with patch("win32com.client.GetActiveObject") as mock_get_obj:
                mock_app = MagicMock()
                mock_app.UserName = "Real M365 User"
                mock_get_obj.return_value = mock_app

                mock_core.return_value = {
                    "applied": 1,
                    "failed": 0,
                    "skipped_details": [],
                    "author_overridden_by_word": "Real M365 User",
                }

                edit = ModifyText(target_text="A", new_text="B")
                res = asyncio.run(process_active_word_batch(ctx, changes=[edit], author_name="Reviewer AI"))

                assert "author_overridden_by_word" in res or "overridden" in res.lower(), (
                    "Missing author override warning in live COM batch response (TECH-1)"
                )

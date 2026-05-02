import io

from docx import Document

from adeu.ingest import extract_text_from_stream
from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine


def test_pure_comment_with_formatting_mismatch():
    """
    Validates that a ModifyText payload with identical target_text and new_text
    is treated purely as a COMMENT_ONLY operation, even if the actual text in the
    document contains formatting markers not present in the payload.
    """
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Important ")
    r1.bold = True
    p.add_run("Text")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Target and new are plain text, omitting the ** markers
    edit = ModifyText(
        target_text="Important Text",
        new_text="Important Text",
        comment="This is a pure comment.",
    )

    engine = RedlineEngine(stream, author="QA")
    engine.apply_edits([edit])

    res_stream = engine.save_to_stream()
    text = extract_text_from_stream(res_stream, clean_view=False)

    # No redlines should be emitted
    assert "{--" not in text
    assert "{++" not in text

    # The formatting should survive, the comment should be attached, and ingest
    # will wrap the differently-styled runs in separate highlight blocks.
    assert "{==**Important **==}" in text
    assert "{==Text==}" in text
    assert "This is a pure comment." in text


def test_pure_comment_no_op():
    """
    Validates that a ModifyText payload with identical target and new text
    and NO comment is safely ignored as a no-op without corrupting the document.
    """
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("Important Text")
    r.bold = True

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit = ModifyText(target_text="Important Text", new_text="Important Text")

    engine = RedlineEngine(stream, author="QA")
    engine.apply_edits([edit])

    res_stream = engine.save_to_stream()
    text = extract_text_from_stream(res_stream, clean_view=False)

    # No redlines
    assert "{--" not in text
    assert "{++" not in text

    # Formatting preserved perfectly
    assert "**Important Text**" in text

import io

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from adeu.ingest import extract_text_from_stream
from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine


def test_edit_on_accepted_view_skipping_deletion():
    """
    Scenario:
    Raw: "A {--deleted--} B"
    Clean: "A B"
    Edit Target: "A B"
    Edit New: "A C"

    The engine should match against Clean view, bridge the gap, and produce:
    "A {--deleted--} {--B--}{++C++}" or similar valid XML.
    """
    doc = Document()
    p = doc.add_paragraph()

    # Run A
    p.add_run("A ")

    # Run Deleted (Manual XML)
    d = OxmlElement("w:del")
    d.set(qn("w:id"), "1")
    d.set(qn("w:author"), "X")
    rd = OxmlElement("w:r")
    rt = OxmlElement("w:delText")
    rt.text = "deleted "
    rd.append(rt)
    d.append(rd)
    p._element.append(d)

    # Run B
    p.add_run("B")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Verify Raw View
    raw_text = extract_text_from_stream(stream, clean_view=False)
    assert "{--deleted --}" in raw_text

    # Verify Clean View
    clean_text = extract_text_from_stream(stream, clean_view=True)
    assert "deleted" not in clean_text
    assert "A B" in clean_text

    # Apply Edit targeting Clean View
    edit = ModifyText(target_text="A B", new_text="A C")

    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])

    assert applied == 1
    assert skipped == 0

    res_stream = engine.save_to_stream()
    final_text = extract_text_from_stream(res_stream)

    # Expect: A {--deleted --}{--B--}{++C++}
    # Or: A {--deleted --}{--B--}{++C++}
    assert "{--deleted" in final_text
    assert "{--deleted B--}" in final_text
    assert "{++C++}" in final_text
    assert "{++C++}" in final_text
    assert "A " in final_text

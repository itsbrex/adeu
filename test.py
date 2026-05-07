# FILE: repro_bug5.py
"""
Reproduction script for Bug 5: rejecting a multi-paragraph insertion leaves
an empty paragraph behind.

The failing test (test_issue_9_reject_multi_paragraph_leaves_break):
  1. Doc has two paragraphs: "Paragraph 1." and "Paragraph 2.".
  2. ModifyText("Paragraph 1.", "Paragraph 1.\\n\\nNew Para.") is applied.
     This is a same-text-prefix insertion that adds a tracked paragraph break
     plus "New Para.".
  3. The resulting Chg ID is found from mapper spans.
  4. RejectChange is applied for that Chg ID.
  5. Test expects [Paragraph 1., Paragraph 2.] — i.e., full revert. The bug:
     an empty paragraph is left behind.

Hypothesis: when track_insert injects a multi-paragraph insertion, it creates
a NEW <w:p> sibling and injects an <w:ins> marker inside that new paragraph's
pPr/rPr to track the paragraph BREAK ITSELF. On reject, the engine's
_reject_change function only removes <w:ins> elements from the body content;
it does not look inside pPr/rPr for paragraph-break markers, and even if it
removed them, it would not remove the orphan <w:p> that those markers
belonged to.

This script:
  1. Reproduces the test scenario.
  2. Dumps the OOXML of every paragraph after process_batch (the "inserted"
     state), so we can confirm the multi-paragraph insertion produced the
     expected pPr/rPr/<w:ins> structure.
  3. Lists all <w:ins> elements in the body, distinguishing inline ones from
     pPr/rPr (paragraph-break) ones, with their w:id values.
  4. Finds and rejects the Chg ID from mapper spans.
  5. Dumps every paragraph after rejection.
  6. Reports the test assertion result.
"""

import io

import lxml.etree as etree
from docx import Document
from docx.oxml.ns import qn

from adeu.models import ModifyText, RejectChange
from adeu.redline.engine import RedlineEngine


def section(title: str) -> None:
    print()
    print("=" * 72)
    print(title)
    print("=" * 72)


def dump_paragraphs(engine: RedlineEngine, label: str) -> None:
    section(label)
    paragraphs = list(engine.doc.paragraphs)
    print(f"paragraph count: {len(paragraphs)}")
    for i, p in enumerate(paragraphs):
        print(f"  [{i}] text={p.text!r}")

    print()
    print("--- raw OOXML for each <w:p> in body ---")
    body = engine.doc.element.body
    p_elements = body.findall(qn("w:p"))
    for i, p_el in enumerate(p_elements):
        print(f"  [{i}]")
        xml = etree.tostring(p_el, pretty_print=True).decode("utf-8")
        # Strip the lxml namespace declaration noise for readability
        for line in xml.splitlines():
            print(f"    {line}")


def list_all_ins_elements(engine: RedlineEngine, label: str) -> None:
    section(label)
    body = engine.doc.element.body
    ins_elements = body.findall(f".//{qn('w:ins')}")
    print(f"total <w:ins> elements in body: {len(ins_elements)}")
    for i, ins in enumerate(ins_elements):
        # Determine context: is this inline content, or a pPr/rPr marker?
        parent = ins.getparent()
        grandparent = parent.getparent() if parent is not None else None
        context = "?"
        if parent is not None:
            if parent.tag == qn("w:rPr") and grandparent is not None and grandparent.tag == qn("w:pPr"):
                context = "pPr/rPr (paragraph-break marker)"
            elif parent.tag == qn("w:p"):
                context = "inline (direct child of w:p)"
            else:
                context = f"parent={parent.tag.split('}')[-1]}"
        wid = ins.get(qn("w:id"))
        author = ins.get(qn("w:author"))
        # Children summary
        child_tags = [c.tag.split("}")[-1] for c in ins]
        print(f"  [{i}]  w:id={wid}  author={author!r}  context={context}  children={child_tags}")


def main() -> None:
    # --- Build the same doc the failing test builds ---
    doc = Document()
    doc.add_paragraph("Paragraph 1.")
    doc.add_paragraph("Paragraph 2.")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)

    # --- Step 1: apply the multi-paragraph insertion ---
    section("STEP 1 — apply multi-paragraph ModifyText")
    edit = ModifyText(
        target_text="Paragraph 1.",
        new_text="Paragraph 1.\n\nNew Para.",
    )
    result = engine.process_batch([edit])
    print(f"process_batch stats: {result}")

    # --- Step 2: state after insertion ---
    dump_paragraphs(engine, "STATE AFTER process_batch (insertion applied)")
    list_all_ins_elements(engine, "ALL <w:ins> ELEMENTS AFTER INSERTION")

    # Rebuild map (the test does this) and find the Chg ID
    engine.mapper._build_map()
    chg_id = None
    for span in engine.mapper.spans:
        if span.ins_id:
            chg_id = f"Chg:{span.ins_id}"
            break
    print()
    print(f"Discovered Chg ID from mapper spans: {chg_id}")

    if chg_id is None:
        print("FATAL: no insertion found in mapper spans; cannot proceed.")
        return

    # --- Step 3: reject the change ---
    section(f"STEP 3 — RejectChange({chg_id})")
    reject_result = engine.process_batch([RejectChange(target_id=chg_id)])
    print(f"process_batch stats: {reject_result}")

    # --- Step 4: state after rejection ---
    dump_paragraphs(engine, "STATE AFTER RejectChange (the bug shows here)")
    list_all_ins_elements(engine, "ALL <w:ins> ELEMENTS AFTER REJECTION")

    # --- Step 5: test assertion preview ---
    section("TEST ASSERTION PREVIEW")
    paragraphs = [p.text for p in engine.doc.paragraphs]
    expected = ["Paragraph 1.", "Paragraph 2."]
    print(f"  paragraphs: {paragraphs}")
    print(f"  expected:   {expected}")
    if paragraphs == expected:
        print("  [PASS]")
    else:
        print("  [FAIL] mismatch")


if __name__ == "__main__":
    main()

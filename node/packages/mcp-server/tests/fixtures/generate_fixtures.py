import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

os.makedirs("node/packages/mcp-server/tests/fixtures", exist_ok=True)

doc = Document()
doc.add_paragraph("Section One").style = doc.styles["Heading 1"]
p = doc.add_paragraph()


def run(text):
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    return r


p._p.append(run("Foo bar "))

d = OxmlElement("w:del")
d.set(qn("w:id"), "10")
d.set(qn("w:author"), "Test Negotiator")
d.set(qn("w:date"), "2026-01-22T12:06:55Z")

rd = OxmlElement("w:r")
dt = OxmlElement("w:delText")
dt.text = "old phrase here."
dt.set(qn("xml:space"), "preserve")
rd.append(dt)
d.append(rd)
p._p.append(d)

ins = OxmlElement("w:ins")
ins.set(qn("w:id"), "11")
ins.set(qn("w:author"), "Test Negotiator")
ins.set(qn("w:date"), "2026-01-22T12:06:55Z")
ins.append(run("new phrase here."))
p._p.append(ins)

doc.save("node/packages/mcp-server/tests/fixtures/gap2_minimal_repro.docx")
print(
    "Successfully generated node/packages/mcp-server/tests/fixtures/gap2_minimal_repro.docx"
)

# Generate GAP 1 deleted row fixture
doc1 = Document()
doc1.add_paragraph("Active Heading").style = doc1.styles["Heading 1"]

# Add a table
table = doc1.add_table(rows=1, cols=1)
row = table.rows[0]
cell = row.cells[0]

# Add a paragraph inside the cell with Heading 1 style
p_cell = cell.paragraphs[0]
p_cell.style = doc1.styles["Heading 1"]
# Clear default text and append the text
p_cell.text = "Deleted Heading"

# Now let's mark the row as deleted (w:del inside w:trPr)
trPr = row._tr.get_or_add_trPr()
del_node = OxmlElement("w:del")
del_node.set(qn("w:id"), "100")
del_node.set(qn("w:author"), "Test Negotiator")
del_node.set(qn("w:date"), "2026-01-22T12:06:55Z")
trPr.append(del_node)

doc1.save("node/packages/mcp-server/tests/fixtures/gap1_deleted_row_repro.docx")
print(
    "Successfully generated node/packages/mcp-server/tests/fixtures/gap1_deleted_row_repro.docx"
)

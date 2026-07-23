# FILE: tests/test_repro_qa_mcp_2026_07_23_engine.py
"""
Python-engine mirror of ADEU-MCP-QA-REPORT.md (2026-07-23, black-box QA of the
Node MCP server, v1.29.0+4bb70f9) — engine-level findings that reproduce on the
Python engine too:

  F1   Reject of a multi-paragraph replacement corrupts the document.
       Python variant (verified 2026-07-23): the replacement's deletion and its
       multi-paragraph insertion carry two UNGROUPED ids (the del and the ins
       live in different paragraphs, so the contiguous-sibling pair walk never
       groups them). Rejecting the deletion id restores the original sentence
       but leaves every inserted paragraph pending — the clean view then shows
       BOTH the restored sentence and its full replacement (the QA report's
       "duplicate content" corruption). Additionally, the apply itself strands
       the original paragraph as a "."-only container (the trailing period is
       trimmed as common context and left behind) while the replacement's last
       sentence loses its final period.
       (The report's w:pPrChange sub-item is Node-specific: the Python engine
       puts the heading style on NEW tracked paragraphs instead of restyling
       the original paragraph in place, so there is no untracked restyle to
       pin here.)
  F8   When a later edit fails validation after earlier edits applied, the
       error note says "N earlier edit(s) in this batch were already applied"
       — while the batch was in fact rolled back and nothing was saved. The
       note must say so instead of implying partial persistence.
       (engine.py sequential-contract hint, ~line 2236)
  F20  Change ids number in REVERSE document order when one edit fans out via
       match_mode="all" (verified: separate sequential edits number ascending;
       the backwards bottom-up sweep assigns ids at mutation time, so a
       fan-out numbers Chg:5/6, 3/4, 1/2 for the 1st/2nd/3rd occurrence).

Every test is written test-first: it fails on current main and passes once the
finding is fixed.
"""

import re
from io import BytesIO

import pytest
from docx import Document

from adeu.ingest import extract_text_from_stream
from adeu.models import ModifyText, RejectChange
from adeu.redline.engine import BatchValidationError, RedlineEngine

INTRO = "Intro paragraph stays untouched."
ORIG_SENTENCE = "The parties shall negotiate disputes in good faith."
TAIL = "Tail paragraph stays untouched."

REPLACEMENT = (
    "## Governing Law\n\n"
    "This Agreement is governed by the laws of Finland.\n\n"
    "Any dispute shall be settled in the courts of Helsinki."
)


def build_three_paragraph_doc() -> BytesIO:
    d = Document()
    d.add_paragraph(INTRO)
    d.add_paragraph(ORIG_SENTENCE)
    d.add_paragraph(TAIL)
    buf = BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf


def apply_multi_paragraph_replacement() -> BytesIO:
    """One ModifyText replacing a body sentence with heading + two sentences."""
    engine = RedlineEngine(build_three_paragraph_doc(), author="QA Agent")
    stats = engine.process_batch([ModifyText(target_text=ORIG_SENTENCE, new_text=REPLACEMENT)])
    assert stats["edits_applied"] == 1, "setup: the replacement edit must apply"
    return engine.save_to_stream()


class TestF1MultiParagraphReplacementReject:
    def test_apply_produces_a_faithful_clean_view(self):
        """The accepted view of the replacement must not strand a '.'-only
        paragraph nor drop the replacement's final period (F1 apply-side)."""
        stream = apply_multi_paragraph_replacement()
        clean = extract_text_from_stream(stream, clean_view=True)

        assert ORIG_SENTENCE not in clean, "setup: the original sentence was replaced"
        assert "the courts of Helsinki." in clean, (
            "the replacement's final period was lost (paired away with the "
            f"original sentence's trailing period):\n{clean}"
        )
        stray = [ln for ln in clean.splitlines() if ln.strip() == "."]
        assert not stray, (
            f"the original paragraph survives as a stranded '.'-only container in the accepted view:\n{clean}"
        )

    def test_reject_unwinds_the_whole_replacement(self):
        """Rejecting the modification (addressed by its deletion id, the
        realistic agent action) must return the document to its pre-edit
        state: no pending insertions, no duplicated content (F1 reject-side).
        """
        stream = apply_multi_paragraph_replacement()
        raw = extract_text_from_stream(stream, clean_view=False)
        del_ids = re.findall(r"\[Chg:(\d+) delete\]", raw)
        assert del_ids, f"setup: expected a tracked deletion bubble in:\n{raw}"

        stream.seek(0)
        engine = RedlineEngine(stream, author="Counterparty")
        engine.apply_review_actions([RejectChange(target_id=del_ids[0])])
        out = engine.save_to_stream()

        raw_after = extract_text_from_stream(out, clean_view=False)
        assert "{++" not in raw_after and "{--" not in raw_after, (
            "rejecting the replacement left part of it pending — the inserted "
            "paragraphs survive as orphan insertions (accepting the remainder "
            f"duplicates the restored sentence):\n{raw_after}"
        )

        out.seek(0)
        clean_after = extract_text_from_stream(out, clean_view=True)
        expected = f"{INTRO}\n\n{ORIG_SENTENCE}\n\n{TAIL}"
        assert clean_after.strip() == expected, (
            "reject did not restore the pre-edit document:\n"
            f"--- expected ---\n{expected}\n--- got ---\n{clean_after.strip()}"
        )


class TestF8RejectedBatchNote:
    def test_note_states_nothing_was_saved(self):
        """The sequential-contract hint reads "1 earlier edit(s) ... were
        already applied" on a batch that was rolled back; it must state that
        nothing was saved."""
        engine = RedlineEngine(build_three_paragraph_doc(), author="QA Agent")
        with pytest.raises(BatchValidationError) as exc_info:
            engine.process_batch(
                [
                    ModifyText(target_text="good faith", new_text="the utmost good faith"),
                    ModifyText(
                        target_text="THIS TEXT DOES NOT EXIST ANYWHERE",
                        new_text="irrelevant",
                    ),
                ]
            )
        message = str(exc_info.value)

        # Sanity (passes on main): the rollback itself does happen.
        after = extract_text_from_stream(engine.save_to_stream(), clean_view=True)
        assert "utmost" not in after, "setup: the batch must be rolled back"

        assert re.search(
            r"nothing (was |has been )?(saved|written)|rolled back|not (been )?saved",
            message,
            re.IGNORECASE,
        ), (
            "the rejection note implies earlier edits persisted; it must state "
            f"the batch was rolled back / nothing was saved. Message was:\n{message}"
        )


class TestF20ChangeIdOrdering:
    def test_match_mode_all_ids_ascend_in_document_order(self):
        """A match_mode='all' fan-out numbers its occurrences bottom-up
        (Chg:5/6, 3/4, 1/2 for the 1st/2nd/3rd occurrence); ids must ascend in
        document order like sequential separate edits already do."""
        d = Document()
        d.add_paragraph("alpha apple one.")
        d.add_paragraph("beta apple two.")
        d.add_paragraph("gamma apple three.")
        buf = BytesIO()
        d.save(buf)
        buf.seek(0)

        engine = RedlineEngine(buf, author="QA Agent")
        stats = engine.process_batch([ModifyText(target_text="apple", new_text="pear", match_mode="all")])
        assert stats["edits_applied"] == 1, "setup: the fan-out edit must apply"

        raw = extract_text_from_stream(engine.save_to_stream(), clean_view=False)
        del_ids = [int(m) for m in re.findall(r"\[Chg:(\d+) delete\]", raw)]
        assert len(del_ids) == 3, f"setup: expected three deletion bubbles in:\n{raw}"
        assert del_ids == sorted(del_ids), (
            "change ids must ascend in document order; projection order was "
            f"{del_ids} (reverse numbering makes ids read as if the last "
            f"occurrence were edited first):\n{raw}"
        )

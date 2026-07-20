# FILE: tests/test_repro_qa_report_v9.py
"""
Repro tests for the 2026-07-19 black-box QA and UX evaluation of 1.27.0
(adeu 1.27.0+7a0a821, "Adeu CLI Black-Box QA and UX Evaluation").

Finding index (report finding -> test class below):
  ADEU-QA-001  sanitize reports `Result: CLEAN` while a `w:docVar` secret
               survives in word/settings.xml (Critical/High)
  ADEU-QA-002  DOCX-to-DOCX `diff --json` output is not reliably applicable
               by `apply`: paragraph deletions and reorderings fail (High).
               Three mechanisms:
                 A. word-level hunks cross paragraph boundaries misaligned
                    ("...arrears.\n\nThe ") and are rightly rejected;
                 B. deleting a full paragraph merges the FOLLOWING paragraph
                    into the deleted one's container, bleeding its style
                    ("## The first body paragraph...");
                 C. later edits in a replayed batch match ambiguously against
                    text introduced by EARLIER edits of the same batch — and
                    against virtual meta-bubble text (timestamps matching "4")
  ADEU-QA-003  a failed post-write verification leaves the requested output
               file behind (exit 1 but file exists) (Medium)
  ADEU-QA-004  a replacement's del+ins pair exposes two IDs that imply
               independent resolution; resolving one side decides both, and a
               follow-up action on the paired ID is still counted "applied"
               (Medium)
  ADEU-QA-005  page semantics vary across modes and outputs never say pages
               are synthetic (Medium/Low)
  ADEU-QA-006  `adeu ... | head` prints a BrokenPipeError traceback (Low)

Every test fails against the commit preceding its fix.
"""

import json
import subprocess
import sys
import zipfile
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from pydantic import TypeAdapter

from adeu.diff import generate_edits_via_paragraph_alignment, generate_structured_edits
from adeu.ingest import _extract_text_from_doc, extract_text_from_stream
from adeu.models import AcceptChange, ModifyText, RejectChange, StrictBatchChanges
from adeu.redline.engine import BatchValidationError, RedlineEngine
from adeu.redline.mapper import DocumentMapper
from adeu.sanitize.core import SanitizeError, sanitize_docx

# ---------------------------------------------------------------------------
# Helpers / fixture builders
# ---------------------------------------------------------------------------

SECRET = "SECRET-MATTER-4711-CONFIDENTIAL"


def run_cli(args, capsys):
    """Invoke the CLI in-process; returns (exit_code, stdout, stderr)."""
    from unittest.mock import patch

    from adeu.cli import main

    code = 0
    with patch.object(sys, "argv", ["adeu"] + [str(a) for a in args]):
        try:
            main()
        except SystemExit as e:
            code = e.code or 0
    captured = capsys.readouterr()
    return code, captured.out, captured.err


def doc_to_stream(doc) -> BytesIO:
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_docx(paragraphs, path: Path = None):
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if path is not None:
        doc.save(path)
    return doc


def clean_text_of(doc) -> str:
    return _extract_text_from_doc(doc, clean_view=True, include_appendix=False)


def build_docvar_docx(path: Path, secret: str = SECRET):
    """A valid DOCX carrying a Word document variable with a secret value in
    word/settings.xml — the report's ADEU-QA-001 fixture."""
    doc = Document()
    doc.add_paragraph("An ordinary contract paragraph.")
    buf = BytesIO()
    doc.save(buf)

    src = zipfile.ZipFile(BytesIO(buf.getvalue()))
    out = BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
        for item in src.infolist():
            data = src.read(item.filename)
            if item.filename == "word/settings.xml":
                text = data.decode("utf-8")
                inject = f'<w:docVars><w:docVar w:name="MatterRef" w:val="{secret}"/></w:docVars>'
                idx = text.index(">", text.index("<w:settings")) + 1
                text = text[:idx] + inject + text[idx:]
                data = text.encode("utf-8")
            z.writestr(item, data)
    path.write_bytes(out.getvalue())


def secret_in_package(path: Path, secret: str = SECRET) -> bool:
    with zipfile.ZipFile(path) as z:
        for name in z.namelist():
            if secret.encode("utf-8") in z.read(name):
                return True
    return False


def build_rich_docx(paragraph_specs, path: Path):
    """spec: list of (kind, content); kind in heading/plain/styled/list.
    'styled' content is a list of (text, style) with style in "", "b", "i"."""
    doc = Document()
    for kind, content in paragraph_specs:
        if kind == "heading":
            doc.add_heading(content, level=2)
        elif kind == "styled":
            p = doc.add_paragraph()
            for text, style in content:
                run = p.add_run(text)
                if "b" in style:
                    run.bold = True
                if "i" in style:
                    run.italic = True
        elif kind == "list":
            doc.add_paragraph(content, style="List Number")
        else:
            doc.add_paragraph(content)
    doc.save(str(path))


RICH_BASE = [
    ("heading", "1. Interest and Payment"),
    (
        "styled",
        [
            ("The outstanding balance accrues interest at a rate of ", ""),
            ("four percent (4%) per annum", "i"),
            (", compounded monthly and payable in arrears.", ""),
        ],
    ),
    ("plain", "The Supplier shall invoice the Customer at the end of each calendar month."),
    (
        "styled",
        [
            ("Late payments accrue a penalty of ", ""),
            ("two percent (2%) per annum", "i"),
            (" above the base rate, as defined by the ", ""),
            ("European Central Bank", "b"),
            (".", ""),
        ],
    ),
    ("plain", "All notices must be delivered in writing to the registered office."),
    ("list", "First list item covering delivery obligations."),
    ("list", "Second list item covering acceptance testing."),
    ("plain", "Thís paragraph — with unicode punctuation ‘quotes’ — stays."),
]


def json_round_trip(edits):
    """Exactly what `adeu diff --json > f.json && adeu apply orig f.json`
    does to an edit batch: serialize without private pins, re-validate."""
    payload = json.dumps([e.model_dump(exclude={"_match_start_index"}) for e in edits])
    return TypeAdapter(StrictBatchChanges).validate_python(json.loads(payload))


def structured_roundtrip(tmp_path: Path, name: str, orig_spec, mod_spec):
    """Runs the full docx-to-docx pipeline: build both fixtures, diff
    structurally, JSON round-trip, apply to the original, return
    (expected_clean_text, actual_clean_text, warnings)."""
    orig = tmp_path / f"{name}_orig.docx"
    mod = tmp_path / f"{name}_mod.docx"
    build_rich_docx(orig_spec, orig)
    build_rich_docx(mod_spec, mod)

    text_o, struct_o = _extract_text_from_doc(
        Document(str(orig)), clean_view=True, include_appendix=False, return_structure=True
    )
    text_m, struct_m = _extract_text_from_doc(
        Document(str(mod)), clean_view=True, include_appendix=False, return_structure=True
    )
    edits, warnings = generate_structured_edits(text_o, struct_o, text_m, struct_m)

    engine = RedlineEngine(BytesIO(orig.read_bytes()), author="QA")
    engine.process_batch(list(json_round_trip(edits)))  # raises on rejection

    final_clean = _extract_text_from_doc(engine.doc, clean_view=True, include_appendix=False)
    return text_m.strip(), final_clean.strip(), warnings


def spec_delete(idx):
    return [s for i, s in enumerate(RICH_BASE) if i != idx]


def spec_swap(i, j):
    out = list(RICH_BASE)
    out[i], out[j] = out[j], out[i]
    return out


def build_tracked_modification_stream() -> BytesIO:
    """A document carrying exactly one tracked modification: the del+ins
    pair (Chg:1 + Chg:2) of `30 days` -> `60 days`."""
    doc = build_docx(["Payment is due in 30 days.", "Second paragraph here."])
    engine = RedlineEngine(doc_to_stream(doc), author="Editor")
    engine.apply_edits([ModifyText(target_text="30 days", new_text="60 days")])
    return engine.save_to_stream()


# ---------------------------------------------------------------------------
# ADEU-QA-001: sanitizer must not report CLEAN while w:docVar secrets remain
# ---------------------------------------------------------------------------


class TestQA001SanitizerDocVars:
    def test_full_sanitize_removes_document_variables(self, tmp_path):
        inp = tmp_path / "docvar.docx"
        outp = tmp_path / "docvar_sanitized.docx"
        build_docvar_docx(inp)
        assert secret_in_package(inp), "fixture must carry the secret"

        result = sanitize_docx(str(inp), str(outp))

        assert not secret_in_package(outp), (
            "sanitize reported %r but the w:docVar secret is still in the package" % result.status
        )

    def test_keep_markup_sanitize_removes_document_variables(self, tmp_path):
        inp = tmp_path / "docvar_km.docx"
        outp = tmp_path / "docvar_km_sanitized.docx"
        build_docvar_docx(inp)

        sanitize_docx(str(inp), str(outp), keep_markup=True)

        assert not secret_in_package(outp), "keep-markup sanitize left the w:docVar secret in the package"

    def test_report_names_variable_but_never_its_value(self, tmp_path):
        inp = tmp_path / "docvar_rep.docx"
        outp = tmp_path / "docvar_rep_sanitized.docx"
        build_docvar_docx(inp)

        result = sanitize_docx(str(inp), str(outp))

        assert "document variable" in result.report_text.lower(), (
            "the report must disclose that document variables were found/removed:\n" + result.report_text
        )
        assert "MatterRef" in result.report_text, "the removed variable should be named for review"
        assert SECRET not in result.report_text, "the report must never echo the secret value"

    def test_verification_fails_closed_if_docvars_survive(self, tmp_path, monkeypatch):
        """The CLEAN verdict must come from re-scanning the SAVED bytes: if the
        transform is (hypothetically) broken, sanitize must raise instead of
        writing output + printing CLEAN."""
        from adeu.sanitize import transforms

        inp = tmp_path / "docvar_fc.docx"
        outp = tmp_path / "docvar_fc_sanitized.docx"
        build_docvar_docx(inp)

        monkeypatch.setattr(transforms, "strip_document_variables", lambda doc: [])

        with pytest.raises(SanitizeError):
            sanitize_docx(str(inp), str(outp))
        assert not outp.exists(), "no output may be written when verification fails"


# ---------------------------------------------------------------------------
# ADEU-QA-002: structured diff output must replay through apply
# ---------------------------------------------------------------------------


class TestQA002StructuredDiffReplay:
    """apply(source, json(diff(source, target))) must reproduce target for
    paragraph deletions and reorderings over realistic content (styled runs,
    headings, lists, unicode) — the report measured 0/8 for these."""

    @pytest.mark.parametrize("idx", range(len(RICH_BASE)))
    def test_paragraph_deletion_replays(self, tmp_path, idx):
        expected, actual, _ = structured_roundtrip(tmp_path, f"del{idx}", RICH_BASE, spec_delete(idx))
        assert actual == expected

    @pytest.mark.parametrize(
        "name,i,j",
        [
            ("swap_styled", 1, 3),
            ("swap_plain", 2, 4),
            ("swap_list_items", 5, 6),
        ],
    )
    def test_paragraph_swap_replays(self, tmp_path, name, i, j):
        expected, actual, _ = structured_roundtrip(tmp_path, name, RICH_BASE, spec_swap(i, j))
        assert actual == expected

    def test_move_styled_paragraph_to_end_replays(self, tmp_path):
        moved = [s for i, s in enumerate(RICH_BASE) if i != 1] + [RICH_BASE[1]]
        expected, actual, _ = structured_roundtrip(tmp_path, "move_end", RICH_BASE, moved)
        assert actual == expected

    def test_move_heading_down_replays(self, tmp_path):
        moved = RICH_BASE[1:3] + [RICH_BASE[0]] + RICH_BASE[3:]
        expected, actual, _ = structured_roundtrip(tmp_path, "move_heading", RICH_BASE, moved)
        assert actual == expected


class TestQA002ParagraphDeletionKeepsFollowingStyle:
    """Mechanism B (both diff paths): deleting a full paragraph must not
    restyle the paragraph that follows it."""

    def test_deleting_heading_paragraph_keeps_body_plain(self, tmp_path):
        doc = Document()
        doc.add_heading("Definitions", level=2)
        doc.add_paragraph("The first body paragraph after the heading.")
        doc.add_paragraph("A second body paragraph.")
        src = tmp_path / "heading.docx"
        doc.save(str(src))

        text_o = clean_text_of(Document(str(src)))
        text_m = text_o.replace("## Definitions\n\n", "")

        edits = generate_edits_via_paragraph_alignment(text_o, text_m)
        engine = RedlineEngine(BytesIO(src.read_bytes()), author="QA")
        engine.process_batch(list(edits))

        final = clean_text_of(engine.doc)
        assert final.strip() == text_m.strip(), (
            "deleting the heading paragraph must not promote the following body text to a heading"
        )

    def test_deleting_list_item_keeps_following_paragraph_plain(self, tmp_path):
        doc = Document()
        doc.add_paragraph("Intro paragraph.")
        doc.add_paragraph("Item one of the list.", style="List Number")
        doc.add_paragraph("Closing plain paragraph.")
        src = tmp_path / "list.docx"
        doc.save(str(src))

        text_o = clean_text_of(Document(str(src)))
        text_m = text_o.replace("1. Item one of the list.\n\n", "")

        edits = generate_edits_via_paragraph_alignment(text_o, text_m)
        engine = RedlineEngine(BytesIO(src.read_bytes()), author="QA")
        engine.process_batch(list(edits))

        final = clean_text_of(engine.doc)
        assert final.strip() == text_m.strip(), (
            "deleting a list item must not turn the following plain paragraph into a list item"
        )

    def test_deleting_plain_paragraph_before_list_keeps_list_numbering(self, tmp_path):
        doc = Document()
        doc.add_paragraph("A paragraph to delete.")
        doc.add_paragraph("Surviving list item.", style="List Number")
        src = tmp_path / "plainlist.docx"
        doc.save(str(src))

        text_o = clean_text_of(Document(str(src)))
        text_m = text_o.replace("A paragraph to delete.\n\n", "")

        edits = generate_edits_via_paragraph_alignment(text_o, text_m)
        engine = RedlineEngine(BytesIO(src.read_bytes()), author="QA")
        engine.process_batch(list(edits))

        final = clean_text_of(engine.doc)
        assert final.strip() == text_m.strip(), "the surviving list item must keep its numbering"


class TestQA002VirtualTextNeverMatches:
    """Mechanism C (matcher): text that exists only in virtual projection
    chrome (meta bubbles, timestamps, author names) is not document text and
    must neither satisfy nor ambiguate a match."""

    def test_bubble_only_target_is_not_found(self):
        stream = build_tracked_modification_stream()
        engine = RedlineEngine(stream, author="Reviewer")
        # "Editor" appears ONLY inside the {>>[Chg:1 delete] Editor ...<<}
        # bubble — never in document text.
        with pytest.raises(BatchValidationError) as exc:
            engine.process_batch([ModifyText(target_text="Editor", new_text="X")])
        assert "not found" in "\n".join(exc.value.errors).lower()

    def test_digit_unique_in_body_resolves_despite_bubble_timestamps(self):
        doc = build_docx(["Clause 7 applies to this agreement.", "Another paragraph."])
        engine = RedlineEngine(doc_to_stream(doc), author="Editor")
        engine.process_batch(
            [ModifyText(target_text="Another", new_text="A different", comment="Reviewed on 2026-07-19")]
        )
        stream = engine.save_to_stream()

        # The comment bubble now carries a timestamp full of digits. "7" is
        # unique in the real document text and must resolve strictly.
        engine2 = RedlineEngine(stream, author="Reviewer")
        stats = engine2.process_batch([ModifyText(target_text="7", new_text="9")])
        assert stats["edits_applied"] == 1
        final = clean_text_of(engine2.doc)
        assert "Clause 9 applies" in final


# ---------------------------------------------------------------------------
# ADEU-QA-003: failed verification must not leave the requested output file
# ---------------------------------------------------------------------------


class TestQA003TransactionalOutput:
    def _build_small_structured(self, path: Path):
        doc = Document()
        doc.add_heading("Small Doc", level=1)
        doc.add_paragraph("Some prose here.")
        t = doc.add_table(rows=1, cols=2)
        t.rows[0].cells[0].text = "K"
        t.rows[0].cells[1].text = "V"
        doc.save(path)

    def test_failed_verification_leaves_no_file_at_requested_path(self, tmp_path, capsys):
        src = tmp_path / "small.docx"
        self._build_small_structured(src)
        tiny = tmp_path / "tiny.txt"
        tiny.write_text("tiny\n", encoding="utf-8")
        out = tmp_path / "result.docx"

        code, stdout, _ = run_cli(["apply", src, tiny, "-o", out, "--json", "--allow-major-deletions"], capsys)

        assert code != 0
        assert not out.exists(), (
            "exit status says the apply failed, but a file exists at the requested output path — "
            "automation checking file existence would consume a wrong document"
        )
        stats = json.loads(stdout)
        assert stats.get("verified") is False
        # The diagnostic copy must live at a clearly-not-the-output path.
        diag = stats.get("unverified_output_path")
        assert diag, "the JSON result must say where the diagnostic copy went"
        assert Path(diag).exists()
        assert Path(diag) != out
        assert ".unverified" in Path(diag).name
        assert stats.get("output_path") is None, "output_path must not point at anything on failure"

    def test_successful_apply_still_writes_requested_path(self, tmp_path, capsys):
        src = tmp_path / "plain.docx"
        build_docx(["Alpha paragraph.", "Beta paragraph."], src)
        modified = tmp_path / "mod.txt"
        modified.write_text("Alpha paragraph.\n\nGamma paragraph.", encoding="utf-8")
        out = tmp_path / "result.docx"

        code, stdout, _ = run_cli(["apply", src, modified, "-o", out, "--json"], capsys)
        assert code == 0
        assert out.exists()
        stats = json.loads(stdout)
        assert stats.get("verified") is True
        assert stats.get("output_path") == str(out)
        assert not stats.get("unverified_output_path")


# ---------------------------------------------------------------------------
# ADEU-QA-004: replacement del+ins pairs — linkage, counts, contradictions
# ---------------------------------------------------------------------------


class TestQA004ReplacementPairSemantics:
    def test_projection_annotates_paired_changes(self):
        stream = build_tracked_modification_stream()
        raw = extract_text_from_stream(stream)
        assert "[Chg:1 delete] Editor (pairs with Chg:2)" in raw, raw
        assert "[Chg:2 insert] Editor (pairs with Chg:1)" in raw, raw

    def test_mapper_projection_matches_ingest_with_pairing(self):
        stream = build_tracked_modification_stream()
        raw = extract_text_from_stream(BytesIO(stream.getvalue()))
        mapper = DocumentMapper(Document(BytesIO(stream.getvalue())))
        assert mapper.full_text == raw, "Virtual Text contract: ingest and mapper must render pairing identically"

    def test_standalone_change_is_not_annotated(self):
        doc = build_docx(["Payment is due in 30 days."])
        engine = RedlineEngine(doc_to_stream(doc), author="Editor")
        engine.apply_edits([ModifyText(target_text="Payment is due in 30 days.", new_text="")])
        raw = extract_text_from_stream(engine.save_to_stream())
        assert "pairs with" not in raw, raw

    def test_accepting_both_sides_counts_one_state_transition(self):
        stream = build_tracked_modification_stream()
        engine = RedlineEngine(stream, author="Reviewer")
        stats = engine.process_batch([AcceptChange(target_id="Chg:1"), AcceptChange(target_id="Chg:2")])
        # The first accept resolves the replacement pair; the second action
        # causes no state transition and must not be reported as applied.
        assert stats["actions_applied"] == 1, stats
        assert stats.get("actions_already_resolved") == 1, stats
        assert stats["actions_skipped"] == 0, stats
        details = "\n".join(stats.get("skipped_details") or [])
        assert "already resolved" in details.lower(), details

        final = clean_text_of(engine.doc)
        assert "60 days" in final and "30 days" not in final

    def test_contradictory_actions_on_pair_are_rejected(self):
        stream = build_tracked_modification_stream()
        engine = RedlineEngine(stream, author="Reviewer")
        with pytest.raises(BatchValidationError) as exc:
            engine.process_batch([AcceptChange(target_id="Chg:1"), RejectChange(target_id="Chg:2")])
        msg = "\n".join(exc.value.errors)
        assert "Chg:1" in msg and "Chg:2" in msg, msg
        assert "pair" in msg.lower(), msg

    def test_reject_then_accept_on_pair_is_rejected(self):
        stream = build_tracked_modification_stream()
        engine = RedlineEngine(stream, author="Reviewer")
        with pytest.raises(BatchValidationError):
            engine.process_batch([RejectChange(target_id="Chg:1"), AcceptChange(target_id="Chg:2")])

    def test_rejecting_both_sides_counts_one_state_transition(self):
        stream = build_tracked_modification_stream()
        engine = RedlineEngine(stream, author="Reviewer")
        stats = engine.process_batch([RejectChange(target_id="Chg:1"), RejectChange(target_id="Chg:2")])
        assert stats["actions_applied"] == 1, stats
        assert stats.get("actions_already_resolved") == 1, stats

        final = clean_text_of(engine.doc)
        assert "30 days" in final and "60 days" not in final

    def test_independent_changes_still_resolve_independently(self):
        doc = build_docx(["First paragraph mentioning alpha.", "Second paragraph mentioning beta."])
        engine = RedlineEngine(doc_to_stream(doc), author="Editor")
        engine.apply_edits([ModifyText(target_text="alpha", new_text="ALPHA")])
        engine.apply_edits([ModifyText(target_text="beta", new_text="BETA")])
        stream = engine.save_to_stream()

        raw = extract_text_from_stream(BytesIO(stream.getvalue()))
        import re as _re

        ids = _re.findall(r"\[Chg:(\d+) delete\]", raw)
        assert len(ids) == 2

        engine2 = RedlineEngine(stream, author="Reviewer")
        stats = engine2.process_batch(
            [AcceptChange(target_id=f"Chg:{ids[0]}"), RejectChange(target_id=f"Chg:{ids[1]}")]
        )
        assert stats["actions_applied"] == 2
        assert stats.get("actions_already_resolved", 0) == 0
        final = clean_text_of(engine2.doc)
        assert "ALPHA" in final and "beta" in final and "BETA" not in final


# ---------------------------------------------------------------------------
# ADEU-QA-005: outputs must say pages are synthetic
# ---------------------------------------------------------------------------


class TestQA005SyntheticPageDisclosure:
    def _build_long_doc(self, path: Path):
        doc = Document()
        for i in range(400):
            doc.add_paragraph(f"Clause {i}: " + ("lorem ipsum dolor sit amet " * 8))
        doc.save(path)

    def test_page_banner_discloses_synthetic_pagination(self, tmp_path, capsys):
        src = tmp_path / "long.docx"
        self._build_long_doc(src)
        code, stdout, _ = run_cli(["extract", src, "--page", "1"], capsys)
        assert code == 0
        banner_line = next((line for line in stdout.splitlines() if "Page 1 of" in line), None)
        assert banner_line, "extract of a multi-page document must carry a page banner:\n" + stdout[:400]
        assert "synthetic" in banner_line.lower(), (
            "the page banner must say the page is a synthetic chunk, not a printed Word page: " + banner_line
        )

    def test_banner_with_disclosure_still_round_trips(self, tmp_path, capsys):
        from adeu.cli import _strip_page_chrome

        src = tmp_path / "long2.docx"
        self._build_long_doc(src)
        code, stdout, _ = run_cli(["extract", src, "--page", "1"], capsys)
        assert code == 0
        stripped, page, total = _strip_page_chrome(stdout)
        assert page == 1 and total and total > 1
        assert "synthetic" not in stripped.splitlines()[0].lower()


# ---------------------------------------------------------------------------
# ADEU-QA-006: broken pipes must not traceback
# ---------------------------------------------------------------------------


@pytest.mark.skipif(sys.platform == "win32", reason="POSIX pipe semantics")
class TestQA006BrokenPipe:
    def test_extract_piped_to_early_close_is_quiet(self, tmp_path):
        src = tmp_path / "big.docx"
        doc = Document()
        for i in range(2000):
            doc.add_paragraph(f"Paragraph {i} " + ("filler text " * 20))
        doc.save(str(src))

        script = (
            "import sys; sys.argv = ['adeu', 'extract', %r, '--page', 'all']; "
            "from adeu.cli import main; main()" % str(src)
        )
        proc = subprocess.Popen(
            [sys.executable, "-c", script],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
        # Read a token amount then close the pipe early — the `| head` shape.
        proc.stdout.read(1024)
        proc.stdout.close()
        stderr = proc.stderr.read().decode("utf-8", "replace")
        proc.stderr.close()
        code = proc.wait(timeout=60)

        assert "Traceback" not in stderr, stderr
        assert "BrokenPipeError" not in stderr, stderr
        assert code in (0, 141), f"expected quiet exit (0) or 128+SIGPIPE (141), got {code}: {stderr}"

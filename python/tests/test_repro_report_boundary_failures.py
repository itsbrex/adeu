import io

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine


def test_boundary_failure():
    """
    TC_BOUNDARY: target_text spans a paragraph boundary with body text on both sides
    """
    doc = Document()
    doc.add_paragraph("First paragraph text.")
    doc.add_paragraph("Second paragraph text.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)

    edit = ModifyText(
        target_text="First paragraph text.\n\nSecond paragraph text.",
        new_text="First paragraph text. New.\n\nSecond paragraph text. New.",
    )

    # On the unpatched codebase, this throws a BatchValidationError.
    # We assert that the batch executes successfully to replicate the failure on unpatched environments.
    result = engine.process_batch([edit], dry_run=False)
    assert result.get("applied", 0) == 1 or result.get("edits_applied", 0) == 1


def test_conflict_failure():
    """
    TC_CONFLICT: Modification targets an active insertion from another author (Supplier's Counsel)
    """
    doc = Document()
    p = doc.add_paragraph("The party shall provide ")

    # Inject active insertion from another author (Supplier's Counsel)
    p_el = p._element
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "201")
    ins.set(qn("w:author"), "Supplier's Counsel")
    ins.set(qn("w:date"), "2026-06-30T08:00:00Z")

    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "written notice"
    r.append(t)
    ins.append(r)
    p_el.append(ins)

    r2 = OxmlElement("w:r")
    t2 = OxmlElement("w:t")
    t2.text = " within 30 days."
    r2.append(t2)
    p_el.append(r2)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Engine is Reviewer AI (different from Supplier's Counsel)
    engine = RedlineEngine(stream, author="Reviewer AI")

    edit = ModifyText(target_text="written notice", new_text="email notification")

    # On the unpatched codebase, this throws a BatchValidationError.
    # We assert that the batch executes successfully to replicate the failure on unpatched environments.
    result = engine.process_batch([edit], dry_run=False)
    assert result.get("applied", 0) == 1 or result.get("edits_applied", 0) == 1


def test_comment_only_on_foreign_insertion_applies_without_crash():
    """
    Regression guard: a COMMENT_ONLY edit (new_text == target_text + comment)
    whose target lies entirely inside another author's pending <w:ins> must
    attach the comment cleanly. Allowing such edits (the foreign-insertion
    relaxation) once routed them to _attach_comment with a run nested inside the
    <w:ins>, raising "Element is not a child of this node" and aborting the whole
    batch. The comment range must wrap the <w:ins>, not crash.
    """
    doc = Document()
    p = doc.add_paragraph("Prefix ")

    p_el = p._element
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "9")
    ins.set(qn("w:author"), "Author A")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "beta"
    r.append(t)
    ins.append(r)
    p_el.append(ins)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="Editor")
    edit = ModifyText(target_text="beta", new_text="beta", comment="flag this clause")

    result = engine.process_batch([edit], dry_run=False)
    assert result["edits_applied"] == 1
    assert result["edits_skipped"] == 0

    out_xml = Document(engine.save_to_stream()).paragraphs[0]._element.xml
    # The comment markers must wrap the foreign insertion, which must survive.
    assert 'w:author="Author A"' in out_xml
    assert out_xml.index("commentRangeStart") < out_xml.index('w:ins w:id="9"')
    assert "commentRangeEnd" in out_xml

import io

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from adeu.models import RejectChange
from adeu.redline.engine import RedlineEngine


def _inject_comment_into_element(engine: RedlineEngine, target_element, comment_text: str) -> str:
    """
    Helper to inject a complete comment structure (Start, End, Reference)
    directly into a specific DOM element (like a w:ins or w:del tag) to mock
    Word's behavior of tracking changes over commented text.
    """
    c_id = engine.comments_manager.add_comment("TestAuthor", comment_text)

    start = OxmlElement("w:commentRangeStart")
    start.set(qn("w:id"), c_id)

    end = OxmlElement("w:commentRangeEnd")
    end.set(qn("w:id"), c_id)

    ref_run = OxmlElement("w:r")
    ref = OxmlElement("w:commentReference")
    ref.set(qn("w:id"), c_id)
    ref_run.append(ref)

    # Append all parts inside the target element
    target_element.append(start)
    target_element.append(end)
    target_element.append(ref_run)

    return c_id


def test_reject_change_cleans_encapsulated_comments():
    """
    Tests that rejecting a w:ins element that contains comment anchors
    successfully cascades the deletion to all 4 comment XML parts.
    """
    # 1. Setup Document
    doc = Document()
    doc.add_paragraph("Base text. ")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # 2. Inject a w:ins element manually
    engine = RedlineEngine(stream, author="TestAuthor")

    ins_tag = OxmlElement("w:ins")
    ins_tag.set(qn("w:id"), "99")
    ins_tag.set(qn("w:author"), "TestAuthor")

    r_tag = OxmlElement("w:r")
    t_tag = OxmlElement("w:t")
    t_tag.text = "Inserted with comment"
    r_tag.append(t_tag)
    ins_tag.append(r_tag)

    p_element = engine.doc.paragraphs[0]._element
    p_element.append(ins_tag)

    # 3. Add a comment INSIDE the w:ins tag
    _inject_comment_into_element(engine, ins_tag, "This comment lives inside the insertion.")

    # Verify pre-condition: Comment exists in parts
    c_mgr = engine.comments_manager
    assert len(c_mgr.comments_part.element.findall(qn("w:comment"))) == 1
    if c_mgr.ids_part:
        assert len(c_mgr.ids_part.element.findall(qn("w16cid:commentId"))) == 1

    # 4. Act: Reject the change (should delete the w:ins and trigger comment cleanup)
    action = RejectChange(target_id="Chg:99")
    applied, _, _ = engine.apply_review_actions([action])

    assert applied == 1, "Action should be successfully applied"

    # 5. Assert: Complete Eradication
    doc_xml = engine.doc.element.xml
    assert "w:commentRangeStart" not in doc_xml, "Start anchor leaked in body"
    assert "w:commentReference" not in doc_xml, "Reference anchor leaked in body"

    assert len(c_mgr.comments_part.element.findall(qn("w:comment"))) == 0, "Leaked in comments.xml"
    if c_mgr.ids_part:
        assert len(c_mgr.ids_part.element.findall(qn("w16cid:commentId"))) == 0, "Leaked in commentsIds.xml"
    if c_mgr.extended_part:
        assert len(c_mgr.extended_part.element.findall(qn("w15:commentEx"))) == 0, "Leaked in commentsExtended.xml"
    if c_mgr.extensible_part:
        assert len(c_mgr.extensible_part.element.findall(qn("w16cex:commentExtensible"))) == 0, (
            "Leaked in commentsExtensible.xml"
        )


def test_accept_all_revisions_total_comment_wipe():
    """
    Tests that accepting all revisions completely cleans the comment subsystem,
    preventing the invisible ghost comments seen in Issue 5.
    """
    # 1. Setup Document
    doc = Document()
    doc.add_paragraph("Some text.")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # 2. Inject a w:del element containing a comment
    engine = RedlineEngine(stream, author="TestAuthor")

    del_tag = OxmlElement("w:del")
    del_tag.set(qn("w:id"), "88")
    del_tag.set(qn("w:author"), "TestAuthor")

    # 3. Add comment inside the w:del tag
    _inject_comment_into_element(engine, del_tag, "This comment lives inside the deletion.")

    # Add to body
    p_element = engine.doc.paragraphs[0]._element
    p_element.append(del_tag)

    # Pre-condition verification
    c_mgr = engine.comments_manager
    assert len(c_mgr.comments_part.element.findall(qn("w:comment"))) == 1

    # 4. Act: Accept all revisions (wipes w:del from body)
    engine.accept_all_revisions()

    # 5. Assert: Total Wipe
    doc_xml = engine.doc.element.xml
    assert "w:commentRangeStart" not in doc_xml
    assert "w:commentReference" not in doc_xml

    assert len(c_mgr.comments_part.element.findall(qn("w:comment"))) == 0
    if c_mgr.ids_part:
        assert len(c_mgr.ids_part.element.findall(qn("w16cid:commentId"))) == 0


def test_accept_all_revisions_preserves_untouched_comments():
    """
    VAL-CRIT-2: accept_all_revisions must NOT perform global normalizations or
    blindly wipe unrelated comments/proofErrs.
    """
    doc = Document()
    doc.add_paragraph("Sentence one. ")
    doc.add_paragraph("Sentence two. ")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="TestAuthor")

    # Add a legitimate unrelated comment that should SURVIVE
    _inject_comment_into_element(engine, engine.doc.paragraphs[0]._element, "Legitimate comment.")

    # Add a fake proofErr to verify normalize_docx isn't called
    proof = OxmlElement("w:proofErr")
    proof.set(qn("w:type"), "spellStart")
    engine.doc.paragraphs[0]._element.append(proof)

    # Add a track-change (w:ins) to paragraph 2
    ins_tag = OxmlElement("w:ins")
    ins_tag.set(qn("w:id"), "99")
    r_tag = OxmlElement("w:r")
    t_tag = OxmlElement("w:t")
    t_tag.text = "Inserted"
    r_tag.append(t_tag)
    ins_tag.append(r_tag)
    engine.doc.paragraphs[1]._element.append(ins_tag)

    # Run acceptance
    engine.accept_all_revisions()

    doc_xml = engine.doc.element.xml
    assert "w:ins" not in doc_xml, "Track change must be accepted/gone"
    assert "w:commentRangeStart" in doc_xml, "Original comment must survive"
    assert "w:proofErr" in doc_xml, "w:proofErr must survive (normalize_docx was wrongly called)"


def test_val_obs_new_5_orphan_comments_spanning_redlines():
    """
    VAL-OBS-NEW-5: Comments that wrap across adjacent redline tags (del + ins)
    must be correctly cleaned up during accept_all_revisions, rather than
    halting prematurely and leaving orphans.
    """
    doc = Document()
    doc.add_paragraph()
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    engine = RedlineEngine(stream)

    p_elem = engine.doc.paragraphs[0]._element
    c_id = engine.comments_manager.add_comment("Test", "Spanning comment")

    # Construct: [Start] <del/> <ins/> [End] [Ref]
    start = OxmlElement("w:commentRangeStart")
    start.set(qn("w:id"), c_id)
    p_elem.append(start)

    del_tag = OxmlElement("w:del")
    del_tag.set(qn("w:id"), "1")
    p_elem.append(del_tag)

    ins_tag = OxmlElement("w:ins")
    ins_tag.set(qn("w:id"), "1")
    p_elem.append(ins_tag)

    end = OxmlElement("w:commentRangeEnd")
    end.set(qn("w:id"), c_id)
    p_elem.append(end)

    engine.accept_all_revisions()

    doc_xml = engine.doc.element.xml
    assert "w:commentRangeStart" not in doc_xml, "Comment start anchor leaked"
    assert "w:commentRangeEnd" not in doc_xml, "Comment end anchor leaked"


def test_val_obs_new_6_empty_ins_tag():
    """
    VAL-OBS-NEW-6: Inserting an empty string should return None,
    preventing the generation of self-closing <w:ins/> ghost tags.
    """
    doc = Document()
    doc.add_paragraph("Target word.")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    engine = RedlineEngine(stream)

    ins = engine._track_insert_inline("")
    assert ins is None, "Empty string should not return an empty <w:ins> tag"

import io

from docx import Document

from adeu.ingest import extract_text_from_stream
from adeu.models import AcceptChange, ModifyText, RejectChange
from adeu.redline.engine import RedlineEngine


def test_batch_accept_does_not_corrupt():
    """
    Verify that accepting multiple changes in one batch does not corrupt the document.
    """
    doc = Document()
    doc.add_paragraph("Para 1")
    doc.add_paragraph("Para 2")
    doc.add_paragraph("Para 3")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Create 3 edits (Modifications) -> 3 Del, 3 Ins = 6 IDs
    # Edit 1 (Para 1): Del(1), Ins(2)
    # Edit 2 (Para 2): Del(3), Ins(4)
    # Edit 3 (Para 3): Del(5), Ins(6)

    edits = [
        ModifyText(target_text="Para 1", new_text="Para One"),
        ModifyText(target_text="Para 2", new_text="Para Two"),
        ModifyText(target_text="Para 3", new_text="Para Three"),
    ]

    engine = RedlineEngine(stream)
    engine.apply_edits(edits)
    stream_redlined = engine.save_to_stream()

    # Verify IDs exist
    text = extract_text_from_stream(stream_redlined)
    assert "[Chg:1 " in text
    assert "[Chg:2 " in text
    assert "[Chg:3 " in text
    assert "[Chg:4 " in text
    assert "[Chg:5 " in text
    assert "[Chg:6 " in text

    # BATCH ACCEPT ALL
    # We accept the Insertions (2, 4, 6) AND the Deletions (1, 3, 5)
    # Order shouldn't matter for correctness, but let's mix them
    actions = [
        AcceptChange(target_id="Chg:1"),
        AcceptChange(target_id="Chg:2"),
        AcceptChange(target_id="Chg:3"),
        AcceptChange(target_id="Chg:4"),
        AcceptChange(target_id="Chg:5"),
        AcceptChange(target_id="Chg:6"),
    ]

    engine2 = RedlineEngine(stream_redlined)
    applied, skipped, already_resolved = engine2.apply_review_actions(actions)

    # Three replacements: each del+ins pair resolves as one unit on the
    # first accept, so 3 actions transition state and 3 are accurate no-ops
    # (QA 2026-07-19 ADEU-QA-004).
    assert applied == 3
    assert already_resolved == 3
    assert skipped == 0

    stream_final = engine2.save_to_stream()
    text_final = extract_text_from_stream(stream_final)

    # Check for corruption / truncation
    assert "Para One" in text_final
    assert "Para Two" in text_final
    assert "Para Three" in text_final

    assert "Para 1" not in text_final
    assert "Para 2" not in text_final
    assert "Para 3" not in text_final

    # Ensure no markup remains
    assert "[Chg:" not in text_final
    assert "{++" not in text_final
    assert "{--" not in text_final


def test_batch_mixed_accept_reject_integrity():
    """
    Scenario:
    Edit 1 (Para 1): Accept (Finalize change)
    Edit 2 (Para 2): Reject (Revert to original)
    """
    doc = Document()
    doc.add_paragraph("Para 1")
    doc.add_paragraph("Para 2")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Edit 1: Del(1), Ins(2)
    # Edit 2: Del(3), Ins(4)
    edits = [
        ModifyText(target_text="Para 1", new_text="Para One"),
        ModifyText(target_text="Para 2", new_text="Para Two"),
    ]

    engine = RedlineEngine(stream)
    engine.apply_edits(edits)
    stream_redlined = engine.save_to_stream()

    # Batch Action:
    # Note: Because edits are processed backwards (bottom to top) to prevent index drift,
    # Edit 2 ("Para 2") is processed first and gets IDs 1 (Del) and 2 (Ins).
    # Edit 1 ("Para 1") is processed second and gets IDs 3 (Del) and 4 (Ins).
    #
    # To Accept Edit 1, we Accept 3 & 4.
    # To Reject Edit 2, we Reject 1 & 2.

    actions = [
        AcceptChange(target_id="Chg:3"),
        AcceptChange(target_id="Chg:4"),
        RejectChange(target_id="Chg:1"),
        RejectChange(target_id="Chg:2"),
    ]

    engine2 = RedlineEngine(stream_redlined)
    applied, skipped, already_resolved = engine2.apply_review_actions(actions)

    # Two replacement pairs, each resolved by its first action; the paired
    # follow-ups are accurate no-ops (QA 2026-07-19 ADEU-QA-004).
    assert applied == 2
    assert already_resolved == 2
    assert skipped == 0

    stream_final = engine2.save_to_stream()
    text_final = extract_text_from_stream(stream_final)

    # Edit 1 Accepted: "Para One" exists, "Para 1" gone
    assert "Para One" in text_final
    assert "Para 1" not in text_final

    # Edit 2 Rejected: "Para 2" exists, "Para Two" gone
    assert "Para 2" in text_final
    assert "Para Two" not in text_final

"""
Reproduction harness for three field observations reported against the engine
(originally seen on the Node backend: playbook-commenting, policy-checklist-review,
party-swap). This file establishes, empirically, whether each pathology is
reproducible on the PYTHON backend.

Summary of findings (see each test for the proof):

  Issue 1 (Sequential batch accept+modify conflict)
      -> NOT reproducible on Python. A single batch that accepts a foreign
         author's change and then edits the now-accepted text SUCCEEDS, because
         Python applies review actions BEFORE validating edits. (Node rejects the
         same batch up-front: node/packages/core/src/repro.feedback.test.ts.)

  Issue 2 (Comment-only run-shredding / comment duplication)
      -> The described MECHANISM (a single match whose physical runs are shredded,
         duplicating the comment per run segment) is NOT reproducible: a
         self-replacement comment over 9 fragmented runs emits exactly ONE comment.
         The 9x SYMPTOM is reproducible only via match_mode="all" fanning one
         logical comment across 9 occurrences. The underlying ergonomic gap is
         real: there is no dedicated comment-only change type.

  Issue 3 (File-path ambiguity and save shuffling)
      -> Reproducible. process_document_batch defaults to *_processed.docx and
         accept_all_changes defaults to *_clean.docx, so the active state
         fragments across files. (Python DOES have an idempotency guard the Node
         tool lacks: re-batching a *_processed.docx reuses the same path.)
"""

import asyncio
import io
import re
import zipfile

from docx import Document

from adeu.ingest import extract_text_from_stream
from adeu.mcp_components.tools.document import (
    accept_all_changes,
    process_document_batch,
)
from adeu.models import AcceptChange, DocumentChange, ModifyText
from adeu.redline.engine import RedlineEngine


class MockContext:
    """Absorbs async logging calls from the FastMCP tool implementations."""

    async def info(self, *a, **k):
        pass

    async def debug(self, *a, **k):
        pass

    async def warning(self, *a, **k):
        pass

    async def error(self, *a, **k):
        pass


def _count_comment_markers(stream: io.BytesIO) -> tuple[int, int, int]:
    zf = zipfile.ZipFile(io.BytesIO(stream.getvalue()))
    xml = zf.read("word/document.xml").decode("utf-8")
    return (
        xml.count("<w:commentRangeStart"),
        xml.count("<w:commentRangeEnd"),
        xml.count("<w:commentReference"),
    )


# ────────────────────────────────────────────────────────────────────────────
# ISSUE 1 — NOT reproducible on Python.
# ────────────────────────────────────────────────────────────────────────────
def test_issue1_accept_then_modify_same_change_in_one_batch_succeeds(tmp_path):
    """
    A prior author ("Supplier's Counsel") tracked an edit. In a SINGLE batch we
    accept that change and then refine the now-accepted text. On Python this is
    one atomic round-trip; the engine applies the accept first, rebuilds its view,
    and the modify lands cleanly. (The identical batch is rejected on Node.)
    """
    doc = Document()
    doc.add_paragraph("The term is 12 months.")
    base = tmp_path / "lease.docx"
    doc.save(base)

    # Round 1: Supplier's Counsel redlines 12 -> 24 months.
    with open(base, "rb") as f:
        eng1 = RedlineEngine(io.BytesIO(f.read()), author="Supplier's Counsel")
    eng1.apply_edits([ModifyText(target_text="12 months", new_text="24 months")])
    mid = tmp_path / "lease_redlined.docx"
    with open(mid, "wb") as f:
        f.write(eng1.save_to_stream().getvalue())

    mid_text = extract_text_from_stream(io.BytesIO(open(mid, "rb").read()))
    chg_ids = re.findall(r"\[Chg:(\d+)", mid_text)
    assert chg_ids, "Round 1 did not produce tracked changes."

    # Round 2: Acme accepts Supplier's change AND edits the accepted text — same batch.
    changes = [AcceptChange(target_id=f"Chg:{cid}") for cid in chg_ids]
    changes.append(ModifyText(target_text="24 months", new_text="36 months"))

    out = tmp_path / "lease_final.docx"
    msg = asyncio.run(
        process_document_batch(
            reasoning="test",
            original_docx_path=str(mid),
            author_name="Acme's Counsel",
            ctx=MockContext(),
            changes=changes,
            output_path=str(out),
        )
    )

    assert "Batch complete" in msg, f"Expected success, got: {msg}"
    assert "Edits: 1 applied, 0 skipped" in msg
    assert "active insertion from another author" not in msg

    final = extract_text_from_stream(io.BytesIO(open(out, "rb").read()))
    # Supplier's insertion was accepted (no longer redlined) and Acme's new
    # redline replaces 24 -> 36.
    assert "{++36++}" in final
    assert "Acme's Counsel" in final


# ────────────────────────────────────────────────────────────────────────────
# ISSUE 2 — mechanism NOT reproducible; symptom reproducible via match_mode=all.
# ────────────────────────────────────────────────────────────────────────────
def test_issue2_self_replacement_comment_over_fragmented_runs_is_single_comment():
    """
    A paragraph split into 9 physical runs. Attaching a margin comment via a
    self-replacement ModifyText (target_text == new_text) produces exactly ONE
    comment, not one per run segment. The run-shredding mechanism is not present.
    """
    doc = Document()
    p = doc.add_paragraph()
    words = [
        "The ",
        "Purchase ",
        "Price ",
        "is ",
        "[   ] ",
        "dollars ",
        "per ",
        "unit ",
        "total.",
    ]
    for w in words:
        p.add_run(w)
    s = io.BytesIO()
    doc.save(s)
    s.seek(0)

    eng = RedlineEngine(s, author="Reviewer")
    sentence = "".join(words).strip()
    eng.apply_edits([ModifyText(target_text=sentence, new_text=sentence, comment="Missing value")])
    starts, ends, refs = _count_comment_markers(eng.save_to_stream())
    assert (starts, ends, refs) == (1, 1, 1)


def test_issue2_match_mode_all_fans_one_comment_across_nine_occurrences():
    """
    The reported 9x duplication. With no comment-only primitive, the agent must
    self-replace; with match_mode="all" over a placeholder appearing 9 times, the
    single logical comment is emitted 9 times. This reproduces the SYMPTOM.
    """
    doc = Document()
    for i in range(9):
        doc.add_paragraph(f"Item {i + 1}: value is [   ] here.")
    s = io.BytesIO()
    doc.save(s)
    s.seek(0)

    eng = RedlineEngine(s, author="Reviewer")
    eng.apply_edits(
        [
            ModifyText(
                target_text="[   ]",
                new_text="[   ]",
                comment="Missing value",
                match_mode="all",
            )
        ]
    )
    _, _, refs = _count_comment_markers(eng.save_to_stream())
    assert refs == 9


def test_issue2_no_dedicated_comment_only_change_type_exists():
    """
    Confirms the ergonomic root cause: the public DocumentChange union has no
    comment-only primitive, so comments can only ride on a (self-replacing)
    ModifyText.
    """
    members = DocumentChange.__args__[0].__args__  # Annotated[Union[...], ...]
    type_literals = {m.model_fields["type"].default for m in members if "type" in m.model_fields}
    assert "comment" not in type_literals
    assert "modify" in type_literals
    # The comment payload lives on ModifyText, not on a first-class comment type.
    assert "comment" in ModifyText.model_fields


# ────────────────────────────────────────────────────────────────────────────
# ISSUE 3 — reproducible: divergent default output stems fragment state.
# ────────────────────────────────────────────────────────────────────────────
def test_issue3_default_output_paths_diverge_and_fragment_state(tmp_path):
    base = tmp_path / "contract.docx"
    d = Document()
    d.add_paragraph("The Provider shall deliver the goods.")
    d.save(base)

    # 1. Redline with NO output_path -> contract_processed.docx
    msg1 = asyncio.run(
        process_document_batch(
            reasoning="test",
            original_docx_path=str(base),
            author_name="Agent",
            ctx=MockContext(),
            changes=[ModifyText(target_text="Provider", new_text="Supplier")],
        )
    )
    assert "contract_processed.docx" in msg1
    assert (tmp_path / "contract_processed.docx").exists()

    # 2. Accept-all on the redlined file, NO output_path -> *_clean.docx (different stem).
    msg2 = asyncio.run(
        accept_all_changes(
            reasoning="test",
            docx_path=str(tmp_path / "contract_processed.docx"),
            ctx=MockContext(),
        )
    )
    assert "contract_processed_clean.docx" in msg2

    # The active state is now fragmented across three files with no single
    # source of truth — the exact "save shuffling" pathology.
    files = sorted(p.name for p in tmp_path.iterdir())
    assert files == [
        "contract.docx",
        "contract_processed.docx",
        "contract_processed_clean.docx",
    ]


def test_issue3_python_idempotency_guard_reuses_processed_path(tmp_path):
    """
    Python (unlike Node) guards against suffix compounding: re-running a batch on
    a *_processed.docx writes back to the SAME file instead of creating
    contract_processed_processed.docx. Documents the Python/Node divergence.
    """
    processed = tmp_path / "contract_processed.docx"
    d = Document()
    d.add_paragraph("The Provider shall deliver the goods.")
    d.save(processed)

    msg = asyncio.run(
        process_document_batch(
            reasoning="test",
            original_docx_path=str(processed),
            author_name="Agent",
            ctx=MockContext(),
            changes=[ModifyText(target_text="goods", new_text="products")],
        )
    )

    assert "contract_processed.docx" in msg
    assert "contract_processed_processed.docx" not in msg
    assert not (tmp_path / "contract_processed_processed.docx").exists()

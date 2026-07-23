import io
import sys
from pathlib import Path

import docx
from docx import Document

from adeu.ingest import _extract_text_from_doc, extract_text_from_stream
from adeu.models import ModifyText
from adeu.outline import extract_outline
from adeu.pagination import paginate
from adeu.redline.engine import RedlineEngine


def test_gfm_table_divider_extraction():
    # 1. Create a document with a table
    doc = Document()
    table = doc.add_table(rows=3, cols=3)

    # Fill first row (header)
    table.cell(0, 0).text = "ID"
    table.cell(0, 1).text = "Name"
    table.cell(0, 2).text = "Description"

    # Fill second row
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "Item A"
    table.cell(1, 2).text = "This is item A"

    # Fill third row
    table.cell(2, 0).text = "2"
    table.cell(2, 1).text = "Item B"
    table.cell(2, 2).text = "This is item B"

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # 2. Extract clean view text
    text = extract_text_from_stream(stream, clean_view=True)

    # 3. Assert correct GFM pipe table formatting with divider row
    expected_lines = [
        "ID | Name | Description",
        "--- | --- | ---",
        "1 | Item A | This is item A",
        "2 | Item B | This is item B",
    ]
    expected_text = "\n".join(expected_lines)
    assert expected_text in text


def test_gfm_table_divider_mapping():
    # Verifies that mapping and editing works flawlessly when the divider is present
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Col1"
    table.cell(0, 1).text = "Col2"
    table.cell(1, 0).text = "Val1"
    table.cell(1, 1).text = "Val2"

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Let's apply a text modification inside a table cell
    # Target "Val1", replace with "NewVal1"
    edit = ModifyText(target_text="Val1", new_text="NewVal1")
    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])

    assert applied == 1
    assert skipped == 0

    res_stream = engine.save_to_stream()
    res_text = extract_text_from_stream(res_stream, clean_view=False)

    # Check that CriticMarkup is correct
    assert "{--Val1--}{++NewVal1++}" in res_text

    # Assert that the divider row is present in the mapped/edited raw view text
    assert "Col1 | Col2\n--- | ---\n" in res_text


def test_read_docx_negative_page_boundary(tmp_path):
    """
    Verifies that calling read_docx with a negative page number (like -1)
    correctly raises a ToolError, instead of silently defaulting to page 1.
    """
    import asyncio

    import docx
    import pytest
    from fastmcp.exceptions import ToolError

    from adeu.mcp_components.tools.document import read_docx

    # Create a 3-page test document
    doc = docx.Document()
    doc.add_paragraph("This is Page 1 content.")
    doc.add_page_break()
    doc.add_paragraph("This is Page 2 content.")
    doc.add_page_break()
    doc.add_paragraph("This is Page 3 content.")

    doc_path = tmp_path / "boundary_test.docx"
    doc.save(str(doc_path))

    class MockContext:
        async def info(self, msg, **kwargs):
            pass

        async def debug(self, msg, **kwargs):
            pass

        async def warning(self, msg, **kwargs):
            pass

        async def error(self, msg, **kwargs):
            pass

    ctx = MockContext()

    # Calling with page=-1 must raise a ToolError indicating that page is out of range
    with pytest.raises(ToolError) as exc_info:
        asyncio.run(read_docx(reasoning="test", file_path=str(doc_path), ctx=ctx, page=-1))

    assert "Page -1 out of range" in str(exc_info.value)

    # Calling with page=0 must also raise a ToolError
    with pytest.raises(ToolError) as exc_info_zero:
        asyncio.run(read_docx(reasoning="test", file_path=str(doc_path), ctx=ctx, page=0))

    assert "Page 0 out of range" in str(exc_info_zero.value)


def create_manual_break_doc(path: Path):
    doc = docx.Document()
    doc.add_heading("Pagination Test Document", level=1)
    for page_num in range(1, 6):
        doc.add_heading(f"Heading on Page {page_num}", level=2)
        doc.add_paragraph(f"This is paragraph content belonging strictly to page {page_num}. " * 5)
        if page_num < 5:
            doc.add_page_break()
    doc.save(str(path))


def test_manual_page_breaks_pagination(tmp_path):
    doc_path = tmp_path / "manual_breaks.docx"
    create_manual_break_doc(doc_path)

    doc = docx.Document(str(doc_path))
    projected_body = _extract_text_from_doc(doc, include_appendix=False)

    # Run paginate
    pag_res = paginate(projected_body)

    # The document must have 5 pages
    assert pag_res.total_pages == 5, f"Expected 5 pages, got {pag_res.total_pages}"

    # Assert each page has the correct heading
    for page_num in range(1, 6):
        page_content = pag_res.pages[page_num - 1].page_content
        assert f"Heading on Page {page_num}" in page_content
        assert f"This is paragraph content belonging strictly to page {page_num}." in page_content
        # Ensure other page content is NOT leaked to this page
        for other_num in range(1, 6):
            if other_num != page_num:
                assert f"Heading on Page {other_num}" not in page_content


def test_manual_page_breaks_outline(tmp_path):
    doc_path = tmp_path / "manual_breaks.docx"
    create_manual_break_doc(doc_path)

    doc = docx.Document(str(doc_path))
    projected_body = _extract_text_from_doc(doc, include_appendix=False)

    pag_res = paginate(projected_body)

    # Get outline
    # extract_outline expects (doc, projected_body, body_pages, body_page_offsets, paragraph_offsets)
    body_pages = [p.page_content for p in pag_res.pages]
    body_page_offsets = pag_res.body_page_offsets

    nodes = extract_outline(doc, projected_body, body_pages, body_page_offsets)

    # We should have headings on pages 1 to 5
    headings = [node for node in nodes if node.level == 2]
    assert len(headings) == 5
    for i, node in enumerate(headings):
        expected_page = i + 1
        assert node.text == f"Heading on Page {expected_page}"
        assert node.page == expected_page, f"Expected {node.text} to be on page {expected_page}, got page {node.page}"


def run_cli(args, capsys):
    """Invoke the CLI in-process; returns (exit_code, stdout, stderr)."""
    from unittest.mock import patch

    from adeu.cli import main

    code = 0
    with patch.object(sys, "argv", ["adeu"] + [str(a) for a in args]):
        try:
            main()
        except SystemExit as e:
            code = e.code or 0
    captured = capsys.readouterr()
    return code, captured.out, captured.err


def test_cli_apply_large_document_major_deletions(tmp_path, capsys):
    # 1. Create the original document with 100 sections
    doc_path = tmp_path / "large_bug.docx"
    doc = docx.Document()
    for i in range(1, 101):
        doc.add_heading(f"Section {i}", level=2)
        doc.add_paragraph(f"This is the body paragraph for section {i}. " * 15)
        if i % 10 == 0:
            doc.add_paragraph("Some special marker for search in section " + str(i))
    doc.save(str(doc_path))

    # 2. Create the truncated modified text file
    txt_path = tmp_path / "large_truncated_bug.txt"
    content = f"> **File Path:** {doc_path.name}\n\n# Large Document Test\n\nOnly Section 1 is here."
    txt_path.write_text(content, encoding="utf-8")

    # 3. Execute apply command
    out_path = tmp_path / "large_applied_bug.docx"
    code, stdout, stderr = run_cli(
        ["apply", str(doc_path), str(txt_path), "-o", str(out_path), "--allow-major-deletions"], capsys
    )

    # The regression test asserts that the bug is fixed and it completes successfully
    assert code == 0, f"apply failed with code {code}\nSTDOUT:\n{stdout}\nSTDERR:\n{stderr}"
    assert out_path.exists(), "Output file was not generated."


def test_accept_all_json_response_enrichment(tmp_path, capsys):
    import json

    import docx

    # 1. Create a simple base document
    doc_path = tmp_path / "base.docx"
    doc = docx.Document()
    doc.add_paragraph("This is a test document.")
    doc.save(str(doc_path))

    # 2. Define a modify edit with a comment
    changes_file = tmp_path / "changes.json"
    changes_file.write_text(
        json.dumps(
            [
                {
                    "type": "modify",
                    "target_text": "document",
                    "new_text": "dossier",
                    "comment": "Review note to be stripped",
                }
            ]
        ),
        encoding="utf-8",
    )

    redlined_path = tmp_path / "redlined.docx"

    # 3. Apply the edit to create tracked changes + comment
    code, stdout, stderr = run_cli(
        ["apply", str(doc_path), str(changes_file), "-o", str(redlined_path), "--json"],
        capsys,
    )
    assert code == 0, f"apply failed with code {code}\nSTDOUT:\n{stdout}\nSTDERR:\n{stderr}"
    assert redlined_path.exists()

    # 4. Run accept-all in JSON mode
    accepted_path = tmp_path / "accepted.docx"
    code, stdout, stderr = run_cli(
        ["accept-all", str(redlined_path), "-o", str(accepted_path), "--json"],
        capsys,
    )
    assert code == 0, f"accept-all failed with code {code}\nSTDOUT:\n{stdout}\nSTDERR:\n{stderr}"

    # 5. Parse JSON output and assert that keys are present and counts are correct
    result = json.loads(stdout.strip())
    assert result.get("status") == "ok"
    assert "accepted_insertions" in result, "accepted_insertions missing from JSON output"
    assert "accepted_deletions" in result, "accepted_deletions missing from JSON output"
    assert "accepted_formatting" in result, "accepted_formatting missing from JSON output"
    assert "removed_comments" in result, "removed_comments missing from JSON output"

    assert result["accepted_insertions"] == 1
    assert result["accepted_deletions"] == 1
    assert result["accepted_formatting"] == 0
    assert result["removed_comments"] == 1


def test_sanitize_baseline_similarity_percent_is_literal(tmp_path):
    """
    The similarity message carries LITERAL percent signs on every surface.

    An error message is data, not a format string: nothing in Adeu %-formats
    it, so pre-escaping to '%%' does not protect a printf host — it simply
    corrupts the text for every consumer that does not un-escape. A previous
    fix escaped these here and then stripped '%%' back out at the CLI print
    sites, which left the SDK, the MCP sanitize tool and the Node engine
    disagreeing about their own message. Format at the point of formatting,
    never at the point of construction.
    """
    import docx
    import pytest

    from adeu.sanitize.core import SanitizeError, sanitize_docx

    # 1. Create two highly different documents to trigger similarity check failure
    working = tmp_path / "normal.docx"
    doc_norm = docx.Document()
    doc_norm.add_paragraph("Agreement between Alpha and Beta.")
    doc_norm.save(str(working))

    baseline = tmp_path / "unicode.docx"
    doc_uni = docx.Document()
    doc_uni.add_paragraph("Different text in Chinese 统一码.")
    doc_uni.save(str(baseline))

    out = tmp_path / "out.docx"

    # 2. The SDK surface (what the MCP sanitize tool also returns verbatim)
    with pytest.raises(SanitizeError) as exc_info:
        sanitize_docx(str(working), str(out), baseline_path=str(baseline))

    err_msg = str(exc_info.value)

    assert "share only 41% of" in err_msg
    assert "differs" in err_msg
    assert "%%" not in err_msg


def test_multi_paragraph_newline_comment(tmp_path):
    import io

    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    from adeu.models import ModifyText
    from adeu.redline.engine import RedlineEngine

    # 1. Create a base document with the target text
    doc = Document()
    doc.add_paragraph("Hello world. 🚀🔥🌟 and some suffix.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # 2. Define the edit with a comment and a newline (paragraph break) in new_text
    edit = ModifyText(
        type="modify",
        target_text="🚀🔥🌟",
        new_text="🚀🔥🌟\n\nAdded: This is an extra line.",
        comment="Diff: Text inserted",
    )

    # 3. Apply edits
    engine = RedlineEngine(stream)
    engine.apply_edits([edit])

    # 4. Save and load document
    result_stream = engine.save_to_stream()
    doc = Document(result_stream)

    # 5. Assert the comment ranges and references exist in the main document XML
    doc_xml = doc.element.xml
    assert "w:commentRangeStart" in doc_xml
    assert "w:commentRangeEnd" in doc_xml
    assert "w:commentReference" in doc_xml

    # 6. Assert the comment actually exists in comments.xml
    comments_part = None
    for rel in doc.part.rels.values():
        if rel.reltype == RT.COMMENTS:
            comments_part = rel.target_part
            break

    assert comments_part is not None, "Comments part was not created"
    comments_xml = comments_part.blob.decode("utf-8")
    assert "Diff: Text inserted" in comments_xml, "Comment text missing from comments.xml"


def test_process_document_batch_flat_schema():
    """
    Asserts that the JSON Schema for process_document_batch.changes is a single flat Change object schema,
    rather than a nested union (oneOf/anyOf) of separate schema definitions.
    This ensures compatibility with 100% of MCP host implementations.
    """
    import asyncio

    from adeu.server import mcp

    tools = asyncio.run(mcp.list_tools())
    process_tool = next(t for t in tools if t.name == "process_document_batch")

    assert process_tool.parameters is not None, "process_document_batch has no parameters"

    # Extract changes property schema
    properties = process_tool.parameters.get("properties", {})
    assert "changes" in properties, "changes parameter missing from process_document_batch tool"

    changes_schema = properties["changes"]
    assert changes_schema.get("type") == "array", "changes must be an array type"

    items_schema = changes_schema.get("items", {})

    # The schema must NOT use oneOf or anyOf for the items
    assert "oneOf" not in items_schema, (
        "changes.items schema uses oneOf, which breaks nested array parsing in some MCP hosts"
    )
    assert "anyOf" not in items_schema, (
        "changes.items schema uses anyOf, which breaks nested array parsing in some MCP hosts"
    )

    # Instead, it must be a single flat object
    assert items_schema.get("type") == "object", "changes.items must be a flat object type"

    # Check that individual properties from all change variants are defined as optional properties of this single model
    item_properties = items_schema.get("properties", {})

    # Must have the type discriminator
    assert "type" in item_properties, "type field missing in flat change schema"

    # Must have all the other specific properties as optional properties
    expected_fields = [
        "target_text",
        "new_text",
        "target_id",
        "text",
        "cells",
        "position",
        "regex",
        "comment",
        "match_mode",
    ]
    for field in expected_fields:
        assert field in item_properties, f"expected field {field} missing from the unified flat Change schema"

    # The other fields should NOT be listed as 'required' at the JSON schema level to ensure they are optional
    required_fields = items_schema.get("required", [])
    # Only type is required (acting as discriminator), or maybe even type is optional.
    # Let's make sure none of the variant-specific fields are required.
    for field in expected_fields:
        assert field not in required_fields, f"field {field} must be optional (not required) in the flat Change schema"


def test_sanitize_blocked_msg_includes_keep_markup(tmp_path, capsys):
    doc_path = tmp_path / "blocked_repro.docx"
    doc = docx.Document()
    p = doc.add_paragraph()
    p.add_run("The ")

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    d = OxmlElement("w:del")
    d.set(qn("w:id"), "1")
    d.set(qn("w:author"), "Opposing Counsel")
    d.set(qn("w:date"), "2025-01-15T10:00:00Z")
    rd = OxmlElement("w:r")
    rt = OxmlElement("w:delText")
    rt.set(qn("xml:space"), "preserve")
    rt.text = "Vendor"
    rd.append(rt)
    d.append(rd)
    p._element.append(d)

    doc.save(str(doc_path))

    # Run CLI sanitize on this doc, which should block because it contains unresolved changes
    code, stdout, stderr = run_cli(["sanitize", str(doc_path)], capsys)

    assert code == 1
    # Check that the block message has our updated guidance suggesting --keep-markup
    assert "use --keep-markup" in stderr, "Validation message should suggest --keep-markup"
    assert "Review in Word first, use --accept-all, or use --keep-markup." in stderr


def test_process_document_batch_changes_payload_shapes(tmp_path):
    """
    There is exactly ONE `changes` parameter, and both engines treat the same
    payload the same way.

    A `changes_json` twin was removed: it made the model choose between two
    spellings of one argument, and Python and Node had drifted into opposite
    rules for which one wins when both were sent. Per-item stringification is
    repaired; a WHOLLY stringified payload is rejected on both engines (the
    Node zod schema cannot accept one without dropping `changes` out of its
    `required` list, and a silent repair on only one engine is how a working
    agent call breaks on a backend switch).
    """
    import asyncio
    import json as _json
    from unittest.mock import patch

    import pytest

    from adeu.ingest import _extract_text_from_doc
    from adeu.server import mcp

    doc_path = tmp_path / "minimal.docx"
    doc = docx.Document()
    doc.add_paragraph("This is a minimal document with some basic text.")
    doc.add_paragraph("It has exactly two paragraphs.")
    doc.save(str(doc_path))

    output_path = tmp_path / "output.docx"

    edit = {
        "type": "modify",
        "target_text": "It has exactly two paragraphs.",
        "new_text": "It has exactly three paragraphs after editing.",
    }
    base_args = {
        "reasoning": "Testing the accepted changes payload shapes.",
        "original_docx_path": str(doc_path),
        "author_name": "Reviewer AI",
        "output_path": str(output_path),
    }

    ctx_patches = (
        patch("fastmcp.server.context.Context.info"),
        patch("fastmcp.server.context.Context.debug"),
        patch("fastmcp.server.context.Context.warning"),
        patch("fastmcp.server.context.Context.error"),
    )

    # 1. Per-item stringification is repaired (the Gemini double-serialize quirk).
    with ctx_patches[0], ctx_patches[1], ctx_patches[2], ctx_patches[3]:
        result = asyncio.run(mcp.call_tool("process_document_batch", {**base_args, "changes": [_json.dumps(edit)]}))

    text = "".join(item.text for item in result.content if item.type == "text")
    assert "Batch complete" in text
    assert output_path.exists()
    clean_text = _extract_text_from_doc(docx.Document(str(output_path)), clean_view=True)
    assert "It has exactly three paragraphs after editing." in clean_text

    # 2. A WHOLLY stringified payload is rejected, not silently repaired and
    #    not silently treated as an empty batch.
    output_path.unlink()
    with (
        patch("fastmcp.server.context.Context.info"),
        patch("fastmcp.server.context.Context.debug"),
        patch("fastmcp.server.context.Context.warning"),
        patch("fastmcp.server.context.Context.error"),
        pytest.raises(Exception) as exc_info,
    ):
        asyncio.run(mcp.call_tool("process_document_batch", {**base_args, "changes": _json.dumps([edit])}))

    assert "changes" in str(exc_info.value)
    assert not output_path.exists(), "a rejected batch must not write a document"

    # 3. `changes` is advertised as REQUIRED — omitting it is a schema error,
    #    never an empty no-op batch reported as success.
    tools = asyncio.run(mcp.list_tools())
    pdb = next(t for t in tools if t.name == "process_document_batch")
    assert "changes" in pdb.parameters["required"]
    assert pdb.parameters["properties"]["changes"]["type"] == "array"
    assert "changes_json" not in pdb.parameters["properties"]


def test_repro_sanitize_double_percent_escaping(tmp_path, capsys):
    """
    Asserts that the similarity guard message printed to stderr uses single percent signs (%)
    instead of double percent signs (%%) when a low similarity baseline is checked, both
    in the blocked block reason and in the low-similarity warning.
    """
    import docx

    # 1. Create two completely different documents to trigger the similarity guard
    doc_a_path = tmp_path / "doc_a.docx"
    doc_a = docx.Document()
    doc_a.add_paragraph("This is document A")
    doc_a.save(str(doc_a_path))

    doc_b_path = tmp_path / "doc_b.docx"
    doc_b = docx.Document()
    doc_b.add_paragraph("Something totally different and unrelated")
    doc_b.save(str(doc_b_path))

    out_path = tmp_path / "out.docx"

    # 2. Run sanitize without allowing low similarity (should block)
    code, stdout, stderr = run_cli(
        ["sanitize", str(doc_a_path), "-o", str(out_path), "--baseline", str(doc_b_path), "--report"], capsys
    )

    assert code == 1
    # The message should contain single percent signs '%' and NOT double percent signs '%%'
    assert "share only 31%" in stderr or "share only 31%" in stdout
    assert "69% differs" in stderr or "69% differs" in stdout
    assert "31%%" not in stderr and "31%%" not in stdout
    assert "69%%" not in stderr and "69%%" not in stdout

    # 3. Run sanitize allowing low similarity (should warning-proceed)
    code, stdout, stderr = run_cli(
        [
            "sanitize",
            str(doc_a_path),
            "-o",
            str(out_path),
            "--baseline",
            str(doc_b_path),
            "--report",
            "--allow-low-similarity-baseline",
        ],
        capsys,
    )

    assert code == 0
    # The warning should contain single percent signs '%' and NOT double percent signs '%%'
    assert "share only 31%" in stderr or "share only 31%" in stdout
    assert "69% differs" in stderr or "69% differs" in stdout
    assert "31%%" not in stderr and "31%%" not in stdout
    assert "69%%" not in stderr and "69%%" not in stdout


def test_process_document_batch_relaxed_validation_repro(tmp_path):
    """
    Verifies that passing a list of changes containing invalid elements (such as `1`)
    to the `changes` parameter of `process_document_batch` does NOT raise a raw
    FastMCP ValidationError (Pydantic validation error) during tool parameter validation,
    but instead successfully executes the tool and returns a friendly validation error message.
    """
    import asyncio
    from unittest.mock import patch

    import docx

    from adeu.server import mcp

    doc_path = tmp_path / "minimal.docx"
    doc = docx.Document()
    doc.add_paragraph("This is a minimal document with some basic text.")
    doc.save(str(doc_path))

    output_path = tmp_path / "output.docx"

    arguments = {
        "reasoning": "Attempting modification via changes parameter with invalid elements to test error handling",
        "original_docx_path": str(doc_path),
        "author_name": "Reviewer AI",
        "output_path": str(output_path),
        "changes": [1],
    }

    # Patch FastMCP Context logging to avoid session-not-established RuntimeError
    with (
        patch("fastmcp.server.context.Context.info"),
        patch("fastmcp.server.context.Context.debug"),
        patch("fastmcp.server.context.Context.warning"),
        patch("fastmcp.server.context.Context.error"),
    ):
        result = asyncio.run(mcp.call_tool("process_document_batch", arguments))

    # The expected behavior is that the tool executes successfully and returns a friendly error message
    # rather than raising a fastmcp.exceptions.ValidationError before execution.
    text = "".join(item.text for item in result.content if item.type == "text")
    assert "Error: No valid changes to apply" in text
    assert "changes[0]:" in text


def test_table_row_match_mode_all(tmp_path):
    import io

    import docx

    from adeu.ingest import _extract_text_from_doc
    from adeu.models import DeleteTableRow, InsertTableRow
    from adeu.redline.engine import RedlineEngine

    doc = docx.Document()
    table = doc.add_table(rows=3, cols=3)
    # Row 0: ID | Name | Notes
    # Row 1: 1 | Alice | First record
    # Row 2: 2 |       | Second record with empty name
    for row, data in zip(
        table.rows,
        [["ID", "Name", "Notes"], ["1", "Alice", "First record"], ["2", "", "Second record with empty name"]],
        strict=True,
    ):
        for cell, text in zip(row.cells, data, strict=True):
            cell.text = text

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # --- Test Delete All ---
    engine_del = RedlineEngine(stream)
    stats_del = engine_del.process_batch([DeleteTableRow(type="delete_row", target_text="record", match_mode="all")])

    assert stats_del["edits_applied"] == 1
    assert stats_del["occurrences_modified"] == 2

    engine_del.accept_all_revisions()
    clean_text_del = _extract_text_from_doc(docx.Document(engine_del.save_to_stream()), clean_view=True)
    assert "First record" not in clean_text_del
    assert "Second record" not in clean_text_del

    # --- Test Insert All ---
    stream.seek(0)
    engine_ins = RedlineEngine(stream)
    stats_ins = engine_ins.process_batch(
        [
            InsertTableRow(
                type="insert_row",
                target_text="record",
                match_mode="all",
                position="below",
                cells=["NEW_ID", "NEW_NAME", "NEW_NOTES"],
            )
        ]
    )

    assert stats_ins["edits_applied"] == 1
    assert stats_ins["occurrences_modified"] == 2

    engine_ins.accept_all_revisions()
    clean_text_ins = _extract_text_from_doc(docx.Document(engine_ins.save_to_stream()), clean_view=True)

    # We should have exactly 2 injected rows
    occurrences = clean_text_ins.count("NEW_ID | NEW_NAME | NEW_NOTES")
    assert occurrences == 2


def test_search_query_paragraph_filtering(tmp_path):
    """
    Ensures that when a document is searched using the MCP read_docx tool,
    the snippet returned strictly bounds to the paragraph the match lives in,
    instead of an arbitrary 100-character window that leaks neighboring lines.
    """
    import asyncio
    from unittest.mock import patch

    import docx

    from adeu.server import mcp

    doc = docx.Document()
    doc.add_paragraph("Unicode Test Document")
    doc.add_paragraph("This is some English text.")
    doc.add_paragraph("Chinese: 一些中文 and 更多文本")
    doc.add_paragraph("Accented: Café, naïve, résumé, garçon, déjà vu, Straße.")
    doc.add_paragraph("Emojis & Symbols: 🌟 🦄 💻 ⚙️ ⛩️ 🎴日本語 🌍")

    doc_path = tmp_path / "search.docx"
    doc.save(str(doc_path))

    arguments = {
        "reasoning": "Filter document to target paragraph.",
        "file_path": str(doc_path),
        "search_query": "Chinese",
    }

    # The tool logs progress through the MCP session, which does not exist when
    # a tool is invoked directly outside a client connection.
    with (
        patch("fastmcp.server.context.Context.info"),
        patch("fastmcp.server.context.Context.debug"),
        patch("fastmcp.server.context.Context.warning"),
        patch("fastmcp.server.context.Context.error"),
    ):
        result = asyncio.run(mcp.call_tool("read_docx", arguments))
    text = "".join(item.text for item in result.content if item.type == "text")

    # The query MUST be matched and returned.
    assert "**Chinese**" in text

    # The search query should filter out non-matching paragraphs to conserve LLM context.
    assert "This is some English text." not in text
    assert "Accented: Café" not in text
    assert "Emojis & Symbols" not in text


def test_read_docx_appendix_mode_schema_compat(tmp_path):
    """
    Verifies that 'appendix' is a valid mode for the read_docx tool,
    allowing the client/LLM to retrieve defined terms and diagnostics
    without triggering a Pydantic schema validation error.
    """
    import asyncio
    from unittest.mock import patch

    import docx

    from adeu.server import mcp

    # 1. Create a minimal document
    doc = docx.Document()
    doc.add_paragraph('The term (the "Agreement") shall mean this contract.')

    doc_path = tmp_path / "test_appendix_mode.docx"
    doc.save(str(doc_path))

    arguments = {
        "reasoning": "Retrieve defined terms and references.",
        "file_path": str(doc_path),
        "mode": "appendix",
    }

    # The tool logs progress through the MCP session, which does not exist when
    # a tool is invoked directly outside a client connection.
    with (
        patch("fastmcp.server.context.Context.info"),
        patch("fastmcp.server.context.Context.debug"),
        patch("fastmcp.server.context.Context.warning"),
        patch("fastmcp.server.context.Context.error"),
    ):
        result = asyncio.run(mcp.call_tool("read_docx", arguments))

    # Get the text content of the result
    text = "".join(item.text for item in result.content if item.type == "text")

    # Assert that the result was successful and contains appendix information
    assert "Agreement" in text

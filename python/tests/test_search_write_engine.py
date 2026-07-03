import asyncio

import pytest
from docx import Document

from adeu.ingest import extract_text_from_stream
from adeu.mcp_components.tools.document import process_document_batch, read_docx
from adeu.models import ModifyText


class MockContext:
    """Mock FastMCP Context to absorb async logging calls during tests."""

    async def info(self, msg, **kwargs):
        pass

    async def debug(self, msg, **kwargs):
        pass

    async def warning(self, msg, **kwargs):
        pass

    async def error(self, msg, **kwargs):
        pass


@pytest.fixture
def repetitive_doc(tmp_path) -> str:
    """Creates a DOCX with duplicate text and multiple headings to test targeted writes and search."""
    doc = Document()
    doc.add_heading("1. Confidentiality", level=1)
    doc.add_paragraph(
        "The Recipient shall maintain confidentiality of all Confidential Information "
        "for a period of two (2) years from the date of disclosure."
    )

    # Pad with some empty paragraphs to simulate a page break in a real doc
    for _ in range(10):
        doc.add_paragraph("")

    doc.add_heading("2. Term & Termination", level=1)
    doc.add_paragraph("This agreement shall remain in effect for a period of two (2) years.")

    path = tmp_path / "repetitive.docx"
    doc.save(path)
    return str(path)


# ---------------------------------------------------------
# read_docx Search Tests
# ---------------------------------------------------------


def test_read_docx_search_query_exact(repetitive_doc):
    ctx = MockContext()
    # Call the tool with the new search_query parameter
    result = asyncio.run(
        read_docx(
            reasoning="test",
            file_path=repetitive_doc,
            ctx=ctx,
            search_query="two (2) years",
        )
    )

    content = str(result.content)

    # Verify the new channel contract formatting
    assert "> **Search Results** — Found 2 matches" in content
    assert "### Match 1" in content
    assert "### Match 2" in content
    assert "**Path:** `1. Confidentiality`" in content
    assert "This exact phrasing appears 2 times" in content


def test_read_docx_search_regex(repetitive_doc):
    ctx = MockContext()
    # Call the tool with regex search
    result = asyncio.run(
        read_docx(
            reasoning="test",
            file_path=repetitive_doc,
            ctx=ctx,
            search_query=r"two \(\d+\) years",
            search_regex=True,
        )
    )

    content = str(result.content)
    assert "> **Search Results** — Found 2 matches" in content


def test_read_docx_outline_page_ranges(repetitive_doc):
    ctx = MockContext()
    result = asyncio.run(read_docx(reasoning="test", file_path=repetitive_doc, ctx=ctx, mode="outline"))

    content = str(result.content)

    # Redundant style labels should be gone (no "Heading 1")
    assert "Heading 1" not in content

    # Should project ranges (e.g. p1-p1 or p1)
    # The spec format requires `# 1. Confidentiality (p1)` or `(p1-p2)`
    assert "# 1. Confidentiality (p" in content
    assert "# 2. Term & Termination (p" in content


# ---------------------------------------------------------
# process_document_batch Targeted Write Tests
# ---------------------------------------------------------


def test_process_batch_match_mode_strict_fails(repetitive_doc, tmp_path):
    ctx = MockContext()

    # Mode defaults to 'strict'. The string "two (2) years" appears twice.
    # It should throw a BatchValidationError / reject the batch.
    edits = [ModifyText(target_text="two (2) years", new_text="five (5) years")]

    result = asyncio.run(
        process_document_batch(
            reasoning="test",
            original_docx_path=repetitive_doc,
            author_name="AI Agent",
            ctx=ctx,
            changes=edits,
            output_path=str(tmp_path / "out_strict.docx"),
        )
    )

    assert "Batch rejected" in result
    assert "ambiguous" in result.lower() or "appears multiple times" in result.lower()


def test_process_batch_match_mode_first(repetitive_doc, tmp_path):
    ctx = MockContext()

    # We bypass Pydantic initialization checks for the test since the model
    # doesn't officially have these fields yet.
    edit = ModifyText(target_text="two (2) years", new_text="five (5) years")
    edit.match_mode = "first"

    result = asyncio.run(
        process_document_batch(
            reasoning="test",
            original_docx_path=repetitive_doc,
            author_name="AI Agent",
            ctx=ctx,
            changes=[edit],
            output_path=str(tmp_path / "out_first.docx"),
        )
    )

    assert "1 applied" in result
    assert "Mode:** `first` (1 occurrence modified" in result  # Enriched Report check

    with open(str(tmp_path / "out_first.docx"), "rb") as f:
        text = extract_text_from_stream(f, clean_view=True)

    # Only the first one should be modified
    assert text.count("five (5) years") == 1
    assert text.count("two (2) years") == 1


def test_process_batch_match_mode_all(repetitive_doc, tmp_path):
    ctx = MockContext()

    edit = ModifyText(target_text="two (2) years", new_text="five (5) years")
    edit.match_mode = "all"

    result = asyncio.run(
        process_document_batch(
            reasoning="test",
            original_docx_path=repetitive_doc,
            author_name="AI Agent",
            ctx=ctx,
            changes=[edit],
            output_path=str(tmp_path / "out_all.docx"),
        )
    )

    assert "2 applied" in result
    assert "Mode:** `all` (2 occurrences modified" in result  # Enriched Report check

    with open(str(tmp_path / "out_all.docx"), "rb") as f:
        text = extract_text_from_stream(f, clean_view=True)

    # Both should be modified
    assert text.count("five (5) years") == 2
    assert text.count("two (2) years") == 0


def test_process_batch_regex(repetitive_doc, tmp_path):
    ctx = MockContext()

    # Using python regex capture group \1
    edit = ModifyText(target_text=r"two \((\d+)\) years", new_text=r"five (\1) years")
    edit.regex = True
    edit.match_mode = "all"

    result = asyncio.run(
        process_document_batch(
            reasoning="test",
            original_docx_path=repetitive_doc,
            author_name="AI Agent",
            ctx=ctx,
            changes=[edit],
            output_path=str(tmp_path / "out_regex.docx"),
        )
    )

    assert "2 applied" in result

    with open(str(tmp_path / "out_regex.docx"), "rb") as f:
        text = extract_text_from_stream(f, clean_view=True)
        assert text.count("five (2) years") == 2

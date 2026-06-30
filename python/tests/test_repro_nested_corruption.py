import io

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from adeu.ingest import extract_text_from_stream
from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine


def test_repro_nested_edit_corruption():
    """
    REPRO: Attempting to edit text that is ALREADY inside a w:ins tag
    should ideally work or fail gracefully, but currently causes corruption/truncation.
    """
    doc = Document()
    p = doc.add_paragraph("Start ")

    # Simulate an existing Tracked Change (Insertion)
    # <w:ins w:id="1" w:author="Other"><w:r><w:t>Existing Insert</w:t></w:r></w:ins>
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "1")
    ins.set(qn("w:author"), "Other")

    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Existing Insert"
    run.append(t)
    ins.append(run)
    p._element.append(ins)

    p_run_end = OxmlElement("w:r")
    t_end = OxmlElement("w:t")
    t_end.text = " End"
    p_run_end.append(t_end)
    p._element.append(p_run_end)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Verify Ingest sees it (currently sees "Existing Insert" as plain text)
    initial_text = extract_text_from_stream(stream)
    assert "Existing Insert" in initial_text

    # Action: Try to edit the "Existing Insert" text
    # This triggers the "edit inside edit" scenario
    edit = ModifyText(target_text="Existing Insert", new_text="Modified Insert")

    engine = RedlineEngine(stream, author="Me")
    engine.apply_edits([edit])

    res_stream = engine.save_to_stream()

    # Check 1: Does output text look right?
    final_text = extract_text_from_stream(res_stream, clean_view=True)

    # Failure condition reported by user: Truncation or missing text
    # We expect "Start Modified Insert End"
    # If corruption happened, we might see "Start  End" or broken XML
    print(f"Final Text: {final_text}")

    assert "Modified Insert" in final_text, "Edit was lost or corrupted"

    # Check 2: XML Validity (Manual Inspection logic)
    # We don't want nested <w:ins><w:ins>...</w:ins></w:ins>
    doc_res = Document(res_stream)
    xml = doc_res.element.xml

    if xml.count("<w:ins") > 1:
        # Check nesting
        # A simple string check isn't perfect but nested tags usually look like:
        # <w:ins ...><w:ins ...>
        import re

        if re.search(r"<w:ins[^>]*>.*?<w:ins", xml, re.DOTALL):
            print("WARNING: Nested w:ins detected!")
            # This is technically what we want to fix, but for this repro
            # we just want to confirm if it breaks the doc content.


def test_repro_surgical_edit_inside_insertion():
    """
    Issue 3: Edits inside insertions replace the entire insertion.

    Scenario:
    Round 1: Author A inserts " The notice period is 60 days."
    Round 2: Author B edits "60" to "45" inside that insertion.

    Expected:
    The engine splits the existing <w:ins> tag and surgically deletes/inserts the numbers.
    It MUST NOT delete the entire "The notice period is 60 days." clause.
    """
    doc = Document()
    doc.add_paragraph("Base text.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # 1. Simulate Round 1 (Author A inserts a new clause)
    engine1 = RedlineEngine(stream, author="Author A")
    edit1 = ModifyText(target_text="Base text.", new_text="Base text. The notice period is 60 days.")
    engine1.apply_edits([edit1])
    mid_stream = engine1.save_to_stream()

    # 2. Simulate Round 2 (Author B changes "60" to "45" inside the insertion)
    engine2 = RedlineEngine(mid_stream, author="Author B")
    edit2 = ModifyText(target_text="60", new_text="45", comment="Standardizing to 45 days")
    applied, skipped = engine2.apply_edits([edit2])

    assert applied == 1, "The engine skipped the surgical edit."

    final_stream = engine2.save_to_stream()
    doc_final = Document(final_stream)

    # 3. VERIFY NO WHOLESALE REPLACEMENT
    # If the bug is present, we will see <w:delText> The notice period is 60 days.</w:delText>
    del_texts = [t.text for t in doc_final.element.xpath("//w:del//w:delText") if t.text]
    full_deleted_text = "".join(del_texts)

    assert "notice period" not in full_deleted_text, "BUG PRESENT: The entire inserted clause was deleted wholesale!"

    # 4. VERIFY SURGICAL DELETION & INSERTION
    assert "60" in del_texts, "The specific target '60' was not deleted."

    ins_texts = [t.text for t in doc_final.element.xpath("//w:ins//w:t") if t.text]
    assert "45" in ins_texts, "The new text '45' was not inserted."

    # 5. VERIFY XML VALIDITY
    # <w:del> nested inside <w:ins> is the CANONICAL OOXML representation of
    # "this author deletes another author's still-pending insertion" (both
    # w:ins and w:del are EG_RunLevelElts and their content models recursively
    # re-admit run-level elements; it is exactly what Word itself authors). It
    # is REQUIRED here so that reject-all reverts the contingent text to nothing
    # instead of promoting it to committed body text. So we now EXPECT it.
    expected_nested_del = doc_final.element.xpath("//w:ins//w:del")
    assert expected_nested_del, "Expected <w:del> nested inside the foreign author's <w:ins>."

    # <w:ins> directly inside <w:ins> is the structure Word does NOT author and
    # normalizes/repairs; the engine must split the enclosing insertion so the
    # new insertion is a sibling. This invariant stays.
    invalid_nested_ins = doc_final.element.xpath("//w:ins//w:ins")
    assert not invalid_nested_ins, "Corrupt XML: Found <w:ins> nested inside <w:ins>."

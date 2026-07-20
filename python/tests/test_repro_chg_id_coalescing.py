# FILE: tests/test_repro_chg_id_coalescing.py
"""
Regression tests for the Chg-ID coalescing + annotation fix.

Two related fixes are protected here:

  PROBLEM B (ID coalescing, primary): A single ModifyText whose target spans
  multiple <w:r> elements (e.g. because a bold word forces OOXML to split the
  text across runs) used to mint a fresh w:id per <w:del> element, surfacing as
  N [Chg:N] entries in the projected bubble. After the fix, all <w:del>
  elements produced by one logical edit share a single w:id; same for <w:ins>.

  PROBLEM A (bubble annotation): The projected metadata bubble now annotates
  Chg entries with their kind ([Chg:N delete] / [Chg:N insert] / [Chg:N format])
  so an agent reading the bubble can tell which side of a substitution each ID
  represents.

If any test in this file regresses, look at:
  - src/adeu/redline/engine.py: _create_track_change_tag(reuse_id=...),
    track_delete_run(reuse_id=...), track_insert(reuse_id=...),
    _apply_single_edit_indexed (allocates one del_id and one ins_id per
    logical operation).
  - src/adeu/redline/mapper.py: _build_merged_meta_block (annotates ins_map
    entries with 'insert' and del_map entries with 'delete').
  - src/adeu/ingest.py: _build_merged_meta_block (same annotation, plus
    'format' for tracked formatting changes via fmt_map).
"""

import io
import re

from docx import Document
from docx.oxml.ns import qn

from adeu.ingest import extract_text_from_stream
from adeu.models import AcceptChange, ModifyText, RejectChange
from adeu.redline.engine import RedlineEngine

# ----------------------------------------------------------------------------
# Helpers
# ----------------------------------------------------------------------------


def _build_doc_with_bold_span() -> io.BytesIO:
    """
    Builds an in-memory DOCX where 'Scotland Limited' is split across two
    <w:r> elements: 'Scotland' (bold) + ' Limited' (plain). A single modify
    targeting both will resolve to two real runs.
    """
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Governed by the laws of ")
    bold_run = p.add_run("Scotland")
    bold_run.bold = True
    p.add_run(" Limited and any disputes shall be resolved there.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def _build_simple_doc(text: str = "Hello world.") -> io.BytesIO:
    doc = Document()
    doc.add_paragraph(text)
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def _collect_revision_ids(doc) -> tuple[list[str], list[str]]:
    """Returns (del_ids_in_document_order, ins_ids_in_document_order)."""
    del_ids = [el.get(qn("w:id")) for el in doc.element.xpath("//w:del")]
    ins_ids = [el.get(qn("w:id")) for el in doc.element.xpath("//w:ins")]
    return del_ids, ins_ids


# ----------------------------------------------------------------------------
# Fix B: multi-run target produces shared w:id per side
# ----------------------------------------------------------------------------


def test_multi_run_target_shares_one_del_id():
    """
    A single ModifyText whose target spans two <w:r> elements (because of a
    bold word) must produce multiple <w:del> elements that share a single w:id.
    """
    stream = _build_doc_with_bold_span()
    engine = RedlineEngine(stream, author="Reviewer AI")

    edit = ModifyText(
        target_text="**Scotland** Limited",
        new_text="England and Wales PLC",
    )
    stats = engine.process_batch([edit])
    assert stats["edits_applied"] == 1
    assert stats["edits_skipped"] == 0

    out_doc = Document(engine.save_to_stream())
    del_ids, ins_ids = _collect_revision_ids(out_doc)

    # Two <w:del> elements (one per real run touched), but they must share one id.
    assert len(del_ids) == 2, f"Expected 2 <w:del> elements, got {len(del_ids)}: {del_ids}"
    assert del_ids[0] == del_ids[1], (
        f"Multi-run delete must share one w:id (Problem B). Got {del_ids}. "
        "If this fails, check that track_delete_run is being called with "
        "reuse_id=del_id from _apply_single_edit_indexed's MODIFICATION branch."
    )

    # One <w:ins> with its own distinct id.
    assert len(ins_ids) == 1, f"Expected 1 <w:ins>, got {len(ins_ids)}: {ins_ids}"
    assert ins_ids[0] != del_ids[0], (
        f"Insert and delete must have distinct ids; got del={del_ids[0]}, ins={ins_ids[0]}. "
        "Coalescing del+ins into a single id was explicitly rejected — Word treats them as "
        "separate review entries, so the projection must too."
    )


def test_multi_run_target_bubble_shows_two_chg_entries():
    """
    The projected bubble for a multi-run modify must show exactly two [Chg:N]
    entries (one delete, one insert) — not three or more. This is the agent-
    visible surface of Fix B.
    """
    stream = _build_doc_with_bold_span()
    engine = RedlineEngine(stream, author="Reviewer AI")
    engine.process_batch([ModifyText(target_text="**Scotland** Limited", new_text="England and Wales PLC")])
    text = extract_text_from_stream(engine.save_to_stream())

    chg_entries = re.findall(r"\[Chg:\d+(?:\s+\w+)?\]", text)
    assert len(chg_entries) == 2, (
        f"Expected exactly 2 [Chg:N] entries in bubble, got {len(chg_entries)}: {chg_entries}. "
        f"Full projected text:\n{text}"
    )


# ----------------------------------------------------------------------------
# Fix A: bubble annotation
# ----------------------------------------------------------------------------


def test_bubble_annotates_delete_and_insert():
    """
    The projected bubble must annotate Chg entries with 'delete' or 'insert'
    so an agent can tell which side of a substitution each ID represents.
    """
    stream = _build_simple_doc("The quick brown fox.")
    engine = RedlineEngine(stream, author="Reviewer AI")
    engine.process_batch([ModifyText(target_text="quick", new_text="fast")])
    text = extract_text_from_stream(engine.save_to_stream())

    delete_entries = re.findall(r"\[Chg:\d+ delete\]", text)
    insert_entries = re.findall(r"\[Chg:\d+ insert\]", text)

    assert len(delete_entries) == 1, (
        f"Expected exactly 1 [Chg:N delete] entry, got {len(delete_entries)}. Text:\n{text}"
    )
    assert len(insert_entries) == 1, (
        f"Expected exactly 1 [Chg:N insert] entry, got {len(insert_entries)}. Text:\n{text}"
    )

    # Belt-and-braces: the unannotated form must NOT appear.
    bare = re.findall(r"\[Chg:\d+\](?!\s)", text)  # [Chg:N] not followed by whitespace
    assert not bare, (
        f"Found unannotated [Chg:N] entries in bubble: {bare}. "
        "Both mapper.py and ingest.py _build_merged_meta_block must annotate."
    )


# ----------------------------------------------------------------------------
# Behavioural guard: shared-ID elements still resolve together on accept
# ----------------------------------------------------------------------------


def test_accept_resolves_all_shared_id_elements_together():
    """
    Accepting a coalesced delete (one w:id, two <w:del> elements) must remove
    BOTH elements and finalize the deletion across the whole logical edit.
    Guards against a future refactor that resolves only one element per call.
    """
    stream = _build_doc_with_bold_span()
    engine = RedlineEngine(stream, author="Reviewer AI")
    engine.process_batch([ModifyText(target_text="**Scotland** Limited", new_text="England and Wales PLC")])
    redlined_stream = engine.save_to_stream()

    # Find the coalesced delete id directly from the XML.
    redlined_doc = Document(redlined_stream)
    del_ids, ins_ids = _collect_revision_ids(redlined_doc)
    assert del_ids[0] == del_ids[1], "Precondition: deletes must share an id"
    coalesced_del_id = del_ids[0]
    coalesced_ins_id = ins_ids[0]

    # Now accept both sides.
    redlined_stream.seek(0)
    engine2 = RedlineEngine(redlined_stream, author="Reviewer AI")
    applied, skipped, already_resolved = engine2.apply_review_actions(
        [
            AcceptChange(target_id=f"Chg:{coalesced_del_id}"),
            AcceptChange(target_id=f"Chg:{coalesced_ins_id}"),
        ]
    )
    # The del+ins pair resolves as ONE unit on the first accept; the second
    # action is an accurate no-op, never a second "applied" transition
    # (QA 2026-07-19 ADEU-QA-004).
    assert applied == 1, f"Expected 1 action applied, got {applied} (skipped={skipped})"
    assert already_resolved == 1
    assert skipped == 0

    final_doc = Document(engine2.save_to_stream())
    final_del_ids, final_ins_ids = _collect_revision_ids(final_doc)
    assert not final_del_ids, f"Deletes should be fully resolved; remaining: {final_del_ids}"
    assert not final_ins_ids, f"Inserts should be fully resolved; remaining: {final_ins_ids}"

    final_text = extract_text_from_stream(engine2.save_to_stream())
    assert "Scotland" not in final_text
    assert " Limited" not in final_text
    assert "England and Wales PLC" in final_text


def test_reject_resolves_all_shared_id_elements_together():
    """
    Symmetric to the accept test: rejecting must revert the full coalesced
    edit, restoring the original text.
    """
    stream = _build_doc_with_bold_span()
    engine = RedlineEngine(stream, author="Reviewer AI")
    engine.process_batch([ModifyText(target_text="**Scotland** Limited", new_text="England and Wales PLC")])
    redlined_stream = engine.save_to_stream()

    redlined_doc = Document(redlined_stream)
    del_ids, ins_ids = _collect_revision_ids(redlined_doc)
    coalesced_del_id = del_ids[0]
    coalesced_ins_id = ins_ids[0]

    redlined_stream.seek(0)
    engine2 = RedlineEngine(redlined_stream, author="Reviewer AI")
    applied, _, already_resolved = engine2.apply_review_actions(
        [
            RejectChange(target_id=f"Chg:{coalesced_del_id}"),
            RejectChange(target_id=f"Chg:{coalesced_ins_id}"),
        ]
    )
    # One resolution unit + one accurate no-op (QA 2026-07-19 ADEU-QA-004).
    assert applied == 1
    assert already_resolved == 1

    final_doc = Document(engine2.save_to_stream())
    final_del_ids, final_ins_ids = _collect_revision_ids(final_doc)
    assert not final_del_ids
    assert not final_ins_ids

    final_text = extract_text_from_stream(engine2.save_to_stream())
    # After reject, original formatting (bold on 'Scotland') must be preserved.
    # The projected Markdown therefore contains '**Scotland**', not 'Scotland'
    # surrounded by space — so we check both sides of the original split
    # independently rather than asserting on the concatenated form.
    assert "**Scotland**" in final_text, f"Bold 'Scotland' must be restored. Got:\n{final_text}"
    assert " Limited" in final_text, f"' Limited' must be restored. Got:\n{final_text}"
    assert "England and Wales PLC" not in final_text, f"Inserted text must be removed by reject. Got:\n{final_text}"


# ----------------------------------------------------------------------------
# Negative guard: don't over-coalesce
# ----------------------------------------------------------------------------


def test_simple_modify_still_produces_two_distinct_ids():
    """
    A 1-run modify must produce 1 <w:del> and 1 <w:ins> with DISTINCT w:ids.
    We deliberately do NOT coalesce del+ins into one synthetic id, because
    Word's Review pane treats them as two separate revisions and the projection
    must match Word's view.
    """
    stream = _build_simple_doc("The quick brown fox.")
    engine = RedlineEngine(stream, author="Reviewer AI")
    engine.process_batch([ModifyText(target_text="quick", new_text="fast")])

    out_doc = Document(engine.save_to_stream())
    del_ids, ins_ids = _collect_revision_ids(out_doc)

    assert len(del_ids) == 1
    assert len(ins_ids) == 1
    assert del_ids[0] != ins_ids[0], (
        "Delete and insert must have distinct ids. Coalescing them was an "
        "explicit non-goal — see the design discussion in the original fix."
    )


def test_three_modifies_in_one_batch_produce_six_distinct_ids():
    """
    Three independent ModifyText edits in a single batch must produce six
    distinct ids (3 del + 3 ins). Guards against a future change that tries
    to share ids across logically-independent edits in a batch.
    """
    doc = Document()
    doc.add_paragraph("Para 1")
    doc.add_paragraph("Para 2")
    doc.add_paragraph("Para 3")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="Reviewer AI")
    stats = engine.process_batch(
        [
            ModifyText(target_text="Para 1", new_text="Para One"),
            ModifyText(target_text="Para 2", new_text="Para Two"),
            ModifyText(target_text="Para 3", new_text="Para Three"),
        ]
    )
    assert stats["edits_applied"] == 3
    assert stats["edits_skipped"] == 0

    out_doc = Document(engine.save_to_stream())
    del_ids, ins_ids = _collect_revision_ids(out_doc)

    assert len(del_ids) == 3, f"Expected 3 <w:del>, got {len(del_ids)}"
    assert len(ins_ids) == 3, f"Expected 3 <w:ins>, got {len(ins_ids)}"
    assert len(set(del_ids + ins_ids)) == 6, (
        f"Expected 6 distinct ids across 3 independent edits, got "
        f"del_ids={del_ids}, ins_ids={ins_ids}, distinct={set(del_ids + ins_ids)}"
    )


# ----------------------------------------------------------------------------
# Multi-paragraph insertion: same ins id across all <w:ins> elements
# ----------------------------------------------------------------------------


def test_multi_paragraph_insertion_shares_one_ins_id():
    """
    A single ModifyText whose new_text spans multiple paragraphs produces
    multiple <w:ins> elements (one per paragraph, plus paragraph-break trackers
    inside <w:pPr><w:rPr>). All of them must share a single w:id, because
    they represent one logical insertion from the agent's point of view.
    """
    stream = _build_simple_doc("Before. Anchor. After.")
    engine = RedlineEngine(stream, author="Reviewer AI")
    engine.process_batch([ModifyText(target_text="Anchor.", new_text="Line one.\n\nLine two.\n\nLine three.")])

    out_doc = Document(engine.save_to_stream())
    _, ins_ids = _collect_revision_ids(out_doc)

    assert len(ins_ids) >= 2, f"Multi-paragraph insert should produce multiple <w:ins>, got {ins_ids}"
    assert len(set(ins_ids)) == 1, (
        f"All <w:ins> elements from a single multi-paragraph ModifyText must share one w:id. "
        f"Got distinct ids: {set(ins_ids)}. Check that track_insert is propagating reuse_id "
        "to every _create_track_change_tag('w:ins', ...) call site (there are several inside "
        "the heading-styled-block path AND the inline+remaining-lines path)."
    )

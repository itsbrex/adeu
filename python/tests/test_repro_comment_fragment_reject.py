"""
Reproduction for the CLI QA report (2026-07-22, bug #1):

    "Silent comment deletion on partial rejection of a fragmented edit.

    When a single `modify` edit gets split by word-level diffing into multiple
    tracked-change fragments — e.g. \"thirty (30) days'\" -> \"sixty (60) days'\"
    splits into two separate Chg pairs — and a comment is attached to that edit,
    the comment anchor (commentRangeStart/commentRangeEnd) wraps only ONE
    fragment, not the full original span. Rejecting that fragment reverts the
    fragment correctly but SILENTLY DELETES the comment (and any reply thread),
    while rejecting the other fragment leaves the comment intact."

Root cause (python/src/adeu/redline/engine.py, _word_diff_sub_edits): a commented
modify was fanned out into one sub-edit per word-level hunk, and the comment was
attached to a single fragment. The reject path then treats that fragment's
insertion cleanup as license to remove the wrapping comment, so rejecting the
fragment that happens to own the comment destroys the whole thread with no
warning — `Actions: 1 applied` gives no hint.

Fix (per operator guidance): do NOT word-split a modify that carries a comment.
Keep it as a single contiguous tracked change (still trimming shared prefix/
suffix so the redline stays minimal at the edges) so the comment anchor always
wraps the whole logical edit. A commented change then has exactly one Chg pair
and rejecting it reverts the entire edit atomically — there is no "other half"
to leave behind and no fragment whose rejection silently orphans the comment.

STYLE: these assert the DESIRED behaviour, so they are RED before the fix and
GREEN after. The final control pins the boundary of the fix: an *uncommented*
modify must still word-split into minimal fragments.
"""

import io
import re
import zipfile

from docx import Document

from adeu.ingest import extract_text_from_stream
from adeu.models import ModifyText, RejectChange
from adeu.redline.engine import RedlineEngine

COMMENT_TEXT = "Client requested a longer notice period; please confirm with legal."


def _comment_survives(stream: io.BytesIO) -> bool:
    """True if the comment body text is still present in any comments part."""
    zf = zipfile.ZipFile(io.BytesIO(stream.getvalue()))
    for name in zf.namelist():
        if re.match(r"word/comments\d*\.xml$", name):
            if "Client requested" in zf.read(name).decode("utf-8"):
                return True
    return False


def _revision_ids(stream: io.BytesIO) -> tuple[set, set]:
    """Distinct w:id values on <w:ins> / <w:del> in the main document part."""
    zf = zipfile.ZipFile(io.BytesIO(stream.getvalue()))
    xml = zf.read("word/document.xml").decode("utf-8")
    ins_ids = set(re.findall(r"<w:ins\b[^>]*\bw:id=\"(\d+)\"", xml))
    del_ids = set(re.findall(r"<w:del\b[^>]*\bw:id=\"(\d+)\"", xml))
    return ins_ids, del_ids


def _fragmenting_modify_doc() -> io.BytesIO:
    """
    A clause whose amendment ("thirty (30) days'" -> "sixty (60) days'") is
    exactly the kind of multi-word change that word-level diffing fragments
    into two Chg pairs. The edit carries a comment.
    """
    doc = Document()
    doc.add_paragraph(
        "Either party may terminate this Agreement upon thirty (30) days' written notice to the other party."
    )
    base = io.BytesIO()
    doc.save(base)
    base.seek(0)

    eng = RedlineEngine(base, author="Reviewer")
    eng.apply_edits(
        [
            ModifyText(
                target_text="thirty (30) days'",
                new_text="sixty (60) days'",
                comment=COMMENT_TEXT,
            )
        ]
    )
    return eng.save_to_stream()


# ────────────────────────────────────────────────────────────────────────────
# RED — the reported bug.
# ────────────────────────────────────────────────────────────────────────────
def test_commented_modify_is_a_single_atomic_change():
    """
    A commented multi-word modify must land as ONE tracked change (one delete
    id + one insert id), not several word-level fragments. Otherwise the comment
    can only anchor to one fragment and rejecting a different fragment orphans
    or destroys it.
    """
    edited = _fragmenting_modify_doc()

    assert _comment_survives(edited)
    marked = extract_text_from_stream(io.BytesIO(edited.getvalue()), clean_view=False)
    assert COMMENT_TEXT in marked

    ins_ids, del_ids = _revision_ids(edited)
    assert len(ins_ids) == 1, f"expected one insertion revision, got {sorted(ins_ids)}"
    assert len(del_ids) == 1, f"expected one deletion revision, got {sorted(del_ids)}"


def test_rejecting_commented_change_reverts_whole_edit_no_fragment_left():
    """
    Because the commented change is atomic, rejecting EITHER paired id reverts
    the entire edit — the clean view returns to the original wording, with no
    surviving 'sixty'/'(60)' fragment. This is the guarantee that makes silent
    partial deletion impossible: there is no 'other half' to keep.
    """
    edited = _fragmenting_modify_doc()
    ins_ids, del_ids = _revision_ids(edited)
    all_ids = sorted(int(x) for x in (ins_ids | del_ids))

    for rid in all_ids:
        eng = RedlineEngine(io.BytesIO(edited.getvalue()), author="Reviewer")
        res = eng.process_batch([RejectChange(target_id=f"Chg:{rid}")])
        assert res["actions_applied"] == 1
        out = eng.save_to_stream()

        clean = extract_text_from_stream(io.BytesIO(out.getvalue()), clean_view=True)
        assert "thirty (30) days'" in clean, f"reject Chg:{rid} did not restore original"
        assert "sixty" not in clean and "(60)" not in clean, f"reject Chg:{rid} left a fragment behind: {clean!r}"


def test_rejecting_commented_change_flags_removed_comment_not_silent():
    """
    The report's other half: the removal must not be *silent*. When a reject
    (or accept) removes a comment because its anchor was inside the resolved
    change, the batch report must say so — an informational note, not a skip,
    so the action still counts as applied and the exit status stays success.
    """
    edited = _fragmenting_modify_doc()
    ins_ids, del_ids = _revision_ids(edited)
    a_change_id = sorted(int(x) for x in (ins_ids | del_ids))[0]

    eng = RedlineEngine(io.BytesIO(edited.getvalue()), author="Reviewer")
    res = eng.process_batch([RejectChange(target_id=f"Chg:{a_change_id}")])

    # Applied, not skipped — the note is informational only.
    assert res["actions_applied"] == 1
    assert res["actions_skipped"] == 0

    details = "\n".join(res["skipped_details"])
    assert "Com:" in details, f"comment removal was not reported: {details!r}"
    assert "removed" in details.lower()
    assert not _comment_survives(eng.save_to_stream())


def test_reject_uncommented_change_emits_no_removal_note():
    """Control: rejecting a change with no comment must not emit a removal note."""
    doc = Document()
    doc.add_paragraph("The color is red today.")
    base = io.BytesIO()
    doc.save(base)
    base.seek(0)
    eng = RedlineEngine(base, author="Reviewer")
    eng.apply_edits([ModifyText(target_text="red", new_text="blue")])
    edited = eng.save_to_stream()

    eng2 = RedlineEngine(io.BytesIO(edited.getvalue()), author="Reviewer")
    res = eng2.process_batch([RejectChange(target_id="Chg:1")])
    assert res["actions_applied"] == 1
    joined = "\n".join(res["skipped_details"]).lower()
    assert "removed comment" not in joined


# ────────────────────────────────────────────────────────────────────────────
# GREEN control — pin the boundary of the fix.
# ────────────────────────────────────────────────────────────────────────────
def test_control_uncommented_modify_still_word_splits():
    """
    The fix must be scoped to *commented* edits only. An identical modify with
    NO comment must still fan out into minimal word-level fragments (two Chg
    pairs here) — the surgical-diff behaviour praised in the QA report.
    """
    doc = Document()
    doc.add_paragraph(
        "Either party may terminate this Agreement upon thirty (30) days' written notice to the other party."
    )
    base = io.BytesIO()
    doc.save(base)
    base.seek(0)

    eng = RedlineEngine(base, author="Reviewer")
    eng.apply_edits([ModifyText(target_text="thirty (30) days'", new_text="sixty (60) days'")])
    out = eng.save_to_stream()

    ins_ids, del_ids = _revision_ids(out)
    assert len(ins_ids) >= 2 and len(del_ids) >= 2, (
        "uncommented modify should still word-split into minimal fragments; "
        f"got ins={sorted(ins_ids)} del={sorted(del_ids)}"
    )

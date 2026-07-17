import io

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine


def test_tc1_sequential_chaining_works_in_both_modes():
    """
    TC1 [report F1], extended 2026-07-17: sequential chaining is a first-class
    batch semantic in BOTH modes. Each edit is validated and applied against
    the document state produced by the edits before it, so a second modify may
    target the first modify's output — and the real run behaves exactly like
    the dry-run (QA M1 parity).
    """
    doc = Document()
    doc.add_paragraph("As defined in Section 1, the Recipient shall maintain confidentiality of all materials.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    def batch():
        return [
            ModifyText(target_text="the Recipient", new_text="Receiving Party"),
            ModifyText(target_text="Receiving Party", new_text="Disclosee"),
        ]

    engine = RedlineEngine(stream)
    res_dry = engine.process_batch(batch(), dry_run=True)
    assert res_dry["edits_applied"] == 2
    assert res_dry["edits_skipped"] == 0
    assert all(r["status"] == "applied" for r in res_dry["edits"])

    # Real run chains identically and the final accepted text reflects both edits.
    res_wet = engine.process_batch(batch(), dry_run=False)
    assert res_wet["edits_applied"] == 2
    assert res_wet["edits_skipped"] == 0

    engine.accept_all_revisions(remove_comments=True)
    from adeu.ingest import extract_text_from_stream

    final_text = extract_text_from_stream(engine.save_to_stream(), filename="x.docx", clean_view=True)
    assert "Disclosee" in final_text
    assert "Receiving Party" not in final_text


def test_tc2_dry_run_write_parity_active_insertion_edit():
    """
    TC2: dry-run outcome must equal write outcome.

    A single (strict) modification whose target lies entirely inside another
    author's pending insertion is now applied: track_delete_run splits the
    enclosing <w:ins> and nests the change. Dry-run and write must agree that
    the edit applies. (Partial overlaps and match_mode="all" fan-outs are still
    refused — see TestQaReportV2EngineSafety.)
    """
    doc = Document()
    p = doc.add_paragraph("The party shall provide ")

    # In python-docx, inject an active insertion from "Original Drafter"
    p_el = p._element
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "101")
    ins.set(qn("w:author"), "Original Drafter")
    ins.set(qn("w:date"), "2026-06-29T12:00:00Z")

    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "five (5)"
    r.append(t)
    ins.append(r)
    p_el.append(ins)

    r2 = OxmlElement("w:r")
    t2 = OxmlElement("w:t")
    t2.text = " years."
    r2.append(t2)
    p_el.append(r2)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="QA Tester")

    # Run A: dry_run = True — the edit applies. (Fresh edit object per call: a
    # dry-run mutates the input edit's internal resolution state.)
    res_dry = engine.process_batch([ModifyText(target_text="five (5)", new_text="seven (7)")], dry_run=True)
    assert res_dry["edits_applied"] == 1
    assert res_dry["edits_skipped"] == 0
    assert res_dry["edits"][0]["status"] == "applied"

    # Run B: dry_run = False — same outcome, no rejection.
    res_wet = engine.process_batch([ModifyText(target_text="five (5)", new_text="seven (7)")], dry_run=False)
    assert res_wet["edits_applied"] == 1
    assert res_wet["edits_skipped"] == 0


def test_tc3_heading_targeted_by_markdown_hash():
    """
    TC3: Heading targeted by markdown '#' corrupts instead of failing [report F4, F5]
    Intended behavior: targeting a heading by its rendered '#' form should NOT produce
    a redline containing literal '#' characters in the preview snippet.
    Asserting that the critic_markup does NOT contain "{--#" or "{++#" will fail precisely due to this bug.
    """
    doc = Document()
    doc.add_heading("3. Pending Review", level=1)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)

    res = engine.process_batch([ModifyText(target_text="# 3. Pending Review", new_text="# 3. Final Review")])

    assert res["edits_applied"] == 1

    # The buggy engine produces a preview that contains hashes inside deletion/insertion markup
    critic_markup = res["edits"][0]["critic_markup"]
    assert "{--#" not in critic_markup
    assert "{++#" not in critic_markup


def test_tc5_w16du_namespace_untouched_parts():
    """
    TC5: editing document.xml should not add namespaces/attributes to parts that weren't modified [report F9]
    Intended behavior: parts like headers should remain untouched without being stamped with w16du namespace.
    """
    doc = Document()
    doc.add_paragraph("This is untouched body text.")

    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = "Header Text"

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)
    engine.process_batch([ModifyText(target_text="untouched body", new_text="changed body")], dry_run=False)

    stream_edited = engine.save_to_stream()
    doc_edited = Document(stream_edited)

    saved_header_xml = doc_edited.sections[0].header._element.xml
    # The untouched header should not contain word16du namespace
    assert "word16du" not in saved_header_xml


def test_tc8_error_message_mislabeled_index():
    """
    TC8: Error message mislabels edit index [report F7]
    Intended behavior: in dry_run mode, when a subsequent edit fails, the error message
    should be correctly labeled with the 1-based index of that edit.
    """
    doc = Document()
    doc.add_paragraph("First paragraph text.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)

    res = engine.process_batch(
        [
            ModifyText(target_text="First paragraph", new_text="Updated first paragraph"),
            ModifyText(target_text="Non-existent text", new_text="Failed update"),
        ],
        dry_run=True,
    )

    # Since QA 2026-07-17 M1, dry-run mirrors the real run's transactional
    # semantics: a batch with any invalid edit applies nothing. The valid edit
    # is reported as blocked by the batch, the invalid one keeps its own error.
    assert res["edits_applied"] == 0
    assert res["edits_skipped"] == 2
    assert res["edits"][0]["status"] == "failed"
    assert "transactional" in res["edits"][0]["error"]
    assert res["edits"][1]["status"] == "failed"

    error_msg = res["edits"][1]["error"]
    # Under correct behavior, the inner error should label it as Edit 2.
    # The buggy unpatched codebase says "Edit 1 Failed:" inside the error message for Edit 2.
    assert "Edit 2 Failed:" in error_msg
    assert "Edit 1 Failed:" not in error_msg

# FILE: tests/test_repro_qa_report_v6.py
"""
Repro tests for the 2026-07-18 black-box QA and UX assessment
(adeu 1.23.0+8102b64).

Finding index (report section -> test class below):
  C1  sanitize reports "Result: CLEAN" while docProps/custom.xml (custom
      document properties: matter numbers, client secrets) survives in the
      output package, and dc:identifier / dc:language / cp:version are
      silently retained in docProps/core.xml
  C2  `adeu extract victim.docx -o victim.docx` replaces the source DOCX
      with extracted UTF-8 text and exits 0 — silent destruction of the
      only source document (same for path aliases and symlinks)
  H1  `adeu diff --json` output is not replayable by `adeu apply`: a table
      row modification followed by an insert_row anchored on the modified
      row either fails to apply or (worse) inserts the row at the wrong
      position — the clean-view fallback resolved an index in clean-view
      coordinates but the row lookup ran against the raw mapper
  H2  the extract → edit text → apply round trip drops the paragraph
      separator of a newly inserted paragraph, concatenating it with the
      following (or previous) paragraph; the defect survives accept-all
  M1  DOCX-writing commands accept misleading output extensions and
      silently replace existing outputs; writes are not atomic
  M3  batch sanitize writes some outputs and exits 1 — partial results
  M4  changes-file kind decided by filename suffix, not content
  M5  invalid --search-regex silently degrades to a literal search, exit 0
  L1  usage-shaped errors exit 1 instead of 2
  L2  apply --help does not document the JSON change schema
  L4  init prints "Config found" when creating a new config

Every C/H test (and the M1/M3/M4/M5 tests) fails against the commit
preceding its fix.
"""

import json
import os
import sys
import zipfile
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from pydantic import TypeAdapter

from adeu.diff import generate_edits_via_paragraph_alignment, generate_structured_edits, make_edits_self_contained
from adeu.ingest import _extract_text_from_doc, extract_text_from_stream
from adeu.models import BatchChanges
from adeu.redline.engine import RedlineEngine
from adeu.sanitize.core import sanitize_docx

# ---------------------------------------------------------------------------
# Helpers / fixture builders
# ---------------------------------------------------------------------------


def run_cli(args, capsys):
    """Invoke the CLI in-process; returns (exit_code, stdout, stderr)."""
    from unittest.mock import patch

    from adeu.cli import main

    code = 0
    with patch.object(sys, "argv", ["adeu"] + [str(a) for a in args]):
        try:
            main()
        except SystemExit as e:
            code = e.code or 0
    captured = capsys.readouterr()
    return code, captured.out, captured.err


CUSTOM_PROPS_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
    '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties" '
    'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">'
    '<property fmtid="{D5CDD505-2E9C-101B-9397-08002B2CF9AE}" pid="2" name="ClientSecret">'
    "<vt:lpwstr>TOP-SECRET-ORCHID</vt:lpwstr></property>"
    '<property fmtid="{D5CDD505-2E9C-101B-9397-08002B2CF9AE}" pid="3" name="MatterNumber">'
    "<vt:lpwstr>MAT-998877</vt:lpwstr></property>"
    "</Properties>"
)

CUSTOM_PROPS_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.custom-properties+xml"
CUSTOM_PROPS_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/custom-properties"


def inject_custom_properties(path: Path):
    """Adds a docProps/custom.xml part (with secrets) to an existing DOCX."""
    src = path.read_bytes()
    out = BytesIO()
    with ZipFile(BytesIO(src)) as zin, ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "[Content_Types].xml":
                text = data.decode("utf-8")
                override = f'<Override PartName="/docProps/custom.xml" ContentType="{CUSTOM_PROPS_CONTENT_TYPE}"/>'
                text = text.replace("</Types>", override + "</Types>")
                data = text.encode("utf-8")
            elif item.filename == "_rels/.rels":
                text = data.decode("utf-8")
                rel = f'<Relationship Id="rIdCustomProps" Type="{CUSTOM_PROPS_REL_TYPE}" Target="docProps/custom.xml"/>'
                text = text.replace("</Relationships>", rel + "</Relationships>")
                data = text.encode("utf-8")
            zout.writestr(item, data)
        zout.writestr("docProps/custom.xml", CUSTOM_PROPS_XML)
    path.write_bytes(out.getvalue())


def build_metadata_doc(path: Path):
    """The QA report's metadata-rich fixture: core + custom properties."""
    doc = Document()
    doc.add_paragraph("Body content that is perfectly fine to share.")
    core = doc.core_properties
    core.author = "Alice Example"
    core.last_modified_by = "Bob Reviewer"
    core.comments = "Internal note: project codename ORCHID"
    core.identifier = "CLIENT-12345"
    core.language = "fi-FI"
    core.version = "9.9-internal"
    doc.save(path)
    inject_custom_properties(path)


def zip_member_names(path: Path) -> list:
    with ZipFile(path) as z:
        return z.namelist()


def whole_package_text(path: Path) -> str:
    """Every member of the ZIP decoded permissively, concatenated."""
    chunks = []
    with ZipFile(path) as z:
        for name in z.namelist():
            chunks.append(z.read(name).decode("utf-8", errors="ignore"))
    return "\n".join(chunks)


def build_table_doc(path: Path, rows):
    doc = Document()
    doc.add_paragraph("Pricing schedule below.")
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for r, row in enumerate(rows):
        for c, cell in enumerate(row):
            table.rows[r].cells[c].text = cell
    doc.add_paragraph("Terms follow the table.")
    doc.save(path)


def build_paragraph_doc(path: Path):
    doc = Document()
    doc.add_paragraph("Intro paragraph one.")
    doc.add_paragraph("Middle paragraph two.")
    doc.add_paragraph("Final sentinel: END-OF-DOCUMENT-9f2c.")
    doc.save(path)


def clean_text_of(path: Path) -> str:
    return extract_text_from_stream(BytesIO(path.read_bytes()), filename=path.name, clean_view=True)


ORIG_ROWS = [["Seats", "5", "€100"], ["Support", "1", "€500"]]
MOD_ROWS = [["Seats", "10", "€125"], ["Support", "1", "Included"], ["Storage", "100 GB", "€50"]]


# ---------------------------------------------------------------------------
# C1 — sanitize leaves custom document properties + core identifier behind
# ---------------------------------------------------------------------------


class TestC1SanitizeCustomProperties:
    def test_full_sanitize_removes_custom_properties_part(self, tmp_path):
        src = tmp_path / "metadata.docx"
        build_metadata_doc(src)
        assert "docProps/custom.xml" in zip_member_names(src), "fixture must carry custom props"

        out = tmp_path / "metadata_sanitized.docx"
        result = sanitize_docx(str(src), str(out))

        assert "docProps/custom.xml" not in zip_member_names(out), (
            "sanitize reported success but docProps/custom.xml survived in the package"
        )
        package = whole_package_text(out)
        assert "TOP-SECRET-ORCHID" not in package
        assert "MAT-998877" not in package
        # The removal must be disclosed in the report.
        assert "custom" in result.report_text.lower()

    def test_full_sanitize_scrubs_core_identifier_language_version(self, tmp_path):
        src = tmp_path / "metadata.docx"
        build_metadata_doc(src)

        out = tmp_path / "metadata_sanitized.docx"
        result = sanitize_docx(str(src), str(out))

        core_xml = ZipFile(out).read("docProps/core.xml").decode("utf-8")
        assert "CLIENT-12345" not in core_xml, "dc:identifier silently retained"
        assert "9.9-internal" not in core_xml, "cp:version silently retained"
        assert "fi-FI" not in core_xml, "dc:language silently retained"
        # Retained-or-scrubbed fields must be enumerated, not silent.
        assert "CLIENT-12345" in result.report_text

    def test_keep_markup_sanitize_also_removes_custom_properties(self, tmp_path):
        src = tmp_path / "metadata.docx"
        build_metadata_doc(src)

        out = tmp_path / "metadata_km.docx"
        sanitize_docx(str(src), str(out), keep_markup=True)

        assert "docProps/custom.xml" not in zip_member_names(out)
        assert "TOP-SECRET-ORCHID" not in whole_package_text(out)

    def test_sanitized_output_still_loads(self, tmp_path):
        src = tmp_path / "metadata.docx"
        build_metadata_doc(src)
        out = tmp_path / "metadata_sanitized.docx"
        sanitize_docx(str(src), str(out))

        reopened = Document(str(out))
        assert "perfectly fine to share" in "\n".join(p.text for p in reopened.paragraphs)

    def test_cli_sanitize_report_never_claims_clean_with_retained_secrets(self, tmp_path, capsys):
        """The QA report's exact command shape: sanitize --report."""
        src = tmp_path / "metadata.docx"
        build_metadata_doc(src)
        out = tmp_path / "metadata_sanitized.docx"

        code, _, err = run_cli(["sanitize", src, "-o", out, "--report"], capsys)
        assert code == 0, err
        if "Result: CLEAN" in err:
            assert "TOP-SECRET-ORCHID" not in whole_package_text(out)
            assert "docProps/custom.xml" not in zip_member_names(out)


# ---------------------------------------------------------------------------
# C2 — extract destroys its own source when output == input
# ---------------------------------------------------------------------------


class TestC2ExtractSameFileGuard:
    def _victim(self, tmp_path) -> Path:
        victim = tmp_path / "victim.docx"
        build_paragraph_doc(victim)
        return victim

    def test_extract_refuses_output_equal_to_input(self, tmp_path, capsys):
        victim = self._victim(tmp_path)
        before = victim.read_bytes()

        code, _, err = run_cli(["extract", victim, "-o", victim], capsys)

        assert code != 0, "extract must refuse to overwrite its own input"
        assert victim.read_bytes() == before, "source DOCX was destroyed"
        assert "input" in err.lower() or "overwrite" in err.lower()

    def test_extract_refuses_relative_alias_of_input(self, tmp_path, capsys):
        victim = self._victim(tmp_path)
        before = victim.read_bytes()
        (tmp_path / "sub").mkdir()
        alias = tmp_path / "sub" / ".." / "victim.docx"

        code, _, _ = run_cli(["extract", victim, "-o", alias], capsys)

        assert code != 0
        assert victim.read_bytes() == before

    @pytest.mark.skipif(sys.platform == "win32", reason="symlinks not reliable on Windows CI")
    def test_extract_refuses_symlink_alias_of_input(self, tmp_path, capsys):
        victim = self._victim(tmp_path)
        before = victim.read_bytes()
        link = tmp_path / "link.docx"
        link.symlink_to(victim)

        code, _, _ = run_cli(["extract", victim, "-o", link], capsys)

        assert code != 0
        assert victim.read_bytes() == before

    def test_extract_refuses_docx_extension_for_text_output(self, tmp_path, capsys):
        """Extracted text may never masquerade as a DOCX (QA C2/M1)."""
        victim = self._victim(tmp_path)
        other = tmp_path / "other.docx"

        code, _, err = run_cli(["extract", victim, "-o", other], capsys)

        assert code != 0
        assert not other.exists()
        assert ".docx" in err or "docx" in err.lower()

    def test_extract_json_refuses_output_equal_to_input(self, tmp_path, capsys):
        victim = self._victim(tmp_path)
        before = victim.read_bytes()

        code, _, _ = run_cli(["extract", victim, "--json", "-o", victim], capsys)

        assert code != 0
        assert victim.read_bytes() == before

    def test_extract_to_fresh_text_path_still_works(self, tmp_path, capsys):
        victim = self._victim(tmp_path)
        out = tmp_path / "extracted.md"

        code, _, err = run_cli(["extract", victim, "--page", "all", "-o", out], capsys)

        assert code == 0, err
        assert "Final sentinel" in out.read_text(encoding="utf-8")

    def test_markup_refuses_output_equal_to_docx_input(self, tmp_path, capsys):
        victim = self._victim(tmp_path)
        before = victim.read_bytes()
        edits = tmp_path / "edits.json"
        edits.write_text(
            json.dumps([{"type": "modify", "target_text": "Middle paragraph two.", "new_text": "Changed."}])
        )

        code, _, _ = run_cli(["markup", victim, edits, "-o", victim], capsys)

        assert code != 0, "markup must refuse to overwrite its own input"
        assert victim.read_bytes() == before

    def test_sanitize_report_file_refuses_input_path(self, tmp_path, capsys):
        victim = self._victim(tmp_path)
        before = victim.read_bytes()
        out = tmp_path / "sanitized.docx"

        code, _, _ = run_cli(["sanitize", victim, "-o", out, "--report-file", victim], capsys)

        assert code != 0, "a text report must never overwrite the input DOCX"
        assert victim.read_bytes() == before

    def test_sanitize_report_file_refuses_output_path(self, tmp_path, capsys):
        victim = self._victim(tmp_path)
        out = tmp_path / "sanitized.docx"

        code, _, _ = run_cli(["sanitize", victim, "-o", out, "--report-file", out], capsys)

        assert code != 0, "the text report must never overwrite the sanitized DOCX output"

    def test_apply_in_place_docx_output_still_allowed(self, tmp_path, capsys):
        """DOCX-in DOCX-out commands legitimately support in-place operation."""
        victim = self._victim(tmp_path)
        edits = tmp_path / "edits.json"
        edits.write_text(
            json.dumps([{"type": "modify", "target_text": "Middle paragraph two.", "new_text": "Edited middle."}])
        )

        code, _, err = run_cli(["apply", victim, edits, "-o", victim], capsys)

        assert code == 0, err
        assert "Edited middle." in extract_text_from_stream(BytesIO(victim.read_bytes()), clean_view=True)


# ---------------------------------------------------------------------------
# H1 — diff --json row operations must be replayable by apply
# ---------------------------------------------------------------------------


class TestH1DiffRowOpsReplay:
    def _structured_edits(self, orig: Path, mod: Path):
        doc_o = Document(str(orig))
        doc_m = Document(str(mod))
        text_o, struct_o = _extract_text_from_doc(doc_o, clean_view=True, include_appendix=False, return_structure=True)
        text_m, struct_m = _extract_text_from_doc(doc_m, clean_view=True, include_appendix=False, return_structure=True)
        return generate_structured_edits(text_o, struct_o, text_m, struct_m)

    def test_row_modify_plus_insert_replays_from_json(self, tmp_path):
        """The QA report's exact failing composition, replayed WITHOUT pins."""
        orig = tmp_path / "original.docx"
        mod = tmp_path / "modified.docx"
        build_table_doc(orig, ORIG_ROWS)
        build_table_doc(mod, MOD_ROWS)

        edits, warnings = self._structured_edits(orig, mod)
        assert not warnings, f"unexpected diff warnings: {warnings}"
        # JSON round trip drops the private positional pins — exactly what
        # `adeu diff --json > changes.json && adeu apply … changes.json` does.
        dumped = json.loads(json.dumps([e.model_dump() for e in edits]))
        changes = TypeAdapter(BatchChanges).validate_python(dumped)

        engine = RedlineEngine(BytesIO(orig.read_bytes()), author="QA")
        stats = engine.process_batch(list(changes))
        assert stats["edits_skipped"] == 0, f"diff output not replayable: {stats['skipped_details']}"

        engine.accept_all_revisions(remove_comments=True)
        final = extract_text_from_stream(engine.save_to_stream(), clean_view=True)
        assert final == clean_text_of(mod), (
            "replayed diff did not reproduce the modified document (row order or content wrong)"
        )

    def test_inserted_row_lands_below_modified_anchor(self, tmp_path):
        """Regression for the silent wrong-position variant of H1."""
        orig = tmp_path / "original.docx"
        mod = tmp_path / "modified.docx"
        build_table_doc(orig, ORIG_ROWS)
        build_table_doc(mod, MOD_ROWS)

        edits, _ = self._structured_edits(orig, mod)
        dumped = json.loads(json.dumps([e.model_dump() for e in edits]))
        changes = TypeAdapter(BatchChanges).validate_python(dumped)

        engine = RedlineEngine(BytesIO(orig.read_bytes()), author="QA")
        engine.process_batch(list(changes))
        engine.accept_all_revisions(remove_comments=True)
        final = extract_text_from_stream(engine.save_to_stream(), clean_view=True)

        support = final.index("Support | 1 | Included")
        storage = final.index("Storage | 100 GB | €50")
        assert support < storage, f"inserted row landed at the wrong position:\n{final}"

    def test_cli_diff_apply_accept_roundtrip(self, tmp_path, capsys):
        """End-to-end over the public CLI: diff --json → apply → accept-all."""
        orig = tmp_path / "original.docx"
        mod = tmp_path / "modified.docx"
        build_table_doc(orig, ORIG_ROWS)
        build_table_doc(mod, MOD_ROWS)

        code, out, err = run_cli(["diff", orig, mod, "--json"], capsys)
        assert code == 0, err
        changes_file = tmp_path / "changes.json"
        changes_file.write_text(out, encoding="utf-8")

        redlined = tmp_path / "redlined.docx"
        code, _, err = run_cli(["apply", orig, changes_file, "-o", redlined], capsys)
        assert code == 0, f"apply could not consume diff's own JSON output:\n{err}"

        accepted = tmp_path / "accepted.docx"
        code, _, err = run_cli(["accept-all", redlined, "-o", accepted], capsys)
        assert code == 0, err

        assert clean_text_of(accepted) == clean_text_of(mod)

    def test_delete_and_modify_composition_replays_from_json(self, tmp_path):
        """Row deletion combined with a modification of the surviving row."""
        orig = tmp_path / "original.docx"
        mod = tmp_path / "modified.docx"
        build_table_doc(orig, [["Alpha", "1", "a"], ["Beta", "2", "b"], ["Gamma", "3", "c"]])
        build_table_doc(mod, [["Alpha", "9", "z"], ["Gamma", "3", "c"]])

        edits, _ = self._structured_edits(orig, mod)
        dumped = json.loads(json.dumps([e.model_dump() for e in edits]))
        changes = TypeAdapter(BatchChanges).validate_python(dumped)

        engine = RedlineEngine(BytesIO(orig.read_bytes()), author="QA")
        stats = engine.process_batch(list(changes))
        assert stats["edits_skipped"] == 0, f"not replayable: {stats['skipped_details']}"
        engine.accept_all_revisions(remove_comments=True)
        final = extract_text_from_stream(engine.save_to_stream(), clean_view=True)
        assert final == clean_text_of(mod)


# ---------------------------------------------------------------------------
# H2 — text round trip must preserve paragraph boundaries
# ---------------------------------------------------------------------------


class TestH2ParagraphBoundaryRoundTrip:
    def _roundtrip_in_process(self, orig: Path, text_mod: str) -> str:
        text_orig = clean_text_of(orig)
        edits = generate_edits_via_paragraph_alignment(text_orig, text_mod)
        engine = RedlineEngine(BytesIO(orig.read_bytes()), author="QA")
        stats = engine.process_batch(list(edits))
        assert stats["edits_skipped"] == 0, f"skipped: {stats['skipped_details']}"
        engine.accept_all_revisions(remove_comments=True)
        return extract_text_from_stream(engine.save_to_stream(), clean_view=True)

    def test_insert_paragraph_before_final_marker(self, tmp_path):
        """The QA report's exact repro: new paragraph before the sentinel."""
        orig = tmp_path / "original.docx"
        build_paragraph_doc(orig)
        text_orig = clean_text_of(orig)
        text_mod = text_orig.replace(
            "Final sentinel", "Additional sentence inserted before the final marker.\n\nFinal sentinel", 1
        )

        final = self._roundtrip_in_process(orig, text_mod)

        assert "marker.Final" not in final, "paragraph separator lost: sentences concatenated"
        assert final == text_mod

    def test_append_paragraph_at_end(self, tmp_path):
        orig = tmp_path / "original.docx"
        build_paragraph_doc(orig)
        text_mod = clean_text_of(orig) + "\n\nAppended final paragraph."

        final = self._roundtrip_in_process(orig, text_mod)

        assert "9f2c.Appended" not in final, "appended paragraph glued to the last paragraph"
        assert final == text_mod

    def test_insert_paragraph_at_start(self, tmp_path):
        orig = tmp_path / "original.docx"
        build_paragraph_doc(orig)
        text_mod = "Brand new first paragraph.\n\n" + clean_text_of(orig)

        final = self._roundtrip_in_process(orig, text_mod)

        assert "paragraph.Intro" not in final, "inserted paragraph glued to the first paragraph"
        assert final == text_mod

    def test_insert_multiple_paragraphs_mid_document(self, tmp_path):
        orig = tmp_path / "original.docx"
        build_paragraph_doc(orig)
        text_orig = clean_text_of(orig)
        text_mod = text_orig.replace(
            "Final sentinel", "First inserted paragraph.\n\nSecond inserted paragraph.\n\nFinal sentinel", 1
        )

        final = self._roundtrip_in_process(orig, text_mod)

        assert final == text_mod

    def test_delete_paragraph_still_roundtrips(self, tmp_path):
        """Deletions already worked before the fix — keep them working."""
        orig = tmp_path / "original.docx"
        build_paragraph_doc(orig)
        text_orig = clean_text_of(orig)
        text_mod = text_orig.replace("\n\nMiddle paragraph two.", "", 1)

        final = self._roundtrip_in_process(orig, text_mod)

        assert final == text_mod

    def test_cli_text_apply_roundtrip(self, tmp_path, capsys):
        """End-to-end over the public CLI: extract → edit → apply → accept-all."""
        orig = tmp_path / "original.docx"
        build_paragraph_doc(orig)

        edited = tmp_path / "edited.md"
        code, _, err = run_cli(["extract", orig, "--page", "all", "--clean-view", "-o", edited], capsys)
        assert code == 0, err

        text = edited.read_text(encoding="utf-8")
        text_mod = text.replace(
            "Final sentinel", "Additional sentence inserted before the final marker.\n\nFinal sentinel", 1
        )
        edited.write_text(text_mod, encoding="utf-8")
        # apply strips extract's "> **File Path:** …" header on ingestion;
        # compare against the same document-content view.
        import re as _re

        text_mod = _re.sub(r"^> \*\*File Path:\*\*[^\n]*\n+", "", text_mod)

        redlined = tmp_path / "redlined.docx"
        code, _, err = run_cli(["apply", orig, edited, "-o", redlined], capsys)
        assert code == 0, err

        accepted = tmp_path / "accepted.docx"
        code, _, err = run_cli(["accept-all", redlined, "-o", accepted], capsys)
        assert code == 0, err

        final = clean_text_of(accepted)
        assert "marker.Final" not in final, "paragraph separator lost through the CLI round trip"
        assert final == text_mod

    def test_diff_json_insertion_edits_are_replayable(self, tmp_path):
        """The diff --json text path must carry the separator too."""
        orig = tmp_path / "original.docx"
        build_paragraph_doc(orig)
        text_orig = clean_text_of(orig)
        text_mod = text_orig.replace(
            "Final sentinel", "Additional sentence inserted before the final marker.\n\nFinal sentinel", 1
        )

        edits = generate_edits_via_paragraph_alignment(text_orig, text_mod)
        edits = make_edits_self_contained(edits, text_orig)
        dumped = json.loads(json.dumps([e.model_dump() for e in edits]))
        changes = TypeAdapter(BatchChanges).validate_python(dumped)

        engine = RedlineEngine(BytesIO(orig.read_bytes()), author="QA")
        stats = engine.process_batch(list(changes))
        assert stats["edits_skipped"] == 0, f"skipped: {stats['skipped_details']}"
        engine.accept_all_revisions(remove_comments=True)
        final = extract_text_from_stream(engine.save_to_stream(), clean_view=True)

        assert final == text_mod


# ---------------------------------------------------------------------------
# M1 — output extension validation, overwrite visibility, atomic writes
# ---------------------------------------------------------------------------


class TestM1OutputSafety:
    def test_apply_refuses_txt_output_extension(self, tmp_path, capsys):
        orig = tmp_path / "original.docx"
        build_paragraph_doc(orig)
        edits = tmp_path / "edits.json"
        edits.write_text(json.dumps([{"type": "modify", "target_text": "Middle paragraph two.", "new_text": "X."}]))

        out = tmp_path / "result.txt"
        code, _, err = run_cli(["apply", orig, edits, "-o", out], capsys)

        assert code != 0, "apply must refuse to write a DOCX package to a .txt path"
        assert not out.exists()
        assert ".docx" in err

    def test_accept_all_refuses_txt_output_extension(self, tmp_path, capsys):
        orig = tmp_path / "original.docx"
        build_paragraph_doc(orig)

        out = tmp_path / "accepted.txt"
        code, _, err = run_cli(["accept-all", orig, "-o", out], capsys)

        assert code != 0
        assert not out.exists()
        assert ".docx" in err

    def test_sanitize_refuses_txt_output_extension(self, tmp_path, capsys):
        orig = tmp_path / "original.docx"
        build_paragraph_doc(orig)

        out = tmp_path / "sanitized.txt"
        code, _, err = run_cli(["sanitize", orig, "-o", out], capsys)

        assert code != 0
        assert not out.exists()
        assert ".docx" in err

    def test_overwriting_existing_output_is_announced(self, tmp_path, capsys):
        orig = tmp_path / "original.docx"
        build_paragraph_doc(orig)
        edits = tmp_path / "edits.json"
        edits.write_text(json.dumps([{"type": "modify", "target_text": "Middle paragraph two.", "new_text": "X."}]))
        out = tmp_path / "result.docx"
        out.write_bytes(b"sentinel-not-a-docx")

        code, _, err = run_cli(["apply", orig, edits, "-o", out], capsys)

        assert code == 0, err
        assert "Overwriting existing" in err

    def test_failed_write_leaves_existing_output_untouched(self, tmp_path, capsys, monkeypatch):
        """Atomicity: a write failure must never truncate the previous output."""
        orig = tmp_path / "original.docx"
        build_paragraph_doc(orig)
        edits = tmp_path / "edits.json"
        edits.write_text(json.dumps([{"type": "modify", "target_text": "Middle paragraph two.", "new_text": "X."}]))
        out = tmp_path / "result.docx"
        out.write_bytes(b"previous-good-output")

        import adeu.cli as cli_mod

        real_replace = os.replace

        def failing_replace(src, dst, *a, **k):
            if str(dst) == str(out):
                raise OSError(28, "No space left on device")
            return real_replace(src, dst, *a, **k)

        monkeypatch.setattr(cli_mod.os, "replace", failing_replace)
        code, _, err = run_cli(["apply", orig, edits, "-o", out], capsys)

        assert code != 0
        assert out.read_bytes() == b"previous-good-output", "existing output was corrupted by a failed write"


# ---------------------------------------------------------------------------
# M3 — batch sanitize is all-or-nothing
# ---------------------------------------------------------------------------


class TestM3TransactionalBatchSanitize:
    def test_blocked_input_means_no_outputs_at_all(self, tmp_path, capsys):
        clean = tmp_path / "cleanable.docx"
        build_paragraph_doc(clean)

        blocked = tmp_path / "blocked.docx"
        doc = Document()
        p = doc.add_paragraph("Base text ")
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn as _qn

        ins = OxmlElement("w:ins")
        ins.set(_qn("w:id"), "1")
        ins.set(_qn("w:author"), "Reviewer")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "pending insertion"
        r.append(t)
        ins.append(r)
        p._element.append(ins)
        doc.save(blocked)

        outdir = tmp_path / "out"
        code, _, err = run_cli(["sanitize", clean, blocked, "--outdir", outdir], capsys)

        assert code != 0
        assert not (outdir / "cleanable.docx").exists(), "batch sanitize wrote a partial result despite a blocked input"
        assert not (outdir / "blocked.docx").exists()
        assert list(outdir.glob("*.tmp")) == [], "staging temp files were left behind"

    def test_all_good_batch_writes_all_outputs(self, tmp_path, capsys):
        a = tmp_path / "a.docx"
        b = tmp_path / "b.docx"
        build_paragraph_doc(a)
        build_paragraph_doc(b)

        outdir = tmp_path / "out"
        code, _, err = run_cli(["sanitize", a, b, "--outdir", outdir], capsys)

        assert code == 0, err
        assert (outdir / "a.docx").exists()
        assert (outdir / "b.docx").exists()
        assert list(outdir.glob("*.tmp")) == []


# ---------------------------------------------------------------------------
# M4 — changes-file kind is decided by content, not filename
# ---------------------------------------------------------------------------


class TestM4ContentSniffedChangesFile:
    def test_json_batch_with_txt_suffix_is_applied_as_batch(self, tmp_path, capsys):
        orig = tmp_path / "original.docx"
        build_paragraph_doc(orig)
        changes = tmp_path / "changes.txt"
        changes.write_text(
            json.dumps([{"type": "modify", "target_text": "Middle paragraph two.", "new_text": "Edited middle."}])
        )

        out = tmp_path / "result.docx"
        code, _, err = run_cli(["apply", orig, changes, "-o", out], capsys)

        assert code == 0, f"a JSON batch must be recognized by content, not suffix:\n{err}"
        final = clean_text_of(out)
        assert "Edited middle." in final
        assert "Intro paragraph one." in final, "the batch was diffed as document text, deleting everything else"

    def test_extensionless_json_batch_is_applied_as_batch(self, tmp_path, capsys):
        orig = tmp_path / "original.docx"
        build_paragraph_doc(orig)
        changes = tmp_path / "changes"
        changes.write_text(
            json.dumps([{"type": "modify", "target_text": "Middle paragraph two.", "new_text": "Edited middle."}])
        )

        out = tmp_path / "result.docx"
        code, _, err = run_cli(["apply", orig, changes, "-o", out], capsys)

        assert code == 0, err
        assert "Edited middle." in clean_text_of(out)

    def test_invalid_json_with_json_suffix_is_an_error_not_a_text_diff(self, tmp_path, capsys):
        orig = tmp_path / "original.docx"
        build_paragraph_doc(orig)
        changes = tmp_path / "changes.json"
        changes.write_text('[{"type": "modify", "target_text": "broken...')

        out = tmp_path / "result.docx"
        code, _, err = run_cli(["apply", orig, changes, "-o", out], capsys)

        assert code != 0, "a broken .json file must be a parse error, never a text diff"
        assert not out.exists()

    def test_text_starting_with_bracket_is_still_text(self, tmp_path, capsys):
        """A document whose text begins with '[' must stay on the text path."""
        orig = tmp_path / "original.docx"
        build_paragraph_doc(orig)
        text_orig = clean_text_of(orig)
        changes = tmp_path / "edited.md"
        changes.write_text("[DRAFT] Cover note.\n\n" + text_orig, encoding="utf-8")

        out = tmp_path / "result.docx"
        code, _, err = run_cli(["apply", orig, changes, "-o", out], capsys)

        assert code == 0, err
        assert "[DRAFT] Cover note." in clean_text_of(out)


# ---------------------------------------------------------------------------
# M5 / L1 / L2 / L4 — strict regex, exit codes, discoverable schema, wording
# ---------------------------------------------------------------------------


class TestM5StrictCliRegex:
    def test_invalid_search_regex_is_a_hard_error(self, tmp_path, capsys):
        orig = tmp_path / "original.docx"
        build_paragraph_doc(orig)

        code, _, err = run_cli(["extract", orig, "--search-query", "[", "--search-regex"], capsys)

        assert code != 0, "an invalid regex must not silently degrade to a literal search on the CLI"
        assert "regular expression" in err

    def test_valid_search_regex_still_works(self, tmp_path, capsys):
        orig = tmp_path / "original.docx"
        build_paragraph_doc(orig)

        code, out, err = run_cli(["extract", orig, "--search-query", "sentinel: END.*9f2c", "--search-regex"], capsys)

        assert code == 0, err
        assert "END-OF-DOCUMENT-9f2c" in out
        assert "appears 1 time" in out


class TestLowSeverityContracts:
    def test_missing_argument_paths_exit_2(self, tmp_path, capsys):
        code, _, _ = run_cli(["extract"], capsys)
        assert code == 2

        orig = tmp_path / "original.docx"
        build_paragraph_doc(orig)
        code, _, _ = run_cli(["apply", orig], capsys)
        assert code == 2

    def test_invalid_page_value_exits_2(self, tmp_path, capsys):
        orig = tmp_path / "original.docx"
        build_paragraph_doc(orig)
        code, _, _ = run_cli(["extract", orig, "--page", "abc"], capsys)
        assert code == 2
        code, _, _ = run_cli(["extract", orig, "--page", "-1"], capsys)
        assert code == 2

    def test_apply_help_documents_the_change_schema(self, capsys):
        code, out, err = run_cli(["apply", "--help"], capsys)
        assert code == 0
        help_text = out + err
        for token in ("insert_row", "delete_row", "match_mode", "target_text", "reply"):
            assert token in help_text, f"apply --help must document '{token}'"

    def test_init_wording_distinguishes_new_config(self, tmp_path, capsys, monkeypatch):
        import adeu.cli as cli_mod

        config = tmp_path / "claude" / "claude_desktop_config.json"
        monkeypatch.setattr(cli_mod, "_get_claude_config_path", lambda: config)
        monkeypatch.setattr(cli_mod.shutil, "which", lambda name: "/usr/bin/uvx")

        code, _, err = run_cli(["init"], capsys)
        assert code == 0
        assert "Config will be created" in err
        assert "Config found" not in err

        code, _, err = run_cli(["init"], capsys)
        assert code == 0
        assert "Config found" in err


# ---------------------------------------------------------------------------
# L3 — markup renders structural row operations in the preview
# ---------------------------------------------------------------------------


class TestL3MarkupStructuralPreview:
    def _diff_changes(self, tmp_path, capsys):
        orig = tmp_path / "original.docx"
        mod = tmp_path / "modified.docx"
        build_table_doc(orig, ORIG_ROWS)
        build_table_doc(mod, MOD_ROWS)
        code, out, err = run_cli(["diff", orig, mod, "--json"], capsys)
        assert code == 0, err
        changes = tmp_path / "changes.json"
        changes.write_text(out, encoding="utf-8")
        return orig, changes

    def test_insert_row_renders_in_preview(self, tmp_path, capsys):
        orig, changes = self._diff_changes(tmp_path, capsys)
        out = tmp_path / "preview.md"

        code, _, err = run_cli(["markup", orig, changes, "-o", out], capsys)

        assert code == 0, err
        preview = out.read_text(encoding="utf-8")
        assert "{++Storage | 100 GB | €50++}" in preview, "inserted row missing from the rendered preview"
        support = preview.index("Support | 1 |")
        storage = preview.index("{++Storage | 100 GB | €50++}")
        assert support < storage, "inserted row rendered at the wrong position"
        assert "ignored" not in err, "row operations must render, not be ignored"

    def test_delete_row_renders_in_preview(self, tmp_path, capsys):
        orig = tmp_path / "original.docx"
        build_table_doc(orig, [["Alpha", "1", "a"], ["Beta", "2", "b"]])
        changes = tmp_path / "changes.json"
        changes.write_text(json.dumps([{"type": "delete_row", "target_text": "Beta | 2 | b"}]))
        out = tmp_path / "preview.md"

        code, _, err = run_cli(["markup", orig, changes, "-o", out], capsys)

        assert code == 0, err
        assert "{--Beta | 2 | b--}" in out.read_text(encoding="utf-8")

    def test_row_substring_anchor_renders_in_preview(self, tmp_path, capsys):
        orig = tmp_path / "original.docx"
        build_table_doc(orig, [["ID", "Product", "Price"], ["1", "Apples", "$1.50"], ["2", "Bananas", "$0.80"]])
        changes = tmp_path / "changes.json"
        changes.write_text(
            json.dumps(
                [
                    {
                        "type": "insert_row",
                        "target_text": "Bananas",
                        "cells": ["3", "Cherries", "$2.50"],
                        "position": "below",
                    },
                    {"type": "delete_row", "target_text": "Apples"},
                ]
            )
        )
        out = tmp_path / "preview.md"

        code, _, err = run_cli(["markup", orig, changes, "-o", out], capsys)

        assert code == 0, err
        preview_content = out.read_text(encoding="utf-8")
        assert "{--1 | Apples | $1.50--}" in preview_content
        assert "{++3 | Cherries | $2.50++}" in preview_content

    def test_missing_row_anchor_fails_without_output(self, tmp_path, capsys):
        orig = tmp_path / "original.docx"
        build_table_doc(orig, [["Alpha", "1", "a"]])
        changes = tmp_path / "changes.json"
        changes.write_text(
            json.dumps([{"type": "insert_row", "target_text": "No | Such | Row", "cells": ["x", "y", "z"]}])
        )
        out = tmp_path / "preview.md"

        code, _, _ = run_cli(["markup", orig, changes, "-o", out], capsys)

        assert code != 0
        assert not out.exists()

    def test_markup_json_result_mode(self, tmp_path, capsys):
        orig, changes = self._diff_changes(tmp_path, capsys)
        out = tmp_path / "preview.md"

        code, stdout, err = run_cli(["markup", orig, changes, "-o", out, "--json"], capsys)

        assert code == 0, err
        payload = json.loads(stdout)
        assert payload["status"] == "ok"
        assert payload["failed"] == 0
        assert payload["applied"] >= 3
        assert payload["output_path"] == str(out)


# ---------------------------------------------------------------------------
# M6 — extract must not pay the MCP framework's import cost
# ---------------------------------------------------------------------------


class TestM6ExtractStartup:
    def test_extract_never_imports_fastmcp(self, tmp_path):
        """fastmcp's import chain costs ~0.7 s — 2× the rest of the command.

        Guarded structurally (no wall-clock flakiness): a successful extract
        subprocess must finish without fastmcp in sys.modules.
        """
        import subprocess
        import sys as _sys

        orig = tmp_path / "original.docx"
        build_paragraph_doc(orig)

        probe = (
            "import sys\n"
            f"sys.argv = ['adeu', 'extract', {str(orig)!r}, '--page', 'all']\n"
            "from adeu.cli import main\n"
            "try:\n"
            "    main()\n"
            "except SystemExit as e:\n"
            "    assert (e.code or 0) == 0, e.code\n"
            "assert 'fastmcp' not in sys.modules, 'extract imported the MCP framework'\n"
        )
        result = subprocess.run([_sys.executable, "-c", probe], capture_output=True, text=True)
        assert result.returncode == 0, result.stderr


# ---------------------------------------------------------------------------
# Hypothesis-found regressions (2026-07-18 fuzz hunt) — deterministic pins
# ---------------------------------------------------------------------------


class TestFuzzFoundRegressions:
    """Each case is a minimized hypothesis counterexample; the property suite
    (test_property_invariants.py) hunts the families, these pin the fixes."""

    def _text_roundtrip(self, paras, mod, via_json=False):
        from adeu.diff import make_edits_self_contained

        orig = BytesIO()
        doc = Document()
        for p in paras:
            doc.add_paragraph(p)
        doc.save(orig)
        text_orig = extract_text_from_stream(BytesIO(orig.getvalue()), clean_view=True)
        text_mod = "\n\n".join(mod)
        edits = generate_edits_via_paragraph_alignment(text_orig, text_mod)
        if via_json:
            edits = make_edits_self_contained(edits, text_orig)
            dumped = json.loads(json.dumps([e.model_dump() for e in edits]))
            edits = TypeAdapter(BatchChanges).validate_python(dumped)
        engine = RedlineEngine(BytesIO(orig.getvalue()), author="Fuzz")
        stats = engine.process_batch(list(edits))
        assert stats["edits_skipped"] == 0, stats["skipped_details"]
        engine.accept_all_revisions(remove_comments=True)
        final = extract_text_from_stream(engine.save_to_stream(), clean_view=True)
        assert final == text_mod

    def test_paragraph_split_keeps_suffix(self):
        """track_insert suffix relocation: 'alpha beta.' split into two."""
        self._text_roundtrip(["alpha beta."], ["alpha.", "beta."])

    def test_paragraph_split_with_changed_halves(self):
        self._text_roundtrip(["alpha beta gamma."], ["alpha beta.", "gamma delta."])

    def test_three_way_paragraph_split(self):
        self._text_roundtrip(["p q."], ["p.", "q.", "r."])

    def test_repetitive_prefix_insertion_via_json(self):
        """Positional vs style anchor: insertion whose trim ends with a space."""
        self._text_roundtrip(["0 00."], ["0 0.", "0 00."], via_json=True)

    def test_inserted_plain_paragraph_with_number_lead_stays_plain(self):
        """'2024. Year…' must not become a numbered-list item."""
        self._text_roundtrip(["Intro."], ["Intro.", "2024. Year in review."], via_json=True)

    def test_projected_list_marker_still_creates_list(self):
        """The projection's own '1. ' shape keeps converting to a list."""
        from adeu.models import ModifyText

        orig = BytesIO()
        doc = Document()
        doc.add_paragraph("Intro.")
        doc.save(orig)
        engine = RedlineEngine(BytesIO(orig.getvalue()), author="Fuzz")
        edit = ModifyText(type="modify", target_text="Intro.", new_text="Intro.\n\n1. first item")
        stats = engine.process_batch([edit])
        assert stats["edits_skipped"] == 0
        raw = engine.save_to_stream().getvalue()
        assert b"ListNumber" in ZipFile(BytesIO(raw)).read("word/document.xml").replace(b" ", b"")

    def _table_roundtrip(self, rows, mod_rows, drop_pins):
        orig = BytesIO()
        doc = Document()
        doc.add_paragraph("Lead paragraph.")
        table = doc.add_table(rows=len(rows), cols=3)
        for r, row in enumerate(rows):
            for c, cell in enumerate(row):
                table.rows[r].cells[c].text = cell
        doc.add_paragraph("Trailing paragraph.")
        doc.save(orig)

        mod = BytesIO()
        doc2 = Document()
        doc2.add_paragraph("Lead paragraph.")
        table2 = doc2.add_table(rows=len(mod_rows), cols=3)
        for r, row in enumerate(mod_rows):
            for c, cell in enumerate(row):
                table2.rows[r].cells[c].text = cell
        doc2.add_paragraph("Trailing paragraph.")
        doc2.save(mod)

        doc_o = Document(BytesIO(orig.getvalue()))
        doc_m = Document(BytesIO(mod.getvalue()))
        t_o, s_o = _extract_text_from_doc(doc_o, clean_view=True, include_appendix=False, return_structure=True)
        t_m, s_m = _extract_text_from_doc(doc_m, clean_view=True, include_appendix=False, return_structure=True)
        edits, warnings = generate_structured_edits(t_o, s_o, t_m, s_m)
        assert not warnings, warnings
        if drop_pins:
            dumped = json.loads(json.dumps([e.model_dump() for e in edits]))
            edits = TypeAdapter(BatchChanges).validate_python(dumped)
        engine = RedlineEngine(BytesIO(orig.getvalue()), author="Fuzz")
        stats = engine.process_batch(list(edits))
        assert stats["edits_skipped"] == 0, stats["skipped_details"]
        engine.accept_all_revisions(remove_comments=True)
        final = extract_text_from_stream(engine.save_to_stream(), clean_view=True)
        want = extract_text_from_stream(BytesIO(mod.getvalue()), clean_view=True)
        assert final == want

    def test_cell_edit_lands_in_the_right_cell_json(self):
        """Middle-cell change on repetitive content: [0,0,0] -> [0,'0 0',0]."""
        self._table_roundtrip([["0", "0", "0"]], [["0", "0 0", "0"]], drop_pins=True)

    def test_cell_edit_lands_in_the_right_cell_in_process(self):
        """The in-process (sanitize --baseline) shape of the same edit."""
        self._table_roundtrip([["0", "0", "0"]], [["0", "0 0", "0"]], drop_pins=False)

    def test_multi_cell_row_modification_json(self):
        self._table_roundtrip([["0 0", "0 0", "00"]], [["0 0", "0", "0"]], drop_pins=True)

    def test_cell_shrink_brushing_separator_json(self):
        self._table_roundtrip([["0", "0 00", "0"]], [["0", "0", "0"]], drop_pins=True)

    def test_text_path_cell_edit_roundtrip(self):
        """Cell edits through the extract → edit text → apply path."""
        orig = BytesIO()
        doc = Document()
        doc.add_paragraph("Lead paragraph.")
        table = doc.add_table(rows=2, cols=3)
        for r, row in enumerate([["Seats", "5", "€100"], ["Support", "1", "€500"]]):
            for c, cell in enumerate(row):
                table.rows[r].cells[c].text = cell
        doc.add_paragraph("Trailing paragraph.")
        doc.save(orig)
        text_orig = extract_text_from_stream(BytesIO(orig.getvalue()), clean_view=True)
        text_mod = text_orig.replace("Seats | 5 | €100", "Seats | 10 | €125")

        edits = generate_edits_via_paragraph_alignment(text_orig, text_mod)
        engine = RedlineEngine(BytesIO(orig.getvalue()), author="Fuzz")
        stats = engine.process_batch(list(edits))
        assert stats["edits_skipped"] == 0, stats["skipped_details"]
        engine.accept_all_revisions(remove_comments=True)
        final = extract_text_from_stream(engine.save_to_stream(), clean_view=True)
        assert final == text_mod

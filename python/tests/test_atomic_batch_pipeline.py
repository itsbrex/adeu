# FILE: tests/test_atomic_batch_pipeline.py
import asyncio
import io
import re

from docx import Document

from adeu.ingest import extract_text_from_stream
from adeu.mcp_components.tools.document import process_document_batch
from adeu.models import AcceptChange, ModifyText
from adeu.redline.engine import RedlineEngine


class MockContext:
    """Mock FastMCP Context to absorb async logging calls during tests."""

    async def info(self, msg, **kwargs):
        pass

    async def debug(self, msg, **kwargs):
        pass

    async def warning(self, msg, **kwargs):
        pass

    async def error(self, msg, **kwargs):
        pass


def test_atomic_batch_prevents_cascading_misanchor(tmp_path):
    """
    Validates Issue #2 Fix:
    Ensures that processing AcceptChanges (which mutate the XML DOM and shift text lengths)
    does not cause subsequent ModifyTexts in the SAME batch to misanchor.
    """
    # 1. Setup initial doc
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.add_paragraph("Third paragraph.")

    orig_path = tmp_path / "original.docx"
    doc.save(orig_path)

    # 2. Make an initial tracked change (Simulating Round 1)
    with open(orig_path, "rb") as f:
        engine = RedlineEngine(io.BytesIO(f.read()), author="Round1")

    # Edit: "First" -> "1st" (Creates a w:del and w:ins)
    engine.apply_edits([ModifyText(target_text="First", new_text="1st")])

    mid_path = tmp_path / "mid.docx"
    with open(mid_path, "wb") as f:
        f.write(engine.save_to_stream().getvalue())

    # Verify intermediate state (Round 1)
    with open(mid_path, "rb") as f:
        mid_text = extract_text_from_stream(io.BytesIO(f.read()))

    assert "{--First--}" in mid_text
    assert "{++1st++}" in mid_text

    # Extract dynamically generated Change IDs for the Accept action
    chg_ids = set(re.findall(r"\[Chg:(\d+)(?:\s+\w+)?\]", mid_text))
    assert len(chg_ids) > 0, "Tracked changes were not generated."

    # 3. Execute the Atomic Batch (Simulating Round 2)
    # We ACCEPT the previous changes. This removes the w:del and w:ins wrappers,
    # shrinking the XML and shifting the text indices of everything below it.
    actions = [AcceptChange(target_id=f"Chg:{i}") for i in chg_ids]

    # We edit text further down the document.
    # If the mapper is not rebuilt, "Third" will look for the wrong index and fail.
    edits = [ModifyText(target_text="Third", new_text="3rd")]
    changes = actions + edits

    out_path = tmp_path / "final.docx"

    # Run the new server tool asynchronously with the mock context
    result_msg = asyncio.run(
        process_document_batch(
            reasoning="test",
            original_docx_path=str(mid_path),
            author_name="Round2",
            ctx=MockContext(),
            changes=changes,
            output_path=str(out_path),
        )
    )

    # 4. Assertions on the Tool Execution
    assert "Batch complete" in result_msg
    assert f"Actions: {len(actions)} applied, 0 skipped" in result_msg
    assert "Edits: 1 applied, 0 skipped" in result_msg, "The edit misanchored and was skipped!"

    # 5. Assertions on the Final Document State
    with open(out_path, "rb") as f:
        final_text = extract_text_from_stream(io.BytesIO(f.read()))

    # The first paragraph should be cleanly accepted
    assert "1st paragraph." in final_text
    assert "{--First--}" not in final_text

    # The third paragraph should have the new tracked change anchored perfectly
    assert "{--Third--}" in final_text
    assert "{++3rd++}" in final_text


def _make_dpa(tmp_path):
    """Build the DPA scenario: an identical placeholder in two clauses."""
    doc = Document()
    doc.add_paragraph("PROVIDER: [official company name] shall process the data.")
    doc.add_paragraph("PROVIDER: [official company name] is the data processor.")
    path = tmp_path / "dpa.docx"
    doc.save(path)
    return path


def test_ambiguous_batch_rejection_then_match_mode_all_through_real_tool(tmp_path):
    """
    End-to-end through the ACTUAL process_document_batch tool (not just the
    engine): an ambiguous strict edit is rejected with guidance naming the
    match_mode escape hatch, and re-running with the suggested match_mode="all"
    actually mutates the document on disk. This is the full "Turn Loop Trap"
    fix verified along the real agent-facing path.
    """
    dpa_path = _make_dpa(tmp_path)

    # 1. The flawed strict edit — exactly what trapped the agent originally.
    strict_edit = ModifyText(
        target_text="PROVIDER: [official company name]",
        new_text="PROVIDER: Acme Corp",
    )
    rejection = asyncio.run(
        process_document_batch(
            reasoning="test",
            original_docx_path=str(dpa_path),
            author_name="Reviewer",
            ctx=MockContext(),
            changes=[strict_edit],
        )
    )

    # The agent-facing message must reject AND show how to re-call.
    assert "Batch rejected" in rejection
    assert "Ambiguous match" in rejection
    assert 'match_mode": "all"' in rejection
    assert 'match_mode": "first"' in rejection

    # 2. Follow the guidance verbatim: re-send with match_mode="all".
    all_edit = ModifyText(
        target_text="PROVIDER: [official company name]",
        new_text="PROVIDER: Acme Corp",
        match_mode="all",
    )
    out_path = tmp_path / "dpa_all.docx"
    result = asyncio.run(
        process_document_batch(
            reasoning="test",
            original_docx_path=str(dpa_path),
            author_name="Reviewer",
            ctx=MockContext(),
            changes=[all_edit],
            output_path=str(out_path),
        )
    )

    assert "Batch complete" in result
    # match_mode="all" fans the single edit out into one applied edit per occurrence.
    assert "Edits: 2 applied" in result

    # 3. The document on disk is genuinely modified in BOTH clauses. (The engine
    # trims the shared "PROVIDER: " prefix, so only the differing tail is redlined.)
    with open(out_path, "rb") as f:
        final_text = extract_text_from_stream(io.BytesIO(f.read()))

    assert final_text.count("{++Acme Corp++}") == 2
    assert final_text.count("{--[official company name]--}") == 2


def test_ambiguous_batch_resolved_by_match_mode_first_through_real_tool(tmp_path):
    """
    Same agent-facing path, but following the match_mode="first" branch of the
    guidance: only the first occurrence is mutated on disk.
    """
    dpa_path = _make_dpa(tmp_path)

    first_edit = ModifyText(
        target_text="PROVIDER: [official company name]",
        new_text="PROVIDER: Acme Corp",
        match_mode="first",
    )
    out_path = tmp_path / "dpa_first.docx"
    result = asyncio.run(
        process_document_batch(
            reasoning="test",
            original_docx_path=str(dpa_path),
            author_name="Reviewer",
            ctx=MockContext(),
            changes=[first_edit],
            output_path=str(out_path),
        )
    )

    assert "Batch complete" in result
    assert "Edits: 1 applied" in result

    with open(out_path, "rb") as f:
        final_text = extract_text_from_stream(io.BytesIO(f.read()))

    # Exactly one clause is changed; the other placeholder is untouched.
    assert final_text.count("{++Acme Corp++}") == 1
    assert final_text.count("{--[official company name]--}") == 1
    assert "PROVIDER: [official company name] is the data processor." in final_text

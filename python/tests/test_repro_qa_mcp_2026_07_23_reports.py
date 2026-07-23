# FILE: tests/test_repro_qa_mcp_2026_07_23_reports.py
"""
Repro tests for ADEU-MCP-QA-REPORT.md (2026-07-23, black-box QA of the Node
MCP server, v1.29.0+4bb70f9) — Python mirror.

Each Node finding below was re-verified against the Python engine / MCP tools
and encoded here only where it reproduces:

  F5   Unsupported markdown handled inconsistently:
       - hyperlink insertion is rejected with dead-end advice ("Use a
         dedicated structural operation." — no such operation exists for MCP
         callers), engine.py validate_edit_strings;
       - "- item" as a new paragraph half-applies: ListParagraph style with
         NO w:numPr (renders indented, bulletless, and the "- " marker is
         silently eaten on round-trip);
       - "* item" — the projection's OWN bullet syntax (AI_CONTEXT §13
         Lists) — also half-applies instead of producing a real bullet that
         re-reads as "* item".
       (The "only '1. ' converts to a list" rule of AI_CONTEXT §11 is
       deliberate and NOT contradicted here.)
  F6   Batch report previews are not faithful to the saved document:
       (1) match_mode="all" previews only the FIRST occurrence changed;
       (2) an older pending insertion renders as plain accepted text in a
           later edit's preview (clean-mapper fallback path);
       (3) same-author re-edit of a pending insertion previews as NESTED
           CriticMarkup {++...{++...++}...++}.
  F7   Dry-run report omits comments entirely (engine report dict AND the
       rendered process_document_batch dry-run response).
  F12  (a) sanitize --accept-all headline counts revision MARKS (deliberate
       unit, AI_CONTEXT 2026-07-22) but the report then LISTS fewer change
       items with no reconciliation — a fragmented revision (multi-paragraph
       insertion: paragraph-mark w:ins elements carry no text) makes the
       headline and the list disagree silently;
       (b) the accept_all_changes MCP tool removes all comment bodies but
       its response never mentions comments at all.
  F14  `adeu diff` on identical files prints only "Found 0 changes:" — no
       explicit no-differences statement.
  F15  compare_clean=False diffs split CriticMarkup tokens across output
       lines: a "+" payload line ends with an unclosed {>>...<<} block and
       bare continuation lines carry closers with no opener.

Verified NON-repros (no test):
  - F6(2) raw-match path: when the later edit's target resolves in the RAW
    view, _tidy_preview_context keeps complete {++...++} blocks intact — the
    defect is specific to the clean-mapper fallback (tested below).
  - F12(b) sanitize_docx: the sanitize report DOES disclose comment removal
    ("COMMENTS (stripped)" section with per-comment lines) — only the
    accept_all_changes tool response is silent (tested below).
  - F14 MCP tool: diff_docx_files already returns the explicit
    "No text differences found between the documents." — only the CLI lacks
    a statement (tested below).

Every test is written test-first: it fails on current main and passes once
the finding is fixed.
"""

import asyncio
import io
import json
import re
import sys

from docx import Document
from docx.oxml.ns import qn

from adeu.ingest import extract_text_from_stream
from adeu.mcp_components.tools.document import (
    accept_all_changes,
    diff_docx_files,
    process_document_batch,
)
from adeu.models import ModifyText
from adeu.redline.engine import BatchValidationError, RedlineEngine
from adeu.sanitize.core import sanitize_docx

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


class MockContext:
    """Mock FastMCP Context to absorb async logging calls during tests."""

    async def info(self, msg, **kwargs):
        pass

    async def debug(self, msg, **kwargs):
        pass

    async def warning(self, msg, **kwargs):
        pass

    async def error(self, msg, **kwargs):
        pass


def run_cli(args, capsys):
    """Invoke the CLI in-process; returns (exit_code, stdout, stderr)."""
    from unittest.mock import patch

    from adeu.cli import main

    code = 0
    with patch.object(sys, "argv", ["adeu"] + [str(a) for a in args]):
        try:
            main()
        except SystemExit as e:
            code = e.code or 0
    captured = capsys.readouterr()
    return code, captured.out, captured.err


def doc_stream(*paragraphs) -> io.BytesIO:
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_docx(path, *paragraphs):
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(str(path))
    return path


def paragraph_containing(doc, needle):
    """The first w:p element of `doc` whose visible text contains `needle`
    (python-docx's Paragraph.text is blind to w:ins content, so walk w:t)."""
    for p_el in doc.element.iter(qn("w:p")):
        text = "".join(t.text or "" for t in p_el.iter(qn("w:t")))
        if needle in text:
            return p_el
    return None


# ---------------------------------------------------------------------------
# F5. Unsupported markdown handled inconsistently
# ---------------------------------------------------------------------------


class TestF5UnsupportedMarkdown:
    def test_hyperlink_insert_rejection_gives_no_dead_end_advice(self):
        """Inserting `[text](url)` must still be rejected — but the error must
        not point the caller at a "dedicated structural operation" that does
        not exist anywhere in the MCP toolset (F5)."""
        engine = RedlineEngine(doc_stream("Visit our website for details."), author="QA")
        errors = engine.validate_edits(
            [ModifyText(target_text="our website", new_text="[our website](https://example.com)")]
        )

        assert len(errors) == 1, "hyperlink insertion must remain a validation rejection"
        assert "dedicated structural operation" not in errors[0], (
            "dead-end advice: no such operation exists for MCP callers — the message must "
            "instead state what IS possible (e.g. edit the text of an existing link, or that "
            "new hyperlinks cannot be created):\n" + errors[0]
        )

    def test_dash_list_item_is_not_half_applied(self):
        """`- item` in a new paragraph currently produces a ListParagraph-styled
        paragraph with NO w:numPr: Word renders indented text with no bullet,
        and the projection round-trips it as plain "item one" (marker silently
        eaten). Acceptable outcomes are literal round-trip, a real bullet, or
        rejection — never the half-applied state (F5)."""
        engine = RedlineEngine(doc_stream("Intro paragraph.", "Closing paragraph."), author="QA")
        try:
            engine.process_batch(
                [ModifyText(target_text="Intro paragraph.", new_text="Intro paragraph.\n\n- item one")]
            )
        except BatchValidationError:
            return  # hard rejection is one of the acceptable consistent policies

        result = Document(engine.save_to_stream())
        p_el = paragraph_containing(result, "item one")
        assert p_el is not None, "the inserted paragraph must exist"

        pPr = p_el.find(qn("w:pPr"))
        pStyle = pPr.find(qn("w:pStyle")) if pPr is not None else None
        style_val = pStyle.get(qn("w:val")) if pStyle is not None else None
        has_numpr = pPr is not None and pPr.find(qn("w:numPr")) is not None

        assert not (style_val == "ListParagraph" and not has_numpr), (
            "half-applied list: the new paragraph got the ListParagraph style but no "
            "<w:numPr> — Word shows indented text with NO bullet, and the '- ' marker "
            "was silently consumed. Either keep '- item one' literal, emit a real "
            "bullet (w:numPr), or reject like hyperlinks are rejected."
        )

    def test_star_bullet_round_trips_as_real_bullet(self):
        """`* item` is the projection's OWN bullet syntax (AI_CONTEXT §13:
        bullets project as "* "). Writing it back must produce a real bullet
        that re-reads as "* item one" — the Virtual Text contract. Currently it
        half-applies exactly like "- item" and re-reads as bare "item one" (F5)."""
        engine = RedlineEngine(doc_stream("Intro paragraph.", "Closing paragraph."), author="QA")
        engine.process_batch([ModifyText(target_text="Intro paragraph.", new_text="Intro paragraph.\n\n* item one")])

        clean = extract_text_from_stream(engine.save_to_stream(), clean_view=True)
        assert "* item one" in clean, (
            "the projection's own bullet syntax must round-trip: writing '* item one' "
            "should yield a real bullet (w:numPr resolved from the paragraph or its "
            "style chain) that re-reads as '* item one', got:\n" + clean
        )


# ---------------------------------------------------------------------------
# F6. Batch report previews are not faithful to the saved document
# ---------------------------------------------------------------------------


class TestF6PreviewFidelity:
    def test_match_mode_all_preview_shows_every_occurrence(self):
        """match_mode="all" over "apple apple apple." reports 3 occurrences
        modified, but both previews show only the FIRST occurrence changed
        while the saved document changes all three (F6.1)."""
        engine = RedlineEngine(doc_stream("apple apple apple."), author="QA")
        stats = engine.process_batch([ModifyText(target_text="apple", new_text="pear", match_mode="all")])
        report = stats["edits"][0]

        # Sanity: the document itself is correct and all 3 occurrences counted.
        assert report["occurrences_modified"] == 3
        final = extract_text_from_stream(engine.save_to_stream(), clean_view=True)
        assert "pear pear pear." in final

        critic = report["critic_markup"] or ""
        clean = report["clean_text"] or ""
        assert critic.count("{++pear++}") == 3, (
            "the CriticMarkup preview must reflect ALL occurrences the edit modified "
            f"(3 reported), but it marks only {critic.count('{++pear++}')}: {critic!r}"
        )
        assert "apple" not in clean, (
            "the Clean preview claims to show the accepted result of an all-occurrences "
            f"edit, yet still contains unmodified 'apple': {clean!r} "
            f"(the saved document reads {final.strip()!r})"
        )

    def test_older_pending_insertion_keeps_markup_in_later_preview(self):
        """An OLDER pending insertion (previous batch) inside a later edit's
        preview window renders as plain accepted text when the later edit
        resolves via the clean-mapper fallback — it must keep its {++...++}
        markup so it never looks silently accepted (F6.2)."""
        engine = RedlineEngine(doc_stream("Payment due in 30 days of invoice receipt at the office."), author="Alice")
        engine.process_batch(
            [
                ModifyText(target_text="30 days", new_text="30 business days"),
                ModifyText(target_text="invoice receipt", new_text="written invoice receipt"),
            ]
        )
        stream = engine.save_to_stream()

        # The later edit's target only exists in the CLEAN view (it spans the
        # first pending insertion), forcing the clean-mapper fallback.
        engine2 = RedlineEngine(stream, author="Alice")
        stats = engine2.process_batch(
            [ModifyText(target_text="in 30 business days", new_text="within 30 business days")]
        )
        report = stats["edits"][0]
        assert report["status"] == "applied"

        critic = report["critic_markup"] or ""
        # The unrelated pending insertion "written " sits inside the preview
        # window. If shown, it must carry its insertion markup.
        assert ("written" not in critic) or ("{++written" in critic), (
            "the preview presents the OLDER pending insertion 'written ' as plain "
            "accepted text — an agent reading this believes it was silently accepted "
            f"(it is still pending): {critic!r}"
        )

    def test_same_author_reedit_preview_never_nests_criticmarkup(self):
        """Same-author re-edit of a pending insertion previews as NESTED
        CriticMarkup ({++... {++...++} ...++}) — invalid notation that no
        CriticMarkup consumer can parse (F6.3)."""
        engine = RedlineEngine(
            doc_stream("The parties shall resolve disputes amicably before litigation."),
            author="Alice",
        )
        engine.process_batch(
            [ModifyText(target_text="resolve disputes amicably", new_text="first negotiate in good faith")]
        )
        stream = engine.save_to_stream()

        engine2 = RedlineEngine(stream, author="Alice")
        stats = engine2.process_batch(
            [ModifyText(target_text="negotiate in good faith", new_text="negotiate in utmost good faith")]
        )
        report = stats["edits"][0]
        assert report["status"] == "applied"

        nested = re.compile(r"\{\+\+[^}]*\{\+\+")
        for field in ("critic_markup", "clean_text"):
            preview = report[field] or ""
            assert not nested.search(preview), (
                f"the {field} preview nests CriticMarkup insertions — invalid notation (F6.3): {preview!r}"
            )


# ---------------------------------------------------------------------------
# F7. Dry-run report omits comments entirely
# ---------------------------------------------------------------------------


COMMENT_TEXT = "Please review this change."


class TestF7DryRunOmitsComments:
    def test_engine_dry_run_report_carries_the_comment(self):
        """A dry-run batch report is the one place an agent can verify a
        comment before committing — the report dict must mention it (F7)."""
        engine = RedlineEngine(doc_stream("The Supplier shall deliver within 10 days."), author="QA")
        stats = engine.process_batch(
            [ModifyText(target_text="10 days", new_text="14 days", comment=COMMENT_TEXT)],
            dry_run=True,
        )

        assert stats["edits"][0]["status"] == "applied"
        assert COMMENT_TEXT in json.dumps(stats), (
            "the dry-run report contains no trace of the edit's comment — the one thing "
            "dry-run exists to let you verify before committing. Report dict fields: "
            + json.dumps(stats["edits"][0], indent=2)
        )

    def test_dry_run_tool_response_mentions_the_comment(self, tmp_path):
        """The rendered process_document_batch dry-run response must surface
        the comment text, not silently drop it (F7)."""
        src = build_docx(tmp_path / "f7.docx", "The Supplier shall deliver within 10 days.")

        result = asyncio.run(
            process_document_batch(
                reasoning="test",
                original_docx_path=str(src),
                author_name="Reviewer",
                ctx=MockContext(),
                changes=[ModifyText(target_text="10 days", new_text="14 days", comment=COMMENT_TEXT)],
                dry_run=True,
            )
        )

        assert "Dry-run simulation complete" in result
        assert COMMENT_TEXT in result, "the dry-run response never mentions the comment the edit carries:\n" + result


# ---------------------------------------------------------------------------
# F12. Finalization/sanitize report inconsistencies
# ---------------------------------------------------------------------------


class TestF12FinalizationReports:
    def test_sanitize_accept_all_headline_matches_listed_change_items(self, tmp_path):
        """The accept-all headline counts revision MARKS (deliberate unit,
        AI_CONTEXT 2026-07-22) — but the report then lists only the marks that
        carry text. A fragmented revision (multi-paragraph insertion whose
        paragraph-mark w:ins elements have no text) makes the headline and the
        listed items disagree with no reconciliation, exactly the F12 shape
        ("auto-accepted: 15" over 13 listed items)."""
        engine = RedlineEngine(doc_stream("First clause text here.", "Second clause stays put."), author="Alice")
        engine.process_batch(
            [
                # One logical change fragmented across multiple w:ins elements:
                # two content paragraphs plus two textless paragraph-mark w:ins.
                ModifyText(
                    target_text="First clause text here.",
                    new_text=("First clause text here.\n\nA second inserted paragraph.\n\nA third inserted paragraph."),
                ),
                # Plus one normal tracked deletion.
                ModifyText(target_text="stays put", new_text=""),
            ]
        )
        inp = tmp_path / "f12_fragmented.docx"
        inp.write_bytes(engine.save_to_stream().getvalue())

        result = sanitize_docx(str(inp), str(tmp_path / "f12_fragmented_sanitized.docx"), accept_all=True)
        report = result.report_text

        headline = re.search(r"Tracked changes auto-accepted: (\d+)", report)
        assert headline, "the accept-all report must carry its headline count:\n" + report
        headline_count = int(headline.group(1))
        listed = len(re.findall(r"Accepted (?:insertion|deletion of): \"", report))
        assert listed > 0, "the report must list the accepted changes:\n" + report

        if headline_count != listed:
            unlisted = headline_count - listed
            reconciled = re.search(
                r"(?i)\b%d\b[^\n]*(paragraph mark|revision mark|no (?:visible )?text|not listed|formatting)" % unlisted,
                report,
            )
            assert reconciled, (
                f"the headline says {headline_count} changes were auto-accepted but only "
                f"{listed} items are listed, and nothing in the report reconciles the "
                f"missing {unlisted} (textless paragraph-mark revisions of the fragmented "
                "insertion). Either list every counted mark or state the difference "
                "explicitly:\n" + report
            )

    def test_accept_all_changes_tool_discloses_comment_removal(self, tmp_path):
        """accept_all_changes deletes every comment body (remove_comments=True)
        yet its response never mentions comments — silent destruction of review
        content (F12). The sanitize report already discloses this; the tool
        response must too."""
        engine = RedlineEngine(doc_stream("Clause with a comment target."), author="Alice")
        engine.process_batch(
            [
                ModifyText(
                    target_text="comment target",
                    new_text="comment target",
                    comment="Check this wording.",
                )
            ]
        )
        # NB: neither path may contain the word "comment" — the response echoes
        # the output path, which would satisfy the assertion vacuously.
        src = tmp_path / "f12_annotated.docx"
        src.write_bytes(engine.save_to_stream().getvalue())
        out = tmp_path / "f12_annotated_clean.docx"

        result = asyncio.run(
            accept_all_changes(
                reasoning="test",
                docx_path=str(src),
                ctx=MockContext(),
                output_path=str(out),
            )
        )

        # Sanity: the comment body really was removed from the output.
        final_raw = extract_text_from_stream(io.BytesIO(out.read_bytes()))
        assert "Check this wording." not in final_raw

        assert re.search(r"comment", result, re.IGNORECASE), (
            "the tool removed a comment but its response never says so — the report "
            "must mention the comment removal:\n" + result
        )


# ---------------------------------------------------------------------------
# F14. Diff on identical files: no explicit "No differences" statement
# ---------------------------------------------------------------------------


class TestF14DiffIdenticalFiles:
    # The MCP diff_docx_files tool already returns the explicit
    # "No text differences found between the documents." — only the CLI
    # leaves the reader to interpret "Found 0 changes:" (F14).

    def test_cli_diff_identical_files_states_no_differences(self, tmp_path, capsys):
        a = build_docx(tmp_path / "same_a.docx", "Clause one text.", "Clause two text.")
        b = build_docx(tmp_path / "same_b.docx", "Clause one text.", "Clause two text.")

        code, out, err = run_cli(["diff", a, b], capsys)

        assert code == 0
        combined = out + err
        assert re.search(r"no (text(ual)? )?differences", combined, re.IGNORECASE), (
            "diffing two identical documents must state explicitly that there are no "
            "differences (the MCP tool already says 'No text differences found between "
            f"the documents.'), got only: {combined!r}"
        )


# ---------------------------------------------------------------------------
# F15. Raw-markup diff splits CriticMarkup tokens across output lines
# ---------------------------------------------------------------------------


_CM_OPENERS = ("{--", "{++", "{>>", "{==")
_CM_CLOSERS = ("--}", "++}", "<<}", "==}")


def _criticmarkup_unbalanced_lines(text: str):
    """Lines whose CriticMarkup delimiters do not pair up within the line —
    the observed breakage: a '+' payload line ends inside an unclosed
    {>>...<<} block and bare continuation lines carry the closer."""
    bad = []
    for line in text.splitlines():
        opens = sum(line.count(tok) for tok in _CM_OPENERS)
        closes = sum(line.count(tok) for tok in _CM_CLOSERS)
        if opens != closes:
            bad.append(f"(open={opens}, close={closes}) {line!r}")
    return bad


class TestF15RawMarkupDiffTokenIntegrity:
    def test_compare_clean_false_keeps_criticmarkup_tokens_atomic_per_line(self, tmp_path):
        """diff_docx_files(compare_clean=False) over a document carrying
        tracked changes + a comment emits Word Patch hunks whose payload lines
        split CriticMarkup tokens: the '+' line ends mid-{>>...<<} (opener
        without closer) and its continuation lines start with bare metadata /
        closers. Every output line must keep its delimiters paired (F15)."""
        orig = build_docx(
            tmp_path / "raw_orig.docx",
            "The Supplier shall deliver the goods within 30 days of the order date.",
            "This agreement is governed by the laws of Finland.",
        )
        engine = RedlineEngine(io.BytesIO(orig.read_bytes()), author="Alice Reviewer")
        engine.process_batch(
            [
                ModifyText(
                    target_text="30 days",
                    new_text="60 days",
                    comment="Extended per client request.",
                ),
                ModifyText(target_text="Finland", new_text="Sweden"),
            ]
        )
        mod = tmp_path / "raw_mod.docx"
        mod.write_bytes(engine.save_to_stream().getvalue())

        result = asyncio.run(
            diff_docx_files(
                reasoning="test",
                original_path=str(orig),
                modified_path=str(mod),
                ctx=MockContext(),
                compare_clean=False,
            )
        )

        # Sanity: the raw comparison actually surfaces the markup.
        assert "{++" in result and "{>>" in result, (
            "compare_clean=False must diff the raw projection including markup:\n" + result
        )

        bad = _criticmarkup_unbalanced_lines(result)
        assert not bad, (
            "raw-markup diff output splits CriticMarkup tokens across lines — each of "
            "these lines opens or closes a {--/{++/{>>/{== block that is completed on "
            "a different line, which is unreadable and unparseable:\n  "
            + "\n  ".join(bad)
            + "\n\nFull output:\n"
            + result
        )

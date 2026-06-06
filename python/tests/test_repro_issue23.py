"""
Regression tests for GitHub Issue #23:
"Malformed comments.xml when creating the comments part from scratch (+ smaller findings)"

All tests in this file are DETECTION tests: they are expected to FAIL until the
described bug is fixed.  They must NOT be changed to accommodate the current
broken behaviour.

Cross-platform parity: matching tests live in
    node/packages/core/src/engine.issue23.test.ts
"""

import io
import re
import shutil
import subprocess
import zipfile

import pytest
from docx import Document

from adeu.ingest import extract_text_from_stream
from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_clean_docx(*paragraphs: str) -> io.BytesIO:
    """Returns a fresh DOCX BytesIO stream with no pre-existing comments part."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _xmllint(xml_bytes: bytes, tmp_path, filename: str = "test.xml") -> subprocess.CompletedProcess:
    """
    Validates *xml_bytes* with xmllint.
    Hard-fails with installation instructions if xmllint is not on PATH.
    """
    if not shutil.which("xmllint"):
        raise RuntimeError(
            "xmllint is required for this test but was not found on PATH.\n"
            "Install it with:\n"
            "  Ubuntu/Debian : sudo apt install libxml2-utils\n"
            "  macOS (Homebrew): brew install libxml2\n"
            "  Alpine        : apk add libxml2-utils\n"
        )
    path = tmp_path / filename
    path.write_bytes(xml_bytes)
    return subprocess.run(
        ["xmllint", "--noout", str(path)],
        capture_output=True,
        text=True,
    )


def _find_comments_part_name(names: list) -> str | None:
    """
    Returns the zip entry name for the main comments part (word/comments.xml or
    word/comments1.xml etc.) — the one that contains <w:comments>, not the
    Extended/Ids/Extensible siblings.
    """
    for name in sorted(names):
        if re.fullmatch(r"word/comments\d*\.xml", name):
            return name
    return None


def _extract_comments_xml(engine: RedlineEngine) -> bytes:
    """Saves the engine output and extracts the main comments part from the zip."""
    buf = engine.save_to_stream()
    buf.seek(0)
    with zipfile.ZipFile(buf) as z:
        name = _find_comments_part_name(z.namelist())
        assert name is not None, f"No word/comments*.xml found in output. Parts: {z.namelist()}"
        return z.read(name)


# ===========================================================================
# Bug #1 (primary) — comments.xml missing xmlns:w14 on freshly created part
# ===========================================================================


class TestCommentsXmlNamespace:
    """
    BUG-23-1: When comments.xml is created from scratch (no pre-existing
    comments part), the root <w:comments> element must declare xmlns:w14 so
    that w14:paraId / w14:textId attributes on child <w:p> elements are valid.

    Without the declaration xmllint emits:
      namespace error: Namespace prefix w14 for paraId on p is not defined
    and read_docx raises:
      NamespaceError: prefix is non-null and namespace is null

    Cross-platform parity: 'BUG-23-1' tests in engine.issue23.test.ts
    """

    def test_comments_xml_declares_w14_on_fresh_doc(self, tmp_path):
        buf = _make_clean_docx("The only paragraph in this document.")
        engine = RedlineEngine(buf, author="Test Author")
        engine.apply_edits(
            [
                ModifyText(
                    target_text="only",
                    new_text="only",
                    comment="Forces creation of comments.xml from scratch",
                )
            ]
        )

        comments_xml = _extract_comments_xml(engine)

        assert b"xmlns:w14=" in comments_xml, (
            "BUG-23-1 (Python, fresh): xmlns:w14 not declared on <w:comments> root.\n"
            f"XML (first 600 bytes):\n{comments_xml[:600].decode('utf-8', errors='replace')}"
        )

    def test_comments_xml_passes_xmllint_on_fresh_doc(self, tmp_path):
        buf = _make_clean_docx("The only paragraph in this document.")
        engine = RedlineEngine(buf, author="Test Author")
        engine.apply_edits(
            [
                ModifyText(
                    target_text="only",
                    new_text="only",
                    comment="Forces creation of comments.xml from scratch",
                )
            ]
        )

        comments_xml = _extract_comments_xml(engine)
        result = _xmllint(comments_xml, tmp_path, "comments_fresh.xml")
        assert result.returncode == 0, (
            f"BUG-23-1 (Python, xmllint): comments.xml fails XML validation:\n{result.stderr}"
        )

    def test_read_docx_roundtrip_does_not_raise_namespace_error(self):
        """
        read_docx (extract_text_from_stream) on the saved output must not
        raise NamespaceError: prefix is non-null and namespace is null.
        """
        buf = _make_clean_docx("Hello world, this is a roundtrip test.")
        engine = RedlineEngine(buf, author="Test Author")
        engine.apply_edits(
            [
                ModifyText(
                    target_text="Hello",
                    new_text="Hello",
                    comment="Roundtrip comment",
                )
            ]
        )

        saved = engine.save_to_stream()
        saved.seek(0)

        # Must not raise
        text = extract_text_from_stream(saved)
        assert "Hello" in text, f"Document unreadable after roundtrip: {text!r}"

    def test_comments_xml_declares_w14_on_doc_with_bare_legacy_part(self, tmp_path):
        """
        When an existing comments.xml has no xmlns:w14 (e.g. a Pandoc-generated
        or older Word document), adding a new comment must still produce a
        valid output — the namespace must be patched onto the root element.

        This is the Node-side blind spot: _ensureNamespaces() is a no-op stub.
        Cross-platform parity: 'BUG-23-1-legacy' in engine.issue23.test.ts
        """
        from docx.opc.constants import CONTENT_TYPE as CT
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.opc.part import XmlPart
        from docx.oxml import parse_xml

        doc_obj = Document()
        doc_obj.add_paragraph("Anchor text for the legacy-part test.")

        # Inject a bare comments.xml that deliberately omits xmlns:w14
        bare_xml = b'<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"></w:comments>'
        pkg = doc_obj.part.package
        pn = pkg.next_partname("/word/comments%d.xml")
        part = XmlPart(pn, CT.WML_COMMENTS, parse_xml(bare_xml), pkg)
        pkg.parts.append(part)
        doc_obj.part.relate_to(part, RT.COMMENTS)

        buf = io.BytesIO()
        doc_obj.save(buf)
        buf.seek(0)

        engine = RedlineEngine(buf, author="Test Author")
        engine.apply_edits(
            [
                ModifyText(
                    target_text="Anchor",
                    new_text="Anchor",
                    comment="Comment added to doc with legacy bare comments part",
                )
            ]
        )

        comments_xml = _extract_comments_xml(engine)

        assert b"xmlns:w14=" in comments_xml, (
            "BUG-23-1 (Python, legacy-part): xmlns:w14 not added to existing "
            "bare comments.xml.\n"
            f"XML (first 600 bytes):\n{comments_xml[:600].decode('utf-8', errors='replace')}"
        )

        result = _xmllint(comments_xml, tmp_path, "comments_legacy.xml")
        assert result.returncode == 0, (
            f"BUG-23-1 (Python, legacy-part, xmllint): comments.xml fails XML validation:\n{result.stderr}"
        )


# ===========================================================================
# Bug #2 — Inserted runs inherit anchor paragraph's character formatting
# ===========================================================================


class TestInsertedRunFormatting:
    """
    BUG-23-2: When modify inserts text into a paragraph whose runs are italic,
    the inserted w:ins/w:r must NOT automatically be italic.  There is currently
    no override mechanism.

    Cross-platform parity: 'BUG-23-2' in engine.issue23.test.ts
    """

    def test_inserted_run_does_not_inherit_italic(self):
        from xml.etree import ElementTree as ET

        doc_obj = Document()
        para = doc_obj.add_paragraph()
        run = para.add_run("italicized anchor text here")
        run.italic = True

        buf = io.BytesIO()
        doc_obj.save(buf)
        buf.seek(0)

        engine = RedlineEngine(buf, author="Test Author")
        engine.apply_edits(
            [
                ModifyText(
                    target_text="anchor",
                    new_text="plain",
                )
            ]
        )

        saved = engine.save_to_stream()
        saved.seek(0)

        with zipfile.ZipFile(saved) as z:
            doc_xml = z.read("word/document.xml")

        root = ET.fromstring(doc_xml)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        inserted_runs = root.findall(".//w:ins/w:r", ns)
        assert len(inserted_runs) > 0, "No w:ins/w:r elements found — edit did not produce a tracked insertion"

        offending = []
        for r in inserted_runs:
            rpr = r.find("w:rPr", ns)
            if rpr is not None and rpr.find("w:i", ns) is not None:
                offending.append(ET.tostring(r, encoding="unicode"))

        assert offending == [], (
            "BUG-23-2: Inserted run(s) inherited italic formatting from the "
            "surrounding paragraph.  Offending runs:\n" + "\n".join(offending)
        )


# ===========================================================================
# Bug #3 — modify diff placement ignores new_text ordering
# ===========================================================================


class TestDiffPlacement:
    """
    BUG-23-3: The diff engine always appends the delta AFTER the common match
    regardless of where it sits in new_text.  Prefixing new text before the
    anchor should produce an insert-before, not an append-after.

    Cross-platform parity: 'BUG-23-3' in engine.issue23.test.ts
    """

    def test_prefix_insertion_lands_before_anchor(self):
        """
        target_text="fox", new_text="red fox"

        Expected markup: "...{++red ++}fox..."   (prefix before, fox kept)
        Bug behaviour:   "...{--fox--}{++red fox++}..." or "...fox{++red ++}..."
        """
        buf = _make_clean_docx("The quick brown fox jumps over the lazy dog.")
        engine = RedlineEngine(buf, author="Test Author")
        engine.apply_edits([ModifyText(target_text="fox", new_text="red fox")])

        saved = engine.save_to_stream()
        saved.seek(0)
        text = extract_text_from_stream(saved)

        # fox must NOT be struck out — it is the preserved anchor
        assert "{--fox--}" not in text, (
            "BUG-23-3: 'fox' was struck out but it should remain as the anchor; "
            "only 'red ' should be inserted before it.\n"
            f"Full text: {text!r}"
        )

        # The inserted prefix must appear somewhere in the text
        inserted = re.search(r"\{\+\+red\s*\+\+\}", text)
        assert inserted is not None, f"BUG-23-3: Expected insertion of 'red' not found in output.\nFull text: {text!r}"

        # And it must appear BEFORE "fox" — guard that fox was not silently deleted
        fox_pos = text.find("fox")
        assert fox_pos >= 0, (
            "BUG-23-3: 'fox' was removed from the output entirely; it must be "
            "preserved as the anchor.\n"
            f"Full text: {text!r}"
        )
        assert inserted.start() < fox_pos, (
            "BUG-23-3: '{++red++}' appears AFTER 'fox' in the output; "
            "the delta should be placed before the anchor.\n"
            f"Full text: {text!r}"
        )

    def test_paragraph_separator_before_anchor_is_preserved(self):
        """
        target_text="Conclusion", new_text="Summary\\n\\nConclusion"

        Expected: a new "Summary" paragraph appears BEFORE "Conclusion",
                  separated by a paragraph break.
        Bug behaviour: "Summary" is dropped entirely, or merged into the
                       Conclusion paragraph, or appended after it.

        Cross-platform parity: 'BUG-23-3b' in engine.issue23.test.ts
        """
        buf = _make_clean_docx("Introduction paragraph.", "Conclusion paragraph.")
        engine = RedlineEngine(buf, author="Test Author")
        engine.apply_edits(
            [
                ModifyText(
                    target_text="Conclusion",
                    new_text="Summary\n\nConclusion",
                )
            ]
        )

        saved = engine.save_to_stream()
        saved.seek(0)
        text = extract_text_from_stream(saved)

        assert "Summary" in text, f"BUG-23-3b: The inserted 'Summary' paragraph was lost entirely.\nFull text: {text!r}"

        summary_pos = text.find("Summary")
        conclusion_pos = text.find("Conclusion")
        assert summary_pos < conclusion_pos, (
            f"BUG-23-3b: 'Summary' appears after 'Conclusion' in the output.\nFull text: {text!r}"
        )

        between = text[summary_pos:conclusion_pos]
        assert "\n" in between, (
            "BUG-23-3b: No paragraph break between 'Summary' and 'Conclusion'; "
            "they appear merged into a single paragraph.\n"
            f"Between: {between!r}\nFull text: {text!r}"
        )


# ===========================================================================
# Bug #4 — Multi-paragraph target_text is silently corrupt or opaque error
# ===========================================================================


class TestMultiParagraphTarget:
    """
    BUG-23-4: A target_text that contains a paragraph break (\\n\\n) currently
    collapses the break in the token stream, causing either silent document
    corruption or an opaque 'Target text not found' error with no mention of
    the multi-paragraph limitation.

    The correct behaviour is to either:
      (a) support multi-paragraph targets correctly, OR
      (b) reject them with a clear, actionable error message.

    Cross-platform parity: 'BUG-23-4' in engine.issue23.test.ts
    """

    def test_multi_paragraph_target_rejected_with_actionable_error(self):
        buf = _make_clean_docx("First paragraph content.", "Second paragraph content.")
        engine = RedlineEngine(buf, author="Test Author")

        raised = None
        try:
            engine.apply_edits(
                [
                    ModifyText(
                        target_text="First paragraph content.\n\nSecond paragraph content.",
                        new_text="Single replacement paragraph.",
                    )
                ]
            )
        except Exception as e:
            raised = e

        if raised is None:
            # No error raised — verify the paragraph boundary wasn't silently collapsed.
            # Bug signature: the two paragraphs are merged into a single deleted token
            # without the \\n\\n separator, e.g.
            #   {--First paragraph content.Second paragraph content.--}
            # A correct implementation either keeps the boundary or raises a clear error.
            saved = engine.save_to_stream()
            saved.seek(0)
            text = extract_text_from_stream(saved)

            collapsed = "First paragraph content.Second paragraph content."
            assert collapsed not in text, (
                "BUG-23-4: Multi-paragraph target_text was silently accepted and the "
                "paragraph boundary was collapsed (the \\n\\n separator is missing from "
                "the deletion token).  Either support multi-paragraph targets correctly "
                "or reject them with a clear error.\n"
                f"Full text: {text!r}"
            )
            space_collapsed = "First paragraph content. Second paragraph content."
            assert space_collapsed not in text, (
                "BUG-23-4: Multi-paragraph target_text was silently accepted and the "
                "paragraph boundary was collapsed with a space separator.\n"
                f"Full text: {text!r}"
            )
        else:
            # An error was raised — it must specifically mention the multi-paragraph issue
            error_msg = str(raised).lower()
            actionable_keywords = ["paragraph", "multi", "boundary", "newline", "cross"]
            assert any(kw in error_msg for kw in actionable_keywords), (
                "BUG-23-4: An error was raised but the message gives no indication "
                "that multi-paragraph target_text is unsupported.\n"
                f"Error: {raised!r}"
            )


# ===========================================================================
# Bug #5 — Ambiguous-match check counts text inside w:del (tracked deletions)
# ===========================================================================


class TestAmbiguousMatchDel:
    """
    BUG-23-5: The ambiguity check counts occurrences of text inside w:del
    elements (tracked deletions) as live matches.  After one duplicate line
    is tracked-deleted, the remaining live copy must be uniquely matchable.

    Cross-platform parity: 'BUG-23-5' in engine.issue23.test.ts
    """

    def test_tracked_deleted_copy_not_counted_as_ambiguous(self):
        """
        1. Doc: "Context A: Dupe"  /  "Context B: Dupe"
        2. Batch 1: delete "Context A: Dupe" via tracked change (unique match via context)
        3. Batch 2: target_text="Dupe" -> only "Context B: Dupe" is live; must succeed
        """
        buf = _make_clean_docx("Context A: Dupe", "Context B: Dupe")

        # Batch 1: uniquely delete the first occurrence
        engine1 = RedlineEngine(buf, author="Test Author")
        engine1.apply_edits(
            [
                ModifyText(
                    target_text="Context A: Dupe",
                    new_text="",
                )
            ]
        )

        saved1 = engine1.save_to_stream()
        saved1.seek(0)

        # Sanity: the first copy is inside a w:del
        text1 = extract_text_from_stream(saved1)
        assert "{--Context A: Dupe--}" in text1, f"Test setup: expected tracked deletion, got: {text1!r}"

        saved1.seek(0)
        engine2 = RedlineEngine(saved1, author="Test Author")

        # Batch 2: only "Context B: Dupe" is live — this must NOT raise ambiguous-match
        try:
            engine2.apply_edits(
                [
                    ModifyText(
                        target_text="Dupe",
                        new_text="Unique",
                    )
                ]
            )
        except Exception as e:
            pytest.fail(
                "BUG-23-5: Modifying the only live 'Dupe' raised an error.\n"
                "The w:del-wrapped copy must not be counted toward the ambiguity "
                f"check.\nError: {e}"
            )

        saved2 = engine2.save_to_stream()
        saved2.seek(0)
        text2 = extract_text_from_stream(saved2)

        # 'Unique' must appear as a TRACKED INSERTION ({++Unique++}),
        # NOT as a tracked deletion ({--Unique--}).
        # The bug manifests as {--Unique--} because the engine edits the
        # w:del-wrapped text instead of the live 'Context B: Dupe' paragraph.
        assert "{++Unique++}" in text2, (
            f"BUG-23-5: 'Unique' was not inserted as a tracked change into the live text.\n"
            f"If the output contains '{{--Unique--}}' instead, the engine modified the "
            f"text inside a w:del (tracked deletion) rather than the live paragraph.\n"
            f"Full text: {text2!r}"
        )
        assert "{--Unique--}" not in text2, (
            "BUG-23-5: The engine modified text inside a w:del element — "
            "{--Unique--} is present, meaning the tracked deletion was edited "
            "in addition to (or instead of) the live paragraph.\n"
            f"Full text: {text2!r}"
        )

# FILE: tests/test_repro_qa_report_v7.py
"""
Repro tests for the 2026-07-19 black-box QA and UX report
(adeu 1.25.0+35c8bb4).

Finding index (report finding -> test class below):
  F-01  `adeu diff --json` emits a batch its own `apply` rejects: context
        widening (make_edits_self_contained) pulls an EARLIER edit's target
        into a LATER edit's anchor text, so the sequential apply matches
        text that is already inside a tracked deletion
  F-02  a formatting-only diff (bold -> italic) reports success but the
        accepted document is bold+italic: marker-carrying replacements
        inherit the replaced span's run properties instead of treating the
        explicit markers as authoritative
  F-03  a bold run ending in a space projects as `**The Supplier **` —
        malformed Markdown that also poisons CriticMarkup output (F-10)
  F-04  image handling: (a) documents whose embedded media bytes differ but
        whose text projections are identical diff to [] with no warning;
        (b) an alt-text-only difference emits a text edit that apply
        categorically rejects (read-only image markers) (also F-14)
  F-05  full-text replacement on a small structured document reports a
        clean_text that re-extraction contradicts; the deletion guard has an
        undocumented 2000-char floor
  F-06  `regex + match_mode=all` rebuilds the document map once per
        occurrence — strongly superlinear apply times
  F-07  `sanitize --keep-markup --accept-all` silently resolves the
        contradiction instead of rejecting it
  F-08  `sanitize -o X --outdir Y` silently ignores -o
  F-09  `sanitize --keep-markup` leaves original comment timestamps in
        word/comments.xml while normalizing every other date
  F-12  match_mode is not validated as an enum ("banana", null accepted)
  F-13  an invalid apply regex is reported as "target text not found"
        instead of as a regex syntax error
  F-16  `adeu init` writes an unpinned `--from adeu` uvx command
  F-17  search results report a heading path truncated at the highlighted
        query ("Master" for a hit inside "Master Services Agreement")
  F-18  `extract --mode outline --page all` first warns --page is ignored,
        then errors because 'all' is invalid
  F-19  an invalid --baseline is reported as "<input>.docx is not a valid
        DOCX", naming the wrong file
  F-21  one match_mode=all edit over two occurrences reports
        edits_applied=2 while the edits array has a single entry

Every test fails against the commit preceding its fix.
"""

import json
import re
import sys
import zipfile
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Inches
from pydantic import TypeAdapter, ValidationError

from adeu.diff import generate_structured_edits
from adeu.ingest import _extract_text_from_doc
from adeu.models import BatchChanges, ModifyText
from adeu.redline.engine import BatchValidationError, RedlineEngine

# ---------------------------------------------------------------------------
# Helpers / fixture builders
# ---------------------------------------------------------------------------


def run_cli(args, capsys):
    """Invoke the CLI in-process; returns (exit_code, stdout, stderr)."""
    from unittest.mock import patch

    from adeu.cli import main

    code = 0
    with patch.object(sys, "argv", ["adeu"] + [str(a) for a in args]):
        try:
            main()
        except SystemExit as e:
            code = e.code or 0
    captured = capsys.readouterr()
    return code, captured.out, captured.err


def doc_to_stream(doc) -> BytesIO:
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_docx(paragraphs, path: Path = None):
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if path is not None:
        doc.save(path)
    return doc


def extract_structured(doc):
    return _extract_text_from_doc(doc, clean_view=True, include_appendix=False, return_structure=True)


def json_round_trip(edits):
    """Serializes diff output exactly like `adeu diff --json` and re-parses it,
    dropping the private position pins JSON consumers never see."""
    payload = json.dumps([e.model_dump(exclude={"_match_start_index"}) for e in edits])
    return TypeAdapter(BatchChanges).validate_json(payload)


def clean_text_of(engine: RedlineEngine) -> str:
    return _extract_text_from_doc(engine.doc, clean_view=True, include_appendix=False)


MINIMAL_PNG_RED = bytes.fromhex(
    "89504e470d0a1a0a0000000d49484452000000010000000108020000009077"
    "53de0000000c4944415408d763f8cfc000000301010018dd8db00000000049"
    "454e44ae426082"
)
MINIMAL_PNG_BLUE = bytes.fromhex(
    "89504e470d0a1a0a0000000d49484452000000010000000108020000009077"
    "53de0000000c4944415408d7636060f80f00010501001a3b5ee90000000049"
    "454e44ae426082"
)


def build_image_doc(png_bytes: bytes, alt: str, path: Path):
    import tempfile

    doc = Document()
    doc.add_paragraph("Before image.")
    pic_para = doc.add_paragraph()
    run = pic_para.add_run()
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as fh:
        fh.write(png_bytes)
        png_path = fh.name
    run.add_picture(png_path, width=Inches(0.5))
    inline = run._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline")[
        0
    ]
    doc_pr = inline.find("{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr")
    doc_pr.set("descr", alt)
    doc.add_paragraph("After image.")
    doc.save(path)


def build_commented_tracked_docx(path: Path):
    """A document with one tracked change carrying a comment (known author/date)."""
    doc = Document()
    doc.add_paragraph("The quick brown fox jumps over the lazy dog.")
    doc.add_paragraph("Second paragraph for context.")
    engine = RedlineEngine(doc_to_stream(doc), author="Reviewer One")
    engine.process_batch(
        [
            ModifyText(
                type="modify",
                target_text="quick",
                new_text="slow",
                comment="Speed change rationale",
            )
        ]
    )
    path.write_bytes(engine.save_to_stream().getvalue())


# ---------------------------------------------------------------------------
# F-01: diff --json output must be applicable by apply (overlap coalescing)
# ---------------------------------------------------------------------------


class TestF01DiffApplyContract:
    def _diff_apply_accept(self, orig_doc, mod_doc):
        text_orig, struct_orig = extract_structured(orig_doc)
        text_mod, struct_mod = extract_structured(mod_doc)
        edits, warnings = generate_structured_edits(text_orig, struct_orig, text_mod, struct_mod)

        replayed = json_round_trip(edits)
        engine = RedlineEngine(doc_to_stream(orig_doc), author="QA")
        engine.process_batch(replayed)  # must NOT raise BatchValidationError
        engine.accept_all_revisions(remove_comments=True)
        return clean_text_of(engine), text_mod, warnings

    def test_overlapping_final_paragraph_edits_apply(self):
        orig = build_docx(["Intro paragraph that stays the same.", "Clause: A for seven years. Token-0001."])
        mod = build_docx(
            [
                "Intro paragraph that stays the same.",
                "Clause: A for eight years. Revised-0001.",
                "Appended paragraph.",
            ]
        )
        final_clean, text_mod, _ = self._diff_apply_accept(orig, mod)
        assert final_clean.strip() == text_mod.strip()

    def test_adjacent_replacement_plus_append_applies(self):
        # Same shape as the QA 2,500-clause failure: the LAST clause is
        # modified AND content is appended after it, so the widened insertion
        # anchor overlaps the replacement's target.
        orig = build_docx([f"Clause {i}: A for seven years. Token-{i:04d}." for i in range(8)])
        mod_paras = [f"Clause {i}: A for seven years. Token-{i:04d}." for i in range(7)]
        mod_paras.append("Clause 7: A for eight years. Revised-0007.")
        mod_paras.append("Appended paragraph.")
        mod = build_docx(mod_paras)
        final_clean, text_mod, _ = self._diff_apply_accept(orig, mod)
        assert final_clean.strip() == text_mod.strip()


# ---------------------------------------------------------------------------
# F-02: formatting-only replacements must realize exactly the markers asked for
# ---------------------------------------------------------------------------


def _runs_with_text(doc):
    runs = []
    for p in doc.paragraphs:
        for r in p.runs:
            if r.text.strip():
                runs.append(r)
    return runs


class TestF02FormattingOnlyDiff:
    def _build_bold_doc(self):
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Normal before ")
        r = p.add_run("Important phrase")
        r.bold = True
        p.add_run(" normal after.")
        return doc

    def test_bold_to_italic_replacement_clears_bold(self):
        engine = RedlineEngine(doc_to_stream(self._build_bold_doc()), author="QA")
        engine.process_batch(
            [
                ModifyText(
                    type="modify",
                    target_text="**Important phrase**",
                    new_text="_Important phrase_",
                )
            ]
        )
        engine.accept_all_revisions(remove_comments=True)
        accepted = clean_text_of(engine)
        assert "_Important phrase_" in accepted
        assert "**" not in accepted, f"accepted text still bold: {accepted!r}"

        phrase_runs = [r for r in _runs_with_text(engine.doc) if "Important" in r.text]
        assert phrase_runs, "phrase run missing after accept"
        for r in phrase_runs:
            assert not r.bold, "explicit italic-only replacement kept inherited bold"
            assert r.italic

    def test_unmarked_replacement_still_inherits_bold(self):
        # The inheritance default is load-bearing: plain replacements inside
        # a styled span must keep the style (no visual data loss).
        engine = RedlineEngine(doc_to_stream(self._build_bold_doc()), author="QA")
        engine.process_batch([ModifyText(type="modify", target_text="Important", new_text="Critical")])
        engine.accept_all_revisions(remove_comments=True)
        accepted = clean_text_of(engine)
        assert "**Critical phrase**" in accepted

    def test_bold_removal_replacement_clears_bold(self):
        engine = RedlineEngine(doc_to_stream(self._build_bold_doc()), author="QA")
        engine.process_batch(
            [
                ModifyText(
                    type="modify",
                    target_text="**Important phrase**",
                    new_text="Important phrase",
                )
            ]
        )
        engine.accept_all_revisions(remove_comments=True)
        accepted = clean_text_of(engine)
        assert "Normal before Important phrase normal after." in accepted
        assert "**" not in accepted

    def test_diff_to_apply_formatting_round_trip(self):
        bold_doc = self._build_bold_doc()
        italic_doc = Document()
        p = italic_doc.add_paragraph()
        p.add_run("Normal before ")
        r = p.add_run("Important phrase")
        r.italic = True
        p.add_run(" normal after.")

        text_orig, struct_orig = extract_structured(bold_doc)
        text_mod, struct_mod = extract_structured(italic_doc)
        edits, _ = generate_structured_edits(text_orig, struct_orig, text_mod, struct_mod)
        replayed = json_round_trip(edits)

        engine = RedlineEngine(doc_to_stream(bold_doc), author="QA")
        engine.process_batch(replayed)
        engine.accept_all_revisions(remove_comments=True)
        assert clean_text_of(engine).strip() == text_mod.strip()


# ---------------------------------------------------------------------------
# F-03: styled-run boundary whitespace must stay outside emphasis markers
# ---------------------------------------------------------------------------


class TestF03BoundaryWhitespace:
    def _build_boundary_doc(self):
        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run("The Supplier ")
        r1.bold = True
        p.add_run("shall deliver")
        r3 = p.add_run(" the Services by 31 December 2026.")
        r3.italic = True
        return doc

    def test_extraction_emits_valid_markdown(self):
        text = _extract_text_from_doc(self._build_boundary_doc(), clean_view=False)
        assert text.strip() == "**The Supplier** shall deliver _the Services by 31 December 2026._"

    def test_mapper_matches_ingest_projection(self):
        # The Virtual Text contract: mapper.full_text and ingest output must
        # agree byte for byte, or edits resolve at wrong offsets.
        from adeu.redline.mapper import DocumentMapper

        doc = self._build_boundary_doc()
        ingest_text = _extract_text_from_doc(doc, clean_view=False)
        mapper = DocumentMapper(doc)
        assert mapper.full_text.strip() == ingest_text.strip()

    def test_plain_text_edit_across_boundary_applies(self):
        engine = RedlineEngine(doc_to_stream(self._build_boundary_doc()), author="QA")
        engine.process_batch(
            [
                ModifyText(
                    type="modify",
                    target_text="shall deliver",
                    new_text="must deliver",
                )
            ]
        )
        engine.accept_all_revisions(remove_comments=True)
        accepted = clean_text_of(engine)
        assert "must deliver" in accepted

    def test_whitespace_only_styled_run_is_not_wrapped(self):
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Alpha")
        ws = p.add_run("   ")
        ws.bold = True
        p.add_run("Omega")
        text = _extract_text_from_doc(doc, clean_view=False)
        assert text.strip() == "Alpha   Omega"


# ---------------------------------------------------------------------------
# F-04: images — media-byte differences and alt-text differences
# ---------------------------------------------------------------------------


class TestF04Images:
    def test_same_alt_different_bytes_warns(self, tmp_path, capsys):
        red = tmp_path / "red.docx"
        blue = tmp_path / "blue.docx"
        build_image_doc(MINIMAL_PNG_RED, "Same alt text", red)
        build_image_doc(MINIMAL_PNG_BLUE, "Same alt text", blue)

        code, out, err = run_cli(["diff", red, blue, "--json"], capsys)
        assert code == 0
        assert json.loads(out) == []
        assert "media" in err.lower(), f"no unsupported-media warning on stderr: {err!r}"

    def test_alt_text_diff_is_not_emitted_as_edit(self, tmp_path):
        red = tmp_path / "red.docx"
        blue = tmp_path / "blue.docx"
        build_image_doc(MINIMAL_PNG_RED, "RED logo", red)
        build_image_doc(MINIMAL_PNG_BLUE, "BLUE logo", blue)

        text_orig, struct_orig = extract_structured(Document(str(red)))
        text_mod, struct_mod = extract_structured(Document(str(blue)))
        edits, warnings = generate_structured_edits(text_orig, struct_orig, text_mod, struct_mod)

        assert edits == [], f"alt-text change must not become an edit: {edits}"
        assert any("alternative text" in w or "image" in w.lower() for w in warnings)

    def test_alt_text_diff_batch_is_applicable(self, tmp_path):
        # The workflow rule (F-14): whatever diff emits, apply must accept.
        red = tmp_path / "red.docx"
        blue = tmp_path / "blue.docx"
        build_image_doc(MINIMAL_PNG_RED, "RED logo", red)
        build_image_doc(MINIMAL_PNG_BLUE, "BLUE logo", blue)

        text_orig, struct_orig = extract_structured(Document(str(red)))
        text_mod, struct_mod = extract_structured(Document(str(blue)))
        edits, _ = generate_structured_edits(text_orig, struct_orig, text_mod, struct_mod)

        engine = RedlineEngine(BytesIO(red.read_bytes()), author="QA")
        engine.process_batch(json_round_trip(edits))  # must not raise


# ---------------------------------------------------------------------------
# F-05: post-apply verification of the text round trip
# ---------------------------------------------------------------------------


class TestF05PostApplyVerification:
    def _build_small_structured(self, path: Path):
        doc = Document()
        doc.add_heading("Small Doc", level=1)
        doc.add_paragraph("Some prose here.")
        t = doc.add_table(rows=1, cols=2)
        t.rows[0].cells[0].text = "K"
        t.rows[0].cells[1].text = "V"
        doc.save(path)

    def test_unfaithful_text_apply_fails_loudly(self, tmp_path, capsys):
        src = tmp_path / "small.docx"
        self._build_small_structured(src)
        tiny = tmp_path / "tiny.txt"
        tiny.write_text("tiny\n", encoding="utf-8")
        out = tmp_path / "result.docx"

        # --allow-major-deletions: since QA 2026-07-19 v8 F-12 the deletion
        # guard also arms below 2,000 chars, and it would refuse this apply
        # before the post-write verification (the behavior under test here)
        # ever runs.
        code, stdout, stderr = run_cli(["apply", src, tiny, "-o", out, "--json", "--allow-major-deletions"], capsys)
        assert code != 0, "apply reported success for an unfaithful replacement"
        stats = json.loads(stdout)
        assert stats.get("verified") is False
        assert (
            "verification" in (stats.get("verification_error") or "").lower()
            or "match" in (stats.get("verification_error") or "").lower()
        )

    def test_faithful_text_apply_verifies(self, tmp_path, capsys):
        src = tmp_path / "plain.docx"
        build_docx(["Alpha paragraph.", "Beta paragraph."], src)
        modified = tmp_path / "mod.txt"
        modified.write_text("Alpha paragraph.\n\nGamma paragraph.", encoding="utf-8")
        out = tmp_path / "result.docx"

        code, stdout, stderr = run_cli(["apply", src, modified, "-o", out, "--json"], capsys)
        assert code == 0, f"faithful apply failed: {stderr}"
        stats = json.loads(stdout)
        assert stats.get("verified") is True

    def test_major_deletion_floor_is_documented(self, capsys):
        code, _, err = run_cli(["apply", "--help"], capsys)
        # argparse --help exits 0 and prints to stdout
        code2, out, _ = run_cli(["apply", "--help"], capsys)
        help_text = out
        assert "2,000" in help_text or "2000" in help_text, (
            "the major-deletion guard's absolute floor must be disclosed in --help"
        )


# ---------------------------------------------------------------------------
# F-06: regex + match_mode=all must not rebuild the map per occurrence
# ---------------------------------------------------------------------------


class TestF06AllMatchScaling:
    def test_map_rebuilds_do_not_scale_with_occurrences(self):
        from adeu.redline import mapper as mapper_mod

        n = 30
        doc = build_docx(
            [f"Clause {i}: The party PLACEHOLDER-{i:04d} shall comply. REPLACEME token." for i in range(n)]
        )
        engine = RedlineEngine(doc_to_stream(doc), author="QA")

        counts = {"build": 0}
        orig_build = mapper_mod.DocumentMapper._build_map

        def counting_build(self):
            counts["build"] += 1
            return orig_build(self)

        mapper_mod.DocumentMapper._build_map = counting_build
        try:
            stats = engine.process_batch(
                [
                    ModifyText(
                        type="modify",
                        target_text="REPLACEME",
                        new_text="SWAPPED",
                        regex=True,
                        match_mode="all",
                    )
                ]
            )
        finally:
            mapper_mod.DocumentMapper._build_map = orig_build

        assert stats["edits_skipped"] == 0
        # A constant number of rebuilds (initial mapping, per-batch refresh,
        # report views) is fine; one rebuild PER OCCURRENCE is the QA finding.
        assert counts["build"] < n, f"{counts['build']} map rebuilds for {n} occurrences"

        engine.accept_all_revisions(remove_comments=True)
        final = clean_text_of(engine)
        assert final.count("SWAPPED") == n
        assert "REPLACEME" not in final


# ---------------------------------------------------------------------------
# F-07 / F-08: contradictory sanitize options must be rejected
# ---------------------------------------------------------------------------


class TestF07F08SanitizeOptionConflicts:
    def test_keep_markup_with_accept_all_is_rejected(self, tmp_path, capsys):
        src = tmp_path / "doc.docx"
        build_commented_tracked_docx(src)
        code, _, err = run_cli(
            ["sanitize", src, "--keep-markup", "--accept-all", "-o", tmp_path / "out.docx"],
            capsys,
        )
        assert code == 2
        assert "--keep-markup" in err and "--accept-all" in err

    def test_output_with_outdir_is_rejected(self, tmp_path, capsys):
        src = tmp_path / "doc.docx"
        build_commented_tracked_docx(src)
        code, _, err = run_cli(
            [
                "sanitize",
                src,
                "--accept-all",
                "-o",
                tmp_path / "explicit.docx",
                "--outdir",
                tmp_path / "outdir",
            ],
            capsys,
        )
        assert code == 2
        assert "--outdir" in err and ("-o" in err or "--output" in err)
        assert not (tmp_path / "outdir" / "doc.docx").exists()


# ---------------------------------------------------------------------------
# F-09: comment timestamps must be normalized when comments are retained
# ---------------------------------------------------------------------------


class TestF09CommentTimestamps:
    def test_keep_markup_normalizes_comment_dates(self, tmp_path, capsys):
        src = tmp_path / "doc.docx"
        build_commented_tracked_docx(src)
        out = tmp_path / "out.docx"
        code, _, err = run_cli(["sanitize", src, "--keep-markup", "-o", out], capsys)
        assert code == 0

        with zipfile.ZipFile(out) as z:
            comment_dates = []
            for name in z.namelist():
                if "comments" in name.lower() and name.endswith(".xml"):
                    data = z.read(name).decode("utf-8")
                    comment_dates.extend(re.findall(r'(?:w:date|w16cex:dateUtc)="([^"]+)"', data))
        assert comment_dates, "fixture must contain comment timestamps"
        for d in comment_dates:
            assert d.startswith("2025-01-01"), f"comment timestamp not normalized: {d}"


# ---------------------------------------------------------------------------
# F-12: match_mode must be validated as an enum
# ---------------------------------------------------------------------------


class TestF12MatchModeValidation:
    def test_unknown_match_mode_is_rejected(self):
        with pytest.raises(ValidationError, match="match_mode"):
            TypeAdapter(BatchChanges).validate_python(
                [{"type": "modify", "target_text": "a", "new_text": "b", "match_mode": "banana"}]
            )

    def test_null_match_mode_is_rejected(self):
        with pytest.raises(ValidationError, match="match_mode"):
            TypeAdapter(BatchChanges).validate_python(
                [{"type": "modify", "target_text": "a", "new_text": "b", "match_mode": None}]
            )

    def test_recognized_synonyms_still_normalize(self):
        changes = TypeAdapter(BatchChanges).validate_python(
            [{"type": "modify", "target_text": "a", "new_text": "b", "match_mode": "first_only"}]
        )
        assert changes[0].match_mode == "first"

    def test_cli_reports_invalid_match_mode(self, tmp_path, capsys):
        src = tmp_path / "doc.docx"
        build_docx(["Some text here."], src)
        edits = tmp_path / "edits.json"
        edits.write_text(
            json.dumps([{"type": "modify", "target_text": "text", "new_text": "words", "match_mode": "banana"}]),
            encoding="utf-8",
        )
        code, out, err = run_cli(["apply", src, edits, "--dry-run", "--json"], capsys)
        assert code != 0
        payload = json.loads(out)
        assert payload["error"] == "invalid_changes_file"
        assert "match_mode" in payload["message"]


# ---------------------------------------------------------------------------
# F-13: invalid apply regex must be diagnosed as a regex error
# ---------------------------------------------------------------------------


class TestF13InvalidRegexDiagnosis:
    def test_invalid_regex_names_the_pattern_problem(self):
        doc = build_docx(["alpha DUPTOKEN one.", "beta DUPTOKEN two."])
        engine = RedlineEngine(doc_to_stream(doc), author="QA")
        with pytest.raises(BatchValidationError) as exc_info:
            engine.process_batch([ModifyText(type="modify", target_text="[", new_text="x", regex=True)])
        msg = "\n".join(exc_info.value.errors)
        assert "regular expression" in msg
        assert "not found" not in msg

    def test_valid_regex_still_applies(self):
        doc = build_docx(["alpha DUPTOKEN one."])
        engine = RedlineEngine(doc_to_stream(doc), author="QA")
        stats = engine.process_batch([ModifyText(type="modify", target_text=r"DUP\w+", new_text="XTOKEN", regex=True)])
        assert stats["edits_applied"] == 1


# ---------------------------------------------------------------------------
# F-16: adeu init must pin the package version
# ---------------------------------------------------------------------------


class TestF16InitPinsVersion:
    def test_generated_uvx_command_is_pinned(self, tmp_path, monkeypatch, capsys):
        import adeu

        # Patch the path getter itself: env-var tricks (HOME) don't reach the
        # Windows %APPDATA% branch, and init must never touch the real config.
        config_path = tmp_path / "Claude" / "claude_desktop_config.json"
        monkeypatch.setattr("adeu.cli._get_claude_config_path", lambda: config_path)
        # shutil.which("uvx") must succeed regardless of test environment.
        monkeypatch.setattr("shutil.which", lambda name: f"/fake/bin/{name}")

        code, _, err = run_cli(["init"], capsys)
        assert code == 0

        data = json.loads(config_path.read_text(encoding="utf-8"))
        args = data["mcpServers"]["adeu"]["args"]
        assert f"adeu=={adeu.__version__}" in args, f"unpinned uvx args: {args}"


# ---------------------------------------------------------------------------
# F-17: search heading path must include the heading containing the match
# ---------------------------------------------------------------------------


class TestF17SearchHeadingPath:
    def test_match_inside_heading_reports_full_heading(self):
        from adeu.mcp_components._response_builders import build_search_response

        text = "# Master Services Agreement\n\nThis agreement covers the Services provided."
        res = build_search_response(text, "Services", False, True, None, "doc.docx", is_cli=True)
        markdown = res.structured_content["markdown"]

        paths = re.findall(r"\*\*Path:\*\* `([^`]+)`", markdown)
        assert paths, f"no heading paths rendered: {markdown}"
        assert paths[0] == "Master Services Agreement", f"truncated path: {paths[0]!r}"


# ---------------------------------------------------------------------------
# F-18: outline --page all must not produce contradictory messages
# ---------------------------------------------------------------------------


class TestF18OutlinePageAll:
    def test_outline_page_all_warns_and_succeeds(self, tmp_path, capsys):
        src = tmp_path / "doc.docx"
        doc = Document()
        doc.add_heading("Master Services Agreement", level=1)
        doc.add_paragraph("This agreement covers the Services provided.")
        doc.save(src)

        code, out, err = run_cli(["extract", src, "--mode", "outline", "--page", "all"], capsys)
        assert code == 0, f"outline --page all must succeed after warning: {err}"
        assert "ignored" in err
        assert "Invalid --page value" not in err
        assert "Master Services Agreement" in out


# ---------------------------------------------------------------------------
# F-19: an invalid baseline error must name the baseline file
# ---------------------------------------------------------------------------


class TestF19BaselineErrorNamesBaseline:
    def test_invalid_baseline_names_baseline_file(self, tmp_path, capsys):
        src = tmp_path / "doc.docx"
        build_commented_tracked_docx(src)
        bogus = tmp_path / "bogus_baseline.docx"
        bogus.write_text("not a docx", encoding="utf-8")

        code, _, err = run_cli(
            ["sanitize", src, "--accept-all", "--baseline", bogus, "-o", tmp_path / "out.docx"],
            capsys,
        )
        assert code != 0
        assert "bogus_baseline.docx" in err, f"error does not name the baseline: {err!r}"
        assert f"'{src.name}' is not a valid DOCX" not in err


# ---------------------------------------------------------------------------
# F-21: edits_applied must count edit objects, not occurrences
# ---------------------------------------------------------------------------


class TestF21EditsAppliedSemantics:
    def test_match_all_reports_one_edit_two_occurrences(self):
        doc = build_docx(["alpha DUPTOKEN one.", "beta DUPTOKEN two."])
        engine = RedlineEngine(doc_to_stream(doc), author="QA")
        stats = engine.process_batch(
            [
                ModifyText(
                    type="modify",
                    target_text="DUPTOKEN",
                    new_text="NEWTOKEN",
                    match_mode="all",
                )
            ]
        )
        assert len(stats["edits"]) == 1
        assert stats["edits_applied"] == 1, f"edits_applied must count change objects (got {stats['edits_applied']})"
        assert stats["occurrences_modified"] == 2
        assert stats["edits"][0]["occurrences_modified"] == 2

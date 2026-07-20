import io
import re

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from adeu.models import ModifyText, ReplyComment
from adeu.redline.engine import RedlineEngine


def test_repro_invisible_reply_xml_structure():
    """
    Simulates the workflow of replying to an existing comment in a document
    that already has a comments part, and inspects the resulting XML.
    """
    # 1. Setup: Create a doc with an initial comment
    doc = Document()
    doc.add_paragraph("Paragraph with comment.")

    stream_initial = io.BytesIO()
    doc.save(stream_initial)
    stream_initial.seek(0)

    # Add initial comment via Engine
    # CHANGE: Make an actual text modification so the comment attaches
    engine1 = RedlineEngine(stream_initial, author="OrigAuthor")
    engine1.apply_edits([ModifyText(target_text="Paragraph", new_text="ParagraphModified", comment="Base Comment")])
    stream_with_comment = engine1.save_to_stream()

    # 2. Workflow: Load existing doc and Reply
    engine2 = RedlineEngine(stream_with_comment, author="ReplyAuthor")

    # Get ID of the existing comment
    comments_data = engine2.comments_manager.extract_comments_data()
    assert len(comments_data) == 1, "Setup failed: Base comment was not created."
    base_id = list(comments_data.keys())[0]

    # Perform Reply
    action = ReplyComment(target_id=f"Com:{base_id}", text="The Reply")
    applied, skipped, _already = engine2.apply_review_actions([action])
    assert applied == 1

    stream_final = engine2.save_to_stream()

    # 3. Inspect XML Structure
    doc_final = Document(stream_final)

    # Get Comments Part
    comments_part = None
    for rel in doc_final.part.rels.values():
        if rel.reltype == RT.COMMENTS:
            comments_part = rel.target_part
            break

    assert comments_part is not None
    xml_str = comments_part.blob.decode("utf-8")

    print("\n--- GENERATED COMMENTS XML ---")
    print(xml_str)
    print("------------------------------")

    # --- CHECKS ---

    # 1. Namespace Declaration
    # Word strongly prefers namespaces declared at the root <w:comments> element.
    root_ns_check = 'xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"'
    if root_ns_check not in xml_str:
        print("POTENTIAL ISSUE: w15 namespace NOT explicitly in root element attrs.")
        # Note: lxml might serialize it differently, but string search helps debug.
    else:
        print("OK: w15 namespace found in root.")

    # 2. Date Format
    dates = re.findall(r'w:date="([^"]+)"', xml_str)
    print(f"Dates found: {dates}")
    for d in dates:
        if "." in d:
            print(f"POTENTIAL ISSUE: Date contains microseconds: {d}")
        if not d.endswith("Z"):
            print(f"POTENTIAL ISSUE: Date might lack UTC 'Z' marker: {d}")

    # 3. Parent Linking
    if f'w15:p="{base_id}"' in xml_str:
        print(f"OK: Reply linked to parent {base_id}")
    else:
        print(f"FAIL: No w15:p link to {base_id} found.")

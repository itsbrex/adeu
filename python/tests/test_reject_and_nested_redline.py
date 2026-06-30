"""Canonical track-change semantics for editing inside another author's
pending insertion, plus reject_all_revisions round-trips.

Decision (cross-engine, Python canonical): when an author edits or deletes text
that lies inside another author's still-pending <w:ins>, the deletion is NESTED
inside that <w:ins> (<w:ins author="A"><w:del author="B">…</w:del></w:ins>) and
the replacement insertion is emitted as a SIBLING after the enclosing <w:ins>
(the enclosing insertion is split if the edit is mid-insertion). This is the
representation Word itself authors; it preserves A's authorship and makes
reject-all revert the contingent text to nothing rather than promoting an
unaccepted proposal to committed body text. <w:ins> is never nested in <w:ins>.
"""
import io

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _foreign_ins_doc():
    """'The party shall provide ' + <ins Supplier's Counsel>written notice</ins> + ' within 30 days.'"""
    doc = Document()
    p = doc.add_paragraph("The party shall provide ")
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "201")
    ins.set(qn("w:author"), "Supplier's Counsel")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "written notice"
    r.append(t)
    ins.append(r)
    p._element.append(ins)
    r2 = OxmlElement("w:r")
    t2 = OxmlElement("w:t")
    t2.text = " within 30 days."
    r2.append(t2)
    p._element.append(r2)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _clean_text(stream):
    d = Document(stream)
    return " ".join(
        "".join(t.text or "" for t in p._element.findall(f".//{{{W}}}t")) for p in d.paragraphs
    )


def test_modify_inside_foreign_ins_produces_canonical_c3_structure():
    eng = RedlineEngine(_foreign_ins_doc(), author="Reviewer AI")
    res = eng.process_batch([ModifyText(target_text="written notice", new_text="email notification")])
    assert res["edits_applied"] == 1 and res["edits_skipped"] == 0

    doc = Document(eng.save_to_stream())
    root = doc.element

    # The deletion is nested inside Supplier's Counsel's insertion.
    nested_del = root.xpath("//w:ins[@w:author=\"Supplier's Counsel\"]//w:del[@w:author='Reviewer AI']")
    assert nested_del, "deletion must be nested inside the foreign author's <w:ins>"
    assert "written notice" in "".join(
        t.text or "" for t in nested_del[0].findall(f".//{{{W}}}delText")
    )

    # The replacement insertion is a sibling authored by Reviewer AI, NOT nested
    # inside Supplier's insertion.
    rev_ins = root.xpath("//w:ins[@w:author='Reviewer AI']")
    assert rev_ins, "Reviewer AI's replacement insertion must exist"
    assert any("email notification" in "".join(t.text or "" for t in i.findall(f".//{{{W}}}t")) for i in rev_ins)

    # Invariant: never <w:ins> directly inside <w:ins>.
    assert not root.xpath("//w:ins//w:ins"), "<w:ins> must never be nested inside <w:ins>"


def test_round_trip_accept_and_reject_for_foreign_ins_modify():
    eng = RedlineEngine(_foreign_ins_doc(), author="Reviewer AI")
    eng.process_batch([ModifyText(target_text="written notice", new_text="email notification")])
    tracked = eng.save_to_stream()

    eng_a = RedlineEngine(tracked, author="Reviewer AI")
    eng_a.accept_all_revisions()
    assert "The party shall provide email notification within 30 days." in _clean_text(eng_a.save_to_stream())

    # Reject-all must revert to the TRUE baseline: 'written notice' was only a
    # pending proposal, so it disappears — it must NOT be promoted to committed
    # text.
    eng_r = RedlineEngine(eng.save_to_stream(), author="Reviewer AI")
    eng_r.reject_all_revisions()
    rejected = _clean_text(eng_r.save_to_stream())
    assert "written notice" not in rejected
    assert "email notification" not in rejected
    assert "The party shall provide  within 30 days." in rejected


def test_delete_inside_foreign_ins_reject_reverts_pending_insertion():
    eng = RedlineEngine(_foreign_ins_doc(), author="Reviewer AI")
    eng.process_batch([ModifyText(target_text="written notice", new_text="")])

    eng_a = RedlineEngine(eng.save_to_stream(), author="Reviewer AI")
    eng_a.accept_all_revisions()
    assert "written notice" not in _clean_text(eng_a.save_to_stream())

    eng_r = RedlineEngine(eng.save_to_stream(), author="Reviewer AI")
    eng_r.reject_all_revisions()
    assert "written notice" not in _clean_text(eng_r.save_to_stream())


def test_mid_insertion_edit_positions_replacement_and_round_trips():
    # Author A inserts a clause containing '60' in the middle.
    doc = Document()
    doc.add_paragraph("Base text.")
    b = io.BytesIO()
    doc.save(b)
    b.seek(0)
    e1 = RedlineEngine(b, author="Author A")
    e1.apply_edits([ModifyText(target_text="Base text.", new_text="Base text. The notice period is 60 days.")])
    mid = e1.save_to_stream()

    # Author B revises '60' -> '45' inside A's pending insertion.
    e2 = RedlineEngine(mid, author="Author B")
    applied, _ = e2.apply_edits([ModifyText(target_text="60", new_text="45")])
    assert applied == 1
    out = Document(e2.save_to_stream()).element
    assert not out.xpath("//w:ins//w:ins"), "no <w:ins> in <w:ins>"
    assert out.xpath("//w:ins//w:del"), "the '60' deletion must be nested in A's insertion"

    eng_a = RedlineEngine(e2.save_to_stream())
    eng_a.accept_all_revisions()
    assert "Base text. The notice period is 45 days." in _clean_text(eng_a.save_to_stream())

    eng_r = RedlineEngine(e2.save_to_stream())
    eng_r.reject_all_revisions()
    # A's whole insertion was pending; rejecting everything reverts to 'Base text.'
    assert _clean_text(eng_r.save_to_stream()).strip() == "Base text."


def test_insertion_inside_foreign_ins_splits_no_ins_in_ins():
    """A modify that trims to a pure INSERTION inside a foreign author's pending
    <w:ins> (new_text == target_text + suffix) must split the enclosing <w:ins>
    so the new insertion is a sibling — never <w:ins> nested in <w:ins>."""
    eng = RedlineEngine(_foreign_ins_doc(), author="Reviewer AI")
    res = eng.process_batch([ModifyText(target_text="written notice", new_text="written notice please")])
    assert res["edits_applied"] == 1
    doc = Document(eng.save_to_stream())
    assert not doc.element.xpath("//w:ins//w:ins"), "<w:ins> must never be nested inside <w:ins>"

    eng_a = RedlineEngine(eng.save_to_stream(), author="Reviewer AI")
    eng_a.accept_all_revisions()
    assert "The party shall provide written notice please within 30 days." in _clean_text(eng_a.save_to_stream())

    eng_r = RedlineEngine(eng.save_to_stream(), author="Reviewer AI")
    eng_r.reject_all_revisions()
    # Both A's pending "written notice" and B's appended " please" revert.
    assert "The party shall provide  within 30 days." in _clean_text(eng_r.save_to_stream())


def test_reject_all_restores_original_for_own_edits():
    doc = Document()
    doc.add_paragraph("The cat sat.")
    b = io.BytesIO()
    doc.save(b)
    b.seek(0)
    eng = RedlineEngine(b, author="Z")
    eng.process_batch([ModifyText(target_text="cat", new_text="dog")])

    eng_a = RedlineEngine(eng.save_to_stream())
    eng_a.accept_all_revisions()
    assert "The dog sat." in _clean_text(eng_a.save_to_stream())

    eng_r = RedlineEngine(eng.save_to_stream())
    eng_r.reject_all_revisions()
    assert "The cat sat." in _clean_text(eng_r.save_to_stream())

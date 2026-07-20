import io
import sys

import pytest
from docx import Document

try:
    from hypothesis import HealthCheck
    from hypothesis import settings as _hyp_settings

    # Property-test profiles (tests/test_property_invariants.py). Registered
    # here so `--hypothesis-profile=hunt` resolves at pytest configure time.
    _hyp_settings.register_profile(
        "default", deadline=None, max_examples=25, suppress_health_check=[HealthCheck.too_slow]
    )
    _hyp_settings.register_profile(
        "hunt", deadline=None, max_examples=300, suppress_health_check=[HealthCheck.too_slow]
    )
    _hyp_settings.load_profile("default")
except ImportError:
    pass


@pytest.fixture(scope="session", autouse=True)
def _isolate_windows_appdata(tmp_path_factory):
    """On Windows, `adeu init` resolves the Claude Desktop config via %APPDATA%.
    A test that runs init without patching _get_claude_config_path would rewrite
    the developer's real claude_desktop_config.json (this happened 2026-07-20:
    two QA-repro tests injected fake uvx entries into a live config). Pointing
    APPDATA at a throwaway directory for the whole session makes that class of
    accident impossible; tests that assert on the config still patch the path
    getter explicitly."""
    if sys.platform != "win32":
        yield
        return
    mp = pytest.MonkeyPatch()
    mp.setenv("APPDATA", str(tmp_path_factory.mktemp("appdata")))
    yield
    mp.undo()


@pytest.fixture
def simple_docx_stream():
    """Returns a BytesIO stream containing a simple DOCX."""
    doc = Document()
    doc.add_heading("Contract Agreement", 0)
    doc.add_paragraph("This is a simple contract.")
    doc.add_paragraph("The party of the first part shall be known as the Seller.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


# Only define COM fixtures on Windows
if sys.platform == "win32":
    import pythoncom
    import win32com.client

    @pytest.fixture
    def active_word_app():
        """
        Creates an ephemeral, visible MS Word instance with a fresh document.
        Ensures it is torn down properly after the test.
        """
        pythoncom.CoInitialize()

        app = None
        try:
            # Dispatch starts a new background instance if one doesn't exist.
            # GetActiveObject will then be able to hook into it in the tool.
            app = win32com.client.Dispatch("Word.Application")
            app.Visible = True  # Needs to be visible/active for GetActiveObject sometimes
            doc = app.Documents.Add()

            # Bring to front so GetActiveObject definitely binds to this instance
            app.Activate()

            # Seed initial content
            doc.Range(0, 0).Text = "Hello world! This is a live testing document.\n"

            yield app, doc

        except Exception as e:
            pytest.skip(f"Could not initialize Word COM for testing: {e}")

        finally:
            if app:
                try:
                    doc.Close(0)  # 0 = wdDoNotSaveChanges
                except Exception:
                    pass
                # We intentionally omit app.Quit() and pythoncom.CoUninitialize()
                # to avoid Windows Access Violations (0x800706be) when Pytest holds COM locals.

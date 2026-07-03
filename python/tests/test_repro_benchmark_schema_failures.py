# FILE: tests/test_repro_benchmark_schema_failures.py
"""
Regression tests for the P2 schema-robustness fixes (A1: infer missing `type`,
A2: coerce/sanitize `match_mode`, A3: per-change validation with partial accept).

The payloads below are the VERBATIM `args the model passed` from the
adeu-python-published failures in the agentic-loop benchmark (loop_analysis.md
§3, error clusters). Every one of these was a HARD `process_document_batch`
rejection in the published engine. They must now validate (or, for the
deliberately-unsalvageable case, be reported per-index without sinking the batch).

These exercise adeu.models.coerce_stringified_changes (A1/A2) and
adeu.mcp_components.tools.document._normalize_changes (A3).
"""

import asyncio

import pytest
from docx import Document

from adeu.mcp_components.tools.document import (
    _normalize_changes,
    process_document_batch,
)
from adeu.models import (
    AcceptChange,
    ModifyText,
    ReplyComment,
    coerce_stringified_changes,
)


class MockContext:
    """Absorbs the async logging calls process_document_batch makes."""

    async def info(self, msg, **kwargs):
        pass

    async def debug(self, msg, **kwargs):
        pass

    async def warning(self, msg, **kwargs):
        pass

    async def error(self, msg, **kwargs):
        pass


@pytest.fixture
def sample_docx(tmp_path) -> str:
    doc = Document()
    # Include the literal placeholders the benchmark payloads target so the
    # full end-to-end path (validate -> apply) can actually match something.
    doc.add_paragraph("Company: [Company Name]")
    doc.add_paragraph("Investor: [Investor Name]")
    doc.add_paragraph("The laws of [fill in state, province, and/or country]")
    path = tmp_path / "sample.docx"
    doc.save(path)
    return str(path)


# ---------------------------------------------------------------------------
# A1 — missing `type` discriminator is inferred as `modify`
#
# These reproduce the four union_tag_not_found clusters (form-fill ×1 line 106,
# policy-checklist-review ×3 lines 121/137/145). Each sub-edit omitted `type`.
# ---------------------------------------------------------------------------


def test_a1_form_fill_missing_type_multi_edit():
    """loop_analysis.md line 106: 10 validation errors, every sub-edit lacked `type`."""
    changes = [
        {
            "new_text": "Acme Robotics, Inc.",
            "target_text": "[Company Name]",
            "match_mode": "all",
        },
        {
            "new_text": "Vertex Seed Fund, L.P.",
            "target_text": "[Investor Name]",
            "match_mode": "strict",
        },
        {"match_mode": "strict", "new_text": "June 22, 2026", "target_text": "[Date]"},
    ]
    valid, rejected = _normalize_changes(changes)
    assert rejected == [], f"expected no rejections, got: {rejected}"
    assert len(valid) == 3
    assert all(isinstance(c, ModifyText) for c in valid)
    # Re-assert per-element so the type checker narrows the union to ModifyText.
    m0, m1, m2 = valid
    assert isinstance(m0, ModifyText) and isinstance(m1, ModifyText) and isinstance(m2, ModifyText)
    assert m0.match_mode == "all"
    assert m1.match_mode == "strict"
    # third had no match_mode -> Pydantic default
    assert m2.match_mode == "strict"


def test_a1_policy_missing_type_single_edit():
    """loop_analysis.md lines 121/137/145: governing-law edit with no `type`."""
    target = "The laws of [fill in state, province, and/or country]"
    changes = [
        {
            "target_text": target,
            "new_text": target,
            "match_mode": "strict",
            "comment": "Checklist Point 1: Governing Law is a placeholder.",
        }
    ]
    valid, rejected = _normalize_changes(changes)
    assert rejected == []
    assert len(valid) == 1
    assert isinstance(valid[0], ModifyText)
    assert valid[0].type == "modify"


# ---------------------------------------------------------------------------
# A2 — `match_mode` help-string echo is dropped to the `strict` default
#
# Reproduces the literal_error cluster (policy-checklist-review ×1 line 129):
# the model passed the field's help text as the value.
# ---------------------------------------------------------------------------


def test_a2_match_mode_help_string_echo_defaults_to_strict():
    """loop_analysis.md line 129: match_mode='strict, first, or all'."""
    changes = [
        {
            "type": "modify",
            "target_text": "The laws of [fill in state, province, and/or country]",
            "new_text": "The laws of the State of Delaware",
            "match_mode": "strict, first, or all",  # the help string, not a value
            "regex": False,
        }
    ]
    valid, rejected = _normalize_changes(changes)
    assert rejected == []
    assert len(valid) == 1
    m = valid[0]
    assert isinstance(m, ModifyText)
    # Unrecognized value dropped -> fail-safe default, NOT 'all'.
    assert m.match_mode == "strict"


@pytest.mark.parametrize(
    "raw,expected",
    [
        ("strict", "strict"),
        ("first", "first"),
        ("all", "all"),
        ("first_only", "first"),
        ("all_occurrences", "all"),
        ("every", "all"),
        ("strict, first, or all", "strict"),  # help-string echo -> default
        ("", "strict"),  # empty -> default
        ("garbage", "strict"),  # unknown -> default
    ],
)
def test_a2_match_mode_synonyms(raw, expected):
    out = coerce_stringified_changes([{"type": "modify", "target_text": "x", "new_text": "y", "match_mode": raw}])
    valid, rejected = _normalize_changes(out)
    assert rejected == []
    m = valid[0]
    assert isinstance(m, ModifyText)
    assert m.match_mode == expected


# ---------------------------------------------------------------------------
# A1 boundary — review actions are NOT guessed
#
# `target_id` alone is ambiguous between accept and reject. We must NOT infer
# (silently picking one could apply the opposite of intent). It stays rejected.
# ---------------------------------------------------------------------------


def test_a1_target_id_alone_is_not_inferred():
    valid, rejected = _normalize_changes([{"target_id": "Chg:12"}])
    assert valid == []
    assert len(rejected) == 1
    assert "changes[0]" in rejected[0]


def test_a1_infers_reply_and_insert_row_unambiguously():
    # reply: text + target_id
    valid, rejected = _normalize_changes([{"target_id": "Com:5", "text": "Acknowledged."}])
    assert rejected == []
    assert isinstance(valid[0], ReplyComment)
    # insert_row: cells present
    valid2, rejected2 = _normalize_changes([{"target_text": "anchor", "position": "below", "cells": ["a", "b"]}])
    assert rejected2 == []
    assert valid2[0].type == "insert_row"


# ---------------------------------------------------------------------------
# A3 — per-change validation: one bad sibling does not sink the batch
# ---------------------------------------------------------------------------


def test_a3_partial_accept_one_bad_sibling():
    """
    A mixed batch: 3 salvageable (2 inferred-type modifies + 1 self-healed
    match_mode) and 1 unsalvageable (target_id alone). The valid ones must
    survive; the bad one reported at its index.
    """
    changes = [
        {"target_text": "X", "new_text": "Y", "match_mode": "all"},  # infer modify
        {"target_id": "Com:5", "text": "ok"},  # infer reply
        {"target_id": "Chg:12"},  # ambiguous -> reject
        {
            "target_text": "A",
            "new_text": "B",
            "match_mode": "strict, first, or all",
        },  # heal -> strict
    ]
    valid, rejected = _normalize_changes(changes)
    assert len(valid) == 3
    assert len(rejected) == 1
    assert "changes[2]" in rejected[0]
    # The self-healed one kept its other intent and defaulted match_mode.
    healed = [c for c in valid if isinstance(c, ModifyText) and c.target_text == "A"][0]
    assert healed.match_mode == "strict"


def test_a3_all_bad_returns_empty_without_raising():
    valid, rejected = _normalize_changes([{"target_id": "Chg:12"}, {"foo": "bar"}])
    assert valid == []
    assert len(rejected) == 2
    assert "changes[0]" in rejected[0]
    assert "changes[1]" in rejected[1]


def test_a3_already_validated_instances_pass_through_untouched():
    """
    Fast path: a list of already-built model instances must skip re-validation
    (preserving engine PrivateAttrs) and report no rejections.
    """
    pre = [
        ModifyText(type="modify", target_text="x", new_text="y"),
        AcceptChange(type="accept", target_id="Chg:1"),
    ]
    valid, rejected = _normalize_changes(pre)
    assert rejected == []
    assert valid is pre  # identity: untouched fast-path return


# ---------------------------------------------------------------------------
# End-to-end — the form-fill payload now applies through the real tool
# ---------------------------------------------------------------------------


def test_end_to_end_form_fill_payload_applies(sample_docx, tmp_path):
    """
    The full benchmark form-fill payload (missing `type` on every sub-edit)
    must now flow through process_document_batch and apply, not 422.
    """
    ctx = MockContext()
    output_path = tmp_path / "out.docx"
    changes = [
        {
            "new_text": "Acme Robotics, Inc.",
            "target_text": "[Company Name]",
            "match_mode": "all",
        },
        {
            "new_text": "Vertex Seed Fund, L.P.",
            "target_text": "[Investor Name]",
            "match_mode": "strict",
        },
    ]
    result = asyncio.run(
        process_document_batch(
            reasoning="test",
            original_docx_path=sample_docx,
            author_name="Contract Editor",
            ctx=ctx,  # type: ignore[arg-type]
            changes=changes,  # type: ignore[arg-type]
            output_path=str(output_path),
        )
    )
    assert "Batch complete" in result
    assert "Edits: 2 applied" in result
    assert output_path.exists()


def test_end_to_end_partial_reports_skips_in_response(sample_docx, tmp_path):
    """
    Mixed batch through the real tool: the valid edit applies, the unsalvageable
    one is surfaced as a skipped/validation note in the response text (so the
    model can fix just that index next turn).
    """
    ctx = MockContext()
    output_path = tmp_path / "out_partial.docx"
    changes = [
        {
            "new_text": "Acme Robotics, Inc.",
            "target_text": "[Company Name]",
            "match_mode": "all",
        },  # ok
        {"target_id": "Chg:99"},  # ambiguous -> rejected pre-engine
    ]
    result = asyncio.run(
        process_document_batch(
            reasoning="test",
            original_docx_path=sample_docx,
            author_name="Contract Editor",
            ctx=ctx,  # type: ignore[arg-type]
            changes=changes,  # type: ignore[arg-type]
            output_path=str(output_path),
        )
    )
    assert "Edits: 1 applied" in result
    assert "changes[1]" in result
    assert "failed validation" in result.lower()
    assert output_path.exists()

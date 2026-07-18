# FILE: tests/test_repro_qa_report_v4.py
"""
Regression tests for the 2026-07-17 CLI QA & UX report.

Finding index (severity as reported):
  C1  diff produced edits apply rejects; suggested match_mode remedies corrupt
  C2  oversized new_text/target_text echoed unbounded into reports
  H1  previews garbled / leak internal scaffolding
  H2  plain targets fail across bold/italic run boundaries
  H3  unhandled traceback on filesystem write errors
  M1  dry-run and real apply disagree on validation results
  M2  overlapping edits within a batch are not detected
  M3  insert_row silently drops extra cells
  M4  insert_row report never shows the inserted content
  M5  malformed edit JSON produces a raw pydantic dump
  L1  --page silently accepts invalid values
  L2  outline label malformed for --outline-max-level <= 0
  L3  markup passes CriticMarkup syntax in new_text through unescaped
  L4  empty edit batch silently writes a no-op copy
"""

import io
import json
import sys
from unittest.mock import patch

import pytest
from docx import Document

from adeu.diff import generate_edits_from_text, make_edits_self_contained
from adeu.ingest import extract_text_from_stream
from adeu.models import InsertTableRow, ModifyText
from adeu.redline.engine import BatchValidationError, RedlineEngine


def _stream(doc) -> io.BytesIO:
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def _make_nda(term="two (2) years") -> io.BytesIO:
    """The report's NDA shape: bare digits repeat in a date, a section number
    and the term clause, so a decomposed '2'->'5' edit is ambiguous."""
    doc = Document()
    doc.add_heading("Mutual Non-Disclosure Agreement", level=1)
    doc.add_paragraph(
        "This Agreement is entered into as of January 1, 2025, by and between "
        "Acme Corp and Beta LLC (each a party and collectively the parties)."
    )
    doc.add_heading("1. Definitions", level=1)
    doc.add_paragraph("Confidential Information means any information disclosed by either party.")
    doc.add_heading("2. Obligations", level=1)
    doc.add_paragraph("The Recipient shall protect the Confidential Information.")
    doc.add_heading("3. Governing Law", level=1)
    doc.add_paragraph("This Agreement shall be governed by the laws of the State of California.")
    doc.add_heading("4. Non-Solicitation", level=1)
    doc.add_paragraph("Neither party shall solicit employees of the other party.")
    doc.add_heading("5. Term", level=1)
    doc.add_paragraph(f"This Agreement shall remain in effect for a period of {term}.")
    return _stream(doc)


def _clean_text(stream: io.BytesIO) -> str:
    stream.seek(0)
    return extract_text_from_stream(io.BytesIO(stream.read()), filename="x.docx", clean_view=True)


# ---------------------------------------------------------------------------
# C1 — diff output must apply cleanly against the document it was diffed from
# ---------------------------------------------------------------------------


class TestC1DiffApplyRoundTrip:
    def test_ambiguous_atomic_edit_gets_context(self):
        original = "Effective January 1, 2025. Section 2 applies. The term is two (2) years."
        modified = "Effective January 1, 2025. Section 2 applies. The term is five (5) years."

        edits = make_edits_self_contained(generate_edits_from_text(original, modified), original)

        for edit in edits:
            assert edit.target_text, "diff output must not contain pure insertions with empty targets"
            assert original.count(edit.target_text) == 1, (
                f"diff emitted an ambiguous target: {edit.target_text!r} "
                f"(occurs {original.count(edit.target_text)} times)"
            )

    def test_unique_targets_left_untouched(self):
        original = "The quick brown fox jumps."
        modified = "The slow brown fox jumps."
        edits = make_edits_self_contained(generate_edits_from_text(original, modified), original)
        assert [(e.target_text, e.new_text) for e in edits] == [("quick", "slow")]

    def test_pure_insertion_becomes_anchored_replacement(self):
        original = "Payment is due upon receipt."
        modified = "Payment is due immediately upon receipt."
        edits = make_edits_self_contained(generate_edits_from_text(original, modified), original)
        assert len(edits) == 1
        edit = edits[0]
        assert edit.target_text, "pure insertion must be rewritten with a text anchor"
        assert edit.target_text in original
        assert original.count(edit.target_text) == 1

    def test_diff_json_roundtrip_reproduces_modified_document(self):
        """The metamorphic invariant from the report: for any (original,
        modified) pair, apply(original, diff(original, modified)) followed by
        accept-all must reproduce `modified` — via JSON, which drops the
        private position index."""
        orig_stream = _make_nda(term="two (2) years")
        mod_stream = _make_nda(term="five (5) years")

        text_orig = _clean_text(orig_stream)
        text_mod = _clean_text(mod_stream)

        edits = make_edits_self_contained(generate_edits_from_text(text_orig, text_mod), text_orig)

        # Simulate `adeu diff --json > edits.json && adeu apply original edits.json`:
        # serialization drops every private attribute, so matching is text-only.
        wire = json.loads(json.dumps([e.model_dump() for e in edits]))
        rehydrated = [ModifyText(**item) for item in wire]

        orig_stream.seek(0)
        engine = RedlineEngine(orig_stream)
        stats = engine.process_batch(rehydrated)
        assert stats["edits_skipped"] == 0
        assert stats["edits_applied"] == len(rehydrated)

        engine.accept_all_revisions(remove_comments=True)
        final_text = _clean_text(engine.save_to_stream())
        assert final_text == text_mod

        # The corruption signatures from the report must be absent.
        assert "5025" not in final_text
        assert "# 5. Term" in final_text
        assert "five (5) years" in final_text


# ---------------------------------------------------------------------------
# H2 — plain targets must match across bold/italic run boundaries
# ---------------------------------------------------------------------------


class TestH2RunBoundaryMatching:
    @staticmethod
    def _boundary_doc():
        doc = Document()
        p = doc.add_paragraph("The word ")
        p.add_run("Al").bold = True
        p.add_run("pha is the bold target.")

        p = doc.add_paragraph("The word ")
        p.add_run("Br").italic = True
        p.add_run("avo is the italic target.")

        p = doc.add_paragraph("The word ")
        p.add_run("Ch").bold = True
        p.add_run("arlie").italic = True
        p.add_run(" is the mixed target.")

        doc.add_paragraph("The word Hotel is the control target.")
        return _stream(doc)

    @pytest.mark.parametrize(
        "target,replacement",
        [
            ("Alpha", "AlphaEdited"),
            ("Bravo", "BravoEdited"),
            ("Charlie", "CharlieEdited"),
            ("Hotel", "HotelEdited"),
        ],
    )
    def test_plain_target_matches_formatted_boundary(self, target, replacement):
        engine = RedlineEngine(self._boundary_doc())
        stats = engine.process_batch([ModifyText(target_text=target, new_text=replacement)])
        assert stats["edits_applied"] == 1, f"plain target {target!r} did not match"

        engine.accept_all_revisions(remove_comments=True)
        final_text = _clean_text(engine.save_to_stream())
        assert replacement in final_text
        assert target not in final_text.replace(replacement, "")

    def test_markdown_inclusive_target_still_matches(self):
        engine = RedlineEngine(self._boundary_doc())
        stats = engine.process_batch([ModifyText(target_text="**Al**pha", new_text="AlphaEdited")])
        assert stats["edits_applied"] == 1

    def test_boundary_match_is_still_ambiguity_checked(self):
        """Two formatted occurrences of the same token: a plain target must be
        rejected as ambiguous under strict mode, not silently pick one."""
        doc = Document()
        for _ in range(2):
            p = doc.add_paragraph("Prefix ")
            p.add_run("Zu").bold = True
            p.add_run("lu suffix.")

        engine = RedlineEngine(_stream(doc))
        errors = engine.validate_edits([ModifyText(target_text="Zulu", new_text="Zebra")])
        assert len(errors) == 1
        assert "Ambiguous match" in errors[0]


# ---------------------------------------------------------------------------
# C2 — oversized edit values must not be echoed unbounded
# ---------------------------------------------------------------------------


class TestC2BoundedReports:
    def test_giant_new_text_is_truncated_in_report_but_fully_applied(self):
        big = "X" * 2_000_000
        engine = RedlineEngine(_make_nda())
        stats = engine.process_batch([ModifyText(target_text="California", new_text=big)])

        assert stats["edits_applied"] == 1
        report = stats["edits"][0]
        assert len(report["new_text"]) < 2_000
        assert "chars omitted" in report["new_text"]
        assert len(report["critic_markup"]) < 2_000
        assert len(report["clean_text"]) < 2_000
        assert len(json.dumps(stats)) < 20_000

        # The truncation is display-only: the document receives the full value.
        engine.accept_all_revisions(remove_comments=True)
        final_text = _clean_text(engine.save_to_stream())
        assert big in final_text

    def test_giant_target_text_is_truncated_in_not_found_error(self):
        big = "Y" * 1_000_000
        engine = RedlineEngine(_make_nda())
        errors = engine.validate_edits([ModifyText(target_text=big, new_text="z")])
        assert len(errors) == 1
        assert len(errors[0]) < 2_000
        assert "not found" in errors[0].lower()


# ---------------------------------------------------------------------------
# H1 — previews must be faithful: no scaffolding, no cross-edit bleed
# ---------------------------------------------------------------------------


class TestH1PreviewIntegrity:
    def test_multi_edit_batch_previews_are_clean_and_localized(self):
        engine = RedlineEngine(_make_nda())
        stats = engine.process_batch(
            [
                ModifyText(target_text="California", new_text="Delaware"),
                ModifyText(target_text="Non-Solicitation", new_text="Indemnification"),
                ModifyText(target_text="two (2) years", new_text="five (5) years"),
            ]
        )
        assert stats["edits_applied"] == 3

        for report in stats["edits"]:
            critic = report["critic_markup"]
            assert critic is not None
            # Internal scaffolding must never leak into previews.
            assert "[Chg:" not in critic
            assert "<<}" not in critic and "{>>" not in critic

        # Each preview marks exactly its own change.
        assert "{--California--}{++Delaware++}" in stats["edits"][0]["critic_markup"]
        assert "{--Non-Solicitation--}{++Indemnification++}" in stats["edits"][1]["critic_markup"]

        # A compound change must preview the COMPLETE logical change, not just
        # its first word-diff sub-edit ("{--two--}{++five++} (2) years").
        assert "{--two (2)--}{++five (5)++} years" in stats["edits"][2]["critic_markup"]
        assert "five (5) years" in stats["edits"][2]["clean_text"]
        assert "five (2)" not in stats["edits"][2]["clean_text"]

    def test_dry_run_previews_match_real_previews(self):
        batch = [
            ModifyText(target_text="California", new_text="Delaware"),
            ModifyText(target_text="two (2) years", new_text="five (5) years"),
        ]
        dry = RedlineEngine(_make_nda()).process_batch([m.model_copy(deep=True) for m in batch], dry_run=True)
        wet = RedlineEngine(_make_nda()).process_batch([m.model_copy(deep=True) for m in batch], dry_run=False)

        assert [r["critic_markup"] for r in dry["edits"]] == [r["critic_markup"] for r in wet["edits"]]
        assert [r["clean_text"] for r in dry["edits"]] == [r["clean_text"] for r in wet["edits"]]


# ---------------------------------------------------------------------------
# M1 — dry-run must report exactly what the real run will do
# ---------------------------------------------------------------------------


class TestM1DryRunParity:
    def test_ambiguity_occurrence_counts_agree(self):
        batch = [
            ModifyText(target_text="party", new_text="side"),  # ambiguous on purpose
            ModifyText(target_text="California", new_text="Delaware"),
        ]

        dry = RedlineEngine(_make_nda()).process_batch([m.model_copy(deep=True) for m in batch], dry_run=True)
        dry_error = next(r["error"] for r in dry["edits"] if r["error"] and "Ambiguous" in r["error"])

        try:
            RedlineEngine(_make_nda()).process_batch([m.model_copy(deep=True) for m in batch], dry_run=False)
            raise AssertionError("real run should have rejected the ambiguous batch")
        except BatchValidationError as e:
            wet_error = next(err for err in e.errors if "Ambiguous" in err)

        import re as _re

        dry_count = _re.search(r"appears (\d+) times", dry_error).group(1)
        wet_count = _re.search(r"appears (\d+) times", wet_error).group(1)
        assert dry_count == wet_count

    def test_dry_run_mirrors_transactional_rejection(self):
        """Real mode rejects the whole batch when any edit fails validation;
        dry-run must not claim other edits 'applied'."""
        batch = [
            ModifyText(target_text="California", new_text="Delaware"),
            ModifyText(target_text="Nonexistent text 123", new_text="x"),
        ]
        dry = RedlineEngine(_make_nda()).process_batch(batch, dry_run=True)
        assert dry["edits_applied"] == 0
        assert dry["edits_skipped"] == 2
        assert all(r["status"] == "failed" for r in dry["edits"])
        assert "transactional" in dry["edits"][0]["error"]
        assert "not found" in dry["edits"][1]["error"].lower()


# ---------------------------------------------------------------------------
# M2 — edits whose targets fight over the same text are rejected, not silently
# misapplied. Batches apply SEQUENTIALLY: each edit sees the document state
# produced by the previous edits, so an overlapping stale target fails with a
# clear error (and a hint about the sequential contract), while a chained
# batch that re-targets the updated text applies cleanly.
# ---------------------------------------------------------------------------


class TestM2OverlapAndChaining:
    def test_stale_overlapping_target_rejected_with_sequential_hint(self):
        batch = [
            ModifyText(target_text="two (2) years", new_text="three (3) years"),
            ModifyText(target_text="(2) years", new_text="(4) years"),
        ]
        engine = RedlineEngine(_make_nda())
        with pytest.raises(BatchValidationError) as exc_info:
            engine.process_batch(batch)
        joined = "\n".join(exc_info.value.errors)
        assert "Edit 2 Failed" in joined
        assert "tracked deletion" in joined
        assert "Batches apply sequentially" in joined
        assert "AFTER the preceding edits" in joined

        # Transactional: the rejected batch left no tracked changes behind.
        final_text = _clean_text(engine.save_to_stream())
        assert "two (2) years" in final_text
        assert "three" not in final_text

    def test_dry_run_reports_the_same_rejection(self):
        batch = [
            ModifyText(target_text="two (2) years", new_text="three (3) years"),
            ModifyText(target_text="(2) years", new_text="(4) years"),
        ]
        dry = RedlineEngine(_make_nda()).process_batch(batch, dry_run=True)
        assert dry["edits_applied"] == 0
        assert dry["edits_skipped"] == 2
        assert all(r["status"] == "failed" for r in dry["edits"])
        assert "tracked deletion" in dry["edits"][1]["error"]
        assert "Batches apply sequentially" in dry["edits"][1]["error"]
        assert "transactional" in dry["edits"][0]["error"]

    def test_chained_batch_targeting_updated_text_applies(self):
        """The sequential-contract way to express the report's overlap example:
        the second edit targets the text as it reads AFTER the first edit."""
        batch = [
            ModifyText(target_text="two (2) years", new_text="three (3) years"),
            ModifyText(target_text="(3) years", new_text="(3) calendar years"),
        ]
        engine = RedlineEngine(_make_nda())
        stats = engine.process_batch(batch)
        assert stats["edits_applied"] == 2
        assert stats["edits_skipped"] == 0

        engine.accept_all_revisions(remove_comments=True)
        final_text = _clean_text(engine.save_to_stream())
        assert "three (3) calendar years" in final_text

    def test_disjoint_edits_unaffected_by_sequential_evaluation(self):
        """Edits on unrelated regions apply exactly as before."""
        batch = [
            ModifyText(target_text="California", new_text="Delaware"),
            ModifyText(target_text="two (2) years", new_text="five (5) years"),
        ]
        engine = RedlineEngine(_make_nda())
        stats = engine.process_batch(batch)
        assert stats["edits_applied"] == 2

        engine.accept_all_revisions(remove_comments=True)
        final_text = _clean_text(engine.save_to_stream())
        assert "Delaware" in final_text
        assert "five (5) years" in final_text


# ---------------------------------------------------------------------------
# M3 / M4 — insert_row cell handling and reporting
# ---------------------------------------------------------------------------


def _table_doc() -> io.BytesIO:
    doc = Document()
    doc.add_paragraph("Pricing tiers:")
    table = doc.add_table(rows=2, cols=3)
    for i, text in enumerate(["Plan", "Price", "Seats"]):
        table.rows[0].cells[i].text = text
    for i, text in enumerate(["Starter", "$10", "5"]):
        table.rows[1].cells[i].text = text
    return _stream(doc)


class TestM3M4InsertRow:
    def test_overfilled_cells_rejected(self):
        engine = RedlineEngine(_table_doc())
        with pytest.raises(BatchValidationError) as exc_info:
            engine.process_batch([InsertTableRow(target_text="Starter", cells=["A", "B", "C", "D", "E"])])
        msg = "\n".join(exc_info.value.errors)
        assert "5 cells" in msg and "3 column" in msg

    def test_underfilled_cells_pad_and_apply(self):
        engine = RedlineEngine(_table_doc())
        stats = engine.process_batch([InsertTableRow(target_text="Starter", cells=["OnlyTwo", "Cells"])])
        assert stats["edits_applied"] == 1

    def test_report_shows_inserted_cells(self):
        engine = RedlineEngine(_table_doc())
        stats = engine.process_batch([InsertTableRow(target_text="Starter", cells=["Pro", "$20", "10"])])
        assert stats["edits"][0]["new_text"] == "Pro | $20 | 10"

    def test_row_op_outside_table_rejected(self):
        engine = RedlineEngine(_table_doc())
        with pytest.raises(BatchValidationError) as exc_info:
            engine.process_batch([InsertTableRow(target_text="Pricing tiers", cells=["A"])])
        assert "not inside a table row" in "\n".join(exc_info.value.errors)


# ---------------------------------------------------------------------------
# CLI-level findings (M5, H3, L1, L2, L4) — exercised through adeu.cli.main
# ---------------------------------------------------------------------------


def _write_docx(tmp_path, name="doc.docx"):
    path = tmp_path / name
    doc = Document()
    doc.add_paragraph("The governing law is the law of California.")
    doc.save(path)
    return path


def _run_cli(argv):
    from adeu.cli import main

    with patch.object(sys, "argv", ["adeu", *argv]):
        try:
            main()
        except SystemExit as e:
            return e.code or 0
    return 0


class TestCliFindings:
    def test_m5_malformed_json_gets_plain_language_error(self, tmp_path, capsys):
        doc = _write_docx(tmp_path)
        bad = tmp_path / "bad.json"
        bad.write_text('[{"foo": "bar"}]', encoding="utf-8")

        code = _run_cli(["apply", str(doc), str(bad), "-o", str(tmp_path / "out.docx")])
        err = capsys.readouterr().err

        assert code == 1
        assert "pydantic" not in err.lower()
        assert "union_tag_not_found" not in err
        assert "missing the required 'type' field" in err
        # The error must teach the valid vocabulary.
        for type_name in ("modify", "accept", "reject", "reply", "insert_row", "delete_row"):
            assert type_name in err

    def test_m5_missing_field_names_the_field(self, tmp_path, capsys):
        doc = _write_docx(tmp_path)
        bad = tmp_path / "bad2.json"
        bad.write_text('[{"type": "reply", "target_id": "Com:1"}]', encoding="utf-8")

        code = _run_cli(["apply", str(doc), str(bad), "-o", str(tmp_path / "out.docx")])
        err = capsys.readouterr().err
        assert code == 1
        assert "missing required field 'text'" in err

    def test_h3_unwritable_output_path_is_clean_error(self, tmp_path, capsys):
        doc = _write_docx(tmp_path)
        edits = tmp_path / "edits.json"
        edits.write_text(
            json.dumps([{"type": "modify", "target_text": "California", "new_text": "Delaware"}]),
            encoding="utf-8",
        )
        long_name = "a" * 300 + ".docx"

        code = _run_cli(["apply", str(doc), str(edits), "-o", str(tmp_path / long_name)])
        err = capsys.readouterr().err
        assert code == 1
        assert "Could not write output file" in err
        assert "Traceback" not in err

    def test_l1_invalid_page_values_error(self, tmp_path, capsys):
        # Exit code 2: invalid argument values are usage errors
        # (QA 2026-07-18 v6 L1).
        doc = _write_docx(tmp_path)

        code = _run_cli(["extract", str(doc), "--page", "-1"])
        assert code == 2
        assert "Invalid --page value" in capsys.readouterr().err

        code = _run_cli(["extract", str(doc), "--page", "abc"])
        assert code == 2
        assert "Invalid --page value" in capsys.readouterr().err

    def test_l1_valid_page_still_works(self, tmp_path, capsys):
        doc = _write_docx(tmp_path)
        code = _run_cli(["extract", str(doc), "--page", "1"])
        assert code == 0
        assert "California" in capsys.readouterr().out

    def test_l2_outline_level_bounds_rejected_at_parse_time(self, tmp_path, capsys):
        doc = _write_docx(tmp_path)
        code = _run_cli(["extract", str(doc), "--mode", "outline", "--outline-max-level", "0"])
        assert code == 2  # argparse usage error
        assert "between 1 and 6" in capsys.readouterr().err

    def test_l4_empty_batch_warns(self, tmp_path, capsys):
        doc = _write_docx(tmp_path)
        empty = tmp_path / "empty.json"
        empty.write_text("[]", encoding="utf-8")

        code = _run_cli(["apply", str(doc), str(empty), "-o", str(tmp_path / "out.docx")])
        err = capsys.readouterr().err
        assert code == 0
        assert "0 changes" in err and "nothing to do" in err


class TestL2OutlineBuilderClamp:
    def test_builder_clamps_out_of_range_levels(self):
        """MCP callers bypass argparse; the shared builder must not render a
        nonsensical 'L1-L0' range label."""
        from adeu.mcp_components._response_builders import build_outline_response

        doc = Document()
        doc.add_heading("Top Heading", level=1)
        doc.add_paragraph("Body text.")
        stream = _stream(doc)
        stream.seek(0)
        projected = extract_text_from_stream(io.BytesIO(stream.read()), filename="x.docx")

        stream.seek(0)
        from docx import Document as load_document

        res = build_outline_response(load_document(stream), projected, "x.docx", outline_max_level=0)
        markdown = res.structured_content["markdown"]
        assert "L1-L0" not in markdown
        assert "L1-L-" not in markdown


# ---------------------------------------------------------------------------
# L3 — markup must not pass CriticMarkup syntax in new_text through
# ---------------------------------------------------------------------------


class TestL3MarkupInjection:
    def test_markup_rejects_criticmarkup_in_new_text(self, tmp_path, capsys):
        doc = _write_docx(tmp_path)
        inj = tmp_path / "inj.json"
        inj.write_text(
            json.dumps(
                [
                    {
                        "type": "modify",
                        "target_text": "California",
                        "new_text": "Delaware {++nested++} {>>fake comment<<}",
                    }
                ]
            ),
            encoding="utf-8",
        )
        out = tmp_path / "inj.md"

        code = _run_cli(["markup", str(doc), str(inj), "-o", str(out)])
        err = capsys.readouterr().err
        assert code == 1
        assert "Do not manually write CriticMarkup tags" in err
        assert not out.exists()

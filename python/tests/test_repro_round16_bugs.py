import io
import re
from unittest.mock import AsyncMock, patch

import pytest
from docx import Document

from adeu.mcp_components.tools.sanitize import sanitize_docx
from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine


@pytest.mark.anyio
async def test_bug16_sanitize_kwargs():
    """
    Bug #16: sanitize_docx tool passes arguments to _sanitize positionally,
    which fails if _sanitize expects keyword-only arguments.
    """

    def mock_sanitize(
        file_path,
        output_path=None,
        *,
        keep_markup=False,
        baseline_path=None,
        author=None,
        accept_all=False,
    ):
        # If called incorrectly, Python raises TypeError: takes X positional arguments but Y were given
        from unittest.mock import MagicMock

        res = MagicMock()
        res.output_path = "out.docx"
        res.status = "Success"
        res.tracked_changes_found = 0
        res.tracked_changes_accepted = 0
        res.comments_removed = 0
        res.comments_kept = 0
        res.metadata_stripped = []
        res.warnings = []
        res.report_text = ""
        return res

    ctx = AsyncMock()
    with (
        patch("pathlib.Path.exists", return_value=True),
        patch(
            "adeu.mcp_components.tools.sanitize._sanitize",
            side_effect=mock_sanitize,
            create=True,
        ),
        patch("adeu.sanitize.core.sanitize_docx", side_effect=mock_sanitize, create=True),
    ):
        # This will raise TypeError if the await asyncio.to_thread call is passing them positionally
        result = await sanitize_docx(
            reasoning="test",
            file_path="dummy.docx",
            ctx=ctx,
            output_path="out.docx",
            keep_markup=True,
            baseline_path=None,
            author="Test",
            accept_all=True,
        )
        assert result["status"] == "Success"


def test_bug11_engine_init_whitespace_mutation():
    """
    Bug #11: RedlineEngine.__init__ should not use serialize_for_reading
    which adds massive amounts of pretty-printed whitespace that then pollutes the saved document.
    """
    doc = Document()
    doc.add_paragraph("Test")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Check baseline newlines
    doc_before = Document(stream)
    xml_before = doc_before.part.blob.decode("utf-8")
    newlines_before = xml_before.count("\n")
    stream.seek(0)

    # Init engine
    engine = RedlineEngine(stream)
    out_stream = engine.save_to_stream()

    out_doc = Document(out_stream)
    xml_after = out_doc.part.blob.decode("utf-8")
    newlines_after = xml_after.count("\n")

    # serialize_for_reading injects ~10-20+ newlines even for a tiny document
    # Without the bug, newlines_after should roughly equal newlines_before (typically 0 or 1).
    assert newlines_after <= newlines_before + 2, f"Whitespace leak! Before: {newlines_before}, After: {newlines_after}"


def test_bug12_replace_inside_bold_inherits_bold():
    """
    Bug #12: Replacing text inside a bold span (without using ** in new_text)
    should inherit the bold styling of the target run, not strip it.
    """
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("BOLD TEXT")
    r.bold = True

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Replacing "TEXT" with "WORD", intentionally not providing **WORD**
    edit = ModifyText(target_text="TEXT", new_text="WORD")

    engine = RedlineEngine(stream)
    engine.apply_edits([edit])

    out_stream = engine.save_to_stream()
    out_doc = Document(out_stream)

    ins_runs = out_doc.element.xpath("//w:ins//w:r")
    assert len(ins_runs) > 0, "No insertion found"

    bolds = ins_runs[0].xpath(".//w:b")
    assert len(bolds) > 0, "Inserted text lost the inherited bold styling of the target!"


def test_bug15_debug_xml_diff_noise():
    """
    Bug #15: The diff output must strip lxml's pretty-printing whitespace between
    tags to reduce noise when comparing live vs disk engines.
    """
    xml_with_noise = "<w:p>\n  <w:r>\n    <w:t>Hello</w:t>\n  </w:r>\n</w:p>"

    cleaned = re.sub(r">\s+<", ">\n<", xml_with_noise)

    assert cleaned == "<w:p>\n<w:r>\n<w:t>Hello</w:t>\n</w:r>\n</w:p>", (
        "Inter-tag whitespace was not correctly normalized."
    )

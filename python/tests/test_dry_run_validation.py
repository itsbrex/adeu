# FILE: tests/test_dry_run_validation.py
import io

from docx import Document

from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine


def test_validation_ambiguous_match_with_context():
    """
    Scenario: The exact bug reported. Target text exists in multiple clauses.
    Validates that the engine catches the duplication and extracts surrounding text for the LLM.
    """
    doc = Document()
    # Padded with introductory text to ensure the clause numbers fall fully within
    # the engine's 30-character backward context window without getting sliced.
    doc.add_paragraph("As specified in Clause 2.4 (Interest), the rate shall not exceed 150% of the base.")
    doc.add_paragraph(
        "As specified in Clause 6.1 (Liability Cap), Seller liability shall not exceed 150% of the total contract."
    )

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)

    # The LLM's flawed, overly generic edit
    edit = ModifyText(target_text="shall not exceed 150%", new_text="shall not exceed 100%")

    errors = engine.validate_edits([edit])

    assert len(errors) == 1
    error_msg = errors[0]

    # Verify the error structure
    assert "Edit 1 Failed: Ambiguous match" in error_msg
    assert "appears 2 times" in error_msg

    # Verify the context extraction is working so the LLM can see the difference
    assert "Clause 2.4" in error_msg
    assert "Clause 6.1" in error_msg
    assert "Please provide more surrounding context" in error_msg


def test_validation_not_found():
    """
    Scenario: Target text does not exist in the document (hallucination or severe typo).
    Validates that the engine catches it and reports it.
    """
    doc = Document()
    doc.add_paragraph("The liability of the Seller is strictly limited.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)
    edit = ModifyText(
        target_text="The liability of the Buyer",
        new_text="The liability of the Purchaser",
    )

    errors = engine.validate_edits([edit])

    assert len(errors) == 1
    assert "Edit 1 Failed: Target text not found" in errors[0]
    assert "The liability of the Buyer" in errors[0]


def test_validation_success_and_pure_insertion():
    """
    Scenario: Perfect targets. Includes a pure insertion (empty target text)
    which should bypass validation safely.
    """
    doc = Document()
    doc.add_paragraph("This is a unique clause.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)

    edit1 = ModifyText(target_text="unique clause", new_text="special clause")

    # Pure insertion (relies on index internally, skipping target_text validation)
    edit2 = ModifyText(target_text="", new_text="Inserted text.")

    errors = engine.validate_edits([edit1, edit2])

    # Should pass perfectly with zero errors
    assert len(errors) == 0


def test_validation_fuzzy_match_ambiguity():
    """
    Scenario: The target text is technically different from the doc, but the fuzzy matcher
    resolves it to multiple places (e.g., varying underscores in legal placeholders).
    """
    doc = Document()
    doc.add_paragraph("Buyer Signature: [__________]")  # 10 underscores
    doc.add_paragraph("Seller Signature: [_______]")  # 7 underscores

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)

    # LLM targets with 3 underscores. The fuzzy matcher should find BOTH signatures.
    edit = ModifyText(target_text="[___]", new_text="John Doe")

    errors = engine.validate_edits([edit])

    assert len(errors) == 1
    assert "Ambiguous match" in errors[0]
    assert "appears 2 times" in errors[0]
    assert "Buyer Signature" in errors[0]
    assert "Seller Signature" in errors[0]


def test_multiple_errors_accumulated():
    """
    Scenario: A batch with multiple bad edits.
    Validates that the dry run is exhaustive and doesn't short-circuit after the first error.
    """
    doc = Document()
    doc.add_paragraph("Duplicate word. Duplicate word.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)

    edit1 = ModifyText(target_text="Duplicate word.", new_text="Changed.")  # Ambiguous
    edit2 = ModifyText(target_text="Missing word.", new_text="Found.")  # Not Found

    errors = engine.validate_edits([edit1, edit2])

    assert len(errors) == 2
    assert "Edit 1 Failed: Ambiguous match" in errors[0]
    assert "Edit 2 Failed: Target text not found" in errors[1]


def test_ambiguity_error_surfaces_match_mode_escape_hatch():
    """
    Regression for the "Turn Loop Trap": when process_document_batch returns an
    ambiguous-match failure, the message MUST tell the agent how to re-call using
    match_mode="all" / match_mode="first". Without this, agents burn dozens of
    turns refining target_text and regexes instead of using the built-in escape
    hatch — exactly the failure that blew through the 80-turn / 6.78M-token ceiling.
    """
    doc = Document()
    # Mirror the DPA scenario: the same placeholder appears in multiple clauses.
    doc.add_paragraph("PROVIDER: [official company name] shall process the data.")
    doc.add_paragraph("PROVIDER: [official company name] is the data processor.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)
    edit = ModifyText(target_text="PROVIDER: [official company name]", new_text="PROVIDER: Acme Corp")

    errors = engine.validate_edits([edit])

    assert len(errors) == 1
    error_msg = errors[0]

    assert "Edit 1 Failed: Ambiguous match" in error_msg
    assert "appears 2 times" in error_msg

    # The actionable remediation: the message names BOTH match_mode options
    # explicitly so the agent can immediately re-call instead of looping.
    assert 'match_mode": "all"' in error_msg
    assert 'match_mode": "first"' in error_msg
    assert "ALL 2 occurrences" in error_msg
    assert "FIRST occurrence" in error_msg

    # The original "add more context" guidance is preserved as a third option.
    assert "Please provide more surrounding context" in error_msg


def test_ambiguity_resolved_by_match_mode_all():
    """
    Goal state: re-sending the SAME target_text with match_mode="all" clears the
    ambiguity (zero validation errors) and modifies every occurrence.
    """
    doc = Document()
    doc.add_paragraph("PROVIDER: [official company name] shall process the data.")
    doc.add_paragraph("PROVIDER: [official company name] is the data processor.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)
    edit = ModifyText(
        target_text="PROVIDER: [official company name]",
        new_text="PROVIDER: Acme Corp",
        match_mode="all",
    )

    # Validation no longer trips on the ambiguity.
    assert engine.validate_edits([edit]) == []

    # And a dry run confirms every occurrence is targeted.
    stats = engine.process_batch([edit], dry_run=True)
    report = stats["edits"][0]
    assert report.get("match_mode") == "all"
    assert report.get("occurrences_modified", 0) == 2


def test_ambiguity_resolved_by_match_mode_first():
    """
    Goal state: re-sending the SAME target_text with match_mode="first" clears the
    ambiguity and modifies only the first occurrence.
    """
    doc = Document()
    doc.add_paragraph("PROVIDER: [official company name] shall process the data.")
    doc.add_paragraph("PROVIDER: [official company name] is the data processor.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)
    edit = ModifyText(
        target_text="PROVIDER: [official company name]",
        new_text="PROVIDER: Acme Corp",
        match_mode="first",
    )

    assert engine.validate_edits([edit]) == []

    stats = engine.process_batch([edit], dry_run=True)
    report = stats["edits"][0]
    assert report.get("match_mode") == "first"
    assert report.get("occurrences_modified", 0) == 1


def test_validation_rejects_hallucinated_criticmarkup():
    """
    Scenario: LLM hallucinates CriticMarkup tags inside the new_text instead of using the comment parameter.
    Validates that the engine explicitly catches and rejects this to prevent DOM pollution.
    """
    doc = Document()
    doc.add_paragraph("Target text.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)
    edit = ModifyText(target_text="Target text", new_text="New text {>>LLM Comment<<}")

    errors = engine.validate_edits([edit])

    assert len(errors) == 1
    assert "Do not manually write CriticMarkup tags" in errors[0]

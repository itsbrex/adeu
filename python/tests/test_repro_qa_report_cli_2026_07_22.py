"""
Reproductions for the Adeu CLI QA report (2026-07-22, v1.29.0), bugs #2 and #3.

Bug #2 (cosmetic): the sanitize --report header printed the invocation flag
token ("--baseline") on its own line directly under the title, reading like a
stray debug/args string appended to the filename. Since --report output is shown
to a counterparty as proof of a clean document, the flags line must be clearly
labelled.

Bug #3 (UX): accept/reject/reply on an invalid target_id emitted only
"Failed to apply action: reply on 99" — no reason, no way to find a valid id —
while every other error in the tool is self-service. The error must name the
expected id kind, list the ids that exist, and point at how to discover them.
"""

import io

from docx import Document

from adeu.models import ModifyText, RejectChange, ReplyComment
from adeu.redline.engine import BatchValidationError, RedlineEngine
from adeu.sanitize.report import SanitizeReport


# ────────────────────────────────────────────────────────────────────────────
# Bug #2 — sanitize --report header labels its flags line.
# ────────────────────────────────────────────────────────────────────────────
def _header_lines(mode: str, **kw) -> list[str]:
    report = SanitizeReport(filename="contract_final.docx", mode=mode, **kw)
    # header is everything up to the second separator line
    rendered = report.render().splitlines()
    sep = "═" * 50
    second_sep = rendered.index(sep, 1)
    return rendered[:second_sep]


def test_baseline_report_header_does_not_print_bare_flag_token():
    lines = _header_lines("baseline")
    assert "Sanitize Report: contract_final.docx" in lines
    # The bare token must be gone; the flag must appear under a clear label.
    assert "--baseline" not in lines, "bare flag token still printed as its own line"
    assert any(line.startswith("Options:") and "--baseline" in line for line in lines)


def test_keep_markup_report_header_labels_flags():
    lines = _header_lines("keep-markup")
    assert "--keep-markup" not in lines
    assert any(line.startswith("Options:") and "--keep-markup" in line for line in lines)


def test_report_header_labels_author_and_accept_all():
    report = SanitizeReport(
        filename="contract_final.docx",
        mode="baseline",
        author="Jane Doe",
        tracked_changes_accepted=3,
    )
    header = report.render().splitlines()
    options_line = next(line for line in header if line.startswith("Options:"))
    assert "--baseline" in options_line
    assert '--author "Jane Doe"' in options_line
    assert "--accept-all" in options_line


def test_full_mode_with_no_flags_has_no_options_line():
    # A plain full sanitize with no author/accept-all still reads cleanly.
    lines = _header_lines("full")
    assert not any(line.startswith("Options:") for line in lines)


# ────────────────────────────────────────────────────────────────────────────
# Bug #3 — self-service errors for invalid accept/reject/reply targets.
# ────────────────────────────────────────────────────────────────────────────
def _doc_with_one_change_and_comment() -> io.BytesIO:
    doc = Document()
    doc.add_paragraph("The fee shall be 100 USD per unit as described in the schedule.")
    base = io.BytesIO()
    doc.save(base)
    base.seek(0)
    eng = RedlineEngine(base, author="Reviewer")
    eng.apply_edits([ModifyText(target_text="100 USD", new_text="150 USD", comment="Confirm currency with finance.")])
    return eng.save_to_stream()


def _error_for(action) -> str:
    eng = RedlineEngine(_doc_with_one_change_and_comment(), author="Reviewer")
    try:
        eng.process_batch([action])
    except BatchValidationError as exc:
        return str(exc)
    raise AssertionError("expected the action to fail")


def test_reply_to_missing_comment_lists_valid_ids_and_how_to_find_them():
    msg = _error_for(ReplyComment(target_id="Com:99", text="ok"))
    assert "no comment with that id exists" in msg
    # It must surface the comment id that DOES exist and how to list ids.
    assert "Com:1" in msg
    assert "adeu markup" in msg or "adeu extract" in msg


def test_reject_missing_change_explains_and_lists_valid_ids():
    msg = _error_for(RejectChange(target_id="Chg:99"))
    assert "no tracked change with that id exists" in msg
    assert "Chg:1" in msg  # a real change id is listed
    assert "adeu markup" in msg or "adeu extract" in msg


def test_reject_on_comment_id_flags_the_kind_mismatch():
    msg = _error_for(RejectChange(target_id="Com:1"))
    # echoes what the caller typed and explains the change/comment mix-up
    assert "reject on Com:1" in msg
    assert "comment id" in msg and "reply" in msg


def test_reply_on_change_id_flags_the_kind_mismatch():
    msg = _error_for(ReplyComment(target_id="Chg:1", text="ok"))
    assert "reply on Chg:1" in msg
    assert "tracked-change id" in msg


def test_error_is_not_the_old_terse_form():
    # The old message was exactly "Failed to apply action: reply on 99".
    msg = _error_for(ReplyComment(target_id="99", text="ok"))
    assert msg.strip() != "Failed to apply action: reply on 99"
    assert "no comment with that id exists" in msg


def test_missing_change_on_clean_doc_says_no_changes():
    doc = Document()
    doc.add_paragraph("plain untouched text")
    base = io.BytesIO()
    doc.save(base)
    base.seek(0)
    eng = RedlineEngine(base, author="Reviewer")
    try:
        eng.process_batch([RejectChange(target_id="Chg:5")])
    except BatchValidationError as exc:
        msg = str(exc)
    else:
        raise AssertionError("expected failure")
    assert "no tracked change" in msg
    assert "This document has no tracked changes." in msg

import io

from docx import Document
from docx.oxml.ns import qn

from adeu.ingest import extract_text_from_stream
from adeu.models import DeleteTableRow, InsertTableRow, ModifyText, RejectChange
from adeu.redline.engine import RedlineEngine


def test_interleaved_tables_and_text():
    """
    Verifies that the extractor respects document order:
    Paragraph -> Table -> Paragraph.
    Previously, all tables were extracted at the end of the section.
    """
    doc = Document()
    doc.add_paragraph("Section 1")

    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "TableContent"

    doc.add_paragraph("Section 2")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    text = extract_text_from_stream(stream)

    # 1. Content check
    assert "Section 1" in text
    assert "TableContent" in text
    assert "Section 2" in text

    # 2. Order check
    p1 = text.find("Section 1")
    tbl = text.find("TableContent")
    p2 = text.find("Section 2")

    assert p1 < tbl < p2, f"Table content out of order! Indicies: P1={p1}, Tbl={tbl}, P2={p2}"


def test_nested_tables_extraction_and_editing():
    """
    Verifies recursive extraction logic.
    Structure: Table -> Cell -> Table -> Cell -> Text
    """
    doc = Document()
    outer_table = doc.add_table(rows=1, cols=1)
    outer_cell = outer_table.cell(0, 0)

    # Add nested table inside the cell
    nested_table = outer_cell.add_table(rows=1, cols=1)
    nested_table.cell(0, 0).text = "InnerSecret"

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # 1. Verify Ingest finds it
    text = extract_text_from_stream(stream)
    assert "InnerSecret" in text, "Nested table content failed to extract"

    # 2. Verify Mapping/Editing can reach it
    edit = ModifyText(target_text="InnerSecret", new_text="OuterSecret")
    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])

    assert applied == 1
    assert skipped == 0

    res_stream = engine.save_to_stream()
    res_text = extract_text_from_stream(res_stream)

    # Expect clean insertion logic
    assert "{--InnerSecret--}{++OuterSecret++}" in res_text


def test_merged_cells_no_duplication():
    """
    Verifies that merged cells are extracted exactly once.
    python-docx iterates a 2-col merged row as [Cell A, Cell A].
    We must deduplicate to avoid "Text | Text".
    """
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    c1 = table.cell(0, 0)
    c2 = table.cell(0, 1)

    c1.merge(c2)
    c1.text = "MergedUnique"

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    text = extract_text_from_stream(stream)

    # Should appear exactly once. If deduplication fails, count will be 2.
    count = text.count("MergedUnique")
    assert count == 1, f"Merged cell content appeared {count} times (expected 1)"

    # Verify Edit targets correctly (coordinates shouldn't be confused by skip)
    edit = ModifyText(target_text="MergedUnique", new_text="ChangedUnique")
    engine = RedlineEngine(stream)
    applied, _ = engine.apply_edits([edit])
    assert applied == 1


def test_empty_row_alignment():
    """
    Verifies that Ingest and Mapper stay synchronized even with empty rows.
    If Ingest skips empty rows but Mapper counts them (or vice versa),
    subsequent edits will drift and target the wrong text.
    """
    doc = Document()
    table = doc.add_table(rows=3, cols=1)

    table.cell(0, 0).text = "RowA"
    table.cell(1, 0).text = ""  # Empty Row
    table.cell(2, 0).text = "RowB"  # Target

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # If alignment is broken, "RowB" index will be calculated wrong
    edit = ModifyText(target_text="RowB", new_text="RowC")

    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])

    assert applied == 1, "Edit failed - likely due to mapping index drift caused by empty row"

    res_stream = engine.save_to_stream()
    text = extract_text_from_stream(res_stream)

    # RowA should still exist
    assert "RowA" in text
    # RowB should be modified
    assert "{--RowB--}{++RowC++}" in text


def test_insert_table_row_below():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "A2"
    table.cell(1, 0).text = "B1"
    table.cell(1, 1).text = "B2"
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    change = InsertTableRow(target_text="A1 | A2", position="below", cells=["New B1", "New B2"])

    engine = RedlineEngine(stream)
    stats = engine.process_batch([change])

    assert stats["edits_applied"] == 1

    engine.accept_all_revisions()
    clean_text = extract_text_from_stream(engine.save_to_stream(), clean_view=True)

    assert "A1 | A2" in clean_text
    assert "New B1 | New B2" in clean_text
    assert "B1 | B2" in clean_text

    # Check order
    lines = [line.strip() for line in clean_text.split("\n") if "|" in line]
    assert lines == ["A1 | A2", "--- | ---", "New B1 | New B2", "B1 | B2"]


def test_delete_table_row():
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "A2"
    table.cell(1, 0).text = "B1"
    table.cell(1, 1).text = "B2"
    table.cell(2, 0).text = "C1"
    table.cell(2, 1).text = "C2"
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    change = DeleteTableRow(target_text="B1")

    engine = RedlineEngine(stream)
    stats = engine.process_batch([change])

    assert stats["edits_applied"] == 1

    engine.accept_all_revisions()
    clean_text = extract_text_from_stream(engine.save_to_stream(), clean_view=True)

    assert "A1 | A2" in clean_text
    assert "B1 | B2" not in clean_text
    assert "C1 | C2" in clean_text


def test_reject_insert_table_row():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "A2"
    table.cell(1, 0).text = "B1"
    table.cell(1, 1).text = "B2"
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    change = InsertTableRow(target_text="A1 | A2", position="below", cells=["New B1", "New B2"])

    engine = RedlineEngine(stream)
    engine.process_batch([change])

    ins_id = engine.doc.element.xpath("//w:tr/w:trPr/w:ins")[0].get(qn("w:id"))

    engine.process_batch([RejectChange(target_id=f"Chg:{ins_id}")])

    clean_text = extract_text_from_stream(engine.save_to_stream(), clean_view=True)
    assert "New B1 | New B2" not in clean_text
    assert "A1 | A2" in clean_text
    assert "B1 | B2" in clean_text


def test_reject_delete_table_row():
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "A2"
    table.cell(1, 0).text = "B1"
    table.cell(1, 1).text = "B2"
    table.cell(2, 0).text = "C1"
    table.cell(2, 1).text = "C2"
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    change = DeleteTableRow(target_text="B1")

    engine = RedlineEngine(stream)
    engine.process_batch([change])

    del_id = engine.doc.element.xpath("//w:tr/w:trPr/w:del")[0].get(qn("w:id"))

    engine.process_batch([RejectChange(target_id=f"Chg:{del_id}")])

    clean_text = extract_text_from_stream(engine.save_to_stream(), clean_view=True)
    assert "A1 | A2" in clean_text
    assert "B1 | B2" in clean_text
    assert "C1 | C2" in clean_text


def test_ingest_structural_row_changes():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "A2"
    table.cell(1, 0).text = "B1"
    table.cell(1, 1).text = "B2"
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)
    engine.process_batch([InsertTableRow(target_text="A1", cells=["New", "Row"]), DeleteTableRow(target_text="B1")])

    raw_text = extract_text_from_stream(engine.save_to_stream(), clean_view=False)
    assert "{++ A1 | A2 |" in raw_text or "{++ New | Row |" in raw_text
    # Batches apply sequentially in batch order, so the insert (edit 1) takes
    # Chg:1 and the delete (edit 2) takes Chg:2.
    assert "{++ New | Row |Chg:1++}" in raw_text
    assert "{-- B1 | B2 |Chg:2--}" in raw_text


def test_clean_view_omits_deleted_row():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "A2"
    table.cell(1, 0).text = "B1"
    table.cell(1, 1).text = "B2"
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    change = DeleteTableRow(target_text="B1")

    engine = RedlineEngine(stream)
    engine.process_batch([change])

    # DO NOT accept revisions. We want to test how clean_view handles the active tracked deletion.
    clean_text = extract_text_from_stream(engine.save_to_stream(), clean_view=True)

    assert "A1 | A2" in clean_text
    assert "B1 | B2" not in clean_text

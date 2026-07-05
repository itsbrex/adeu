"""
Reproduction for the field report (2026-07-03, PII-Shield Desktop / memo review):

    "adeu's modify refuses to touch any text that sits inside another author's
    comment range, misclassifying it as 'an active insertion from another
    author.' This holds even when the target is fully inside the span
    (confirmed via dry-run)."

Scenario: a memo arrives with a colleague's margin comments anchored to plain,
untracked body text. The agent is asked to amend that body text. The engine
refuses with "Modification targets an active insertion from another author"
even though no insertion exists anywhere in the document — only a foreign
COMMENT_ONLY annotation. The user's forced workaround (strip all comments,
edit, re-diff) defeats the point of redlining.

Root cause (python/src/adeu/redline/engine.py, validate_edits): the foreign
<w:ins> overlap check and the foreign comment-range overlap check feed the
same rejection branch. Any foreign comment overlap — with zero insertions —
is rejected, and with the insertion-specific error message.

Word's native behavior: editing text under someone else's comment is a normal
review workflow. The comment anchor persists (shrinking to the surviving
range) and the edit becomes a tracked change. Deleting/replying to the
COMMENT ITSELF is a different operation and stays protected.

STYLE: these tests assert the DESIRED behaviour, so they are RED while the
bug is present and turn GREEN once the engine is fixed (the "isolate the bug
before fixing" pattern from AI_CONTEXT.md > Testing). The GREEN controls pin
down the boundary of the fix: same-author comments must keep working, and the
foreign tracked-INSERTION straddle protection must NOT be loosened.
"""

import io
import zipfile

import pytest
from docx import Document

from adeu.ingest import extract_text_from_stream
from adeu.models import ModifyText
from adeu.redline.engine import BatchValidationError, RedlineEngine


def _count_comment_markers(stream: io.BytesIO) -> tuple[int, int, int]:
    zf = zipfile.ZipFile(io.BytesIO(stream.getvalue()))
    xml = zf.read("word/document.xml").decode("utf-8")
    return (
        xml.count("<w:commentRangeStart"),
        xml.count("<w:commentRangeEnd"),
        xml.count("<w:commentReference"),
    )


def _memo_with_colleague_comment() -> io.BytesIO:
    """
    A memo whose second paragraph carries a colleague's pure comment (no
    tracked changes anywhere) over 'the budget is 40,000 EUR'. This mirrors
    the reported document state: clean body text + foreign comment ranges.
    """
    doc = Document()
    doc.add_paragraph("MEMO")
    doc.add_paragraph("The project deadline is 15 September 2026 and the budget is 40,000 EUR.")
    doc.add_paragraph("Please review the terms above.")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    eng = RedlineEngine(stream, author="Colleague")
    eng.apply_edits(
        [
            ModifyText(
                target_text="the budget is 40,000 EUR",
                new_text="the budget is 40,000 EUR",
                comment="Is this figure still correct?",
            )
        ]
    )
    out = eng.save_to_stream()

    # Fixture sanity: exactly one comment, zero redlines.
    text = extract_text_from_stream(io.BytesIO(out.getvalue()), clean_view=False)
    assert "[Com:1]" in text
    assert "{++" not in text and "{--" not in text
    assert _count_comment_markers(out) == (1, 1, 1)
    out.seek(0)
    return out


# ────────────────────────────────────────────────────────────────────────────
# RED — the reported bug. Desired: the edit applies and the comment survives.
# ────────────────────────────────────────────────────────────────────────────
def test_modify_text_strictly_inside_foreign_comment_range_applies():
    """
    Target '40,000 EUR' sits strictly inside Colleague's comment span. There
    is no <w:ins> anywhere in the document, so refusing this as 'an active
    insertion from another author' is a misclassification. Desired: tracked
    change by the agent, colleague's comment intact.
    """
    engine = RedlineEngine(_memo_with_colleague_comment(), author="Agent")
    res = engine.process_batch([ModifyText(target_text="40,000 EUR", new_text="45,000 EUR")])

    assert res["edits_applied"] == 1
    assert res["edits_skipped"] == 0

    out = engine.save_to_stream()
    clean = extract_text_from_stream(io.BytesIO(out.getvalue()), clean_view=True)
    assert "45,000 EUR" in clean

    marked = extract_text_from_stream(io.BytesIO(out.getvalue()), clean_view=False)
    assert "{++" in marked and "{--" in marked, "Edit must land as a tracked change, not a silent rewrite"
    assert "Agent" in marked

    # The colleague's annotation must survive the edit underneath it.
    assert "Is this figure still correct?" in marked
    starts, ends, refs = _count_comment_markers(out)
    assert (starts, ends, refs) == (1, 1, 1), "Colleague's comment anchor was destroyed or duplicated"


def test_modify_entire_foreign_comment_span_applies():
    """
    Boundary variant: the target coincides exactly with the commented span,
    so the edit's edges touch the commentRangeStart/End markers themselves.
    """
    engine = RedlineEngine(_memo_with_colleague_comment(), author="Agent")
    res = engine.process_batch(
        [
            ModifyText(
                target_text="the budget is 40,000 EUR",
                new_text="the budget is 45,000 EUR excluding VAT",
            )
        ]
    )

    assert res["edits_applied"] == 1
    assert res["edits_skipped"] == 0

    out = engine.save_to_stream()
    clean = extract_text_from_stream(io.BytesIO(out.getvalue()), clean_view=True)
    assert "45,000 EUR excluding VAT" in clean

    marked = extract_text_from_stream(io.BytesIO(out.getvalue()), clean_view=False)
    assert "Is this figure still correct?" in marked
    starts, ends, refs = _count_comment_markers(out)
    assert (starts, ends, refs) == (1, 1, 1)


def test_dry_run_reports_edit_under_foreign_comment_as_applicable():
    """
    The report explicitly says the refusal was 'confirmed via dry-run'.
    Desired: dry-run previews the edit as applicable instead of failing it
    with the insertion misclassification.
    """
    engine = RedlineEngine(_memo_with_colleague_comment(), author="Agent")
    res = engine.process_batch(
        [ModifyText(target_text="40,000 EUR", new_text="45,000 EUR")],
        dry_run=True,
    )

    assert res["edits_applied"] == 1
    assert res["edits_skipped"] == 0
    report = res["edits"][0]
    assert report["status"] == "applied", f"Dry-run failed the edit: {report['error']}"
    assert "active insertion from another author" not in (report["error"] or "")


# ────────────────────────────────────────────────────────────────────────────
# GREEN controls — pin the boundary of the fix.
# ────────────────────────────────────────────────────────────────────────────
def test_control_same_author_comment_range_edit_applies():
    """
    GREEN today: the comment author editing their own commented text is not
    blocked. Proves the rejection above is keyed purely on foreign authorship
    of the comment, not on comment ranges per se.
    """
    engine = RedlineEngine(_memo_with_colleague_comment(), author="Colleague")
    res = engine.process_batch([ModifyText(target_text="40,000 EUR", new_text="45,000 EUR")])
    assert res["edits_applied"] == 1
    assert res["edits_skipped"] == 0


def test_control_foreign_insertion_straddle_still_refused():
    """
    GREEN today and must STAY green after the fix: an edit that partially
    straddles a foreign author's real tracked insertion is still refused —
    that protection is legitimate and must not be loosened while enabling
    comment-range edits.
    """
    doc = Document()
    doc.add_paragraph("The quick brown fox jumps.")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    eng1 = RedlineEngine(stream, author="Colleague")
    eng1.apply_edits([ModifyText(target_text="brown", new_text="red")])
    redlined = eng1.save_to_stream()

    engine = RedlineEngine(redlined, author="Agent")
    # 'red fox' = foreign inserted 'red' + untracked body ' fox' -> straddle.
    with pytest.raises(BatchValidationError) as exc:
        engine.process_batch([ModifyText(target_text="red fox", new_text="crimson wolf")])
    assert "active insertion from another author" in str(exc.value)

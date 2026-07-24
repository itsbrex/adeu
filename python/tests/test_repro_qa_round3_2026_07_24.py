# FILE: tests/test_repro_qa_round3_2026_07_24.py
"""
Python-side repro tests for the Adeu MCP QA Report — Round 3 (2026-07-24,
black-box QA of v1.30.0+1fd5285). Findings verified to reproduce on the
Python engine/server (replication notes per finding):

  1.2  keep-markup sanitize corrupts footer whitespace: the run coalescer
       merges a text run with the FOLLOWING <w:tab/>+text run, gluing
       "…Confidential" + "Page " together and relocating the tab after the
       merged text ("…ConfidentialPage " TAB instead of "…Confidential" TAB
       "Page ").
  1.3  A modify targeting a {#cell:...} anchor of a NON-empty cell silently
       prepends new_text at cell start ("By: /s/ SignerBy: ") instead of
       rejecting or inserting at the anchor's actual position (end of cell).
  2.1  Rejecting every id enumerated by a read (the natural agent loop) hard-
       fails on chained pair groups: some partners are gracefully counted
       already_resolved, but one group member (Chg:6 in the QA topology)
       falls through to "no tracked change with that id exists" and is
       counted as skipped — which rejects the whole batch transactionally at
       the server layer. Reproduces on BOTH engines (the QA report only
       blamed Node).
  2.2  read_docx renders format-only tracked changes as actionable-looking
       "[Chg:N format]" ids, but accept/reject by that id fails with "no
       tracked change with that id exists" — read and write disagree.
  2.4  Replacing a body paragraph with "## Heading\\n\\nBody" markdown leaves
       a stray empty paragraph behind in the accepted view (Node replaces in
       place).
  3.1  A rejected batch is returned as a SUCCESS payload while a BLOCKED
       sanitize is raised as a ToolError — the server is internally
       inconsistent about what travels on the error channel.
  3.4  Purely informational notes ("- Note: … the action itself succeeded")
       are rendered under a header literally called "Skipped Details".
  3.7  The success preview of a modify shows a neighboring window with NO
       trace of the change. Root cause found during replication: triggered
       when the edit's COMMENT contains CriticMarkup-like tokens (3.8) —
       the preview builder's bubble-stripping regexes misparse the nested
       tokens. (The QA report's t5 artifact name "commentinject" matches.)
  3.8  CriticMarkup tokens are validated/blocked in new_text but accepted
       verbatim in the comment field; on read-back the raw tokens nest
       inside the {>>…<<} annotation and break naive CriticMarkup parsing.
  3.10 A regex alternation returns the SAME paragraph once per branch hit
       (3 entries for one paragraph) instead of one deduped entry with all
       hits highlighted.
  3.11 The missing-file error speculates about sandboxed/containerized
       environments and pushes a `uv tool install adeu` CLI migration for a
       plain file-not-found, and appends a relative-path hint even when the
       given path was absolute.
  3.12 Search snippets are sliced at line boundaries, so a match whose line
       ends inside a multi-line {>>…<<} bubble renders an unterminated
       annotation ("{>>[Chg:1 delete] … (pairs with Chg:2") and the agent
       cannot harvest ids/pairings from search results.

Findings replicated but NOT covered here (cross-implementation or Node-only;
see node/packages/core/src/repro_qa_round3_2026_07_24.test.ts):
  1.1 (Node deletes anchored comments on accept — Python keeps them, pinned
  here as a control), 2.3 (Node lacks the \\1 backreference warning), 3.5/3.6
  (Node finalize report wording/omissions), 3.2 (footer part extraction
  parity), 3.3 (Node ENOENT directory dump).

Not reproduced / stale (no tests): 3.9 — the "appears N times" counter counts
occurrences of the MATCHED SUBSTRING, a semantic already pinned by
test_search_write_engine.py; S4 zip-size bloat — current outputs are
deflate-compressed (the report's byte counts came from May-era fixture pairs).

Every test is written test-first: it fails on current main and passes once
the finding is fixed.
"""

import asyncio
import re
import zipfile
from io import BytesIO

import pytest
from docx import Document
from docx.oxml.ns import qn
from fastmcp.exceptions import ToolError

from adeu.ingest import extract_text_from_stream
from adeu.mcp_components._response_builders import build_search_response
from adeu.mcp_components.tools.document import process_document_batch, read_docx
from adeu.models import AcceptChange, ModifyText, RejectChange
from adeu.redline.engine import RedlineEngine
from adeu.sanitize.core import sanitize_docx


class MockContext:
    """Mock FastMCP Context to absorb async logging calls during tests."""

    async def info(self, msg, **kwargs):
        pass

    async def debug(self, msg, **kwargs):
        pass

    async def warning(self, msg, **kwargs):
        pass

    async def error(self, msg, **kwargs):
        pass


def build_doc(*paras) -> BytesIO:
    d = Document()
    for p in paras:
        d.add_paragraph(p)
    buf = BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# 1.2 — keep-markup sanitize corrupts footer whitespace
# ---------------------------------------------------------------------------


class TestQA12SanitizeFooterWhitespace:
    def _footer_token_sequence(self, path) -> list:
        """Rendered-order token stream (text and TAB markers) of the footer."""
        with zipfile.ZipFile(path) as z:
            names = [n for n in z.namelist() if re.match(r"word/footer\d*\.xml", n)]
            assert names, "setup: the fixture must have a footer part"
            xml = z.read(names[0]).decode("utf-8")
        toks = re.findall(r"<w:t[^>]*>([^<]*)</w:t>|<w:(tab)[^>]*/>", xml)
        return [(t if t else "TAB") for t, _tab in toks]

    def test_keep_markup_sanitize_preserves_footer_tab_position(self, tmp_path):
        """Mirrors footer2.xml of the QA fixture: run A "… – Confidential",
        run B "<w:tab/>Page " (tab and text in ONE run, same formatting once
        rsids are stripped). Sanitize must be conservative on rendered text —
        the tab must stay between "Confidential" and "Page "."""
        inp = tmp_path / "footer_fixture.docx"
        outp = tmp_path / "footer_sanitized.docx"

        d = Document()
        d.add_paragraph("Body content.")
        p = d.sections[0].footer.paragraphs[0]
        r1 = p.add_run("NordicTech Solutions Inc. - Confidential")
        r1._element.set(qn("w:rsidRPr"), "00E44903")
        r2 = p.add_run()
        r2._element.set(qn("w:rsidRPr"), "00AB12CD")
        r2._element.append(r2._element.makeelement(qn("w:tab"), {}))
        t = r2._element.makeelement(qn("w:t"), {})
        t.text = "Page "
        t.set(qn("xml:space"), "preserve")
        r2._element.append(t)
        d.save(inp)

        assert self._footer_token_sequence(inp) == [
            "NordicTech Solutions Inc. - Confidential",
            "TAB",
            "Page ",
        ], "setup: fixture footer must render text TAB text"

        sanitize_docx(str(inp), str(outp), keep_markup=True, author="Redacted Reviewer")

        assert self._footer_token_sequence(outp) == [
            "NordicTech Solutions Inc. - Confidential",
            "TAB",
            "Page ",
        ], (
            "keep-markup sanitize reordered the footer: the run coalescer glued "
            '"Confidential" and "Page " together across the intervening tab, so '
            'the footer now renders "…ConfidentialPage " — visible corruption in '
            "the exact artifact sent to a counterparty (QA round 3, finding 1.2)"
        )


# ---------------------------------------------------------------------------
# 1.3 — cell-anchor fill on a NON-empty cell garbles the cell text
# ---------------------------------------------------------------------------


class TestQA13CellAnchorNonEmptyCell:
    def _build_signature_table(self) -> BytesIO:
        d = Document()
        d.add_paragraph("Signature block below.")
        table = d.add_table(rows=1, cols=2)
        c0 = table.cell(0, 0).paragraphs[0]
        c0.add_run("By: ")
        c0._element.set(qn("w14:paraId"), "1A26F0BF")
        c1 = table.cell(0, 1).paragraphs[0]
        c1.add_run("Date: ")
        c1._element.set(qn("w14:paraId"), "62EEA09B")
        buf = BytesIO()
        d.save(buf)
        buf.seek(0)
        return buf

    def test_anchor_write_to_non_empty_cell_does_not_interleave_text(self):
        """The {#cell:} anchor contract is documented for EMPTY cells. On a
        cell already containing "By: ", targeting the anchor must either be
        rejected with a helpful error or insert at the anchor's actual
        position (end of cell) — never silently prepend, which garbles the
        signature line to "By: /s/ Test SignerBy: "."""
        buf = self._build_signature_table()
        raw = extract_text_from_stream(buf, clean_view=False)
        assert "{#cell:1A26F0BF}" in raw, "setup: cell anchor must render"

        buf.seek(0)
        engine = RedlineEngine(buf, author="Adeu AI")
        stats = engine.process_batch(
            [ModifyText(target_text="{#cell:1A26F0BF}", new_text="By: /s/ Test Signer")]
        )

        if stats["edits_applied"]:
            out = engine.save_to_stream()
            clean = extract_text_from_stream(out, clean_view=True)
            assert "By: /s/ Test SignerBy:" not in clean, (
                "anchor-write to a non-empty cell was applied as a silent PREPEND, "
                f"interleaving the existing cell text:\n{clean}\n"
                "(QA round 3, finding 1.3: reject non-empty cells or insert at the "
                "anchor's position)"
            )
        else:
            # Rejection is an acceptable resolution, but it must tell the agent
            # how to proceed on a non-empty cell.
            details = "\n".join(str(s) for s in engine.skipped_details)
            assert "cell" in details.lower(), (
                "the anchor-write was skipped without explaining the non-empty-"
                f"cell situation to the agent: {details!r}"
            )


# ---------------------------------------------------------------------------
# 2.1 — "reject everything I enumerated" must not hard-fail on pair groups
# ---------------------------------------------------------------------------


def apply_chained_edit_batch() -> BytesIO:
    """Reproduces the QA t3 chain topology: pair groups 1<->(2,5) and
    3<->(4,6,7) plus a standalone insertion Chg:8."""
    stream = build_doc(
        "Unless otherwise stated in the Order Form, invoiced charges are due "
        "net ten (10) days from the invoice date."
    )
    engine = RedlineEngine(stream, author="Adeu AI")
    stats = engine.process_batch(
        [
            ModifyText(target_text="ten (10)", new_text="twenty (20)"),
            ModifyText(target_text="twenty (20)", new_text="twenty-one (21)"),
            ModifyText(
                target_text="from the invoice date.",
                new_text="from the invoice date. Final decision noted.",
            ),
        ]
    )
    assert stats["edits_applied"] == 3, "setup: the chained batch must apply"
    return engine.save_to_stream()


class TestQA21RejectAllEnumeratedIds:
    def test_rejecting_every_enumerated_id_never_hard_fails(self):
        """Enumerate ids exactly as a reading agent would, reject them all in
        order. Every pair-group member consumed by an earlier action must get
        the graceful already_resolved treatment; none may fall through to the
        fatal "no tracked change with that id exists" path."""
        stream = apply_chained_edit_batch()
        raw = extract_text_from_stream(stream, clean_view=False)
        ids = sorted({int(m) for m in re.findall(r"\[Chg:(\d+)", raw)})
        assert len(ids) == 8, f"setup: expected the QA topology's 8 ids, got {ids}"

        stream.seek(0)
        engine = RedlineEngine(stream, author="Counterparty")
        applied, skipped, already_resolved = engine.apply_review_actions(
            [RejectChange(target_id=f"Chg:{i}") for i in ids]
        )

        failures = [d for d in engine.skipped_details if "Failed to apply action" in str(d)]
        assert skipped == 0 and not failures, (
            "rejecting the full enumerated id list hard-failed on a pair-group "
            f"member (applied={applied}, skipped={skipped}, "
            f"already_resolved={already_resolved}):\n" + "\n".join(map(str, failures)) + "\n"
            "One code path re-snapshots the id map mid-batch while another does "
            "not, so the most natural agent pattern — reject every id the read "
            "listed — is nondeterministically fatal (QA round 3, finding 2.1)."
        )
        assert applied + already_resolved == len(ids)

    def test_reject_all_restores_the_original_text(self):
        stream = apply_chained_edit_batch()
        raw = extract_text_from_stream(stream, clean_view=False)
        ids = sorted({int(m) for m in re.findall(r"\[Chg:(\d+)", raw)})

        stream.seek(0)
        engine = RedlineEngine(stream, author="Counterparty")
        engine.apply_review_actions([RejectChange(target_id=f"Chg:{i}") for i in ids])
        out = engine.save_to_stream()
        clean = extract_text_from_stream(out, clean_view=True)
        assert clean.strip() == (
            "Unless otherwise stated in the Order Form, invoiced charges are due "
            "net ten (10) days from the invoice date."
        ), f"reject-all did not restore the pre-edit text:\n{clean}"


# ---------------------------------------------------------------------------
# 2.2 — format-only tracked changes are readable but not actionable
# ---------------------------------------------------------------------------


class TestQA22FormatChangeIds:
    def _build_doc_with_format_change(self) -> BytesIO:
        d = Document()
        d.add_paragraph("Plain intro paragraph.")
        p = d.add_paragraph("3.2 Invoicing and Payment.")
        rpr = p.runs[0]._element.get_or_add_rPr()
        rprchange = rpr.makeelement(qn("w:rPrChange"), {})
        rprchange.set(qn("w:id"), "901")
        rprchange.set(qn("w:author"), "Mikko Korpela")
        rprchange.set(qn("w:date"), "2026-01-22T16:16:00Z")
        rprchange.append(rpr.makeelement(qn("w:rPr"), {}))
        rpr.append(rprchange)
        buf = BytesIO()
        d.save(buf)
        buf.seek(0)
        return buf

    def test_advertised_format_id_is_actionable_or_marked_view_only(self):
        """read_docx advertises "[Chg:901 format]" exactly like actionable
        ids. Either accept-by-id must work (accept_all_changes already knows
        how to resolve formatting changes), or the projection must mark the
        id as non-actionable so the documented read-then-act loop cannot die
        on it."""
        buf = self._build_doc_with_format_change()
        raw = extract_text_from_stream(buf, clean_view=False)
        m = re.search(r"\[Chg:(\d+) format([^\]]*)\]", raw)
        assert m, f"setup: expected a format-change bubble in:\n{raw}"
        chg_id, qualifier = m.group(1), m.group(2)

        buf.seek(0)
        engine = RedlineEngine(buf, author="Reviewer")
        applied, skipped, _ = engine.apply_review_actions(
            [AcceptChange(target_id=f"Chg:{chg_id}")]
        )

        marked_view_only = bool(re.search(r"view.?only|not actionable", qualifier, re.I))
        assert applied == 1 or marked_view_only, (
            f"read_docx advertises [Chg:{chg_id} format] with no non-actionable "
            "marker, but accepting that id fails with 'no tracked change with "
            "that id exists' — read and write disagree about what exists "
            f"(applied={applied}, skipped={skipped}):\n"
            + "\n".join(map(str, engine.skipped_details))
        )


# ---------------------------------------------------------------------------
# 2.4 — heading+body replacement leaves a stray empty paragraph
# ---------------------------------------------------------------------------


class TestQA24HeadingBodyReplacement:
    def test_accepted_view_has_no_stray_empty_paragraph(self):
        stream = build_doc(
            "Customer may not assign any of its rights.",
            "8.3 Entire Agreement.",
            "This Agreement is the entire agreement between the parties.",
        )
        engine = RedlineEngine(stream, author="Adeu AI")
        stats = engine.process_batch(
            [
                ModifyText(
                    target_text="8.3 Entire Agreement.",
                    new_text=(
                        "## 8.3 Entire Agreement; Amendments\n\n"
                        "No amendment is effective unless in writing."
                    ),
                )
            ]
        )
        assert stats["edits_applied"] == 1, "setup: the replacement must apply"

        out = engine.save_to_stream()
        clean = extract_text_from_stream(out, clean_view=True)
        expected = (
            "Customer may not assign any of its rights.\n\n"
            "## 8.3 Entire Agreement; Amendments\n\n"
            "No amendment is effective unless in writing.\n\n"
            "This Agreement is the entire agreement between the parties."
        )
        # extract may or may not render the heading marker; normalize it away
        # so the assertion pins the PARAGRAPH STRUCTURE, not heading rendering.
        normalize = lambda s: re.sub(r"^#+\s*", "", s.strip(), flags=re.M)
        assert normalize(clean) == normalize(expected), (
            "the accepted view contains a stray empty paragraph where the "
            "original paragraph used to be (Node replaces in place — parity "
            f"break, QA round 3 finding 2.4):\n{clean!r}"
        )


# ---------------------------------------------------------------------------
# 3.7 / 3.8 — CriticMarkup tokens in the `comment` field
# ---------------------------------------------------------------------------

TOKEN_COMMENT = "watch {>>nested<<} and {--del--} tokens"


class TestQA38CommentCriticMarkupInjection:
    def _apply_with_token_comment(self):
        stream = build_doc(
            "then-current online user guides and manuals.",
            "ARTICLE2 PROPRIETARYRIGHTS",
        )
        engine = RedlineEngine(stream, author="Adeu AI")
        stats = engine.process_batch(
            [
                ModifyText(
                    target_text="and manuals",
                    new_text="and handbooks",
                    comment=TOKEN_COMMENT,
                )
            ]
        )
        return engine, stats

    def test_comment_tokens_do_not_break_the_edit_preview(self):
        """Finding 3.7: the ✅ preview must show the change. With CriticMarkup
        tokens in the comment, the preview builder's bubble-stripping regexes
        misparse the nested tokens and render a neighboring window containing
        no trace of the change ('.\\n\\nARTICLE2 PROPRIETARYRIGHTS')."""
        _, stats = self._apply_with_token_comment()
        assert stats["edits_applied"] == 1, "setup: the edit must apply"
        preview = stats["edits"][0]["critic_markup"] or ""
        assert "{--manuals--}" in preview and "{++handbooks++}" in preview, (
            "the success preview shows the wrong window — no trace of the "
            f"change it reports on: {preview!r} (QA round 3, finding 3.7; "
            "trigger: CriticMarkup tokens in the comment field)"
        )

    def test_comment_tokens_are_rejected_or_escaped_on_render(self):
        """Finding 3.8: new_text blocks CriticMarkup tokens, the comment field
        writes them verbatim; on read-back the raw tokens nest inside the
        {>>…<<} annotation, so any CriticMarkup consumer (including the model)
        misparses the document. Either reject the comment like new_text, or
        escape the tokens when rendering."""
        engine, stats = self._apply_with_token_comment()
        if stats["edits_applied"] == 0:
            return  # rejected at validation — acceptable resolution

        out = engine.save_to_stream()
        raw = extract_text_from_stream(out, clean_view=False)
        assert "{>>nested<<}" not in raw and "{--del--}" not in raw, (
            "the comment's CriticMarkup tokens render VERBATIM nested inside "
            "the {>>…<<} annotation — the inner '<<}' terminates the outer "
            f"bubble early for any CriticMarkup parser:\n{raw!r}"
        )


# ---------------------------------------------------------------------------
# 3.1 — error-channel consistency (batch rejection vs sanitize BLOCKED)
# ---------------------------------------------------------------------------


class TestQA31ErrorChannelConsistency:
    def test_rejected_batch_and_blocked_sanitize_use_the_same_channel(self, tmp_path):
        """A rejected batch is returned as a normal success payload while a
        BLOCKED sanitize is raised as a ToolError. Whatever the choice, a
        refused operation must travel on ONE channel — an orchestrator
        branching on the MCP error flag currently behaves differently for
        the two refusals."""
        from adeu.mcp_components.tools.sanitize import sanitize_docx as sanitize_tool

        # A document with one unresolved tracked change.
        stream = build_doc("The fee is ten dollars.")
        engine = RedlineEngine(stream, author="Adeu AI")
        engine.process_batch([ModifyText(target_text="ten", new_text="twenty")])
        path = tmp_path / "pending.docx"
        path.write_bytes(engine.save_to_stream().getvalue())

        ctx = MockContext()

        batch_error = False
        try:
            result = asyncio.run(
                process_document_batch(
                    reasoning="test",
                    original_docx_path=str(path),
                    changes=[RejectChange(target_id="Chg:99")],
                    ctx=ctx,
                    output_path=str(tmp_path / "batch_out.docx"),
                )
            )
            assert "Batch rejected" in str(result), "setup: the batch must be refused"
        except ToolError:
            batch_error = True

        sanitize_error = False
        try:
            result = asyncio.run(
                sanitize_tool(
                    reasoning="test",
                    file_path=str(path),
                    ctx=ctx,
                    output_path=str(tmp_path / "san_out.docx"),
                )
            )
            assert "BLOCKED" in str(result), "setup: the sanitize must be refused"
        except ToolError:
            sanitize_error = True

        assert batch_error == sanitize_error, (
            "refused operations travel on different channels: a rejected batch "
            f"is a normal payload (error={batch_error}) while a BLOCKED sanitize "
            f"is a ToolError (error={sanitize_error}) — align them "
            "(QA round 3, finding 3.1)"
        )


# ---------------------------------------------------------------------------
# 3.4 — informational notes rendered under a "Skipped Details" header
# ---------------------------------------------------------------------------


class TestQA34InformationalNotesHeader:
    def test_informational_notes_are_not_labeled_skipped(self, tmp_path):
        """Rejecting both ids of a replacement pair succeeds (the partner is
        counted already_resolved) — yet the informational note explaining
        that is printed under a header literally called "Skipped Details"
        while saying "the action itself succeeded". Rename to "Notes"."""
        stream = build_doc("The fee is ten dollars.")
        engine = RedlineEngine(stream, author="Adeu AI")
        engine.process_batch([ModifyText(target_text="ten", new_text="twenty")])
        path = tmp_path / "pair.docx"
        path.write_bytes(engine.save_to_stream().getvalue())

        raw = extract_text_from_stream(BytesIO(path.read_bytes()), clean_view=False)
        ids = sorted({int(m) for m in re.findall(r"\[Chg:(\d+)", raw)})
        assert len(ids) >= 2, f"setup: expected a del/ins pair, got {ids}"

        result = asyncio.run(
            process_document_batch(
                reasoning="test",
                original_docx_path=str(path),
                changes=[RejectChange(target_id=f"Chg:{i}") for i in ids],
                ctx=MockContext(),
                output_path=str(tmp_path / "pair_out.docx"),
            )
        )
        text = str(result)
        assert "had no additional effect" in text, (
            f"setup: expected the already-resolved informational note in:\n{text}"
        )
        assert "Skipped Details" not in text, (
            'purely informational notes ("the action itself succeeded") are '
            'filed under a header called "Skipped Details" — misleading; '
            'rename to "Notes" (QA round 3, finding 3.4)'
        )


# ---------------------------------------------------------------------------
# 3.10 / 3.12 — search response defects (builder layer)
# ---------------------------------------------------------------------------


class TestQA310SearchDuplicateMatches:
    def test_multiple_alternation_hits_in_one_paragraph_render_one_entry(self):
        body = (
            "# Section One\n\n"
            "This deal between Dealfluence and the Com: holder grants limited "
            "rights to the platform.\n\n"
            "Unrelated closing paragraph."
        )
        res = build_search_response(
            body,
            "Com:|Dealfluence|limited rights",
            True,
            True,
            None,
            "dummy.docx",
        )
        text = str(res.content)
        entries = text.count("### Match")
        assert entries == 1, (
            f"one paragraph with 3 alternation-branch hits rendered {entries} "
            "separate match entries (each with different bold placement and "
            "occurrence counts) — dedupe matches per paragraph and highlight "
            "all hits in one entry (QA round 3, finding 3.10)"
        )


class TestQA312SearchSnippetTruncation:
    def test_snippet_never_leaves_an_annotation_unterminated(self):
        """The snippet is the line containing the match, but {>>…<<} bubbles
        are MULTI-line, so the snippet is cut mid-annotation and the agent
        cannot harvest ids/pairings from search results."""
        body = (
            "# Section One\n\n"
            "Invoiced charges are due net {--ten--}{++twenty++}{>>[Chg:1 delete] "
            "Adeu AI (pairs with Chg:2)\n"
            "[Chg:2 insert] Adeu AI (pairs with Chg:1)\n"
            "[Com:0] Reviewer @ 2026-01-22T14:13:00Z: chained change<<} days "
            "from the invoice date.\n\n"
            "Unrelated closing paragraph."
        )
        res = build_search_response(
            body, "Invoiced charges", False, True, None, "dummy.docx"
        )
        text = str(res.content)
        opened = text.count("{>>")
        closed = text.count("<<}")
        assert opened == closed, (
            f"search snippet truncates mid-annotation ({opened} '{{>>' opened, "
            f"{closed} '<<}}' closed) — extend the window to close any markup "
            f"span it opens (QA round 3, finding 3.12):\n{text}"
        )


# ---------------------------------------------------------------------------
# 3.11 — missing-file error speculates and mis-advises
# ---------------------------------------------------------------------------


class TestQA311MissingFileError:
    def test_file_not_found_error_does_not_speculate_about_sandboxes(self, tmp_path):
        missing = tmp_path / "does_not_exist_qa.docx"
        with pytest.raises(ToolError) as exc_info:
            asyncio.run(
                read_docx(
                    reasoning="test",
                    file_path=str(missing),
                    ctx=MockContext(),
                )
            )
        msg = str(exc_info.value)
        assert "not found" in msg.lower(), f"setup: expected a not-found error, got {msg!r}"

        for phrase in ("sandboxed", "containerized", "uv tool install"):
            assert phrase not in msg, (
                f"the missing-file error speculates about {phrase!r} and pushes a "
                "CLI migration for a file that simply does not exist — say "
                '"file not found" first and gate the sandbox essay behind an '
                f"actual access-denied signal (QA round 3, finding 3.11):\n{msg}"
            )

    def test_absolute_path_hint_not_shown_for_absolute_paths(self, tmp_path):
        missing = tmp_path / "does_not_exist_qa.docx"
        with pytest.raises(ToolError) as exc_info:
            asyncio.run(
                read_docx(
                    reasoning="test",
                    file_path=str(missing),
                    ctx=MockContext(),
                )
            )
        msg = str(exc_info.value)
        assert "Provide an absolute path" not in msg, (
            "the relative-path hint fires even though the caller already "
            f"provided an absolute path ({missing}):\n{msg}"
        )


# ---------------------------------------------------------------------------
# 1.1 control — Python KEEPS an anchored comment when the change is accepted.
# This pins the correct behavior that the Node engine must match (the Node
# failing test lives in repro_qa_round3_2026_07_24.test.ts).
# ---------------------------------------------------------------------------


class TestQA11PythonAcceptKeepsComment:
    def test_accepting_a_change_keeps_the_wrapping_comment(self):
        stream = build_doc(
            "Agreement between NordicTech and Adeu.ai, a Delaware corporation."
        )
        engine = RedlineEngine(stream, author="Claude")
        engine.process_batch(
            [
                ModifyText(
                    target_text="Adeu.ai, a Delaware corporation",
                    new_text="Dealfluence Oy, a Finnish corporation",
                    comment="Updated party name and jurisdiction",
                )
            ]
        )
        out = engine.save_to_stream()
        raw = extract_text_from_stream(out, clean_view=False)
        del_id = re.search(r"\[Chg:(\d+) delete\]", raw).group(1)

        out.seek(0)
        engine2 = RedlineEngine(out, author="Reviewer")
        applied, _, _ = engine2.apply_review_actions(
            [AcceptChange(target_id=f"Chg:{del_id}")]
        )
        assert applied == 1, "setup: the accept must apply"

        raw_after = extract_text_from_stream(engine2.save_to_stream(), clean_view=False)
        assert "Updated party name and jurisdiction" in raw_after, (
            "accepting a tracked change deleted the comment anchored on the "
            "surviving text — Word semantics keep it (QA round 3, finding 1.1; "
            f"this is the Python control the Node engine must match):\n{raw_after}"
        )

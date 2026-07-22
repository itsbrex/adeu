# FILE: tests/test_repro_qa_2026_07_18.py
"""
Repro tests for the 2026-07-18 black-box QA report (adeu 1.22.0+4947f97).

Finding index (report section -> test class below):
  C1  final body insertion written into the footer / unreadable DOCX
  C2  DOCX-to-DOCX table row changes do not round-trip
  H1  diff generates read-only appendix edits apply cannot consume
  H2  sanitize batch silently overwrites identical basenames
  H3  sanitize --baseline: bogus similarity %, raw rId crash, invalid output
  H4  tracked footnote edit produces a DOCX LibreOffice cannot open
  M1  markup does not honor apply's edit semantics
  M2  failed apply still writes an output and prints "Batch complete"
  M3  VML watermark survives sanitize unreported
  M4  extraction loses list semantics
  M5  images/alt text disappear from extraction
  M6  phantom reserved footnotes ([^fn--1], [^fn-0])
  M7  only first inline defined term per paragraph detected
  M8  --json modes emit non-JSON errors
  L1  search no-match advice names nonexistent options
  L2  mode-specific flags silently ignored
  L3  .txt input to extract described as bad DOCX zip
  L4  extract --json -o writes file AND prints JSON to stdout
  L5  re-running init backs up even when config unchanged
  L6  init --local reports success from an arbitrary directory

Every LibreOffice interop assertion is skipped when soffice (with the Writer
import filter) is unavailable.
"""

import json
import re
import shutil
import struct
import subprocess
import sys
import zlib
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn

from adeu.ingest import _extract_text_from_doc, extract_text_from_stream

W_NS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'


# ---------------------------------------------------------------------------
# Helpers / fixture builders
# ---------------------------------------------------------------------------


def run_cli(args, capsys):
    """Invoke the CLI in-process; returns (exit_code, stdout, stderr)."""
    from unittest.mock import patch

    from adeu.cli import main

    code = 0
    with patch.object(sys, "argv", ["adeu"] + [str(a) for a in args]):
        try:
            main()
        except SystemExit as e:
            code = e.code or 0
    captured = capsys.readouterr()
    return code, captured.out, captured.err


_SOFFICE_STATE = {"probed": False, "ok": False}


def soffice_can_convert(tmp_dir: Path) -> bool:
    """True when a working LibreOffice Writer import filter is available."""
    if _SOFFICE_STATE["probed"]:
        return _SOFFICE_STATE["ok"]
    _SOFFICE_STATE["probed"] = True
    if shutil.which("soffice") is None:
        return False
    probe = tmp_dir / "probe.docx"
    doc = Document()
    doc.add_paragraph("probe")
    doc.save(probe)
    _SOFFICE_STATE["ok"] = lo_loads(probe, tmp_dir)
    return _SOFFICE_STATE["ok"]


def lo_loads(path: Path, out_dir: Path) -> bool:
    """True when LibreOffice can load `path` (via pdf conversion)."""
    pdf_dir = out_dir / "lo_pdf"
    pdf_dir.mkdir(exist_ok=True)
    expected = pdf_dir / (path.stem + ".pdf")
    if expected.exists():
        expected.unlink()
    try:
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(pdf_dir), str(path)],
            capture_output=True,
            timeout=5,
        )
    except (subprocess.TimeoutExpired, subprocess.SubprocessError, OSError):
        return False
    return expected.exists()


def build_header_footer_doc(path: Path, extra_body_paragraph: bool = False):
    doc = Document()
    sec = doc.sections[0]
    sec.header.paragraphs[0].text = "HEADER MARKER"
    sec.footer.paragraphs[0].text = "FOOTER MARKER"
    doc.add_paragraph("Body paragraph one.")
    if extra_body_paragraph:
        doc.add_paragraph("New final body paragraph.")
    doc.save(path)


def read_part_xml(path: Path, member: str) -> str:
    with ZipFile(path) as z:
        return z.read(member).decode("utf-8")


def footer_members(path: Path):
    with ZipFile(path) as z:
        return [n for n in z.namelist() if n.startswith("word/footer")]


FOOTNOTE_SEP_TYPED = (
    f'<w:footnote w:type="separator" w:id="-1" {W_NS}><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
)
FOOTNOTE_CONT_TYPED = (
    f'<w:footnote w:type="continuationSeparator" w:id="0" {W_NS}>'
    "<w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>"
)
# Reserved notes as emitted by some generators: ids -1/0 but NO w:type attribute.
FOOTNOTE_SEP_UNTYPED = f'<w:footnote w:id="-1" {W_NS}><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
FOOTNOTE_CONT_UNTYPED = f'<w:footnote w:id="0" {W_NS}><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'


def build_footnote_doc(
    path: Path, typed_separators: bool = True, note_text: str = "This is a QA footnote about governing law."
):
    """A body paragraph carrying footnote 1, plus separator/continuation notes."""
    doc = Document()
    p = doc.add_paragraph("The governing law clause")
    run_xml = f'<w:r {W_NS}><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr><w:footnoteReference w:id="1"/></w:r>'
    p._element.append(parse_xml(run_xml))
    doc.add_paragraph("Second body paragraph.")

    sep = FOOTNOTE_SEP_TYPED if typed_separators else FOOTNOTE_SEP_UNTYPED
    cont = FOOTNOTE_CONT_TYPED if typed_separators else FOOTNOTE_CONT_UNTYPED
    footnotes_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f"<w:footnotes {W_NS}>"
        f"{sep}{cont}"
        f'<w:footnote w:id="1"><w:p><w:pPr><w:pStyle w:val="FootnoteText"/></w:pPr>'
        f'<w:r><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr><w:footnoteRef/></w:r>'
        f'<w:r><w:t xml:space="preserve"> {note_text}</w:t></w:r></w:p></w:footnote>'
        "</w:footnotes>"
    )

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    ct_override = (
        '<Override PartName="/word/footnotes.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>'
    )
    rel = (
        '<Relationship Id="rIdFn99" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" '
        'Target="footnotes.xml"/>'
    )
    out = BytesIO()
    with ZipFile(stream) as zin, ZipFile(out, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "[Content_Types].xml":
                data = data.replace(b"</Types>", ct_override.encode() + b"</Types>")
            elif item.filename == "word/_rels/document.xml.rels":
                data = data.replace(b"</Relationships>", rel.encode() + b"</Relationships>")
            zout.writestr(item, data)
        zout.writestr("word/footnotes.xml", footnotes_xml)
    path.write_bytes(out.getvalue())


def build_table_doc(path: Path, extra_row: bool = False, drop_middle_row: bool = False):
    doc = Document()
    doc.add_paragraph("Service Level Agreement")
    rows = [
        ["Service", "Uptime", "Price"],
        ["Standard support", "99.0%", "EUR 1,000"],
        ["Premium support", "99.9%", "EUR 2,000"],
    ]
    if drop_middle_row:
        rows = [rows[0], rows[2]]
    table = doc.add_table(rows=len(rows), cols=3)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
    if extra_row:
        r = table.add_row()
        for j, val in enumerate(["Backup service", "99.5%", "EUR 500"]):
            r.cells[j].text = val
    doc.add_paragraph("Documentation: see appendix.")
    doc.save(path)


def make_png() -> bytes:
    sig = b"\x89PNG\r\n\x1a\n"

    def chunk(t, d):
        c = struct.pack(">I", len(d)) + t + d
        return c + struct.pack(">I", zlib.crc32(t + d) & 0xFFFFFFFF)

    ihdr = struct.pack(">IIBBBBB", 4, 4, 8, 2, 0, 0, 0)
    raw = b"".join(b"\x00" + b"\xff\x00\x00" * 4 for _ in range(4))
    return sig + chunk(b"IHDR", ihdr) + chunk(b"IDAT", zlib.compress(raw)) + chunk(b"IEND", b"")


def build_image_doc(path: Path):
    from docx.shared import Inches

    png = path.parent / "red.png"
    png.write_bytes(make_png())
    doc = Document()
    doc.add_paragraph("Before image.")
    run = doc.add_paragraph().add_run()
    run.add_picture(str(png), width=Inches(1))
    inline = run._element.findall(".//" + qn("wp:inline"))[0]
    doc_pr = inline.find(qn("wp:docPr"))
    doc_pr.set("descr", "Red rectangle QA diagram")
    doc_pr.set("title", "Red rectangle QA diagram")
    doc.add_paragraph("After image.")
    doc.save(path)


VML_WATERMARK_HEADER = f"""<w:p {W_NS} xmlns:v="urn:schemas-microsoft-com:vml" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:pPr><w:pStyle w:val="Header"/></w:pPr>
  <w:r>
    <w:pict>
      <v:shape id="PowerPlusWaterMarkObject" type="#_x0000_t136" style="position:absolute;width:400pt;height:100pt">
        <v:textpath style="font-family:Calibri" string="DRAFT ACME SECRET"/>
      </v:shape>
    </w:pict>
  </w:r>
</w:p>"""


def build_watermark_doc(path: Path):
    doc = Document()
    header = doc.sections[0].header
    header.paragraphs[0].text = "Header text"
    header.paragraphs[0]._element.addnext(parse_xml(VML_WATERMARK_HEADER))
    doc.add_paragraph("Confidential body.")
    doc.save(path)


# ---------------------------------------------------------------------------
# C1 — cross-part (body/footer) edits
# ---------------------------------------------------------------------------


class TestC1FooterBoundary:
    def _diff_json(self, capsys, orig, mod):
        code, out, err = run_cli(["diff", orig, mod, "--json"], capsys)
        assert code == 0, err
        return json.loads(out)

    def test_diff_edits_never_span_body_footer(self, tmp_path, capsys):
        """C1: diff must not emit a target that crosses the body/footer wall."""
        orig = tmp_path / "hf_original.docx"
        mod = tmp_path / "hf_modified.docx"
        build_header_footer_doc(orig)
        build_header_footer_doc(mod, extra_body_paragraph=True)

        edits = self._diff_json(capsys, orig, mod)
        assert edits, "expected at least one edit"
        for e in edits:
            tgt = e.get("target_text", "")
            assert "FOOTER" not in tgt or "one." not in tgt, f"edit spans body/footer boundary: {e}"

    def test_apply_diff_json_keeps_new_paragraph_in_body(self, tmp_path, capsys):
        """C1: the diff→apply round trip must put the paragraph in document.xml."""
        orig = tmp_path / "hf_original.docx"
        mod = tmp_path / "hf_modified.docx"
        build_header_footer_doc(orig)
        build_header_footer_doc(mod, extra_body_paragraph=True)

        edits = self._diff_json(capsys, orig, mod)
        edits_file = tmp_path / "edits.json"
        edits_file.write_text(json.dumps(edits))

        applied = tmp_path / "applied.docx"
        code, out, err = run_cli(["apply", orig, edits_file, "-o", applied], capsys)
        assert code == 0, f"apply failed: {err}"

        body_xml = read_part_xml(applied, "word/document.xml")
        assert "New final body paragraph." in body_xml, "insertion missing from document body"
        for member in footer_members(applied):
            f_xml = read_part_xml(applied, member)
            assert "New final body paragraph." not in f_xml, f"body text leaked into {member}"

    def test_apply_text_roundtrip_keeps_new_paragraph_in_body(self, tmp_path, capsys):
        """C1: extract --clean-view → edit text → apply must anchor in the body."""
        orig = tmp_path / "hf_original.docx"
        mod = tmp_path / "hf_modified.docx"
        build_header_footer_doc(orig)
        build_header_footer_doc(mod, extra_body_paragraph=True)

        md = tmp_path / "modified.md"
        code, out, err = run_cli(["extract", mod, "--page", "all", "--clean-view", "-o", md], capsys)
        assert code == 0

        applied = tmp_path / "applied_text.docx"
        code, out, err = run_cli(["apply", orig, md, "-o", applied], capsys)
        assert code == 0, f"apply failed: {err}"

        body_xml = read_part_xml(applied, "word/document.xml")
        assert "New final body paragraph." in body_xml
        for member in footer_members(applied):
            assert "New final body paragraph." not in read_part_xml(applied, member)

        # The new paragraph must be a paragraph of its own, not inline glue.
        clean = extract_text_from_stream(BytesIO(applied.read_bytes()), clean_view=True)
        assert "Body paragraph one.\n\nNew final body paragraph." in clean
        assert "New final body paragraph.FOOTER" not in clean

    def test_cross_part_replacement_is_rejected(self, tmp_path, capsys):
        """C1: a replacement whose changed text spans body+footer must be refused."""
        orig = tmp_path / "hf_original.docx"
        build_header_footer_doc(orig)
        edits_file = tmp_path / "edits.json"
        edits_file.write_text(
            json.dumps(
                [
                    {
                        "type": "modify",
                        "target_text": "paragraph one.\n\nFOOTER MARKER",
                        "new_text": "paragraph two.\n\nALTERED FOOTER",
                    }
                ]
            )
        )
        applied = tmp_path / "applied.docx"
        code, out, err = run_cli(["apply", orig, edits_file, "-o", applied], capsys)
        assert code != 0, "cross-part replacement must fail"
        assert not applied.exists(), "no output may be written for a rejected batch"
        assert "boundary" in err.lower() or "part" in err.lower()

    def test_diff_detects_content_in_wrong_part(self, tmp_path, capsys):
        """C1: structural comparison — same flattened text, different parts."""
        good = tmp_path / "good.docx"
        build_header_footer_doc(good, extra_body_paragraph=True)

        # Corrupt twin: the same flattened text, but the new paragraph lives in
        # the footer (what 1.22.0 used to produce).
        corrupt = tmp_path / "corrupt.docx"
        doc = Document()
        sec = doc.sections[0]
        sec.header.paragraphs[0].text = "HEADER MARKER"
        sec.footer.paragraphs[0].text = "New final body paragraph."
        p2 = f"<w:p {W_NS}><w:r><w:t>FOOTER MARKER</w:t></w:r></w:p>"
        sec.footer.paragraphs[0]._element.addnext(parse_xml(p2))
        doc.add_paragraph("Body paragraph one.")
        doc.save(corrupt)

        code, out, err = run_cli(["diff", corrupt, good, "--json"], capsys)
        assert code == 0
        edits = json.loads(out)
        assert edits != [], "diff must not report structurally different documents as identical"

    def test_applied_output_loads_in_libreoffice(self, tmp_path, capsys):
        """C1: the applied output must be loadable by an independent implementation."""
        if not soffice_can_convert(tmp_path):
            pytest.skip("no working soffice")
        orig = tmp_path / "hf_original.docx"
        mod = tmp_path / "hf_modified.docx"
        build_header_footer_doc(orig)
        build_header_footer_doc(mod, extra_body_paragraph=True)

        edits = self._diff_json(capsys, orig, mod)
        edits_file = tmp_path / "edits.json"
        edits_file.write_text(json.dumps(edits))
        applied = tmp_path / "applied.docx"
        code, _, err = run_cli(["apply", orig, edits_file, "-o", applied], capsys)
        assert code == 0, err
        assert lo_loads(applied, tmp_path), "LibreOffice could not load the applied output"


# ---------------------------------------------------------------------------
# C2 — table row round trip
# ---------------------------------------------------------------------------


class TestC2TableRows:
    def test_diff_emits_insert_row_for_added_row(self, tmp_path, capsys):
        orig = tmp_path / "tbl_orig.docx"
        mod = tmp_path / "tbl_added.docx"
        build_table_doc(orig)
        build_table_doc(mod, extra_row=True)

        code, out, err = run_cli(["diff", orig, mod, "--json"], capsys)
        assert code == 0
        edits = json.loads(out)
        row_ops = [e for e in edits if e.get("type") == "insert_row"]
        assert row_ops, f"expected an insert_row op, got: {edits}"
        assert row_ops[0]["cells"] == ["Backup service", "99.5%", "EUR 500"]
        # No generic modify may smuggle pipe-delimited row text.
        for e in edits:
            if e.get("type") == "modify":
                assert "Backup service | 99.5%" not in (e.get("new_text") or "")

    def test_row_addition_roundtrip(self, tmp_path, capsys):
        orig = tmp_path / "tbl_orig.docx"
        mod = tmp_path / "tbl_added.docx"
        build_table_doc(orig)
        build_table_doc(mod, extra_row=True)

        code, out, err = run_cli(["diff", orig, mod, "--json"], capsys)
        edits_file = tmp_path / "edits.json"
        edits_file.write_text(out)

        applied = tmp_path / "applied.docx"
        code, out, err = run_cli(["apply", orig, edits_file, "-o", applied], capsys)
        assert code == 0, f"apply of diff output failed: {err}"

        clean = extract_text_from_stream(BytesIO(applied.read_bytes()), clean_view=True)
        # The new row must be a table row (joined by row separator \n, not \n\n)
        assert "Premium support | 99.9% | EUR 2,000\nBackup service | 99.5% | EUR 500" in clean

        # Follow-up clean diff must be empty.
        code, out, err = run_cli(["diff", applied, mod, "--json"], capsys)
        assert code == 0
        assert json.loads(out) == [], f"round trip not converged: {out}"

    def test_diff_emits_delete_row_for_removed_row(self, tmp_path, capsys):
        orig = tmp_path / "tbl_orig.docx"
        mod = tmp_path / "tbl_dropped.docx"
        build_table_doc(orig)
        build_table_doc(mod, drop_middle_row=True)

        code, out, err = run_cli(["diff", orig, mod, "--json"], capsys)
        assert code == 0
        edits = json.loads(out)
        row_ops = [e for e in edits if e.get("type") == "delete_row"]
        assert row_ops, f"expected a delete_row op, got: {edits}"

    def test_row_deletion_roundtrip(self, tmp_path, capsys):
        orig = tmp_path / "tbl_orig.docx"
        mod = tmp_path / "tbl_dropped.docx"
        build_table_doc(orig)
        build_table_doc(mod, drop_middle_row=True)

        code, out, err = run_cli(["diff", orig, mod, "--json"], capsys)
        edits_file = tmp_path / "edits.json"
        edits_file.write_text(out)

        applied = tmp_path / "applied.docx"
        code, out, err = run_cli(["apply", orig, edits_file, "-o", applied], capsys)
        assert code == 0, f"apply of diff output failed: {err}"

        clean = extract_text_from_stream(BytesIO(applied.read_bytes()), clean_view=True)
        assert "Standard support" not in clean

        code, out, err = run_cli(["diff", applied, mod, "--json"], capsys)
        assert code == 0
        assert json.loads(out) == []

    def test_generic_modify_cannot_inject_row_into_cell(self, tmp_path, capsys):
        """C2: the 1.22.0 failure shape — pipe row injected into a cell — must be refused."""
        orig = tmp_path / "tbl_orig.docx"
        build_table_doc(orig)
        edits_file = tmp_path / "edits.json"
        edits_file.write_text(
            json.dumps(
                [
                    {
                        "type": "modify",
                        "target_text": "2,000\n\nDocumentation:",
                        "new_text": "2,000\nBackup service | 99.5% | EUR 500\n\nDocumentation:",
                    }
                ]
            )
        )
        applied = tmp_path / "applied.docx"
        code, out, err = run_cli(["apply", orig, edits_file, "-o", applied], capsys)
        assert code != 0, "row-shaped text injection must be rejected"
        assert "insert_row" in err, "error should point the agent at insert_row"
        assert not applied.exists()


# ---------------------------------------------------------------------------
# H1 — appendix leakage into diff
# ---------------------------------------------------------------------------


def build_defined_terms_doc(path: Path, extra_uses: bool = False):
    doc = Document()
    doc.add_paragraph('"Agreement" means this service agreement between the parties.')
    doc.add_paragraph("The Agreement enters into force upon signature.")
    if extra_uses:
        doc.add_paragraph("Termination of the Agreement requires notice.")
        doc.add_paragraph("Amendments to the Agreement must be written.")
    doc.save(path)


class TestH1AppendixLeak:
    def test_diff_json_contains_no_appendix_edits(self, tmp_path, capsys):
        orig = tmp_path / "a.docx"
        mod = tmp_path / "b.docx"
        build_defined_terms_doc(orig)
        build_defined_terms_doc(mod, extra_uses=True)

        code, out, err = run_cli(["diff", orig, mod, "--json"], capsys)
        assert code == 0
        edits = json.loads(out)
        for e in edits:
            joined = (e.get("target_text") or "") + (e.get("new_text") or "")
            assert "— used " not in joined, f"appendix text leaked into diff: {e}"
            assert "Document Structure" not in joined

    def test_diff_output_passes_dry_run_apply(self, tmp_path, capsys):
        """QA regression invariant 1: diff output must be applicable to the original."""
        orig = tmp_path / "a.docx"
        mod = tmp_path / "b.docx"
        build_defined_terms_doc(orig)
        build_defined_terms_doc(mod, extra_uses=True)

        code, out, err = run_cli(["diff", orig, mod, "--json"], capsys)
        assert code == 0
        edits_file = tmp_path / "edits.json"
        edits_file.write_text(out)

        code, out, err = run_cli(["apply", orig, edits_file, "--dry-run", "--json"], capsys)
        assert code == 0, f"diff output failed its own dry-run apply: {err} {out}"
        stats = json.loads(out)
        assert stats["edits_skipped"] == 0


# ---------------------------------------------------------------------------
# H2 — sanitize batch basename collision
# ---------------------------------------------------------------------------


class TestH2SanitizeBatchCollision:
    def test_collision_is_rejected_before_processing(self, tmp_path, capsys):
        (tmp_path / "a").mkdir()
        (tmp_path / "b").mkdir()
        d = Document()
        d.add_paragraph("CONTENT FROM A")
        d.save(tmp_path / "a" / "same.docx")
        d = Document()
        d.add_paragraph("CONTENT FROM B")
        d.save(tmp_path / "b" / "same.docx")

        outdir = tmp_path / "out"
        code, out, err = run_cli(
            ["sanitize", tmp_path / "a" / "same.docx", tmp_path / "b" / "same.docx", "--outdir", outdir],
            capsys,
        )
        assert code != 0, "identical output basenames must not be reported as success"
        assert "same.docx" in err
        assert "collide" in err.lower() or "collision" in err.lower() or "overwrite" in err.lower()
        # Nothing silently overwritten.
        assert not (outdir / "same.docx").exists()


# ---------------------------------------------------------------------------
# H3 — sanitize --baseline
# ---------------------------------------------------------------------------


class TestH3Baseline:
    def _make_pair(self, tmp_path):
        paras = [
            f"Clause {i}: The party shall perform obligation number {i} in accordance with the agreement."
            for i in range(40)
        ]
        base = tmp_path / "base.docx"
        work = tmp_path / "work.docx"
        d = Document()
        for p in paras:
            d.add_paragraph(p)
        d.save(base)
        d = Document()
        d.add_paragraph("Preamble inserted at top.")
        for p in paras:
            d.add_paragraph(p)
        d.save(work)
        return base, work

    def test_similarity_warning_uses_sequence_similarity(self, tmp_path, capsys):
        """H3: a ~2% real difference must not be reported as 93%."""
        base, work = self._make_pair(tmp_path)
        code, out, err = run_cli(
            ["sanitize", work, "--baseline", base, "-o", tmp_path / "san.docx", "--report"],
            capsys,
        )
        assert code == 0, err
        assert "differ by 9" not in err, f"positional similarity metric still in use: {err}"

    def test_incompatible_baseline_never_leaks_raw_key_error(self, tmp_path, capsys):
        """H3: missing relationship keys must not surface as `Error: 'rIdN'`."""
        mini = tmp_path / "mini.docx"
        d = Document()
        d.add_paragraph("Totally different tiny doc.")
        d.save(mini)

        agreement = tmp_path / "agreement.docx"
        doc = Document()
        sec = doc.sections[0]
        # Hyperlink in the FOOTER: its relationship lives on the footer part,
        # which the flattened baseline flow used to look up on the main part.
        footer_p = sec.footer.paragraphs[0]
        footer_part = sec.footer.part
        r_id = footer_part.relate_to(
            "https://example.com/terms",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hl = parse_xml(
            f"<w:hyperlink {W_NS} "
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
            f'r:id="{r_id}"><w:r><w:t>terms link</w:t></w:r></w:hyperlink>'
        )
        footer_p._element.append(hl)
        doc.add_paragraph("Agreement body text here about the obligations.")
        doc.save(agreement)

        code, out, err = run_cli(
            ["sanitize", mini, "--baseline", agreement, "-o", tmp_path / "san.docx", "--report"],
            capsys,
        )
        assert not re.search(r"Error: 'rId\d+'", err), f"raw KeyError leaked: {err}"

    def test_baseline_recomputation_keeps_body_paragraph_out_of_footer(self, tmp_path, capsys):
        """H3/C1: baseline recompute with an appended final paragraph must stay valid."""
        base = tmp_path / "base.docx"
        work = tmp_path / "work.docx"
        build_header_footer_doc(base)
        build_header_footer_doc(work, extra_body_paragraph=True)

        san = tmp_path / "san.docx"
        code, out, err = run_cli(["sanitize", work, "--baseline", base, "-o", san], capsys)
        assert code == 0, err
        assert san.exists()
        body_xml = read_part_xml(san, "word/document.xml")
        assert "New final body paragraph." in body_xml
        for member in footer_members(san):
            assert "New final body paragraph." not in read_part_xml(san, member)
        if soffice_can_convert(tmp_path):
            assert lo_loads(san, tmp_path), "baseline output must load in LibreOffice"


# ---------------------------------------------------------------------------
# H4 — footnote redlines and comments outside the main story
# ---------------------------------------------------------------------------


class TestH4FootnoteInterop:
    def test_footnote_edit_with_comment_produces_no_comment_anchor_in_footnotes(self, tmp_path, capsys):
        orig = tmp_path / "fn.docx"
        build_footnote_doc(orig)
        edits_file = tmp_path / "edits.json"
        edits_file.write_text(
            json.dumps(
                [
                    {
                        "type": "modify",
                        "target_text": "QA footnote about governing law",
                        "new_text": "QA footnote about applicable law",
                        "comment": "Diff: Replacement",
                    }
                ]
            )
        )
        applied = tmp_path / "fn_applied.docx"
        code, out, err = run_cli(["apply", orig, edits_file, "-o", applied], capsys)
        assert code == 0, err

        fn_xml = read_part_xml(applied, "word/footnotes.xml")
        assert "applicable" in fn_xml, "tracked replacement must still be applied"
        assert "commentReference" not in fn_xml, "comment anchors are not valid inside footnotes"
        assert "commentRangeStart" not in fn_xml

        # The agent must be told the comment was dropped.
        assert "comment" in err.lower() and (
            "skip" in err.lower() or "dropped" in err.lower() or "not supported" in err.lower()
        )

    def test_footnote_edit_output_loads_in_libreoffice(self, tmp_path, capsys):
        if not soffice_can_convert(tmp_path):
            pytest.skip("no working soffice")
        orig = tmp_path / "fn.docx"
        build_footnote_doc(orig)
        edits_file = tmp_path / "edits.json"
        edits_file.write_text(
            json.dumps(
                [
                    {
                        "type": "modify",
                        "target_text": "QA footnote about governing law",
                        "new_text": "QA footnote about applicable law",
                        "comment": "Diff: Replacement",
                    }
                ]
            )
        )
        applied = tmp_path / "fn_applied.docx"
        code, out, err = run_cli(["apply", orig, edits_file, "-o", applied], capsys)
        assert code == 0, err
        assert lo_loads(applied, tmp_path), "tracked footnote edit must produce a loadable DOCX"

    def test_footer_edit_comment_is_dropped_but_body_comment_kept(self, tmp_path, capsys):
        orig = tmp_path / "hf.docx"
        build_header_footer_doc(orig)
        edits_file = tmp_path / "edits.json"
        edits_file.write_text(
            json.dumps(
                [
                    {
                        "type": "modify",
                        "target_text": "FOOTER MARKER",
                        "new_text": "FOOTER MARK",
                        "comment": "footer comment",
                    },
                    {
                        "type": "modify",
                        "target_text": "Body paragraph one.",
                        "new_text": "Body paragraph 1.",
                        "comment": "body comment",
                    },
                ]
            )
        )
        applied = tmp_path / "applied.docx"
        code, out, err = run_cli(["apply", orig, edits_file, "-o", applied], capsys)
        assert code == 0, err
        for member in footer_members(applied):
            f_xml = read_part_xml(applied, member)
            assert "commentReference" not in f_xml
        body_xml = read_part_xml(applied, "word/document.xml")
        assert "commentRangeStart" in body_xml, "body comments must keep working"

    def test_pure_comment_on_footnote_text_is_rejected(self, tmp_path, capsys):
        orig = tmp_path / "fn.docx"
        build_footnote_doc(orig)
        edits_file = tmp_path / "edits.json"
        edits_file.write_text(
            json.dumps(
                [
                    {
                        "type": "modify",
                        "target_text": "QA footnote about governing law",
                        "new_text": "QA footnote about governing law",
                        "comment": "please review",
                    }
                ]
            )
        )
        applied = tmp_path / "applied.docx"
        code, out, err = run_cli(["apply", orig, edits_file, "-o", applied], capsys)
        assert code != 0, "a comment-only edit in a footnote cannot be honored and must fail"
        assert not applied.exists()


# ---------------------------------------------------------------------------
# M1 — markup semantics parity with apply
# ---------------------------------------------------------------------------


class TestM1MarkupParity:
    def _md(self, tmp_path):
        md = tmp_path / "doc.md"
        md.write_text("Alpha fee applies.\n\nBeta fee applies.\n\nGamma fee applies.")
        return md

    def test_match_mode_all_marks_every_occurrence(self, tmp_path, capsys):
        md = self._md(tmp_path)
        edits = tmp_path / "e.json"
        edits.write_text(
            json.dumps([{"type": "modify", "target_text": "fee", "new_text": "charge", "match_mode": "all"}])
        )
        out_md = tmp_path / "out.md"
        code, _, err = run_cli(["markup", md, edits, "-o", out_md], capsys)
        assert code == 0, err
        result = out_md.read_text()
        assert result.count("{--fee--}{++charge++}") == 3, result

    def test_regex_flag_is_honored(self, tmp_path, capsys):
        md = self._md(tmp_path)
        edits = tmp_path / "e.json"
        edits.write_text(
            json.dumps([{"type": "modify", "target_text": r"(?<=Alpha )fee", "new_text": "charge", "regex": True}])
        )
        out_md = tmp_path / "out.md"
        code, _, err = run_cli(["markup", md, edits, "-o", out_md], capsys)
        assert code == 0, err
        result = out_md.read_text()
        assert "Alpha {--fee--}{++charge++}" in result, result

    def test_ambiguous_strict_target_fails_like_apply(self, tmp_path, capsys):
        md = self._md(tmp_path)
        edits = tmp_path / "e.json"
        edits.write_text(json.dumps([{"type": "modify", "target_text": "fee", "new_text": "charge"}]))
        out_md = tmp_path / "out.md"
        code, _, err = run_cli(["markup", md, edits, "-o", out_md], capsys)
        assert code != 0, "ambiguous strict target must fail in markup exactly like apply"
        assert "Ambiguous" in err or "ambiguous" in err

    def test_missing_target_fails_nonzero(self, tmp_path, capsys):
        md = self._md(tmp_path)
        edits = tmp_path / "e.json"
        edits.write_text(json.dumps([{"type": "modify", "target_text": "does not exist", "new_text": "x"}]))
        out_md = tmp_path / "out.md"
        code, _, err = run_cli(["markup", md, edits, "-o", out_md], capsys)
        assert code != 0, "a missing target must not exit 0"
        assert "not found" in err.lower()

    def test_stats_reflect_actual_outcomes(self, tmp_path, capsys):
        md = self._md(tmp_path)
        edits = tmp_path / "e.json"
        edits.write_text(
            json.dumps(
                [
                    {"type": "modify", "target_text": "Alpha fee", "new_text": "Alpha charge"},
                    {"type": "modify", "target_text": "does not exist", "new_text": "x"},
                ]
            )
        )
        out_md = tmp_path / "out.md"
        code, _, err = run_cli(["markup", md, edits, "-o", out_md], capsys)
        assert code != 0
        assert "1 edits processed" not in err, "stats must not count skipped edits as processed"
        assert re.search(r"1\s+applied", err) and re.search(r"1\s+(failed|skipped)", err), err

    def test_non_text_actions_are_reported_not_silently_dropped(self, tmp_path, capsys):
        md = self._md(tmp_path)
        edits = tmp_path / "e.json"
        edits.write_text(
            json.dumps(
                [
                    {"type": "accept", "target_id": "Chg:1"},
                    {"type": "modify", "target_text": "Alpha fee", "new_text": "Alpha charge"},
                ]
            )
        )
        out_md = tmp_path / "out.md"
        code, _, err = run_cli(["markup", md, edits, "-o", out_md], capsys)
        assert "accept" in err.lower(), "ignored non-text actions must be called out"


# ---------------------------------------------------------------------------
# M2 — failed apply must not write output
# ---------------------------------------------------------------------------


class TestM2NoOutputOnFailure:
    def test_empty_target_fails_and_writes_nothing(self, tmp_path, capsys):
        orig = tmp_path / "doc.docx"
        d = Document()
        d.add_paragraph("Hello world.")
        d.save(orig)
        edits = tmp_path / "e.json"
        edits.write_text(json.dumps([{"type": "modify", "target_text": "", "new_text": "inserted"}]))
        out = tmp_path / "out.docx"
        code, stdout, err = run_cli(["apply", orig, edits, "-o", out], capsys)
        assert code != 0
        assert not out.exists(), "failed batch must not write an output file"
        assert "Batch complete" not in err


# ---------------------------------------------------------------------------
# M3 — VML watermark reporting
# ---------------------------------------------------------------------------


class TestM3Watermark:
    def test_watermark_is_reported(self, tmp_path, capsys):
        src = tmp_path / "wm.docx"
        build_watermark_doc(src)
        out = tmp_path / "wm_san.docx"
        code, stdout, err = run_cli(["sanitize", src, "-o", out, "--report"], capsys)
        assert code == 0, err
        assert "watermark" in err.lower(), f"watermark object must be surfaced in the report: {err}"
        assert "DRAFT ACME SECRET" in err


# ---------------------------------------------------------------------------
# M4 — list semantics
# ---------------------------------------------------------------------------


class TestM4Lists:
    def _build(self, path: Path):
        doc = Document()
        doc.add_paragraph("Requirements:")
        doc.add_paragraph("Maintain ISO 27001 controls", style="List Bullet")
        doc.add_paragraph("Notify incidents without undue delay", style="List Bullet")
        doc.add_paragraph("Escalation:")
        doc.add_paragraph("First escalation", style="List Number")
        doc.add_paragraph("Second escalation", style="List Number")
        doc.save(path)

    def test_style_based_lists_get_markers(self, tmp_path):
        p = tmp_path / "lists.docx"
        self._build(p)
        text = extract_text_from_stream(BytesIO(p.read_bytes()))
        assert "* Maintain ISO 27001 controls" in text, text
        assert "* Notify incidents without undue delay" in text
        assert "1. First escalation" in text, "ordered lists must be distinguishable from bullets"
        assert "1. Second escalation" in text

    def test_mapper_projection_stays_in_sync(self, tmp_path):
        """The Virtual Text contract: mapper and ingest must produce identical text."""
        from adeu.redline.mapper import DocumentMapper

        p = tmp_path / "lists.docx"
        self._build(p)
        doc = Document(str(p))
        mapper = DocumentMapper(doc)
        ingest_text = _extract_text_from_doc(Document(str(p)), include_appendix=False)
        assert mapper.full_text == ingest_text

    def test_list_item_edit_still_applies(self, tmp_path, capsys):
        p = tmp_path / "lists.docx"
        self._build(p)
        edits = tmp_path / "e.json"
        edits.write_text(
            json.dumps([{"type": "modify", "target_text": "ISO 27001 controls", "new_text": "ISO 27002 controls"}])
        )
        out = tmp_path / "out.docx"
        code, stdout, err = run_cli(["apply", p, edits, "-o", out], capsys)
        assert code == 0, err
        clean = extract_text_from_stream(BytesIO(out.read_bytes()), clean_view=True)
        assert "ISO 27002 controls" in clean


# ---------------------------------------------------------------------------
# M5 — image markers
# ---------------------------------------------------------------------------


class TestM5Images:
    def test_inline_image_projects_alt_marker(self, tmp_path):
        p = tmp_path / "img.docx"
        build_image_doc(p)
        text = extract_text_from_stream(BytesIO(p.read_bytes()))
        assert "![Red rectangle QA diagram](docx-image:" in text, text

    def test_mapper_projection_stays_in_sync(self, tmp_path):
        from adeu.redline.mapper import DocumentMapper

        p = tmp_path / "img.docx"
        build_image_doc(p)
        mapper = DocumentMapper(Document(str(p)))
        ingest_text = _extract_text_from_doc(Document(str(p)), include_appendix=False)
        assert mapper.full_text == ingest_text

    def test_image_marker_is_write_protected(self, tmp_path, capsys):
        p = tmp_path / "img.docx"
        build_image_doc(p)
        edits = tmp_path / "e.json"
        edits.write_text(
            json.dumps(
                [
                    {
                        "type": "modify",
                        "target_text": "![Red rectangle QA diagram](docx-image:1)",
                        "new_text": "some replacement prose",
                    }
                ]
            )
        )
        out = tmp_path / "out.docx"
        code, stdout, err = run_cli(["apply", p, edits, "-o", out], capsys)
        assert code != 0, "image markers are read-only projections and must reject edits"
        assert not out.exists()


# ---------------------------------------------------------------------------
# M6 — reserved footnotes
# ---------------------------------------------------------------------------


class TestM6ReservedFootnotes:
    def test_untyped_separator_footnotes_are_filtered(self, tmp_path):
        p = tmp_path / "fn.docx"
        build_footnote_doc(p, typed_separators=False)
        text = extract_text_from_stream(BytesIO(p.read_bytes()))
        assert "[^fn--1]" not in text, text
        assert "[^fn-0]" not in text, text
        assert "[^fn-1]:" in text
        assert "QA footnote about governing law" in text


# ---------------------------------------------------------------------------
# M7 — inline defined terms
# ---------------------------------------------------------------------------


class TestM7DefinedTerms:
    def test_multiple_sentence_leading_terms_in_one_paragraph(self, tmp_path):
        from adeu.domain import extract_all_domain_metadata

        doc = Document()
        doc.add_paragraph(
            '"Alpha" means the first party. "Beta" means the second party. "Gamma" means the third party.'
        )
        doc.add_paragraph("Alpha, Beta and Gamma perform the obligations.")
        base_text = _extract_text_from_doc(doc, include_appendix=False)
        defs, diags, anchors = extract_all_domain_metadata(doc, base_text)
        assert "Alpha" in defs
        assert "Beta" in defs, f"only got {sorted(defs)}"
        assert "Gamma" in defs, f"only got {sorted(defs)}"

    def test_multiple_paren_terms_in_one_paragraph(self, tmp_path):
        from adeu.domain import extract_all_domain_metadata

        doc = Document()
        doc.add_paragraph(
            'This agreement is between Acme Oy (the "Customer") and Beta Oy '
            '(the "Provider") for the provision of services (the "Services").'
        )
        doc.add_paragraph("The Customer pays the Provider for the Services.")
        base_text = _extract_text_from_doc(doc, include_appendix=False)
        defs, diags, anchors = extract_all_domain_metadata(doc, base_text)
        assert {"Customer", "Provider", "Services"} <= set(defs), sorted(defs)


# ---------------------------------------------------------------------------
# M8 — machine-readable errors in --json modes
# ---------------------------------------------------------------------------


class TestM8JsonErrors:
    def test_apply_json_malformed_changes_file(self, tmp_path, capsys):
        orig = tmp_path / "doc.docx"
        d = Document()
        d.add_paragraph("Hello")
        d.save(orig)
        bad = tmp_path / "bad.json"
        bad.write_text("{not json")
        code, out, err = run_cli(["apply", orig, bad, "--json"], capsys)
        assert code != 0
        payload = json.loads(out)
        assert "error" in payload and payload["error"], payload

    def test_extract_json_missing_file(self, tmp_path, capsys):
        code, out, err = run_cli(["extract", tmp_path / "missing.docx", "--json"], capsys)
        assert code != 0
        payload = json.loads(out)
        assert payload.get("error") == "file_not_found", payload

    def test_extract_json_invalid_docx(self, tmp_path, capsys):
        fake = tmp_path / "fake.docx"
        fake.write_text("this is not a zip")
        code, out, err = run_cli(["extract", fake, "--json"], capsys)
        assert code != 0
        payload = json.loads(out)
        assert payload.get("error") == "invalid_docx", payload


# ---------------------------------------------------------------------------
# L1-L6 — UX paper cuts
# ---------------------------------------------------------------------------


class TestLowSeverity:
    def test_l1_search_no_match_advice_names_real_flags(self, tmp_path, capsys):
        doc = tmp_path / "d.docx"
        d = Document()
        d.add_paragraph("Some content here.")
        d.save(doc)
        code, out, err = run_cli(["extract", doc, "--search-query", "zzz-no-match"], capsys)
        assert code == 0
        combined = out + err
        assert "search_case_sensitive" not in combined
        assert "--search-case-insensitive" in combined
        assert "--search-regex" in combined

    def test_l2_ignored_flags_warn(self, tmp_path, capsys):
        doc = tmp_path / "d.docx"
        d = Document()
        d.add_heading("Title", 0)
        d.add_paragraph("Some content here.")
        d.save(doc)

        code, out, err = run_cli(["extract", doc, "--mode", "outline", "--page", "2"], capsys)
        assert code == 0
        assert "--page" in err and ("ignored" in err.lower() or "has no effect" in err.lower())

        code, out, err = run_cli(["extract", doc, "--outline-verbose"], capsys)
        assert code == 0
        assert "--outline-verbose" in err and ("ignored" in err.lower() or "has no effect" in err.lower())

        code, out, err = run_cli(["extract", doc, "--mode", "outline", "--search-query", "content"], capsys)
        assert code == 0
        assert "search" in err.lower() and "outline" in err.lower()

    def test_l3_txt_input_gets_extension_error(self, tmp_path, capsys):
        f = tmp_path / "notes.txt"
        f.write_text("plain text")
        code, out, err = run_cli(["extract", f], capsys)
        assert code != 0
        assert ".txt" in err and "DOCX" in err
        assert "zip signature" not in err

    def test_l4_extract_json_with_output_keeps_stdout_quiet(self, tmp_path, capsys):
        doc = tmp_path / "d.docx"
        d = Document()
        d.add_paragraph("Some content here.")
        d.save(doc)
        out_file = tmp_path / "out.json"
        code, out, err = run_cli(["extract", doc, "--json", "-o", out_file], capsys)
        assert code == 0
        assert out.strip() == "", "with -o the payload belongs in the file, stdout stays quiet"
        payload = json.loads(out_file.read_text())
        assert isinstance(payload, dict)

    def test_l5_init_noop_rerun_creates_no_backup(self, tmp_path, capsys, monkeypatch):
        # Patch the path getter itself: HOME/Path.home don't reach the Windows
        # %APPDATA% branch, and this test used to rewrite the real config there
        # (while passing, because the .bak assertions looked at an empty dir).
        cfg_dir = tmp_path / "Claude"
        monkeypatch.setattr("adeu.cli._get_claude_config_path", lambda: cfg_dir / "claude_desktop_config.json")
        monkeypatch.setattr(shutil, "which", lambda name: "/usr/bin/uvx" if name == "uvx" else None)

        code, out, err = run_cli(["init"], capsys)
        assert code == 0, err
        baks = list(cfg_dir.glob("*.bak"))
        assert len(baks) == 0

        code, out, err = run_cli(["init"], capsys)
        assert code == 0, err
        baks = list(cfg_dir.glob("*.bak"))
        assert len(baks) == 0, "re-running init with no changes must not create backups"
        assert "unchanged" in out.lower() or "already" in out.lower()

    def test_l6_init_local_refuses_non_source_directory(self, tmp_path, capsys, monkeypatch):
        monkeypatch.setattr(Path, "home", classmethod(lambda cls: tmp_path))
        empty = tmp_path / "somewhere-else"
        empty.mkdir()
        monkeypatch.chdir(empty)
        code, out, err = run_cli(["init", "--local"], capsys)
        assert code != 0, "--local from a non-checkout directory must not claim success"
        assert "source" in err.lower() or "checkout" in err.lower()


# ---------------------------------------------------------------------------
# Post-review regressions (found while reviewing the fixes themselves)
# ---------------------------------------------------------------------------


class TestReviewRegressions:
    def test_diff_of_added_image_warns_instead_of_self_rejecting(self, tmp_path, capsys):
        """An added inline image must never produce an edit apply refuses."""
        orig = tmp_path / "noimg.docx"
        d = Document()
        d.add_paragraph("Before image.")
        d.add_paragraph("After image.")
        d.save(orig)

        mod = tmp_path / "img.docx"
        build_image_doc(mod)  # Before image. / [image] / After image.

        code, out, err = run_cli(["diff", orig, mod, "--json"], capsys)
        assert code == 0
        edits = json.loads(out)
        for e in edits:
            t_imgs = re.findall(r"!\[[^\]]*\]\(docx-image:[^)]*\)", e.get("target_text") or "")
            n_imgs = re.findall(r"!\[[^\]]*\]\(docx-image:[^)]*\)", e.get("new_text") or "")
            assert sorted(t_imgs) == sorted(n_imgs), f"diff emitted an unappliable image edit: {e}"
        assert "image" in err.lower(), "the skipped image difference must be surfaced as a warning"

        # Whatever the diff DID emit must still pass its own dry-run apply.
        edits_file = tmp_path / "edits.json"
        edits_file.write_text(out)
        code, out, err = run_cli(["apply", orig, edits_file, "--dry-run", "--json"], capsys)
        assert code == 0, f"diff output failed its own dry-run: {err} {out}"

    def test_baseline_fails_closed_on_apply_stage_skips(self, tmp_path, capsys, monkeypatch):
        """Apply-stage skips (no exception) must not produce a partial baseline output."""
        from adeu.redline.engine import RedlineEngine
        from adeu.sanitize.core import SanitizeError, sanitize_docx

        base = tmp_path / "base.docx"
        work = tmp_path / "work.docx"
        d = Document()
        d.add_paragraph("Alpha paragraph.")
        d.save(base)
        d = Document()
        d.add_paragraph("Alpha paragraph revised.")
        d.save(work)

        real_process_batch = RedlineEngine.process_batch

        def fake_process_batch(self, changes, dry_run=False):
            stats = real_process_batch(self, changes, dry_run=dry_run)
            stats["edits_skipped"] = stats.get("edits_skipped", 0) + 1
            stats.setdefault("skipped_details", []).append("- simulated apply-stage skip")
            return stats

        monkeypatch.setattr(RedlineEngine, "process_batch", fake_process_batch)
        out_path = tmp_path / "san.docx"
        with pytest.raises(SanitizeError, match="could not be applied"):
            sanitize_docx(str(work), str(out_path), baseline_path=str(base))
        assert not out_path.exists()

    def test_diff_json_projection_failure_emits_json_error(self, tmp_path, capsys, monkeypatch):
        """Extraction blowups inside diff must honor the --json error contract (M8)."""
        import adeu.cli as cli_mod

        a = tmp_path / "a.docx"
        b = tmp_path / "b.docx"
        for p in (a, b):
            d = Document()
            d.add_paragraph("content")
            d.save(p)

        def boom(*args, **kwargs):
            raise RuntimeError("simulated projection failure")

        monkeypatch.setattr(cli_mod, "_extract_text_from_doc", boom, raising=False)
        # handle_diff imports the symbol lazily from adeu.ingest — patch there.
        import adeu.ingest as ingest_mod

        monkeypatch.setattr(ingest_mod, "_extract_text_from_doc", boom)

        code, out, err = run_cli(["diff", a, b, "--json"], capsys)
        assert code != 0
        payload = json.loads(out)
        assert payload.get("error") == "invalid_docx", payload

    def test_footer_prefix_edit_stays_in_footer(self, tmp_path, capsys):
        """A target-anchored insertion at the footer's first character must not
        be hijacked into the body by the C1 boundary re-anchor."""
        orig = tmp_path / "hf.docx"
        build_header_footer_doc(orig)
        edits_file = tmp_path / "edits.json"
        edits_file.write_text(
            json.dumps(
                [
                    {
                        "type": "modify",
                        "target_text": "FOOTER MARKER",
                        "new_text": "DRAFT FOOTER MARKER",
                    }
                ]
            )
        )
        applied = tmp_path / "applied.docx"
        code, out, err = run_cli(["apply", orig, edits_file, "-o", applied], capsys)
        assert code == 0, err
        body_xml = read_part_xml(applied, "word/document.xml")
        assert "DRAFT" not in body_xml, "footer prefix must not be re-anchored into the body"
        assert any("DRAFT" in read_part_xml(applied, m) for m in footer_members(applied))
        clean = extract_text_from_stream(BytesIO(applied.read_bytes()), clean_view=True)
        assert "DRAFT FOOTER MARKER" in clean

    def test_extensionless_docx_still_extracts(self, tmp_path, capsys):
        """Content beats extension: a valid DOCX without .docx must load."""
        src = tmp_path / "sample_noext"
        d = Document()
        d.add_paragraph("Extension-less content.")
        d.save(src)
        code, out, err = run_cli(["extract", src, "--page", "all"], capsys)
        assert code == 0, err
        assert "Extension-less content." in out

    def test_two_consecutive_added_rows_roundtrip_in_process(self, tmp_path, capsys):
        """Consecutive inserts share one pinned anchor; they must not be
        flagged as overlapping edits (zero-width insert ranges)."""
        orig = tmp_path / "t1.docx"
        mod = tmp_path / "t3.docx"
        d = Document()
        t = d.add_table(rows=2, cols=2)
        for i, row in enumerate([["A", "B"], ["C", "D"]]):
            for j, v in enumerate(row):
                t.rows[i].cells[j].text = v
        d.save(orig)
        d = Document()
        t = d.add_table(rows=4, cols=2)
        for i, row in enumerate([["A", "B"], ["C", "D"], ["X1", "Y1"], ["X2", "Y2"]]):
            for j, v in enumerate(row):
                t.rows[i].cells[j].text = v
        d.save(mod)

        # In-process pinned path (the sanitize --baseline shape).
        san = tmp_path / "san.docx"
        code, out, err = run_cli(["sanitize", mod, "--baseline", orig, "-o", san], capsys)
        assert code == 0, err
        clean = extract_text_from_stream(BytesIO(san.read_bytes()), clean_view=True)
        assert "A | B\n--- | ---\nC | D\nX1 | Y1\nX2 | Y2" in clean, clean

        # JSON round trip too.
        code, out, err = run_cli(["diff", orig, mod, "--json"], capsys)
        edits_file = tmp_path / "edits.json"
        edits_file.write_text(out)
        applied = tmp_path / "applied.docx"
        code, out, err = run_cli(["apply", orig, edits_file, "-o", applied], capsys)
        assert code == 0, err
        clean = extract_text_from_stream(BytesIO(applied.read_bytes()), clean_view=True)
        assert "A | B\n--- | ---\nC | D\nX1 | Y1\nX2 | Y2" in clean, clean

    def test_numbering_disabled_override_suppresses_style_list(self, tmp_path):
        """ECMA-376 §17.9.15: a direct numId=0 removes the style's numbering."""
        from docx.oxml import parse_xml

        doc = Document()
        p = doc.add_paragraph("Not a list item", style="List Number")
        pPr = p._element.get_or_add_pPr()
        pPr.append(
            parse_xml(
                '<w:numPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:numId w:val="0"/></w:numPr>'
            )
        )
        doc.add_paragraph("Real item", style="List Number")
        stream = BytesIO()
        doc.save(stream)
        text = extract_text_from_stream(BytesIO(stream.getvalue()))
        assert "1. Not a list item" not in text, text
        assert "Not a list item" in text
        assert "1. Real item" in text

    def test_duplicate_anchor_rows_across_tables_roundtrip(self, tmp_path, capsys):
        """Row ops anchored on text duplicated in another table warn, and the
        in-process baseline path (pinned offsets) still succeeds."""
        orig = tmp_path / "two_tables.docx"
        mod = tmp_path / "two_tables_mod.docx"
        for path, extra in ((orig, False), (mod, True)):
            d = Document()
            for _ in range(2):
                t = d.add_table(rows=2, cols=2)
                t.rows[0].cells[0].text = "Name"
                t.rows[0].cells[1].text = "Value"
                t.rows[1].cells[0].text = "A"
                t.rows[1].cells[1].text = "1"
            if extra:
                r = d.tables[1].add_row()
                r.cells[0].text = "B"
                r.cells[1].text = "2"
            d.save(path)

        code, out, err = run_cli(["diff", orig, mod, "--json"], capsys)
        assert code == 0
        assert "identical" in err.lower() or "ambiguous" in err.lower(), (
            "duplicate-row ambiguity must be surfaced as a warning"
        )

        # The in-process pipeline (sanitize --baseline) rides pinned offsets
        # and must succeed despite the duplicated anchor text.
        san = tmp_path / "san.docx"
        code, out, err = run_cli(["sanitize", mod, "--baseline", orig, "-o", san], capsys)
        assert code == 0, err
        clean = extract_text_from_stream(BytesIO(san.read_bytes()), clean_view=True)
        assert "B | 2" in clean

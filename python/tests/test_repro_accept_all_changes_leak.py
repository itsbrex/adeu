# FILE: python/tests/test_repro_accept_all_changes_leak.py
import asyncio
import io
import zipfile

from docx import Document

from adeu.mcp_components.tools.document import accept_all_changes
from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine


class MockContext:
    """Mock FastMCP Context to absorb async logging calls during tests."""

    async def info(self, msg, **kwargs):
        print(f"INFO: {msg} {kwargs}")

    async def debug(self, msg, **kwargs):
        print(f"DEBUG: {msg} {kwargs}")

    async def warning(self, msg, **kwargs):
        print(f"WARNING: {msg} {kwargs}")

    async def error(self, msg, **kwargs):
        print(f"ERROR: {msg} {kwargs}")


def test_repro_accept_all_changes_comment_leak(tmp_path):
    """
    Reproduction test for accept_all_changes comment leak.
    It verifies that when a document containing tracked changes and comments is finalized via
    accept_all_changes, the output package contains ZERO comment parts (word/comments*.xml)
    and ZERO in-body comment anchors (w:commentRangeStart, w:commentRangeEnd, w:commentReference).
    """
    # 1. Create a simple base document
    doc = Document()
    doc.add_paragraph("This is the original text of the agreement.")
    base_path = tmp_path / "base.docx"
    doc.save(base_path)

    # 2. Add a tracked change with a comment
    with open(base_path, "rb") as f:
        engine = RedlineEngine(io.BytesIO(f.read()), author="Reviewer AI")
        engine.apply_edits(
            [
                ModifyText(
                    target_text="original text",
                    new_text="updated text",
                    comment="Should this be updated or kept as original?",
                )
            ]
        )
        tracked_stream = engine.save_to_stream()

    tracked_path = tmp_path / "tracked.docx"
    with open(tracked_path, "wb") as f:
        f.write(tracked_stream.getvalue())

    # Ensure pre-condition: tracked document has comment in body and comment parts
    with zipfile.ZipFile(tracked_path, "r") as zf:
        namelist = zf.namelist()
        comment_parts = [name for name in namelist if "comments" in name.lower()]
        assert len(comment_parts) > 0, "Pre-condition failed: no comment parts found in tracked.docx"

    tracked_doc = Document(str(tracked_path))
    tracked_xml = tracked_doc.element.xml
    assert "w:commentRangeStart" in tracked_xml, "Pre-condition failed: no commentRangeStart in body"
    assert "w:commentReference" in tracked_xml, "Pre-condition failed: no commentReference in body"

    # 3. Act: Finalize the document using accept_all_changes
    output_path = tmp_path / "finalized.docx"
    ctx = MockContext()

    asyncio.run(accept_all_changes(docx_path=str(tracked_path), ctx=ctx, output_path=str(output_path)))

    assert output_path.exists(), "Finalized document was not created"

    # 4. Assert: Total comment wipe across body, parts, rels, and content-types
    # Check body XML for anchors
    final_doc = Document(str(output_path))
    final_xml = final_doc.element.xml

    # Assert NO in-body comment anchors survive (anchors must be completely gone)
    assert "w:commentRangeStart" not in final_xml, "Found leaked w:commentRangeStart in body XML"
    assert "w:commentRangeEnd" not in final_xml, "Found leaked w:commentRangeEnd in body XML"
    assert "w:commentReference" not in final_xml, "Found leaked w:commentReference in body XML"

    # Inspect the zip package contents directly
    with zipfile.ZipFile(output_path, "r") as zf:
        namelist = zf.namelist()

        # Assert NO comment parts exist (e.g. word/comments.xml, word/commentsExtended.xml, etc.)
        comment_files = [name for name in namelist if "comments" in name.lower()]
        assert len(comment_files) == 0, f"Found leaked comment parts in zip: {comment_files}"

        # Assert word/_rels/document.xml.rels contains ZERO references to comments
        if "word/_rels/document.xml.rels" in namelist:
            rels_xml = zf.read("word/_rels/document.xml.rels").decode("utf-8")
            assert "comments" not in rels_xml.lower(), "Found leaked comment relationships in document.xml.rels"

        # Assert [Content_Types].xml contains ZERO Override entries referencing comments
        if "[Content_Types].xml" in namelist:
            content_types_xml = zf.read("[Content_Types].xml").decode("utf-8")
            assert "comments" not in content_types_xml.lower(), "Found leaked comment overrides in [Content_Types].xml"

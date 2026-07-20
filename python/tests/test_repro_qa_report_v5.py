# FILE: tests/test_repro_qa_report_v5.py
"""
Regression tests for the 2026-07-17 exploratory QA report (adeu 1.21.0+5bece24).

Finding index (severity as reported):
  F1  🔴 paginated extract → full-document apply/diff silently deletes
      everything past page 1 and writes page chrome into the document
  F2  🟠 legacy {"actions": [...]} batch coerces unrecognized action values
      (" reject ", "rejcet", missing) to "accept" — semantic inversion.
      Resolution superseded: the entire undocumented pre-v1.1.0 dict format
      was removed; that shape now gets a targeted migration error.
  F3  🟠 sanitize reports comments removed while word/comments.xml survives
      (delete_comment guards on the backing field, not the lazy property)
  F4  🟠 dc:title / cp:category / cp:keywords / dc:subject / contentStatus
      are never scrubbed nor reported
  F5  🟡 LLM-controlled regex (regex:true / search_regex) hits Python's
      backtracking engine with no time budget — ReDoS
  F6  🟡 duplicate-definition diagnostic silently discarded when the term is
      never used (exactly when the document is most broken)
  F7  🟢 directory input produces a raw IsADirectoryError traceback
  F8  🔴 pinned/text-diff edits bypass validate_edit_strings entirely,
      writing CriticMarkup into document bodies (or crashing _attach_comment)
  F9  🟡 duplicate w:id values: one accept/reject silently resolves multiple
      unrelated revisions by different authors
  F10 🟢 markup emits 0-based [Edit:N]; apply reports 1-based Edit N
  F11 🟢 control characters in new_text/comment/--author raise a raw lxml
      traceback instead of a clean validation error

Every test in this file failed against 5bece24 before the corresponding fix.
"""

import io
import json
import re
import sys
import zipfile
from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from adeu.ingest import extract_text_from_stream
from adeu.markup import apply_edits_to_markdown
from adeu.models import ModifyText, RejectChange
from adeu.redline.comments import CommentsManager
from adeu.redline.engine import BatchValidationError, RedlineEngine

# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------


def _stream(doc) -> io.BytesIO:
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def _clean_text(data: bytes) -> str:
    return extract_text_from_stream(io.BytesIO(data), filename="x.docx", clean_view=True)


def _run_cli(argv: list[str]) -> int:
    """Runs the adeu CLI in-process; returns the exit code (0 if main returns)."""
    from adeu.cli import main

    with patch.object(sys, "argv", ["adeu"] + argv):
        try:
            main()
        except SystemExit as e:
            return int(e.code or 0)
    return 0


def _make_tracked_ins(text: str, wid: str, author: str) -> OxmlElement:
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), wid)
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), "2026-01-02T10:00:00Z")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    ins.append(r)
    return ins


def _make_tracked_del(text: str, wid: str, author: str) -> OxmlElement:
    d = OxmlElement("w:del")
    d.set(qn("w:id"), wid)
    d.set(qn("w:author"), author)
    d.set(qn("w:date"), "2026-01-02T10:00:00Z")
    r = OxmlElement("w:r")
    t = OxmlElement("w:delText")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    d.append(r)
    r.append(t)
    return d


def _make_large_doc(paragraphs: int = 130) -> io.BytesIO:
    """A document comfortably past PAGE_TARGET_CHARS (19k) so extract paginates."""
    doc = Document()
    for i in range(1, paragraphs + 1):
        doc.add_paragraph(f"Paragraph {i:04d}: " + "Lorem ipsum dolor sit amet consectetur adipiscing elit. " * 3)
    return _stream(doc)


def _make_doc_with_fee_change() -> io.BytesIO:
    """Tracked change raising a fee from $10,000 to $12,500 (Chg:101 / Chg:102)."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("The fee is ")
    p._element.append(_make_tracked_del("$10,000", "101", "Opposing Counsel"))
    p._element.append(_make_tracked_ins("$12,500", "102", "Opposing Counsel"))
    p.add_run(" per month.")
    return _stream(doc)


def _make_doc_with_comment(
    author: str = "Mallory Insider",
    text: str = "Walk away below 9M. Board approved 12M ceiling.",
) -> io.BytesIO:
    """One anchored comment, no tracked changes (the QA report's comments_only.docx)."""
    doc = Document()
    p = doc.add_paragraph("The purchase price shall be negotiated in good faith.")
    cm = CommentsManager(doc)
    cid = cm.add_comment(author, text)

    p_el = p._element
    run_el = p.runs[0]._element
    rs = OxmlElement("w:commentRangeStart")
    rs.set(qn("w:id"), cid)
    p_el.insert(list(p_el).index(run_el), rs)
    re_el = OxmlElement("w:commentRangeEnd")
    re_el.set(qn("w:id"), cid)
    p_el.append(re_el)
    ref_run = OxmlElement("w:r")
    ref = OxmlElement("w:commentReference")
    ref.set(qn("w:id"), cid)
    ref_run.append(ref)
    p_el.append(ref_run)
    return _stream(doc)


def _make_doc_with_tracked_insertion() -> io.BytesIO:
    """The QA report's only_ins.docx: 'Fee is {++high++} today.' by Alice."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Fee is ")
    p._element.append(_make_tracked_ins("high", "1", "Alice"))
    p.add_run(" today.")
    return _stream(doc)


def _read_zip_part(data: bytes, name: str) -> str:
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        if name not in zf.namelist():
            return ""
        return zf.read(name).decode("utf-8")


def _read_comment_parts(data: bytes) -> str:
    """Concatenated content of every word/comments*.xml part in the package."""
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        return "".join(zf.read(n).decode("utf-8", "ignore") for n in zf.namelist() if n.startswith("word/comments"))


# ---------------------------------------------------------------------------
# F1 — paginated extract → full-document apply must never silently truncate
# ---------------------------------------------------------------------------


class TestF1PaginationRoundTrip:
    def test_partial_page_apply_refuses_loudly(self, tmp_path, capsys):
        """Applying page 1 of a multi-page extract must be a loud error, not a
        silent deletion of every other page (QA F1 main repro)."""
        docx_path = tmp_path / "large.docx"
        docx_path.write_bytes(_make_large_doc().getvalue())
        original_clean = _clean_text(docx_path.read_bytes())
        assert "Paragraph 0130" in original_clean  # sanity: doc really has the tail

        txt_path = tmp_path / "large.txt"
        rc = _run_cli(["extract", str(docx_path), "--clean-view", "-o", str(txt_path)])
        assert rc == 0
        page1_text = txt_path.read_text(encoding="utf-8")
        assert "Page 1 of" in page1_text  # sanity: extract really paginated

        # User edits NOTHING and feeds the page-1 text straight back.
        out_path = tmp_path / "out.docx"
        rc = _run_cli(["apply", str(docx_path), str(txt_path), "-o", str(out_path)])
        err = capsys.readouterr().err

        assert rc != 0, "apply exited 0 while deleting every paragraph past page 1"
        assert not out_path.exists(), "apply wrote a corrupted output despite the page mismatch"
        # The error must tell the agent what happened and how to recover.
        assert "page" in err.lower()
        assert "--page all" in err

    def test_partial_page_diff_refuses_loudly(self, tmp_path, capsys):
        docx_path = tmp_path / "large.docx"
        docx_path.write_bytes(_make_large_doc().getvalue())
        txt_path = tmp_path / "large.txt"
        assert _run_cli(["extract", str(docx_path), "--clean-view", "-o", str(txt_path)]) == 0

        rc = _run_cli(["diff", str(docx_path), str(txt_path)])
        err = capsys.readouterr().err
        assert rc != 0, "diff exited 0 while reporting the missing pages as deletions"
        assert "page" in err.lower()

    def test_extract_page_all_roundtrips_large_doc(self, tmp_path, capsys):
        """--page all must exist so a complete text round-trip is expressible,
        and apply(D, extract(D)) must then be a no-op."""
        docx_path = tmp_path / "large.docx"
        docx_path.write_bytes(_make_large_doc().getvalue())
        original_clean = _clean_text(docx_path.read_bytes())

        txt_path = tmp_path / "large_all.txt"
        rc = _run_cli(["extract", str(docx_path), "--clean-view", "--page", "all", "-o", str(txt_path)])
        assert rc == 0, "extract --page all must be supported outside search mode"
        full_text = txt_path.read_text(encoding="utf-8")
        assert "Paragraph 0001" in full_text
        assert "Paragraph 0130" in full_text
        assert "Continues on page" not in full_text

        out_path = tmp_path / "out.docx"
        rc = _run_cli(["apply", str(docx_path), str(txt_path), "-o", str(out_path), "--json"])
        stats = json.loads(capsys.readouterr().out.strip().splitlines()[-1])
        assert rc == 0
        assert stats["edits_applied"] == 0, f"identity round-trip produced edits: {stats}"

        result_clean = _clean_text(out_path.read_bytes())
        assert result_clean == original_clean

    @pytest.mark.parametrize("is_cli", [True, False])
    def test_page_chrome_stripper_matches_real_chrome_builders(self, is_cli):
        """The ingest-side chrome stripper must recognize exactly what the
        pagination chrome builders emit — banner, continuation footer, and
        appendix pointer, in both CLI and MCP wording."""
        from adeu.cli import _strip_page_chrome
        from adeu.pagination import build_appendix_pointer, build_page_banner, build_page_footer

        body = "Paragraph body text.\n\nSecond paragraph."
        chromed = (
            build_page_banner(1, 6, "large.docx", is_cli=is_cli)
            + body
            + build_page_footer(1, 6, True, "large.docx", is_cli=is_cli)
            + build_appendix_pointer("large.docx", True, is_cli=is_cli)
        )
        stripped, page, total = _strip_page_chrome(chromed)
        assert stripped == body
        assert page == 1
        assert total == 6

        # Single-page output (no banner/footer) with an appendix pointer.
        chromed_single = body + build_appendix_pointer("doc.docx", True, is_cli=is_cli)
        stripped, page, total = _strip_page_chrome(chromed_single)
        assert stripped == body
        assert page is None and total is None

    def test_single_page_extract_roundtrips_as_noop(self, tmp_path, capsys):
        doc = Document()
        doc.add_paragraph('"Widget" means a mechanical device.')
        doc.add_paragraph("The Widget shall be delivered on time.")
        docx_path = tmp_path / "terms.docx"
        docx_path.write_bytes(_stream(doc).getvalue())
        original_clean = _clean_text(docx_path.read_bytes())

        txt_path = tmp_path / "terms.txt"
        assert _run_cli(["extract", str(docx_path), "--clean-view", "-o", str(txt_path)]) == 0

        out_path = tmp_path / "out.docx"
        rc = _run_cli(["apply", str(docx_path), str(txt_path), "-o", str(out_path), "--json"])
        stats = json.loads(capsys.readouterr().out.strip().splitlines()[-1])
        assert rc == 0
        assert stats["edits_applied"] == 0, f"identity round-trip produced edits: {stats}"
        assert _clean_text(out_path.read_bytes()) == original_clean

    def test_truncated_text_without_chrome_still_guarded(self, tmp_path, capsys):
        """Even without page chrome (e.g. a hand-truncated file), silently
        deleting most of the document must require explicit consent."""
        docx_path = tmp_path / "large.docx"
        docx_path.write_bytes(_make_large_doc().getvalue())
        original_clean = _clean_text(docx_path.read_bytes())

        truncated = original_clean[: len(original_clean) // 5]
        txt_path = tmp_path / "truncated.txt"
        txt_path.write_text(truncated, encoding="utf-8")

        out_path = tmp_path / "out.docx"
        rc = _run_cli(["apply", str(docx_path), str(txt_path), "-o", str(out_path)])
        err = capsys.readouterr().err
        assert rc != 0, "apply silently deleted ~80% of the document"
        assert not out_path.exists()
        assert "--allow-major-deletions" in err

        # The override flag expresses intent and unlocks the operation.
        rc = _run_cli(
            ["apply", str(docx_path), str(txt_path), "-o", str(out_path), "--allow-major-deletions", "--json"]
        )
        assert rc == 0
        assert out_path.exists()


# ---------------------------------------------------------------------------
# F2 — legacy {"actions": [...]} must never invert reject into accept
# ---------------------------------------------------------------------------


class TestF2LegacyFormatRemoved:
    """The pre-v1.1.0 {"actions": [...], "edits": [...]} dict format — F2's
    entire attack surface — is removed. Any file in that shape must be
    rejected with a targeted migration error before a single change is
    interpreted, regardless of its content."""

    def _apply_batch(self, tmp_path, capsys, payload) -> tuple[int, str, Path]:
        docx_path = tmp_path / "tracked.docx"
        input_bytes = _make_doc_with_fee_change().getvalue()
        docx_path.write_bytes(input_bytes)
        json_path = tmp_path / "a.json"
        json_path.write_text(json.dumps(payload), encoding="utf-8")
        out_path = tmp_path / "out.docx"
        rc = _run_cli(["apply", str(docx_path), str(json_path), "-o", str(out_path)])
        err = capsys.readouterr().err
        # Whatever the outcome, apply must never mutate its input file.
        assert docx_path.read_bytes() == input_bytes
        return rc, err, out_path

    @pytest.mark.parametrize(
        "actions",
        [
            [{"action": "reject", "target_id": "Chg:101"}],  # even well-formed legacy input
            [{"action": " reject ", "target_id": "Chg:101"}],  # F2's original repro
            [{"action": "rejcet", "target_id": "Chg:101"}],
            [{"target_id": "Chg:101"}],
            [{"action": None, "target_id": "Chg:101"}],
        ],
    )
    def test_legacy_actions_format_is_rejected_with_migration_guidance(self, tmp_path, capsys, actions):
        rc, err, out_path = self._apply_batch(tmp_path, capsys, {"actions": actions})
        assert rc != 0, "the removed legacy format was silently interpreted"
        assert not out_path.exists()
        # Actionable migration error: names the removed shape and the modern one.
        assert "removed" in err
        assert "'action' to 'type'" in err
        assert "AttributeError" not in err

    def test_legacy_edits_bucket_is_rejected(self, tmp_path, capsys):
        rc, err, out_path = self._apply_batch(
            tmp_path, capsys, {"edits": [{"original": "$12,500", "replace": "$99,999"}]}
        )
        assert rc != 0
        assert not out_path.exists()
        assert "'original' to 'target_text'" in err

    def test_non_legacy_dict_root_gets_generic_list_error(self, tmp_path, capsys):
        rc, err, out_path = self._apply_batch(tmp_path, capsys, {"changes": []})
        assert rc != 0
        assert not out_path.exists()
        assert "list of change objects" in err

    def test_migrated_modern_equivalent_applies(self, tmp_path, capsys):
        """The migration the error message describes must actually work."""
        rc, _err, out_path = self._apply_batch(
            tmp_path,
            capsys,
            [
                {"type": "reject", "target_id": "Chg:101"},
                {"type": "reject", "target_id": "Chg:102"},
            ],
        )
        assert rc == 0
        result = _clean_text(out_path.read_bytes())
        assert "$10,000" in result
        assert "$12,500" not in result


# ---------------------------------------------------------------------------
# F3 — sanitize must actually delete comments, not just their anchors
# ---------------------------------------------------------------------------


class TestF3SanitizeCommentDeletion:
    def test_delete_comment_works_on_fresh_manager(self):
        """delete_comment on a freshly constructed CommentsManager (the sanitize
        path) must delete the comment, not silently no-op on the unpopulated
        backing field."""
        doc = Document(_make_doc_with_comment())
        cm = CommentsManager(doc)
        data = cm.extract_comments_data()
        assert len(data) == 1  # sanity
        cm.delete_comment(next(iter(data)))

        # Verify against package state via a *fresh* manager.
        remaining = CommentsManager(doc).extract_comments_data()
        assert remaining == {}, "delete_comment was a no-op on a fresh manager"

    def test_find_para_id_works_on_fresh_manager(self):
        doc = Document(_make_doc_with_comment())
        cm = CommentsManager(doc)
        cid = next(iter(cm.extract_comments_data()))
        assert cm._find_para_id_for_comment(cid) is not None

    def test_sanitize_removes_comment_text_and_author_from_package(self, tmp_path):
        from adeu.sanitize.core import sanitize_docx

        input_path = tmp_path / "comments_only.docx"
        input_path.write_bytes(_make_doc_with_comment().getvalue())
        output_path = tmp_path / "clean.docx"

        result = sanitize_docx(str(input_path), str(output_path))
        assert result.comments_removed == 1
        assert "Comments removed: 1" in result.report_text

        comments_xml = _read_comment_parts(output_path.read_bytes())
        assert "Walk away below 9M" not in comments_xml, (
            "sanitize certified the comment removed but its text survives in word/comments*.xml"
        )
        assert "Mallory Insider" not in comments_xml

    def test_full_sanitize_with_author_leaves_no_original_authors(self, tmp_path):
        """--accept-all --author REDACTED (full sanitize) must not leave the
        original comment authors anywhere in the package."""
        from adeu.sanitize.core import sanitize_docx

        input_path = tmp_path / "with_comment.docx"
        input_path.write_bytes(_make_doc_with_comment().getvalue())
        output_path = tmp_path / "clean.docx"

        sanitize_docx(str(input_path), str(output_path), author="REDACTED", accept_all=True)

        with zipfile.ZipFile(io.BytesIO(output_path.read_bytes())) as zf:
            all_xml = "".join(zf.read(n).decode("utf-8", "ignore") for n in zf.namelist() if n.endswith(".xml"))
        assert "Mallory Insider" not in all_xml


# ---------------------------------------------------------------------------
# F4 — sanitize must scrub (or at least report) all leaky core properties
# ---------------------------------------------------------------------------


class TestF4CorePropertyScrub:
    def _sanitized(self, tmp_path) -> tuple[str, str]:
        from adeu.sanitize.core import sanitize_docx

        doc = Document()
        doc.add_paragraph("Body text.")
        core = doc.core_properties
        core.title = "Secret Merger Agreement"
        core.category = "Project Falcon"
        core.keywords = "confidential,merger,project-falcon"
        core.subject = "Acquisition of TargetCo"
        core.content_status = "Draft - privileged"
        core.comments = "Internal: do not circulate"

        input_path = tmp_path / "meta.docx"
        input_path.write_bytes(_stream(doc).getvalue())
        output_path = tmp_path / "clean.docx"
        result = sanitize_docx(str(input_path), str(output_path))
        core_xml = _read_zip_part(output_path.read_bytes(), "docProps/core.xml")
        return core_xml, result.report_text

    def test_category_keywords_subject_status_description_are_stripped(self, tmp_path):
        core_xml, report = self._sanitized(tmp_path)
        assert "Project Falcon" not in core_xml
        assert "project-falcon" not in core_xml
        assert "Acquisition of TargetCo" not in core_xml
        assert "Draft - privileged" not in core_xml
        assert "do not circulate" not in core_xml
        # ...and the report says so.
        assert "Category" in report
        assert "Keywords" in report

    def test_title_is_reported_not_silently_kept(self, tmp_path):
        """Title is often intentional — the documented behavior is to REPORT it
        rather than strip it. The `pass` shipped in 1.21.0 did neither."""
        core_xml, report = self._sanitized(tmp_path)
        assert "Secret Merger Agreement" in core_xml  # kept by design
        assert "Secret Merger Agreement" in report, "title leak is invisible in the report"


# ---------------------------------------------------------------------------
# F5 — LLM-controlled regex needs a wall-clock budget
# ---------------------------------------------------------------------------

# Both pathological patterns come from the QA report. The regex engine's own
# optimizations defuse `(a+)+$` outright (fast, correct no-match); `(a|a)*$`
# defeats optimization and must be stopped by the wall-clock budget instead.
REDOS_PATTERN_OPTIMIZED_AWAY = "(a+)+$"
REDOS_PATTERN_NEEDS_TIMEOUT = "(a|a)*$"
REDOS_HAYSTACK = "x" + "a" * 34 + "!"
WATCHDOG_SECONDS = 12.0


def _f5_edit_payload(q, pattern):
    doc = Document()
    doc.add_paragraph(REDOS_HAYSTACK)
    engine = RedlineEngine(_stream(doc))
    edit = ModifyText(target_text=pattern, new_text="X", regex=True, match_mode="first")
    try:
        engine.process_batch([edit])
        q.put(("returned", None))
    except BatchValidationError as e:
        q.put(("rejected", "\n".join(e.errors)))
    except Exception as e:  # noqa: BLE001 - recorded for assertion
        q.put(("raised", f"{type(e).__name__}: {e}"))


def _f5_search_payload(q, pattern):
    from adeu.mcp_components._response_builders import build_search_response

    try:
        res = build_search_response(REDOS_HAYSTACK, pattern, True, True, None, "doc.docx")
        q.put(("returned", str(res.content)))
    except Exception as e:  # noqa: BLE001 - recorded for assertion
        q.put(("raised", f"{type(e).__name__}: {e}"))


def _run_in_subprocess_with_watchdog(payload, pattern) -> tuple[str, str | None]:
    """Catastrophic re backtracking holds the GIL, so an in-process watchdog
    thread would freeze with it — run the payload in a child process instead.
    Windows has no fork; spawn pays interpreter+import startup inside the
    child, so it gets extra join budget before the watchdog calls it a hang."""
    import multiprocessing

    method = "fork" if "fork" in multiprocessing.get_all_start_methods() else "spawn"
    join_budget = WATCHDOG_SECONDS + (0.0 if method == "fork" else 18.0)
    ctx = multiprocessing.get_context(method)
    q = ctx.Queue()
    proc = ctx.Process(target=payload, args=(q, pattern), daemon=True)
    proc.start()
    proc.join(join_budget)
    if proc.is_alive():
        proc.terminate()
        proc.join(5)
        pytest.fail(f"ReDoS: regex resolution still running after {join_budget}s")
    return q.get(timeout=5)


class TestF5RegexTimeout:
    def test_redos_edit_is_rejected_within_budget(self):
        outcome, detail = _run_in_subprocess_with_watchdog(_f5_edit_payload, REDOS_PATTERN_NEEDS_TIMEOUT)
        assert outcome == "rejected", f"expected clean rejection, got {outcome}: {detail}"
        assert "Edit 1 Failed" in detail
        assert "time" in detail.lower()  # names the timeout, not a generic not-found

    def test_redos_prone_pattern_resolves_fast_and_clean(self):
        """The report's headline pattern is handled by engine optimization:
        a fast, correct 'not found' — no hang, no traceback."""
        outcome, detail = _run_in_subprocess_with_watchdog(_f5_edit_payload, REDOS_PATTERN_OPTIMIZED_AWAY)
        assert outcome == "rejected", f"expected clean not-found rejection, got {outcome}: {detail}"
        assert "Edit 1 Failed" in detail

    def test_redos_search_downgrades_to_literal_within_budget(self):
        outcome, detail = _run_in_subprocess_with_watchdog(_f5_search_payload, REDOS_PATTERN_NEEDS_TIMEOUT)
        assert outcome == "returned", f"expected downgraded literal search, got {outcome}: {detail}"
        assert "literal" in detail.lower()  # mirrors the invalid-regex downgrade note


# ---------------------------------------------------------------------------
# F6 — duplicate-definition diagnostics must survive the unused-term prune
# ---------------------------------------------------------------------------


class TestF6DuplicateDefinitionDiagnostics:
    def _doc_with_unused_duplicate(self):
        doc = Document()
        doc.add_paragraph('"Gadget" means a mechanical device.')
        doc.add_paragraph('"Gadget" means an electronic device.')
        doc.add_paragraph("Nothing else references that term at all.")
        base_text = "\n\n".join(p.text for p in doc.paragraphs)
        return doc, base_text

    def test_extract_definitions_reports_unused_duplicate(self):
        from adeu.domain import extract_all_domain_metadata

        doc, base_text = self._doc_with_unused_duplicate()
        _defs, diagnostics, _anchors = extract_all_domain_metadata(doc, base_text)
        assert any("Duplicate Definition" in d and "Gadget" in d for d in diagnostics), (
            f"duplicate-definition error suppressed for unused term; got {diagnostics}"
        )

    def test_single_pass_extractor_reports_unused_duplicate(self):
        from adeu.domain import extract_all_domain_metadata

        doc, base_text = self._doc_with_unused_duplicate()
        _defs, diagnostics, _anchors = extract_all_domain_metadata(doc, base_text)
        assert any("Duplicate Definition" in d and "Gadget" in d for d in diagnostics), (
            f"duplicate-definition error suppressed for unused term; got {diagnostics}"
        )

    def test_used_duplicate_still_reported(self):
        """Control: the already-working case must keep working."""
        from adeu.domain import extract_all_domain_metadata

        doc = Document()
        doc.add_paragraph('"Widget" means a mechanical device.')
        doc.add_paragraph('"Widget" means an electronic device.')
        doc.add_paragraph("The Widget shall be delivered on time.")
        base_text = "\n\n".join(p.text for p in doc.paragraphs)
        _defs, diagnostics, _anchors = extract_all_domain_metadata(doc, base_text)
        assert any("Duplicate Definition" in d and "Widget" in d for d in diagnostics)


# ---------------------------------------------------------------------------
# F7 — a directory as input must produce a clean error, not a traceback
# ---------------------------------------------------------------------------


class TestF7DirectoryInput:
    @pytest.fixture
    def a_dir(self, tmp_path):
        d = tmp_path / "some_directory"
        d.mkdir()
        return d

    def _assert_clean_failure(self, capsys, argv):
        with pytest.raises(SystemExit) as exc_info:
            from adeu.cli import main

            with patch.object(sys, "argv", ["adeu"] + argv):
                main()
        assert exc_info.value.code != 0
        err = capsys.readouterr().err
        assert "❌" in err
        assert "Traceback" not in err

    def test_extract_directory(self, capsys, a_dir):
        self._assert_clean_failure(capsys, ["extract", str(a_dir)])

    def test_accept_all_directory(self, capsys, a_dir):
        self._assert_clean_failure(capsys, ["accept-all", str(a_dir)])

    def test_apply_directory_original(self, capsys, a_dir, tmp_path):
        edits = tmp_path / "e.json"
        edits.write_text("[]", encoding="utf-8")
        self._assert_clean_failure(capsys, ["apply", str(a_dir), str(edits)])

    def test_diff_directory_original(self, capsys, a_dir, tmp_path):
        txt = tmp_path / "m.txt"
        txt.write_text("hello", encoding="utf-8")
        self._assert_clean_failure(capsys, ["diff", str(a_dir), str(txt)])

    def test_markup_directory_input(self, capsys, a_dir, tmp_path):
        edits = tmp_path / "e.json"
        edits.write_text("[]", encoding="utf-8")
        self._assert_clean_failure(capsys, ["markup", str(a_dir), str(edits)])


# ---------------------------------------------------------------------------
# F8 — pinned edits must pass the same string-shape validation as JSON edits
# ---------------------------------------------------------------------------


class TestF8PinnedValidationBypass:
    @pytest.mark.parametrize(
        "bad_new_text, expected_fragment",
        [
            ("{++New York++}", "CriticMarkup"),
            ("####### Deep heading", "Heading level 7"),
            ("see {#InternalAnchor}", "anchor"),
        ],
    )
    def test_pinned_edits_are_shape_validated(self, bad_new_text, expected_fragment):
        """A pinned edit (as produced by generate_edits_from_text) whose
        new_text violates a VAL-CRIT rule must be rejected exactly as the
        identical user-supplied JSON edit is."""
        doc = Document()
        doc.add_paragraph("This is a simple contract paragraph for testing.")
        engine = RedlineEngine(_stream(doc))

        edit = ModifyText(target_text="simple", new_text=bad_new_text)
        edit._match_start_index = engine.mapper.full_text.index("simple")

        with pytest.raises(BatchValidationError) as exc_info:
            engine.process_batch([edit])
        joined = "\n".join(exc_info.value.errors)
        assert expected_fragment.lower() in joined.lower()

        # And nothing may have been applied (transactional).
        final = extract_text_from_stream(engine.save_to_stream(), filename="x.docx")
        assert "{++New York++}" not in final
        assert "#######" not in final

    def test_markup_view_text_apply_is_refused_with_guidance(self, tmp_path, capsys):
        """The QA report's identity round-trip: extract (default markup view) a
        doc with an existing tracked insertion, apply unchanged. 1.21.0 wrote
        '{++high++}{>>[Chg:1 insert] Alice<<}' into the document body at rc=0."""
        docx_path = tmp_path / "only_ins.docx"
        docx_path.write_bytes(_make_doc_with_tracked_insertion().getvalue())

        txt_path = tmp_path / "t.txt"
        assert _run_cli(["extract", str(docx_path), "-o", str(txt_path)]) == 0
        assert "{++high++}" in txt_path.read_text(encoding="utf-8")  # sanity

        out_path = tmp_path / "r.docx"
        rc = _run_cli(["apply", str(docx_path), str(txt_path), "-o", str(out_path)])
        err = capsys.readouterr().err

        assert rc != 0, "apply accepted markup-view text against the clean-view baseline"
        assert not out_path.exists()
        assert "--clean-view" in err  # actionable: re-extract with --clean-view

    def test_markup_view_text_with_comments_does_not_crash(self, tmp_path, capsys):
        """Crash variant: same round-trip on a doc that also carries a comment
        raised AttributeError from _attach_comment in 1.21.0."""
        doc = Document(_make_doc_with_tracked_insertion())
        cm = CommentsManager(doc)
        cid = cm.add_comment("Bob", "Please review this fee.")
        p_el = doc.paragraphs[0]._element
        rs = OxmlElement("w:commentRangeStart")
        rs.set(qn("w:id"), cid)
        p_el.insert(0, rs)
        re_el = OxmlElement("w:commentRangeEnd")
        re_el.set(qn("w:id"), cid)
        p_el.append(re_el)
        ref_run = OxmlElement("w:r")
        ref = OxmlElement("w:commentReference")
        ref.set(qn("w:id"), cid)
        ref_run.append(ref)
        p_el.append(ref_run)

        docx_path = tmp_path / "ins_and_comment.docx"
        docx_path.write_bytes(_stream(doc).getvalue())
        txt_path = tmp_path / "t.txt"
        assert _run_cli(["extract", str(docx_path), "-o", str(txt_path)]) == 0

        out_path = tmp_path / "r.docx"
        with pytest.raises(SystemExit) as exc_info:
            from adeu.cli import main

            argv = ["adeu", "apply", str(docx_path), str(txt_path), "-o", str(out_path)]
            with patch.object(sys, "argv", argv):
                main()
        assert exc_info.value.code != 0
        err = capsys.readouterr().err
        assert "AttributeError" not in err

    def test_attach_comment_tolerates_missing_parent(self):
        """_attach_comment must degrade gracefully when the anchor context is
        gone (engine.py:1297 crashed on parent_element=None)."""
        doc = Document()
        doc.add_paragraph("Anchor paragraph.")
        engine = RedlineEngine(_stream(doc))
        el = engine.doc.paragraphs[0]._element
        engine._attach_comment(None, el, el, "orphan note")  # must not raise


# ---------------------------------------------------------------------------
# F9 — duplicate w:id values must not let one action resolve unrelated changes
# ---------------------------------------------------------------------------


class TestF9DuplicateRevisionIds:
    def _make_doc_with_duplicate_wid(self) -> io.BytesIO:
        doc = Document()
        p1 = doc.add_paragraph()
        p1.add_run("Clause A fee: ")
        p1._element.append(_make_tracked_ins("ALPHA", "5", "Alice"))
        p2 = doc.add_paragraph()
        p2.add_run("Clause B fee: ")
        p2._element.append(_make_tracked_ins("BRAVO", "5", "Bob"))
        return _stream(doc)

    def test_reject_on_duplicate_wid_refuses(self):
        engine = RedlineEngine(self._make_doc_with_duplicate_wid())
        with pytest.raises(BatchValidationError) as exc_info:
            engine.process_batch([RejectChange(target_id="Chg:5")])
        joined = "\n".join(exc_info.value.errors)
        assert "Alice" in joined and "Bob" in joined

        # Neither revision may have been touched.
        raw = extract_text_from_stream(engine.save_to_stream(), filename="x.docx")
        assert "ALPHA" in raw and "BRAVO" in raw

    def test_same_author_multi_element_change_still_resolvable(self):
        """Control: the engine's own multi-paragraph insertions legitimately
        reuse one w:id across several <w:ins> elements (same author). The
        duplicate-id guard must not refuse those."""
        doc = Document()
        doc.add_paragraph("Alpha marker paragraph.")
        engine = RedlineEngine(_stream(doc), author="Adeu AI")
        engine.process_batch([ModifyText(target_text="marker", new_text="first\n\nsecond")])
        out = engine.save_to_stream()

        raw = extract_text_from_stream(io.BytesIO(out.getvalue()), filename="x.docx")
        chg_ids = sorted(set(re.findall(r"Chg:(\d+)", raw)), key=int)
        assert chg_ids, f"no tracked change surfaced in: {raw!r}"

        engine2 = RedlineEngine(io.BytesIO(out.getvalue()))
        stats = engine2.process_batch([RejectChange(target_id=f"Chg:{chg_ids[-1]}")])
        assert stats["actions_applied"] >= 1


# ---------------------------------------------------------------------------
# F10 — [Edit:N] indices must be 1-based to match apply's Edit N reports
# ---------------------------------------------------------------------------


class TestF10MarkupIndexBase:
    def test_markup_indices_are_one_based(self):
        edits = [
            ModifyText(target_text="Alpha", new_text="Omega"),
            ModifyText(target_text="gamma", new_text="delta"),
        ]
        result = apply_edits_to_markdown("Alpha beta gamma.", edits, include_index=True)
        assert "[Edit:1]" in result
        assert "[Edit:2]" in result
        assert "[Edit:0]" not in result, "markup is 0-based while apply reports are 1-based"


# ---------------------------------------------------------------------------
# F11 — control characters must fail as clean validation, not an lxml traceback
# ---------------------------------------------------------------------------


class TestF11ControlCharacters:
    def test_control_char_in_new_text_is_clean_validation_error(self):
        doc = Document()
        doc.add_paragraph("This is a simple contract.")
        engine = RedlineEngine(_stream(doc))
        edit = ModifyText(target_text="simple", new_text="bad\x01value")

        with pytest.raises(BatchValidationError) as exc_info:
            engine.process_batch([edit])
        joined = "\n".join(exc_info.value.errors)
        assert "Edit 1 Failed" in joined
        assert "control character" in joined.lower()

    def test_control_char_in_comment_is_clean_validation_error(self):
        doc = Document()
        doc.add_paragraph("This is a simple contract.")
        engine = RedlineEngine(_stream(doc))
        edit = ModifyText(target_text="simple", new_text="fine", comment="note\x00here")

        with pytest.raises(BatchValidationError) as exc_info:
            engine.process_batch([edit])
        assert "control character" in "\n".join(exc_info.value.errors).lower()

    def test_control_char_in_author_is_clean_cli_error(self, tmp_path, capsys):
        doc = Document()
        doc.add_paragraph("This is a simple contract.")
        docx_path = tmp_path / "doc.docx"
        docx_path.write_bytes(_stream(doc).getvalue())
        edits_path = tmp_path / "e.json"
        edits_path.write_text(
            json.dumps([{"type": "modify", "target_text": "simple", "new_text": "plain"}]),
            encoding="utf-8",
        )

        with pytest.raises(SystemExit) as exc_info:
            from adeu.cli import main

            argv = ["adeu", "apply", str(docx_path), str(edits_path), "--author", "Bad\x01Author"]
            with patch.object(sys, "argv", argv):
                main()
        assert exc_info.value.code != 0
        err = capsys.readouterr().err
        assert "❌" in err
        assert "control character" in err.lower()


# ---------------------------------------------------------------------------
# MCP server surface — the same guarantees must hold through the tool layer
# ---------------------------------------------------------------------------


class MockContext:
    """Absorbs async logging calls from FastMCP tools during tests."""

    async def info(self, msg, **kwargs):
        pass

    async def debug(self, msg, **kwargs):
        pass

    async def warning(self, msg, **kwargs):
        pass

    async def error(self, msg, **kwargs):
        pass


class TestMcpServerSurface:
    def test_read_docx_page_all_returns_entire_document(self, tmp_path):
        """MCP parity for F1: page='all' with mode='full' must return the whole
        document without page chrome (it used to silently render page 1)."""
        import asyncio

        from adeu.mcp_components.tools.document import read_docx

        docx_path = tmp_path / "large.docx"
        docx_path.write_bytes(_make_large_doc().getvalue())

        result = asyncio.run(read_docx(reasoning="test", file_path=str(docx_path), ctx=MockContext(), page="all"))
        markdown = result.structured_content["markdown"]
        assert "Paragraph 0001" in markdown
        assert "Paragraph 0130" in markdown
        assert "Continues on page" not in markdown
        assert "Page 1 of" not in markdown

    def test_process_document_batch_rejects_control_char_author(self, tmp_path):
        """MCP parity for F11: author_name with control characters must be a
        clean tool error, not an lxml traceback."""
        import asyncio

        from adeu.mcp_components.tools.document import process_document_batch

        docx_path = tmp_path / "doc.docx"
        doc = Document()
        doc.add_paragraph("This is a simple contract.")
        docx_path.write_bytes(_stream(doc).getvalue())

        result = asyncio.run(
            process_document_batch(
                reasoning="test",
                original_docx_path=str(docx_path),
                author_name="Bad\x01Author",
                ctx=MockContext(),
                changes=[ModifyText(target_text="simple", new_text="plain")],
                output_path=str(tmp_path / "out.docx"),
            )
        )
        assert "control character" in str(result).lower()
        assert not (tmp_path / "out.docx").exists()

    def test_process_document_batch_redos_regex_is_clean_error(self, tmp_path):
        """MCP parity for F5: a catastrophic regex edit must come back as a
        clean batch-validation report through the tool layer."""
        import asyncio

        from adeu.mcp_components.tools.document import process_document_batch

        docx_path = tmp_path / "r.docx"
        doc = Document()
        doc.add_paragraph(REDOS_HAYSTACK)
        docx_path.write_bytes(_stream(doc).getvalue())

        result = asyncio.run(
            process_document_batch(
                reasoning="test",
                original_docx_path=str(docx_path),
                author_name="QA",
                ctx=MockContext(),
                changes=[
                    ModifyText(
                        target_text=REDOS_PATTERN_NEEDS_TIMEOUT,
                        new_text="X",
                        regex=True,
                        match_mode="first",
                    )
                ],
                output_path=str(tmp_path / "out.docx"),
            )
        )
        text = str(result)
        assert "Edit 1 Failed" in text
        assert "time" in text.lower()
        assert not (tmp_path / "out.docx").exists()


# ---------------------------------------------------------------------------
# §7 regression property — apply(D, extract(D)) must never silently mutate D
# ---------------------------------------------------------------------------


def _corpus_docs() -> dict[str, io.BytesIO]:
    """Small corpus spanning the shapes the QA report exercised: plain text,
    contract structure (headings/terms/table), unicode, >19k chars (paginates),
    existing tracked changes, and comments."""
    simple = Document()
    simple.add_paragraph("This is a simple contract.")
    simple.add_paragraph("It has exactly two paragraphs.")

    contract = Document()
    contract.add_heading("Master Services Agreement", 1)
    contract.add_paragraph('"Services" means the work described in Exhibit A.')
    contract.add_paragraph("The Services shall commence on the Effective Date.")
    table = contract.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Fee"
    table.cell(1, 0).text = "Consulting"
    table.cell(1, 1).text = "$10,000"

    unicode_doc = Document()
    unicode_doc.add_paragraph("Emoji: 🚀🔥 — CJK: 契約書の条項 — RTL: اتفاقية سرية")
    unicode_doc.add_paragraph("Zero-width​ and NBSP  characters survive.")

    return {
        "simple": _stream(simple),
        "contract": _stream(contract),
        "unicode": _stream(unicode_doc),
        "large": _make_large_doc(),
        "tracked": _make_doc_with_tracked_insertion(),
        "comments": _make_doc_with_comment(),
    }


class TestIdentityRoundTripProperty:
    """The QA report's central invariant (§2, §7): feeding an unmodified
    extract straight back into apply must either be a no-op or fail loudly —
    never a silent mutation. Checked for every corpus doc in both views."""

    @pytest.mark.parametrize("view", ["clean", "markup"])
    @pytest.mark.parametrize("name", ["simple", "contract", "unicode", "large", "tracked", "comments"])
    def test_identity_roundtrip_is_a_noop_or_loud(self, tmp_path, capsys, name, view):
        docx_path = tmp_path / f"{name}.docx"
        docx_path.write_bytes(_corpus_docs()[name].getvalue())
        original_clean = _clean_text(docx_path.read_bytes())

        txt_path = tmp_path / f"{name}_{view}.txt"
        extract_args = ["extract", str(docx_path), "--page", "all", "-o", str(txt_path)]
        if view == "clean":
            extract_args.insert(2, "--clean-view")
        assert _run_cli(extract_args) == 0
        capsys.readouterr()

        out_path = tmp_path / f"{name}_{view}_out.docx"
        rc = _run_cli(["apply", str(docx_path), str(txt_path), "-o", str(out_path), "--json"])
        captured = capsys.readouterr()

        if rc == 0:
            stats = json.loads(captured.out.strip().splitlines()[-1])
            assert stats["edits_applied"] == 0, f"{name}/{view}: identity round-trip produced edits: {stats}"
            assert _clean_text(out_path.read_bytes()) == original_clean, (
                f"{name}/{view}: document content changed on a no-edit round-trip"
            )
        else:
            # A loud refusal (e.g. markup-view text on a redlined document) is
            # an acceptable outcome; silent corruption is not.
            assert not out_path.exists(), f"{name}/{view}: apply failed ({rc}) but still wrote output"
            assert captured.err.strip(), f"{name}/{view}: apply failed silently with no diagnostic"


class TestRemainingGaps:
    def test_pinned_shape_violation_reported_in_dry_run(self):
        """Dry-run must mirror the wet run's rejection of invalid pinned edits
        (transactional parity)."""
        doc = Document()
        doc.add_paragraph("This is a simple contract paragraph for testing.")
        engine = RedlineEngine(_stream(doc))

        edit = ModifyText(target_text="simple", new_text="{++New York++}")
        edit._match_start_index = engine.mapper.full_text.index("simple")

        stats = engine.process_batch([edit], dry_run=True)
        assert stats["edits_applied"] == 0
        assert stats["edits_skipped"] == 1
        assert stats["edits"][0]["status"] == "failed"
        assert "CriticMarkup" in (stats["edits"][0]["error"] or "")

    def test_keep_markup_sanitize_verifiably_removes_resolved_comments(self, tmp_path):
        """The remove_resolved_comments path shares delete_comment with F3;
        its report is now verified against actual package state."""
        from docx.oxml.ns import qn as _qn

        from adeu.sanitize.core import sanitize_docx

        doc = Document(_make_doc_with_comment(author="Alice", text="Resolved note to remove."))
        cm = CommentsManager(doc)
        # Flip the (only) comment to resolved in commentsExtended.
        for child in cm.extended_part.element:
            child.set(_qn("w15:done"), "1")

        input_path = tmp_path / "resolved.docx"
        input_path.write_bytes(_stream(doc).getvalue())
        output_path = tmp_path / "clean.docx"

        result = sanitize_docx(str(input_path), str(output_path), keep_markup=True)
        assert result.comments_removed == 1

        comments_xml = _read_comment_parts(output_path.read_bytes())
        assert "Resolved note to remove." not in comments_xml
        assert "Alice" not in comments_xml

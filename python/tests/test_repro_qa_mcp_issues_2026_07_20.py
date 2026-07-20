import io
import re

from docx import Document

from adeu.ingest import _extract_text_from_doc, extract_text_from_stream
from adeu.mcp_components._response_builders import build_search_response
from adeu.models import DeleteTableRow, InsertTableRow, ModifyText
from adeu.outline import extract_outline
from adeu.pagination import paginate
from adeu.redline.engine import RedlineEngine
from adeu.sanitize.core import sanitize_docx


class TestReproQaMcpIssues:
    def test_tc1_2_heading_prefix_leak(self):
        """
        TC 1.2: multi-paragraph insert with heading does not leak ### prefix to other paragraphs
        """
        doc = Document()
        doc.add_paragraph("Replace me.")
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        engine = RedlineEngine(stream)
        edit = ModifyText(
            target_text="Replace me.", new_text="### Title\n\nBody with bold and italic.\n\nSecond paragraph."
        )
        engine.process_batch([edit])
        engine.accept_all_revisions()

        res_stream = engine.save_to_stream()
        clean_text = extract_text_from_stream(res_stream, clean_view=True)

        # Subsequent paragraphs must not have literal "###" prepended
        assert "Body with bold and italic." in clean_text
        assert "Second paragraph" in clean_text
        assert "### Body" not in clean_text
        assert "### Second" not in clean_text

    def test_tc2_1_bold_table_cell_classification(self):
        """
        TC 2.1: bold table cell with uppercase/numeric text is NOT classified as a heading
        """
        doc = Document()
        tbl = doc.add_table(rows=1, cols=1)
        cell = tbl.cell(0, 0)
        p = cell.paragraphs[0]
        p.text = "TOTAL $ 88,136"
        r = p.runs[0]
        r.font.bold = True

        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        doc_obj = Document(stream)
        body = _extract_text_from_doc(doc_obj)
        pagination_result = paginate(body)
        nodes = extract_outline(
            doc_obj,
            body,
            pagination_result.body_pages,
            pagination_result.body_page_offsets,
        )

        # The table cell must not be present in the outline nodes
        fake_heading = any("TOTAL" in n.text for n in nodes)
        assert not fake_heading

    def test_tc3_1_insert_row_anchor_only(self):
        """
        TC 3.1: insert_row succeeds when targeted with a {#cell:paraId} anchor alone
        """
        doc = Document()
        tbl = doc.add_table(rows=1, cols=2)
        tbl.cell(0, 0).paragraphs[0].text = "A1"
        tbl.cell(0, 1).paragraphs[0].text = "A2"

        # Explicitly set paraId to the paragraph of cell(0,0)
        p_el = tbl.cell(0, 0).paragraphs[0]._element
        p_el.set("{http://schemas.microsoft.com/office/word/2010/wordml}paraId", "DEADBEEF")

        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        engine = RedlineEngine(stream)
        edit = InsertTableRow(target_text="{#cell:DEADBEEF}", cells=["B1", "B2"], position="below")
        engine.process_batch([edit])
        assert len(engine.save_to_stream().getvalue()) > 0

    def test_tc3_2_delete_row_anchor_only(self):
        """
        TC 3.2: delete_row succeeds when targeted with a {#cell:paraId} anchor alone
        """
        doc = Document()
        tbl = doc.add_table(rows=2, cols=2)
        tbl.cell(0, 0).paragraphs[0].text = "A1"
        tbl.cell(0, 1).paragraphs[0].text = "A2"
        tbl.cell(1, 0).paragraphs[0].text = "B1"
        tbl.cell(1, 1).paragraphs[0].text = "B2"

        p_el = tbl.cell(0, 0).paragraphs[0]._element
        p_el.set("{http://schemas.microsoft.com/office/word/2010/wordml}paraId", "DEADBEEF")

        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        engine = RedlineEngine(stream)
        edit = DeleteTableRow(target_text="{#cell:DEADBEEF}")
        engine.process_batch([edit])
        assert len(engine.save_to_stream().getvalue()) > 0

    def test_tc4_1_unbounded_search_results(self):
        """
        TC 4.1: search results are capped to at most 20 matches
        """
        body = "\n\n".join(["The quick brown fox jumps over the lazy dog."] * 50)
        res = build_search_response(
            text=body,
            search_query="fox",
            search_regex=False,
            search_case_sensitive=True,
            page=None,
            file_path="doc.docx",
        )
        md = res.content
        assert isinstance(md, str)

        matches = re.findall(r"### Match \d+", md)
        assert len(matches) <= 20

    def test_tc5_1_empty_cell_fill_spacing(self):
        """
        TC 5.1: modifying cell anchor next to a label preserves or inserts a separator space
        """
        doc = Document()
        tbl = doc.add_table(rows=1, cols=1)
        p = tbl.cell(0, 0).paragraphs[0]
        p.text = "Nimi"
        p_el = p._element
        p_el.set("{http://schemas.microsoft.com/office/word/2010/wordml}paraId", "DEADBEEF")

        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        engine = RedlineEngine(stream)
        edit = ModifyText(target_text="{#cell:DEADBEEF}", new_text="Testi Testinen")
        engine.process_batch([edit])
        engine.accept_all_revisions()

        clean_text = extract_text_from_stream(engine.save_to_stream(), clean_view=True)
        # It should contain a separator (e.g. "Nimi Testi Testinen"), not "NimiTesti Testinen"
        assert "NimiTesti Testinen" not in clean_text

    def test_tc7_1_keep_markup_comments_stripped_label(self, tmp_path):
        """
        TC 7.1: keep-markup with open comments does NOT label them as COMMENTS (stripped)
        """
        doc = Document()
        doc.add_paragraph("This is a normal paragraph.")
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        engine = RedlineEngine(stream)
        edit = ModifyText(
            target_text="This is a normal paragraph.",
            new_text="This is a normal paragraph.",
            comment="This is a comment",
        )
        engine.process_batch([edit])

        input_path = tmp_path / "test.docx"
        output_path = tmp_path / "test_out.docx"
        with open(input_path, "wb") as f:
            f.write(engine.save_to_stream().getvalue())

        res = sanitize_docx(input_path=str(input_path), output_path=str(output_path), keep_markup=True)

        assert "COMMENTS (stripped)" not in res.report_text

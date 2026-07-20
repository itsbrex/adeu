"""
Sanitize orchestrator.

Coordinates the transform pipeline based on mode (full / keep-markup / baseline)
and produces the sanitized DOCX + report.
"""

import enum
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Optional

import structlog
from docx import Document

from adeu.diff import generate_structured_edits
from adeu.ingest import _extract_text_from_doc
from adeu.redline.comments import CommentsManager
from adeu.redline.engine import BatchValidationError, RedlineEngine
from adeu.sanitize import transforms
from adeu.sanitize.report import SanitizeReport
from adeu.utils.docx import strip_bom_from_docx_bytes

logger = structlog.get_logger(__name__)


class SanitizeMode(enum.Enum):
    FULL = "full"
    KEEP_MARKUP = "keep-markup"
    BASELINE = "baseline"


@dataclass
class SanitizeResult:
    """Structured result returned by sanitize_docx."""

    output_path: str
    status: str  # "clean", "clean_with_warnings", "blocked"
    tracked_changes_found: int = 0
    tracked_changes_accepted: int = 0
    comments_removed: int = 0
    comments_kept: int = 0
    metadata_stripped: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    report_text: str = ""


class SanitizeError(Exception):
    """Raised when sanitization is blocked (e.g., unresolved track changes)."""

    pass


def _atomic_write(path: Path, payload: bytes) -> None:
    """
    Stages to a same-directory temporary file and os.replace()s it into
    place: a failed or interrupted write never truncates or corrupts an
    existing file at `path`.

    Missing parent directories are created, and any remaining filesystem
    failure surfaces as a SanitizeError naming the OUTPUT path — never the
    raw `Errno 2 ... '.sanitize.docx.<random>.tmp'` the temporary file used
    to leak (QA 2026-07-19 v8 F-09/F-13).
    """
    import os
    import tempfile

    tmp_path: Optional[Path] = None
    try:
        if path.parent and str(path.parent) not in ("", "."):
            path.parent.mkdir(parents=True, exist_ok=True)
        fd, tmp_name = tempfile.mkstemp(prefix=f".{path.name}.", suffix=".tmp", dir=str(path.parent or "."))
        tmp_path = Path(tmp_name)
        with os.fdopen(fd, "wb") as f:
            f.write(payload)
        os.replace(tmp_path, path)
        tmp_path = None
    except OSError as e:
        raise SanitizeError(f"❌ Could not write output file '{path}': {e.strerror or e}") from e
    finally:
        if tmp_path is not None:
            try:
                tmp_path.unlink()
            except OSError:
                pass


def sanitize_docx(
    input_path: str,
    output_path: Optional[str] = None,
    *,
    keep_markup: bool = False,
    baseline_path: Optional[str] = None,
    author: Optional[str] = None,
    accept_all: bool = False,
    allow_low_similarity_baseline: bool = False,
) -> SanitizeResult:
    """
    Sanitize a DOCX file.
    """
    input_p = Path(input_path)
    if not input_p.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    if baseline_path and not Path(baseline_path).exists():
        raise FileNotFoundError(f"Baseline file not found: {baseline_path}")

    # Determine mode
    if baseline_path:
        mode = SanitizeMode.BASELINE
    elif keep_markup:
        mode = SanitizeMode.KEEP_MARKUP
    else:
        mode = SanitizeMode.FULL

    # Default output path
    if not output_path:
        output_path = str(input_p.parent / f"{input_p.stem}_sanitized{input_p.suffix}")

    report = SanitizeReport(
        filename=input_p.name,
        mode=mode.value,
        author=author,
    )

    # Load document
    with open(input_p, "rb") as f:
        sanitized_bytes = strip_bom_from_docx_bytes(f.read())
        doc = Document(BytesIO(sanitized_bytes))

    # --- Mode-specific logic ---

    if mode == SanitizeMode.FULL:
        _sanitize_full(doc, report, accept_all=accept_all)
    elif mode == SanitizeMode.KEEP_MARKUP:
        _sanitize_keep_markup(doc, report, author=author)
    elif mode == SanitizeMode.BASELINE:
        assert baseline_path is not None
        try:
            doc = _sanitize_baseline(
                doc,
                input_path,
                baseline_path,
                report,
                author=author,
                allow_low_similarity=allow_low_similarity_baseline,
            )
        except SanitizeError:
            raise
        except KeyError as e:
            # A relationship/key lookup failure means the documents are
            # structurally incompatible — surface guidance, never a raw
            # `Error: 'rId9'` (QA 2026-07-18 H3).
            raise SanitizeError(
                f"❌ Baseline recomputation failed: the documents reference incompatible "
                f"internal structures (missing key {e}). Verify that --baseline points to an "
                "earlier version of THIS document (same headers/footers, links and images), "
                "not an unrelated file."
            ) from e

    if report.status == "blocked":
        report_text = report.render()
        raise SanitizeError(report_text)

    # --- Common transforms (applied in all modes) ---
    _apply_common_transforms(doc, report)

    # --- Author replacement and date normalization ---
    if mode in (SanitizeMode.KEEP_MARKUP, SanitizeMode.BASELINE):
        if author:
            report.add_transform_lines(transforms.replace_comment_authors(doc, author))
            report.add_transform_lines(transforms.replace_change_authors(doc, author))
        # Always normalize change AND comment dates on outbound docs —
        # prevents counterparty from inferring when edits were made. Retained
        # comments carry the same signal in word/comments.xml as tracked
        # changes do in the body (QA 2026-07-19 F-09).
        report.add_transform_lines(transforms.normalize_change_dates(doc))
        report.add_transform_lines(transforms.normalize_comment_dates(doc))

    # --- Save (verify BEFORE anything reaches disk; write atomically) ---
    output = BytesIO()
    doc.save(output)
    _verify_sanitized_package(output.getvalue())
    _atomic_write(Path(output_path), output.getvalue())

    # Finalize report
    if report.warnings:
        report.status = "clean_with_warnings"
    report_text = report.render()
    report.status = report.status  # ensure set

    logger.info("Sanitization complete", output_path=output_path, status=report.status)

    return SanitizeResult(
        output_path=output_path,
        status=report.status,
        tracked_changes_found=report.tracked_changes_found,
        tracked_changes_accepted=report.tracked_changes_accepted,
        comments_removed=report.comments_removed,
        comments_kept=report.comments_kept,
        metadata_stripped=[line for line in report.metadata_lines],
        warnings=report.warnings,
        report_text=report_text,
    )


def _sanitize_full(doc, report: SanitizeReport, *, accept_all: bool):
    """Full sanitize: strip everything."""
    # Check for unresolved track changes
    ins_count, del_count, fmt_count = transforms.count_tracked_changes(doc)
    total = ins_count + del_count + fmt_count
    report.tracked_changes_found = total

    if total > 0 and not accept_all:
        report.status = "blocked"
        report.blocked_reason = (
            f"Document contains {total} unresolved tracked changes "
            f"({ins_count} insertions, {del_count} deletions, {fmt_count} formatting). "
            f"Review in Word first, or use --accept-all."
        )
        return

    # Accept all tracked changes
    if total > 0:
        # VAL-OBS-NEW-9: Warn if there are multiple authors to prevent silent smuggle
        authors = transforms.get_track_change_authors(doc)
        if len(authors) > 1:
            report.warnings.append(
                f"Multiple authors detected in tracked changes: {', '.join(sorted(authors))}. "
                f"Review per-change list before sending."
            )

        lines = transforms.accept_all_tracked_changes(doc)
        report.tracked_changes_accepted = total
        report.add_transform_lines(lines)

    # Remove all comments
    comments_summary = transforms.get_comments_summary(doc)
    report.comments_removed = comments_summary["total"]
    lines = transforms.remove_all_comments(doc)
    report.add_transform_lines(lines)


def _sanitize_keep_markup(doc, report: SanitizeReport, *, author: Optional[str]):
    """Keep existing track changes and open comments, strip the rest."""
    # Count what's there
    ins_count, del_count, fmt_count = transforms.count_tracked_changes(doc)
    total_changes = ins_count + del_count + fmt_count
    report.tracked_changes_found = total_changes
    report.tracked_changes_kept = total_changes

    # Warn if no markup found
    comments_summary = transforms.get_comments_summary(doc)
    if total_changes == 0 and comments_summary["total"] == 0:
        report.warnings.append(
            "Document contains no tracked changes or comments. "
            "Output will be identical to a full sanitize. "
            "If you edited without Track Changes, use --baseline to reconstruct the redline."
        )

    # Remove resolved comments, keep open
    lines = transforms.remove_resolved_comments(doc)
    report.comments_removed = comments_summary["resolved"]
    report.comments_kept = comments_summary["open"]
    report.add_transform_lines(lines)

    # Collect kept comment info for report
    remaining = transforms.get_comments_summary(doc)
    for c in remaining["comments"]:
        if not c["resolved"]:
            report.kept_comment_lines.append(f'"{transforms._truncate(c["text"], 60)}" ({c["author"]})')


# Below this line-level sequence similarity, a baseline almost certainly is
# not an earlier version of the working document — proceeding would emit the
# baseline's text as the "sanitized" output (QA 2026-07-19 v8 F-01).
_BASELINE_MIN_SIMILARITY = 0.5


def _sanitize_baseline(
    doc,
    input_path: str,
    baseline_path: str,
    report: SanitizeReport,
    *,
    author: Optional[str],
    allow_low_similarity: bool = False,
):
    """
    Recompute delta against a baseline document.

    Returns the Document to continue sanitizing/saving: the BASELINE package
    with the working document's changes applied as tracked changes. Working
    on the baseline package keeps every relationship id (hyperlinks, images,
    headers) resolvable; grafting the recomputed body into the working
    document's package leaves dangling rIds and produces files other
    applications refuse to open.
    """
    # Step 1: Extract structured projections from both documents
    with open(input_path, "rb") as f:
        working_stream = BytesIO(strip_bom_from_docx_bytes(f.read()))
    try:
        with open(baseline_path, "rb") as f:
            baseline_stream = BytesIO(strip_bom_from_docx_bytes(f.read()))
        baseline_doc_view = Document(BytesIO(baseline_stream.getvalue()))
    except Exception as e:
        # The generic CLI error handler blames the INPUT document for package
        # failures; when the file that failed to open is the baseline, the
        # error must say so (QA 2026-07-19 F-19).
        raise SanitizeError(
            f"❌ Baseline file '{Path(baseline_path).name}' is not a valid DOCX file ({e}). "
            f"The input document '{Path(input_path).name}' itself is fine — fix or replace "
            "the --baseline argument."
        ) from e

    working_doc = Document(BytesIO(working_stream.getvalue()))

    # The appendix is generated metadata, not document content — diffing it
    # writes phantom "used N times" edits into the output (QA H1/H3).
    working_text, working_struct = _extract_text_from_doc(
        working_doc, clean_view=True, include_appendix=False, return_structure=True
    )
    baseline_text, baseline_struct = _extract_text_from_doc(
        baseline_doc_view, clean_view=True, include_appendix=False, return_structure=True
    )

    # Divergence check: a real sequence similarity over paragraphs. The old
    # positional character comparison reported a one-paragraph insertion at
    # the top of the document as "93% different" (QA H3).
    #
    # Below _BASELINE_MIN_SIMILARITY this is no longer a warning: recomputing
    # against a clearly unrelated baseline REPLACES the working document with
    # the baseline's content while exiting 0 and printing "Result: CLEAN"
    # (QA 2026-07-19 v8 F-01). Block before any diff is computed; the
    # explicit allow_low_similarity override downgrades the block to the
    # warning for the rare legitimate near-rewrite.
    if baseline_text or working_text:
        import difflib

        ratio = difflib.SequenceMatcher(
            None, baseline_text.split("\n"), working_text.split("\n"), autojunk=False
        ).ratio()
        if ratio < _BASELINE_MIN_SIMILARITY and len(baseline_text) + len(working_text) <= 100_000:
            # Line-level similarity counts only IDENTICAL lines, so on a
            # small document a legitimate edit to its only paragraph reads
            # as 0% similar. Confirm with a character-level alignment before
            # blocking; bounded to small inputs because SequenceMatcher is
            # quadratic (large related documents share plenty of verbatim
            # lines, so the line-level ratio is reliable there).
            ratio = max(
                ratio,
                difflib.SequenceMatcher(None, baseline_text, working_text, autojunk=False).ratio(),
            )
        difference_pct = round((1 - ratio) * 100)
        if ratio < _BASELINE_MIN_SIMILARITY:
            if not allow_low_similarity:
                report.status = "blocked"
                report.blocked_reason = (
                    f"Baseline and working document share only {round(ratio * 100)}% of their "
                    f"content ({difference_pct}% differs) — '{Path(baseline_path).name}' does not "
                    f"look like an earlier version of '{Path(input_path).name}'. Proceeding would "
                    "replace the document's content with the baseline's. Verify the --baseline "
                    "argument; if this near-total rewrite is intentional, re-run with "
                    "--allow-low-similarity-baseline."
                )
                return doc
            report.warnings.append(
                f"Baseline and working document share only {round(ratio * 100)}% of their content "
                f"({difference_pct}% differs). This may indicate the wrong baseline file was "
                "selected; proceeding because --allow-low-similarity-baseline is set."
            )

    # Step 2: Compute the structured diff (part-aware, table-row-aware).
    edits, diff_warnings = generate_structured_edits(baseline_text, baseline_struct, working_text, working_struct)
    report.warnings.extend(diff_warnings)

    # Step 3: Apply edits to the baseline as tracked changes — sequentially
    # and transactionally, exactly like `adeu apply`.
    baseline_stream.seek(0)
    engine_author = author or "Author"
    engine = RedlineEngine(baseline_stream, author=engine_author)

    if edits:
        try:
            stats = engine.process_batch(list(edits))
        except BatchValidationError as e:
            details = "\n".join(e.errors)
            raise SanitizeError(
                "❌ Baseline recomputation failed — the computed changes could not be applied "
                f"to the baseline:\n{details}"
            ) from e
        # Apply-stage skips return in stats instead of raising. A partial
        # redline silently reverts the skipped working-document changes to
        # baseline text — the same false-success shape as QA M2, so fail
        # closed here too.
        if stats.get("edits_skipped", 0) > 0 or stats.get("actions_skipped", 0) > 0:
            details = "\n".join(stats.get("skipped_details") or [])
            raise SanitizeError(
                "❌ Baseline recomputation failed — "
                f"{stats.get('edits_skipped', 0)} computed change(s) could not be applied to the "
                "baseline, so the output would silently miss part of the working document's "
                f"changes:\n{details}"
            )
        # Non-fatal engine notices (e.g. a comment dropped because it landed
        # in a footer part) must reach the sanitize report, not vanish with
        # the discarded stats.
        for detail in stats.get("skipped_details") or []:
            if detail.startswith("- Warning:"):
                report.warnings.append(detail[2:])

    # Reload from the engine's serialized output so every later transform and
    # the final save operate on a self-consistent package.
    result_doc = Document(engine.save_to_stream())

    # Accurately count the generated track changes XML nodes
    ins_count, del_count, fmt_count = transforms.count_tracked_changes(result_doc)
    report.tracked_changes_found = ins_count + del_count + fmt_count
    report.tracked_changes_kept = report.tracked_changes_found

    # Step 4: Handle comments from working doc (keep those not in baseline)
    # Comments are attached to the working doc's XML, not the baseline.
    # The baseline-reconstructed doc won't have any comments from the working
    # version. This is a known limitation — comments require XML-level
    # transplanting. We note this in the report.
    working_cm = CommentsManager(working_doc)
    working_comments = working_cm.extract_comments_data()

    baseline_cm = CommentsManager(baseline_doc_view)
    baseline_comments = baseline_cm.extract_comments_data()

    # Identify comments unique to working doc
    baseline_texts = {info["text"] for info in baseline_comments.values()}
    new_comments = [
        info for info in working_comments.values() if info["text"] not in baseline_texts and not info.get("resolved")
    ]
    removed_comments = [
        info for info in working_comments.values() if info["text"] in baseline_texts or info.get("resolved")
    ]

    report.comments_kept = len(new_comments)
    report.comments_removed = len(removed_comments)

    for c in new_comments:
        report.kept_comment_lines.append(f'"{transforms._truncate(c["text"], 60)}" ({c["author"]})')
    for c in removed_comments:
        status = "[Resolved]" if c.get("resolved") else "[Baseline]"
        report.removed_comment_lines.append(f'{status} "{transforms._truncate(c["text"], 60)}" ({c["author"]})')

    return result_doc


def _apply_common_transforms(doc, report: SanitizeReport):
    """Apply transforms that run in every mode."""
    report.add_transform_lines(transforms.strip_rsid(doc))
    report.add_transform_lines(transforms.strip_para_ids(doc))
    report.add_transform_lines(transforms.strip_proof_errors(doc))
    report.add_transform_lines(transforms.strip_empty_properties(doc))
    report.add_transform_lines(transforms.strip_hidden_text(doc))
    report.add_transform_lines(transforms.coalesce_runs(doc))
    report.add_transform_lines(transforms.scrub_doc_properties(doc))
    report.add_transform_lines(transforms.scrub_timestamps(doc))
    report.add_transform_lines(transforms.strip_custom_xml(doc))
    report.add_transform_lines(transforms.strip_custom_properties(doc))
    report.add_transform_lines(transforms.strip_document_variables(doc))
    report.add_transform_lines(transforms.strip_image_alt_text(doc))

    # Audit (non-destructive — just warnings)
    hyperlink_warnings = transforms.audit_hyperlinks(doc)
    report.warnings.extend(hyperlink_warnings)
    report.warnings.extend(transforms.detect_watermarks(doc))


# Core-property elements the pipeline claims to scrub; the post-sanitize
# verification re-checks each one in the SAVED bytes.
_DC_NS = "http://purl.org/dc/elements/1.1/"
_CP_NS = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
_VERIFIED_CORE_FIELDS = (
    (f"{{{_DC_NS}}}creator", "author (dc:creator)"),
    (f"{{{_CP_NS}}}lastModifiedBy", "last modified by (cp:lastModifiedBy)"),
    (f"{{{_DC_NS}}}identifier", "identifier (dc:identifier)"),
    (f"{{{_DC_NS}}}description", "description (dc:description)"),
    (f"{{{_CP_NS}}}keywords", "keywords (cp:keywords)"),
    (f"{{{_CP_NS}}}category", "category (cp:category)"),
    (f"{{{_DC_NS}}}subject", "subject (dc:subject)"),
    (f"{{{_CP_NS}}}contentStatus", "content status (cp:contentStatus)"),
    (f"{{{_DC_NS}}}language", "language (dc:language)"),
    (f"{{{_CP_NS}}}version", "version (cp:version)"),
)


def _verify_sanitized_package(output_bytes: bytes) -> None:
    """
    Post-sanitize package scan: before any output is
    written or a report rendered, re-open the SAVED bytes — bypassing every
    python-docx caching layer — and verify the claims the report is about to
    make. A "Result: CLEAN" verdict over a package that still carries custom
    properties or an identifier is worse than no sanitizer at all.
    """
    import zipfile

    from lxml import etree

    problems = []
    with zipfile.ZipFile(BytesIO(output_bytes)) as z:
        names = set(z.namelist())
        if "docProps/custom.xml" in names:
            problems.append("docProps/custom.xml (custom document properties) is still in the package")
        if any(n.startswith("customXml/") for n in names):
            problems.append("customXml/* parts are still in the package")
        if "docProps/core.xml" in names:
            root = etree.fromstring(z.read("docProps/core.xml"))
            for tag, label in _VERIFIED_CORE_FIELDS:
                for el in root.iter(tag):
                    if (el.text or "").strip():
                        problems.append(f"core property {label} still contains a value")
        if "word/settings.xml" in names:
            # Document variables are invisible metadata (QA ADEU-QA-001): a
            # surviving w:docVar means the strip transform silently failed.
            settings_root = etree.fromstring(z.read("word/settings.xml"))
            for el in settings_root.iter():
                if isinstance(el.tag, str) and el.tag.endswith("}docVar"):
                    problems.append("word/settings.xml still contains document variables (w:docVar)")
                    break

    if problems:
        raise SanitizeError(
            "❌ Sanitize integrity check failed — the saved package still contains metadata "
            "this run claims to remove:\n  - "
            + "\n  - ".join(problems)
            + "\nNo output was written. Refusing to report a clean document."
        )
